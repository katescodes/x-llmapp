#!/usr/bin/env python3
"""
分析招标文档内容，评估提取要求的全面性
"""
import sys
from docx import Document
import re

def analyze_document(docx_path):
    """分析docx文档"""
    print(f"\n{'='*100}")
    print(f"📄 招标文档分析")
    print(f"{'='*100}")
    print(f"文档路径: {docx_path}\n")
    
    # 读取文档
    doc = Document(docx_path)
    
    # 提取所有段落
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    print(f"文档总段落数：{len(paragraphs)}")
    
    # 提取所有表格
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables.append(table_data)
    
    print(f"文档总表格数：{len(tables)}")
    
    # 统计表格总行数
    total_table_rows = sum(len(table) for table in tables)
    print(f"表格总行数：{total_table_rows}")
    
    # 关键词统计
    full_text = '\n'.join(paragraphs)
    
    keywords_stats = {
        '废标': len(re.findall(r'废标', full_text)),
        '否决': len(re.findall(r'否决', full_text)),
        '无效': len(re.findall(r'无效', full_text)),
        '取消资格': len(re.findall(r'取消资格', full_text)),
        '不得': len(re.findall(r'不得', full_text)),
        '禁止': len(re.findall(r'禁止', full_text)),
        '必须': len(re.findall(r'必须', full_text)),
        '应当': len(re.findall(r'应当', full_text)),
        '投标人须知': len(re.findall(r'投标人须知', full_text)),
        '评审办法': len(re.findall(r'评审办法', full_text)),
        '评分标准': len(re.findall(r'评分标准', full_text)),
        '资格条件': len(re.findall(r'资格条件', full_text)),
        '技术要求': len(re.findall(r'技术要求', full_text)),
        '采购需求': len(re.findall(r'采购需求', full_text)),
        '▲': len(re.findall(r'▲', full_text)),
        '★': len(re.findall(r'★', full_text)),
        '*': full_text.count('*'),
        '投标保证金': len(re.findall(r'投标保证金', full_text)),
        '最高限价': len(re.findall(r'最高限价', full_text)),
        '控制价': len(re.findall(r'控制价', full_text)),
    }
    
    print(f"\n🔍 关键词频次统计：")
    for keyword, count in sorted(keywords_stats.items(), key=lambda x: x[1], reverse=True):
        if count > 0:
            print(f"  - '{keyword}': {count} 次")
    
    # 查找主要章节
    print(f"\n📚 主要章节结构：")
    chapter_pattern = r'^第[一二三四五六七八九十]+章|^第[0-9]+章|^[一二三四五六七八九十]+、|^[0-9]+\.'
    chapters = []
    
    for i, para in enumerate(paragraphs):
        if re.match(chapter_pattern, para) or any(key in para for key in ['投标人须知', '评审办法', '合同条款', '技术要求', '采购需求', '评分标准']):
            chapters.append((i, para[:60]))  # 只保留前60字符
    
    for idx, (para_idx, chapter) in enumerate(chapters[:30], 1):  # 显示前30个章节
        print(f"  {idx:2d}. [Para {para_idx:4d}] {chapter}")
    
    if len(chapters) > 30:
        print(f"  ... 还有 {len(chapters) - 30} 个章节未显示")
    
    # 分析表格
    print(f"\n📋 表格分析：")
    
    if len(tables) > 0:
        print(f"\n表格详情（前5个）：")
        for i, table in enumerate(tables[:5], 1):
            print(f"\n表格 {i}:")
            print(f"  - 行数：{len(table)}")
            print(f"  - 列数：{len(table[0]) if table else 0}")
            
            # 显示表头（第一行）
            if table:
                header = table[0]
                print(f"  - 表头：{' | '.join(header[:5])}")  # 只显示前5列
                
                # 判断表格类型
                header_text = ' '.join(header).lower()
                if '评分' in header_text or '分值' in header_text:
                    print(f"  - 类型：⭐ 评分标准表")
                elif '参数' in header_text or '技术' in header_text:
                    print(f"  - 类型：⚙️ 技术参数表")
                elif '资格' in header_text or '条件' in header_text:
                    print(f"  - 类型：📜 资格条件表")
                elif '商务' in header_text:
                    print(f"  - 类型：💼 商务要求表")
                else:
                    print(f"  - 类型：📄 其他表格")
                
                # 检查是否有特殊标记
                table_text = '\n'.join([' '.join(row) for row in table])
                special_marks = {'▲': 0, '★': 0, '*': 0}
                for mark in special_marks:
                    special_marks[mark] = table_text.count(mark)
                
                if any(special_marks.values()):
                    print(f"  - 特殊标记：", end='')
                    for mark, count in special_marks.items():
                        if count > 0:
                            print(f"{mark}:{count}次 ", end='')
                    print()
                
                # 显示部分内容（前3行）
                print(f"  - 内容示例（前3行）：")
                for row_idx, row in enumerate(table[1:4], 1):  # 跳过表头
                    print(f"    行{row_idx}: {' | '.join(str(cell)[:30] for cell in row[:3])}")  # 只显示前3列，每列最多30字符
    
    # 估算应提取的要求数量
    print(f"\n{'='*100}")
    print(f"📊 提取数量估算")
    print(f"{'='*100}\n")
    
    # 基于关键词和表格估算
    estimated_veto = keywords_stats['废标'] + keywords_stats['否决'] + keywords_stats['取消资格']
    estimated_veto = max(20, min(estimated_veto * 2, 50))  # 估算废标项数量
    
    estimated_scoring = max(30, min(total_table_rows * 0.3, 80))  # 估算评分项数量（假设30%的表格行是评分项）
    
    estimated_tech = keywords_stats['▲'] + keywords_stats['★'] + max(20, min(total_table_rows * 0.2, 60))  # 估算技术要求
    
    estimated_total = estimated_veto + estimated_scoring + estimated_tech + 40  # 加上其他要求
    
    print(f"基于文档内容估算：")
    print(f"  - 废标项：约 {int(estimated_veto)} 条")
    print(f"  - 评分项：约 {int(estimated_scoring)} 条（含细分项）")
    print(f"  - 技术要求：约 {int(estimated_tech)} 条")
    print(f"  - 其他要求：约 40 条")
    print(f"  - **预估总计：约 {int(estimated_total)} 条**")
    
    print(f"\n说明：")
    print(f"  - 如果实际提取数量远低于预估，说明有大量遗漏")
    print(f"  - 特别注意表格内容，{total_table_rows}行表格数据应逐行提取")
    print(f"  - 注意▲/★标记的条款（文档中共{keywords_stats['▲'] + keywords_stats['★']}处）")
    
    # 提取部分关键内容示例
    print(f"\n{'='*100}")
    print(f"📝 关键内容示例")
    print(f"{'='*100}\n")
    
    # 查找包含"废标"或"否决"的段落
    veto_paragraphs = []
    for i, para in enumerate(paragraphs):
        if '废标' in para or '否决' in para:
            veto_paragraphs.append((i, para))
    
    if veto_paragraphs:
        print(f"包含'废标'/'否决'的段落（前5个）：")
        for idx, (para_idx, para) in enumerate(veto_paragraphs[:5], 1):
            print(f"\n  {idx}. [段落 {para_idx}]")
            print(f"     {para[:200]}..." if len(para) > 200 else f"     {para}")
    
    # 查找包含"评分"的段落
    scoring_paragraphs = []
    for i, para in enumerate(paragraphs):
        if '评分' in para and len(para) > 20:  # 过滤掉太短的段落
            scoring_paragraphs.append((i, para))
    
    if scoring_paragraphs:
        print(f"\n\n包含'评分'的段落（前5个）：")
        for idx, (para_idx, para) in enumerate(scoring_paragraphs[:5], 1):
            print(f"\n  {idx}. [段落 {para_idx}]")
            print(f"     {para[:200]}..." if len(para) > 200 else f"     {para}")
    
    return {
        'paragraphs_count': len(paragraphs),
        'tables_count': len(tables),
        'table_rows_count': total_table_rows,
        'keywords_stats': keywords_stats,
        'estimated_total': int(estimated_total),
        'estimated_veto': int(estimated_veto),
        'estimated_scoring': int(estimated_scoring),
    }


if __name__ == "__main__":
    docx_path = "/aidata/x-llmapp1/data/tender_assets/tp_f379d279606a4ff89a6aa2cfabc0a6c5/tender_a6320484adca479caaa83b71ff1de9de_【GC1818TP】 储能技术公司金坛、刘庄储气库控制系统国产化升级改造工程施工项目20260104.docx"
    
    try:
        result = analyze_document(docx_path)
        
        print(f"\n{'='*100}")
        print(f"✅ 分析完成")
        print(f"{'='*100}\n")
        
        print(f"【重要提示】")
        print(f"  如果V3.4版本提取结果少于 {result['estimated_total']} 条，")
        print(f"  说明还有很大的优化空间！")
        print(f"  特别要检查：")
        print(f"    1. 表格是否逐行提取（文档中有{result['table_rows_count']}行表格）")
        print(f"    2. 废标项是否全部提取（文档中'废标'/'否决'共{result['keywords_stats']['废标'] + result['keywords_stats']['否决']}次）")
        print(f"    3. 评分项是否细化（估算约{result['estimated_scoring']}条）")
        print(f"    4. ▲/★标记的条款是否全部提取（共{result['keywords_stats']['▲'] + result['keywords_stats']['★']}处）")
        
    except Exception as e:
        print(f"\n❌ 分析失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

