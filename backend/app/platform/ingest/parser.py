"""
平台级文档解析器
支持文本、HTML、PDF、DOCX、Excel 和音频文件解析
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from typing import Tuple, Optional

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    title: str
    text: str
    metadata: dict


TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json"}
HTML_EXTS = {".html", ".htm"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xls"}
AUDIO_EXTS = {".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".wav", ".webm", ".ogg", ".flac"}


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="ignore")


def _parse_text(data: bytes) -> Tuple[str, dict]:
    text = _decode_text(data)
    return text, {"chars": len(text)}


def _parse_html(data: bytes) -> Tuple[str, dict, str]:
    html = _decode_text(data)
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else ""
    text = soup.get_text("\n", strip=True)
    return text, {"chars": len(text)}, title


def _parse_pdf(data: bytes) -> Tuple[str, dict]:
    pdf = PdfReader(io.BytesIO(data))
    pages = []
    for page in pdf.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n".join(pages)
    return text, {"pages": len(pdf.pages), "chars": len(text)}


def _parse_docx(data: bytes) -> Tuple[str, dict]:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [para.text for para in doc.paragraphs if para.text]
    text = "\n".join(paragraphs)
    return text, {"paragraphs": len(paragraphs), "chars": len(text)}


def _parse_excel(data: bytes) -> Tuple[str, dict]:
    """
    解析 Excel 文件（.xlsx 或 .xls）
    
    提取所有工作表的文本内容，按行列组织
    
    Args:
        data: Excel 文件二进制数据
        
    Returns:
        (text, metadata) 元组
    """
    try:
        import openpyxl
        from openpyxl.utils.exceptions import InvalidFileException
    except ImportError:
        # 如果没有安装 openpyxl，fallback 到纯文本
        return "", {"error": "openpyxl not installed", "chars": 0}
    
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        total_rows = 0
        total_cells = 0
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            lines.append(f"=== {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                # 过滤空行
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    lines.append(" | ".join(row_values))
                    total_rows += 1
                    total_cells += len(row_values)
            
            lines.append("")  # 工作表之间空行
        
        text = "\n".join(lines)
        return text, {
            "sheets": len(wb.sheetnames),
            "rows": total_rows,
            "cells": total_cells,
            "chars": len(text)
        }
    
    except Exception as e:
        # 解析失败，返回错误信息
        return "", {"error": str(e), "chars": 0}


async def parse_document(
    filename: str,
    data: bytes,
    transcribe_audio_func: Optional[callable] = None,
) -> ParsedDocument:
    """
    解析文档，支持文本、HTML、PDF、DOCX 和音频文件
    
    Args:
        filename: 文件名
        data: 文件二进制数据
        transcribe_audio_func: 可选的音频转文字函数
        
    Returns:
        ParsedDocument 对象
    """
    ext = os.path.splitext(filename)[1].lower()
    title = os.path.splitext(os.path.basename(filename))[0] or filename
    metadata = {"filename": filename, "size": len(data)}

    if ext in TEXT_EXTS or not ext:
        text, meta = _parse_text(data)
        metadata.update(meta)
        return ParsedDocument(title=title, text=text, metadata=metadata)

    if ext in HTML_EXTS:
        text, meta, html_title = _parse_html(data)
        metadata.update(meta)
        return ParsedDocument(title=html_title or title, text=text, metadata=metadata)

    if ext in PDF_EXTS:
        text, meta = _parse_pdf(data)
        metadata.update(meta)
        return ParsedDocument(title=title, text=text, metadata=metadata)

    if ext in DOCX_EXTS:
        text, meta = _parse_docx(data)
        metadata.update(meta)
        return ParsedDocument(title=title, text=text, metadata=metadata)

    if ext in EXCEL_EXTS:
        text, meta = _parse_excel(data)
        metadata.update(meta)
        return ParsedDocument(title=title, text=text, metadata=metadata)

    if ext in AUDIO_EXTS:
        if transcribe_audio_func is None:
            raise ValueError(f"音频文件 {filename} 需要提供 transcribe_audio_func 参数")
        text = await transcribe_audio_func(data, filename)
        metadata.update({
            "chars": len(text),
            "type": "audio_transcription",
            "original_format": ext,
        })
        return ParsedDocument(title=title, text=text, metadata=metadata)

    # fallback to plain text
    text, meta = _parse_text(data)
    metadata.update(meta)
    metadata["note"] = "fallback-text"
    return ParsedDocument(title=title, text=text, metadata=metadata)

