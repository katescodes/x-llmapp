"""
确定性模板解析 + 预览文档生成（PDF 优先）

目标：
- 不做 OCR / AI：页眉页脚图片必须从 docx 关系里直接提取 blob bytes 原样保存
- 解析 section 的 page size/orientation，归类为 A4_PORTRAIT / A4_LANDSCAPE / A3_LANDSCAPE
- 生成一份示范 docx（2-3 页，含目录/Heading1~5/正文），并插入对应 variant 的页眉/页脚图片
- 可选：调用 LibreOffice headless 转 pdf 供前端 iframe 预览
"""

from __future__ import annotations

import hashlib
import os
import re
import subprocess
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm

from app.services.docx_style_utils import guess_heading_level


EMU_PER_PX = 9525  # 1 px @ 96dpi ~= 9525 EMU（近似，用于回显）


def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _safe_mkdir(p: str) -> None:
    os.makedirs(p, exist_ok=True)


def _length_to_mm(x) -> float:
    """python-docx Length -> mm"""
    try:
        # Length is EMU-backed, .mm exists
        return float(x.mm)  # type: ignore[attr-defined]
    except Exception:
        try:
            return float(x)  # best effort
        except Exception:
            return 0.0


def _classify_variant(width_mm: float, height_mm: float) -> Tuple[str, float]:
    """
    根据页面尺寸（mm）分类 variant。
    返回 (variant, confidence)
    """
    w = float(width_mm or 0.0)
    h = float(height_mm or 0.0)
    if w <= 0 or h <= 0:
        return ("DEFAULT", 0.0)

    landscape = w > h
    short = min(w, h)
    long = max(w, h)

    # 容差 6mm（模板里常有边距/微小偏差）
    def _near(a: float, b: float, tol: float = 6.0) -> bool:
        return abs(a - b) <= tol

    if _near(short, 210) and _near(long, 297):
        return ("A4_LANDSCAPE" if landscape else "A4_PORTRAIT", 1.0)
    if _near(short, 297) and _near(long, 420):
        return ("A3_LANDSCAPE" if landscape else "DEFAULT", 1.0)

    # 尺寸不标准：保底 DEFAULT
    return ("DEFAULT", 0.3)


def _clear_container_paragraphs(container) -> None:
    """清空 header/footer 的段落内容（保留空段落结构，避免某些 Word 渲染异常）"""
    try:
        for p in list(getattr(container, "paragraphs", []) or []):
            # 清空 runs（保留段落节点）
            try:
                r_elms = list(p._p)  # type: ignore[attr-defined]
                for child in r_elms:
                    p._p.remove(child)  # type: ignore[attr-defined]
            except Exception:
                pass
    except Exception:
        pass


def _insert_toc_field(paragraph) -> None:
    """
    在指定段落插入 TOC 域（LibreOffice 转 pdf 时通常会展开生成目录）。
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-5" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)  # type: ignore[attr-defined]


def _remove_all_body_content_keep_last_sectpr(doc: Document) -> None:
    """
    清空文档 body 内容但保留最后的 sectPr（这样 styles/parts 仍在，且至少保留一个 section）。
    """
    body = doc._body._element  # type: ignore[attr-defined]
    children = list(body)
    for child in children:
        # 保留最后一个 sectPr（通常在 body 最末）
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _style_exists(doc: Document, style_name: str) -> bool:
    try:
        _ = doc.styles[style_name]
        return True
    except Exception:
        return False


def _pick_heading_style(doc: Document, level: int, preferred: Optional[str] = None) -> Optional[str]:
    """
    尽量选择模板里真实存在的 heading 样式名。
    """
    cands: List[str] = []
    if preferred and isinstance(preferred, str) and preferred.strip():
        cands.append(preferred.strip())
        # LLM 有时会输出 '+标题1' 之类
        cands.append(preferred.strip().lstrip("+").strip())
    # 常见内置样式名
    cands.extend(
        [
            f"Heading {level}",
            f"标题 {level}",
            f"标题{level}",
            f"Heading{level}",
        ]
    )
    for n in cands:
        if n and _style_exists(doc, n):
            return n

    # 兜底：扫描 styles，按 guess_heading_level 判定
    try:
        for st in doc.styles:
            try:
                if st.type != WD_STYLE_TYPE.PARAGRAPH:
                    continue
                lvl = guess_heading_level(getattr(st, "name", None) or "")
                if lvl == level:
                    return getattr(st, "name", None) or None
            except Exception:
                continue
    except Exception:
        pass
    return None


def _pick_body_style(doc: Document, preferred: Optional[str] = None) -> Optional[str]:
    cands: List[str] = []
    if preferred and isinstance(preferred, str) and preferred.strip():
        cands.append(preferred.strip())
        cands.append(preferred.strip().lstrip("+").strip())
    cands.extend(["正文", "Body Text", "Normal"])
    for n in cands:
        if n and _style_exists(doc, n):
            return n
    return "Normal" if _style_exists(doc, "Normal") else None


@dataclass
class ExtractedImage:
    where: str  # header|footer
    variant: str
    content_type: str
    file_name: str
    sha256: str
    blob: bytes
    width_px: Optional[int] = None
    height_px: Optional[int] = None


class DocxTemplateDeterministicParser:
    """
    解析 docx：
    - section 的 page size / orientation / margins
    - header/footer 的图片（原样 blob bytes）
    - 样式体系摘要（heading1~5 是否存在、正文样式候选）
    """

    _NS = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    def parse(self, doc: Document) -> Tuple[Dict[str, Any], List[ExtractedImage]]:
        extracted_images: List[ExtractedImage] = []

        # 1) styles: heading 1..5 presence + body style
        heading_present = {str(i): False for i in range(1, 6)}
        try:
            for st in doc.styles:
                try:
                    if st.type != WD_STYLE_TYPE.PARAGRAPH:
                        continue
                    nm = getattr(st, "name", "") or ""
                    lvl = guess_heading_level(nm)
                    if lvl and 1 <= lvl <= 5:
                        heading_present[str(lvl)] = True
                except Exception:
                    continue
        except Exception:
            pass

        # 2) sections: classify variant + margins + header/footer images
        variants: List[str] = []
        sections: List[Dict[str, Any]] = []
        confidence_min = 1.0

        for idx, sec in enumerate(list(doc.sections) or []):
            w_mm = _length_to_mm(getattr(sec, "page_width", 0))
            h_mm = _length_to_mm(getattr(sec, "page_height", 0))
            variant, conf = _classify_variant(w_mm, h_mm)
            confidence_min = min(confidence_min, conf)
            variants.append(variant)

            sec_info: Dict[str, Any] = {
                "index": idx,
                "variant": variant,
                "page_width_mm": w_mm,
                "page_height_mm": h_mm,
                "orientation": "LANDSCAPE" if w_mm > h_mm else "PORTRAIT",
                "margins_mm": {
                    "top": _length_to_mm(getattr(sec, "top_margin", 0)),
                    "bottom": _length_to_mm(getattr(sec, "bottom_margin", 0)),
                    "left": _length_to_mm(getattr(sec, "left_margin", 0)),
                    "right": _length_to_mm(getattr(sec, "right_margin", 0)),
                },
                "header_distance_mm": _length_to_mm(getattr(sec, "header_distance", 0)),
                "footer_distance_mm": _length_to_mm(getattr(sec, "footer_distance", 0)),
                "link_to_previous": {
                    "header": bool(getattr(getattr(sec, "header", None), "is_linked_to_previous", False)),
                    "footer": bool(getattr(getattr(sec, "footer", None), "is_linked_to_previous", False)),
                },
                "header_images": 0,
                "footer_images": 0,
            }

            # header/footer images
            header = getattr(sec, "header", None)
            footer = getattr(sec, "footer", None)

            extracted_images.extend(self._extract_images_from_hf(doc, header, "header", variant, sec_info))
            extracted_images.extend(self._extract_images_from_hf(doc, footer, "footer", variant, sec_info))

            sections.append(sec_info)

        variants_unique = []
        for v in variants:
            if v not in variants_unique:
                variants_unique.append(v)

        parse_result: Dict[str, Any] = {
            "heading_levels": heading_present,
            "variants": variants_unique,
            "sections": sections,
            "header_footer_images": {
                # variant -> {"header": n, "footer": n}
            },
            "confidence": float(confidence_min),
        }

        # aggregate header/footer counts
        agg: Dict[str, Dict[str, int]] = {}
        for sec in sections:
            v = sec.get("variant") or "DEFAULT"
            agg.setdefault(v, {"header": 0, "footer": 0})
            agg[v]["header"] += int(sec.get("header_images") or 0)
            agg[v]["footer"] += int(sec.get("footer_images") or 0)
        parse_result["header_footer_images"] = agg

        return parse_result, extracted_images

    def _extract_images_from_hf(
        self,
        doc: Document,
        header_or_footer,
        where: str,
        variant: str,
        sec_info: Dict[str, Any],
    ) -> List[ExtractedImage]:
        if header_or_footer is None:
            return []

        out: List[ExtractedImage] = []

        try:
            el = getattr(header_or_footer, "_element", None)
            if el is None:
                return []

            # 找到 blip embed rid 列表
            rids = el.xpath(".//a:blip/@r:embed", namespaces=self._NS) or []
            for rid in rids:
                rid = str(rid)
                try:
                    # 注意：图片关系在 header/footer part 上，而不是 document.part
                    part = header_or_footer.part.related_parts.get(rid)  # type: ignore[attr-defined]
                except Exception:
                    part = None
                if part is None:
                    continue

                try:
                    blob = bytes(getattr(part, "blob", b""))
                except Exception:
                    blob = b""
                if not blob:
                    continue

                content_type = str(getattr(part, "content_type", "") or "")
                # 尽量推断扩展名
                ext = "png"
                if "jpeg" in content_type:
                    ext = "jpg"
                elif "png" in content_type:
                    ext = "png"
                elif "gif" in content_type:
                    ext = "gif"
                elif "bmp" in content_type:
                    ext = "bmp"

                sha = _sha256_bytes(blob)
                file_name = f"{where}_{variant}_{sha[:12]}.{ext}"

                width_px, height_px = self._try_get_extent_px(el, rid)

                out.append(
                    ExtractedImage(
                        where=where,
                        variant=variant,
                        content_type=content_type or f"image/{ext}",
                        file_name=file_name,
                        sha256=sha,
                        blob=blob,
                        width_px=width_px,
                        height_px=height_px,
                    )
                )

                if where == "header":
                    sec_info["header_images"] = int(sec_info.get("header_images") or 0) + 1
                else:
                    sec_info["footer_images"] = int(sec_info.get("footer_images") or 0) + 1

        except Exception:
            return out

        return out

    def _try_get_extent_px(self, hf_element, rid: str) -> Tuple[Optional[int], Optional[int]]:
        """
        从 drawing/wp:extent 拿尺寸（EMU），近似转换为 px。
        解析失败返回 (None, None)
        """
        try:
            # 找到 embed=rid 的 blip 节点
            blips = hf_element.xpath(f'.//a:blip[@r:embed="{rid}"]', namespaces=self._NS) or []
            if not blips:
                return (None, None)
            blip = blips[0]
            # 向上找 wp:inline / wp:anchor
            parent = blip
            for _ in range(10):
                parent = parent.getparent()
                if parent is None:
                    break
                if parent.tag.endswith("}inline") or parent.tag.endswith("}anchor"):
                    break
            if parent is None:
                return (None, None)
            extent = parent.xpath(".//wp:extent", namespaces=self._NS) or []
            if not extent:
                return (None, None)
            ex = extent[0]
            cx = int(ex.get("cx") or 0)
            cy = int(ex.get("cy") or 0)
            if cx <= 0 or cy <= 0:
                return (None, None)
            return (int(cx / EMU_PER_PX), int(cy / EMU_PER_PX))
        except Exception:
            return (None, None)


class TemplatePreviewGenerator:
    """
    基于模板 docx 生成示范 docx，并可选转 pdf。
    """

    def __init__(self, work_dir: str):
        self.work_dir = work_dir
        _safe_mkdir(self.work_dir)

    def generate_sample_docx(
        self,
        template_docx_path: str,
        parse_result: Dict[str, Any],
        images_by_variant: Dict[str, Dict[str, List[str]]],
        spec_style_hints: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        返回生成的 sample.docx 路径
        images_by_variant: {variant: {"header":[img_path...], "footer":[...]}}
        """
        doc = Document(template_docx_path)
        _remove_all_body_content_keep_last_sectpr(doc)

        style_hints = spec_style_hints or {}
        h1 = _pick_heading_style(doc, 1, style_hints.get("heading1_style"))
        h2 = _pick_heading_style(doc, 2, style_hints.get("heading2_style"))
        h3 = _pick_heading_style(doc, 3, style_hints.get("heading3_style"))
        h4 = _pick_heading_style(doc, 4, style_hints.get("heading4_style"))
        h5 = _pick_heading_style(doc, 5, style_hints.get("heading5_style"))
        body_style = _pick_body_style(doc, style_hints.get("body_style"))

        def _apply_section_params(sec, variant: str) -> None:
            # 基于 parse_result 里对应 variant 的第一条 section 配置套用边距/页眉页脚距离
            try:
                sec_cfgs = [s for s in (parse_result.get("sections") or []) if (s.get("variant") == variant)]
                cfg = sec_cfgs[0] if sec_cfgs else None
                if not cfg:
                    return
                m = (cfg.get("margins_mm") or {}) if isinstance(cfg, dict) else {}
                if m:
                    sec.top_margin = Mm(float(m.get("top") or 0))  # type: ignore[attr-defined]
                    sec.bottom_margin = Mm(float(m.get("bottom") or 0))  # type: ignore[attr-defined]
                    sec.left_margin = Mm(float(m.get("left") or 0))  # type: ignore[attr-defined]
                    sec.right_margin = Mm(float(m.get("right") or 0))  # type: ignore[attr-defined]
                hd = cfg.get("header_distance_mm")
                fd = cfg.get("footer_distance_mm")
                if hd is not None:
                    sec.header_distance = Mm(float(hd))  # type: ignore[attr-defined]
                if fd is not None:
                    sec.footer_distance = Mm(float(fd))  # type: ignore[attr-defined]
            except Exception:
                return

        def _set_section_variant(sec, variant: str) -> None:
            # 设置 page size + orientation（A4/A3 + 横竖）
            try:
                if variant == "A4_PORTRAIT":
                    sec.orientation = WD_ORIENTATION.PORTRAIT
                    sec.page_width = Mm(210)  # type: ignore[attr-defined]
                    sec.page_height = Mm(297)  # type: ignore[attr-defined]
                elif variant == "A4_LANDSCAPE":
                    sec.orientation = WD_ORIENTATION.LANDSCAPE
                    sec.page_width = Mm(297)  # type: ignore[attr-defined]
                    sec.page_height = Mm(210)  # type: ignore[attr-defined]
                elif variant == "A3_LANDSCAPE":
                    sec.orientation = WD_ORIENTATION.LANDSCAPE
                    sec.page_width = Mm(420)  # type: ignore[attr-defined]
                    sec.page_height = Mm(297)  # type: ignore[attr-defined]
                else:
                    # DEFAULT: 不强制
                    pass
            except Exception:
                pass

        def _apply_header_footer(sec, variant: str) -> None:
            # 横版 section 必须断开 linkToPrevious，避免继承上一节
            try:
                sec.header.is_linked_to_previous = False  # type: ignore[attr-defined]
                sec.footer.is_linked_to_previous = False  # type: ignore[attr-defined]
            except Exception:
                pass

            hf = images_by_variant.get(variant) or images_by_variant.get("DEFAULT") or {}
            header_imgs = list(hf.get("header") or [])
            footer_imgs = list(hf.get("footer") or [])

            # 清空原有页眉页脚内容，再把图片插入
            try:
                _clear_container_paragraphs(sec.header)
                _clear_container_paragraphs(sec.footer)
            except Exception:
                pass

            if header_imgs:
                p = sec.header.add_paragraph()
                r = p.add_run()
                for img_path in header_imgs[:2]:
                    try:
                        r.add_picture(img_path)
                    except Exception:
                        continue
            if footer_imgs:
                p = sec.footer.add_paragraph()
                r = p.add_run()
                for img_path in footer_imgs[:2]:
                    try:
                        r.add_picture(img_path)
                    except Exception:
                        continue

        # section 1: A4 portrait
        sec1 = doc.sections[0]
        _set_section_variant(sec1, "A4_PORTRAIT")
        _apply_section_params(sec1, "A4_PORTRAIT")
        _apply_header_footer(sec1, "A4_PORTRAIT")

        # cover + toc
        p = doc.add_paragraph("示范文档（模板解析预览）")
        if h1:
            try:
                p.style = h1
            except Exception:
                pass
        doc.add_paragraph("目录").style = h1 or doc.styles["Normal"]  # type: ignore[index]
        toc_p = doc.add_paragraph()
        _insert_toc_field(toc_p)

        # heading 1..5 + body paragraphs
        def _add_heading(text: str, style_name: Optional[str]) -> None:
            pp = doc.add_paragraph(text)
            if style_name:
                try:
                    pp.style = style_name
                except Exception:
                    pass

        def _add_body(text: str) -> None:
            pp = doc.add_paragraph(text)
            if body_style:
                try:
                    pp.style = body_style
                except Exception:
                    pass

        _add_heading("第一章 示例：项目概述（Heading 1）", h1)
        _add_body("这是示范正文段落，用于检查模板正文样式（字体/行距/段前段后/缩进）是否生效。")
        _add_body("第二段示例正文：Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt.")
        _add_heading("1.1 示例：背景（Heading 2）", h2)
        _add_body("这里是二级标题下的正文内容示例。")
        _add_heading("1.1.1 示例：范围（Heading 3）", h3)
        _add_body("这里是三级标题下的正文内容示例。")
        _add_heading("1.1.1.1 示例：术语（Heading 4）", h4)
        _add_body("这里是四级标题下的正文内容示例。")
        _add_heading("1.1.1.1.1 示例：补充（Heading 5）", h5)
        _add_body("这里是五级标题下的正文内容示例。")

        # 增加一点内容撑到 2-3 页
        for i in range(18):
            _add_body(f"正文填充段落 {i+1}：用于撑页，确保预览至少 2-3 页，便于观察页眉页脚是否贯穿。")

        # section 2: A4 landscape
        sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_section_variant(sec2, "A4_LANDSCAPE")
        _apply_section_params(sec2, "A4_LANDSCAPE")
        _apply_header_footer(sec2, "A4_LANDSCAPE")
        _add_heading("第二章 示例：A4 横版 Section", h1)
        for i in range(6):
            _add_body(f"A4 横版内容示例段落 {i+1}。")

        # section 3: A3 landscape
        sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_section_variant(sec3, "A3_LANDSCAPE")
        _apply_section_params(sec3, "A3_LANDSCAPE")
        _apply_header_footer(sec3, "A3_LANDSCAPE")
        _add_heading("第三章 示例：A3 横版 Section", h1)
        for i in range(6):
            _add_body(f"A3 横版内容示例段落 {i+1}。")

        out_path = os.path.join(self.work_dir, "sample.docx")
        doc.save(out_path)
        return out_path

    def convert_to_pdf(self, docx_path: str) -> Optional[str]:
        """
        用 LibreOffice headless 把 docx 转成 pdf。
        返回 pdf_path 或 None（表示转换不可用/失败）
        """
        out_dir = self.work_dir
        _safe_mkdir(out_dir)
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(out_dir, f"{base}.pdf")

        # 清理旧文件，避免拿到过期输出
        try:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
        except Exception:
            pass

        cmd = [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            out_dir,
            docx_path,
        ]
        try:
            env = os.environ.copy()
            # LibreOffice 有时需要 HOME 可写
            env.setdefault("HOME", "/tmp")
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)
        except FileNotFoundError:
            return None
        except Exception:
            return None

        return pdf_path if os.path.exists(pdf_path) else None


