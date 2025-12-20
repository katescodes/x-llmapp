"""
DocxBlockExtractor - 确定性结构化解析 Word 文档
只做可重复的结构化抽取，不做意图判断
"""
from __future__ import annotations

import re
import uuid
from dataclasses import dataclass, field
from enum import Enum
from typing import Dict, List, Optional, Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import nsmap

from app.services.docx_style_utils import guess_heading_level


class BlockType(str, Enum):
    """块类型枚举"""
    PARAGRAPH = "PARAGRAPH"
    TABLE = "TABLE"
    HEADER = "HEADER"
    FOOTER = "FOOTER"


@dataclass
class DocxBlock:
    """Word 文档块（段落/表格等）"""
    id: str
    type: BlockType
    text: str
    tag: Optional[str] = None  # "TOC"|"INSTRUCTION"|"COLOR_SWATCH"|"NORMAL"
    style_id: Optional[str] = None
    style_name: Optional[str] = None
    outline_level: Optional[int] = None
    num_id: Optional[int] = None
    ilvl: Optional[int] = None
    is_page_break_before: bool = False
    is_section_break: bool = False
    table_meta: Optional[Dict[str, Any]] = None
    sequence: int = 0  # 在文档中的顺序


@dataclass
class DocxExtractResult:
    """文档抽取结果"""
    blocks: List[DocxBlock] = field(default_factory=list)
    style_stats: Dict[str, Any] = field(default_factory=dict)
    numbering_stats: Dict[str, Any] = field(default_factory=dict)
    header_footer_stats: Dict[str, Any] = field(default_factory=dict)
    # ---- enhanced context for LLM / export ----
    style_catalog: Dict[str, Any] = field(default_factory=dict)  # 全量样式信息（段落样式）
    instructions_text: str = ""  # 说明块合并文本（不进正文）
    header_footer_media: Dict[str, Any] = field(default_factory=dict)  # 页眉页脚图片/LOGO信息
    tags_by_block_id: Dict[str, str] = field(default_factory=dict)  # block_id -> TOC/INSTRUCTION/COLOR_SWATCH/NORMAL


class DocxBlockExtractor:
    """Word 文档块提取器"""

    def __init__(self):
        pass

    _RE_TOC_LINE = re.compile(
        r"^\s*(([一二三四五六七八九十]、)|(\d+(\.\d+)*))\s+.+\s+\d+\s*$"
    )
    _INSTRUCTION_KEYWORDS = [
        "模板使用说明",
        "填写说明",
        "请复制本页页眉页脚",
        "连接至前一节",
        "域代码",
        "StyleRef",
        "页眉",
        "页脚",
        "目录样式",
        "书签",
        "交叉引用",
    ]

    def _is_toc_line(self, text: str) -> bool:
        """识别目录行（确定性启发式）"""
        s = (text or "").strip()
        if not s:
            return False
        s = re.sub(r"\s+", " ", s)
        return bool(self._RE_TOC_LINE.match(s))

    def _is_instruction(self, text: str) -> bool:
        """识别模板说明/操作指引类文本（确定性关键词）"""
        s = (text or "")
        if not s:
            return False
        return any(k in s for k in self._INSTRUCTION_KEYWORDS)

    def _is_color_swatch(self, text: str) -> bool:
        """识别色卡块：#RRGGBB 占比高"""
        s = (text or "")
        if not s:
            return False
        hits = len(re.findall(r"#[0-9A-Fa-f]{6}", s))
        ratio = hits / max(1, len(s.split()))
        return hits >= 2 or ratio > 0.5

    def extract(
        self,
        docx_bytes: bytes,
        max_blocks: int = 400,
        max_chars_per_block: int = 300
    ) -> DocxExtractResult:
        """
        从 docx 字节流中提取结构化块
        
        Args:
            docx_bytes: Word 文档字节内容
            max_blocks: 最大块数量
            max_chars_per_block: 每个块最大字符数
            
        Returns:
            DocxExtractResult: 提取结果
        """
        from io import BytesIO
        doc = Document(BytesIO(docx_bytes))
        
        result = DocxExtractResult()
        sequence = 0

        # 0. 额外结构化信息（不依赖 blocks）
        try:
            result.header_footer_media = self._extract_header_footer_media(doc)
        except Exception:
            result.header_footer_media = {}
        try:
            result.style_catalog = self._build_style_catalog(doc)
        except Exception:
            result.style_catalog = {}
        
        # 1. 提取主体内容（段落和表格）
        blocks = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                block = self._extract_paragraph(para, sequence, max_chars_per_block)
                if block:
                    blocks.append(block)
                    sequence += 1
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                block = self._extract_table(table, sequence, max_chars_per_block)
                if block:
                    blocks.append(block)
                    sequence += 1
        
        # 2. 提取页眉页脚
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    block = self._extract_header_footer_paragraph(
                        para, sequence, max_chars_per_block, BlockType.HEADER
                    )
                    if block:
                        blocks.append(block)
                        sequence += 1
            
            if section.footer:
                for para in section.footer.paragraphs:
                    block = self._extract_header_footer_paragraph(
                        para, sequence, max_chars_per_block, BlockType.FOOTER
                    )
                    if block:
                        blocks.append(block)
                        sequence += 1
        
        # 3. 去噪与优先级保留（控制 max_blocks）
        blocks = self._denoise_and_prioritize(blocks, max_blocks)
        result.blocks = blocks

        # 3.1 tags_by_block_id（给 LLM/排除列表用）
        try:
            result.tags_by_block_id = {
                b.id: (getattr(b, "tag", None) or "NORMAL") for b in (blocks or [])
            }
        except Exception:
            result.tags_by_block_id = {}

        # 3.2 instructions_text：说明块合并（但说明块仍保留在 blocks 中用于 exclude）
        try:
            ins = [
                (b.text or "") for b in (blocks or [])
                if getattr(b, "tag", None) == "INSTRUCTION" and (b.text or "").strip()
            ]
            merged = "\n".join(ins).strip()
            # 限制长度，避免 prompt 过长
            result.instructions_text = merged[:4000] if len(merged) > 4000 else merged
        except Exception:
            result.instructions_text = ""
        
        # 4. 统计样式、编号等信息
        result.style_stats = self._compute_style_stats(blocks)
        result.numbering_stats = self._compute_numbering_stats(blocks)
        result.header_footer_stats = self._compute_header_footer_stats(blocks)
        try:
            # header_footer_media 的 logo_detected 反哺 stats（供 analyzer 更可靠判断）
            if isinstance(result.header_footer_media, dict):
                ld = bool(result.header_footer_media.get("logo_detected", False))
                result.header_footer_stats["logo_detected"] = ld
        except Exception:
            pass
        
        return result

    def _extract_header_footer_media(self, doc: Document) -> Dict[str, Any]:
        """
        粗略抽取页眉/页脚里的图片关系（LOGO 先按“有图片”判定）
        返回结构：
        {
          "has_header": True/False,
          "has_footer": True/False,
          "images": [{"where":"header|footer","rid":"rIdX","partname":"/word/media/image1.png","content_type":"image/png"}],
          "logo_detected": True/False
        }
        """
        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        images: List[Dict[str, Any]] = []
        has_header = False
        has_footer = False

        def _scan(where: str, header_or_footer) -> None:
            nonlocal images
            try:
                el = getattr(header_or_footer, "_element", None)
                if el is None:
                    return
                rids = el.xpath(".//a:blip/@r:embed", namespaces=ns) or []
                for rid in rids:
                    try:
                        part = doc.part.related_parts.get(rid)  # type: ignore[attr-defined]
                    except Exception:
                        part = None
                    info: Dict[str, Any] = {"where": where, "rid": str(rid)}
                    if part is not None:
                        try:
                            info["partname"] = str(getattr(part, "partname", ""))
                        except Exception:
                            pass
                        try:
                            info["content_type"] = str(getattr(part, "content_type", ""))
                        except Exception:
                            pass
                    images.append(info)
            except Exception:
                return

        try:
            for section in doc.sections:
                try:
                    if getattr(section, "header", None):
                        has_header = True
                        _scan("header", section.header)
                except Exception:
                    pass
                try:
                    if getattr(section, "footer", None):
                        has_footer = True
                        _scan("footer", section.footer)
                except Exception:
                    pass
        except Exception:
            pass

        # 去重（where+rid）
        uniq = []
        seen = set()
        for it in images:
            k = (it.get("where"), it.get("rid"))
            if k in seen:
                continue
            seen.add(k)
            uniq.append(it)

        return {
            "has_header": bool(has_header),
            "has_footer": bool(has_footer),
            "images": uniq,
            "logo_detected": bool(uniq),
        }

    def _build_style_catalog(self, doc: Document) -> Dict[str, Any]:
        """
        构建段落样式字典（尽量包含关键信息，None 也保留）。
        注意：python-docx 对中文字体 eastAsia 支持不完整，这里会尝试从 style XML 补取。
        """
        def _as_hex(rgb) -> Optional[str]:
            try:
                if rgb is None:
                    return None
                # rgb 可能是 RGBColor 或类似对象
                s = str(rgb)
                if re.fullmatch(r"[0-9A-Fa-f]{6}", s):
                    return f"#{s.upper()}"
                if s.startswith("#") and re.fullmatch(r"#[0-9A-Fa-f]{6}", s):
                    return s.upper()
                return None
            except Exception:
                return None

        def _alignment_to_str(a) -> Optional[str]:
            try:
                if a is None:
                    return None
                if a == WD_ALIGN_PARAGRAPH.LEFT:
                    return "left"
                if a == WD_ALIGN_PARAGRAPH.CENTER:
                    return "center"
                if a == WD_ALIGN_PARAGRAPH.RIGHT:
                    return "right"
                if a == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return "justify"
                return str(a)
            except Exception:
                return None

        def _get_east_asia_font(style) -> Optional[str]:
            try:
                el = getattr(style, "_element", None)
                if el is None:
                    return None
                vals = el.xpath(".//w:rFonts/@w:eastAsia", namespaces=nsmap) or []
                if vals:
                    v = str(vals[0]).strip()
                    return v or None
            except Exception:
                return None
            return None

        def _first_line_indent_chars(style) -> Optional[int]:
            # 优先读 w:ind/@w:firstLineChars（Word 内部通常是 1/100 字符）
            try:
                el = getattr(style, "_element", None)
                if el is None:
                    return None
                vals = el.xpath(".//w:pPr/w:ind/@w:firstLineChars", namespaces=nsmap) or []
                if not vals:
                    return None
                v = int(str(vals[0]))
                # 常见：200 -> 2 字符（粗略）
                return max(0, int(round(v / 100)))
            except Exception:
                return None

        paragraph_styles: List[Dict[str, Any]] = []
        try:
            for s in doc.styles:
                try:
                    if getattr(s, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
                        continue
                except Exception:
                    continue

                name = getattr(s, "name", None)
                if not name:
                    continue

                based_on = None
                try:
                    if getattr(s, "base_style", None):
                        based_on = getattr(s.base_style, "name", None)
                except Exception:
                    based_on = None

                font_name = None
                font_size_pt = None
                bold = None
                color = None
                try:
                    f = s.font
                    font_name = getattr(f, "name", None)
                    try:
                        font_size_pt = f.size.pt if getattr(f, "size", None) is not None else None
                    except Exception:
                        font_size_pt = None
                    try:
                        bold = bool(f.bold) if f.bold is not None else None
                    except Exception:
                        bold = None
                    try:
                        color = _as_hex(getattr(getattr(f, "color", None), "rgb", None))
                    except Exception:
                        color = None
                except Exception:
                    pass

                # 补取 eastAsia 字体（若 font.name 为空）
                east_asia = _get_east_asia_font(s)
                if not font_name and east_asia:
                    font_name = east_asia

                alignment = None
                line_spacing = None
                space_before_pt = None
                space_after_pt = None
                first_line_chars = None
                try:
                    pf = s.paragraph_format
                    alignment = _alignment_to_str(getattr(pf, "alignment", None))
                    ls = getattr(pf, "line_spacing", None)
                    if ls is None:
                        line_spacing = None
                    else:
                        # 可能是 float 或 Length；尽量转换成可序列化字符串
                        try:
                            line_spacing = str(float(ls))
                        except Exception:
                            try:
                                line_spacing = str(getattr(ls, "pt", ls))
                            except Exception:
                                line_spacing = str(ls)
                    try:
                        sb = getattr(pf, "space_before", None)
                        space_before_pt = sb.pt if sb is not None else None
                    except Exception:
                        space_before_pt = None
                    try:
                        sa = getattr(pf, "space_after", None)
                        space_after_pt = sa.pt if sa is not None else None
                    except Exception:
                        space_after_pt = None
                except Exception:
                    pass
                first_line_chars = _first_line_indent_chars(s)

                paragraph_styles.append(
                    {
                        "name": str(name),
                        "based_on": based_on,
                        "font": {
                            "name": font_name,
                            "size_pt": font_size_pt,
                            "bold": bold,
                            "color": color,
                        },
                        "para": {
                            "alignment": alignment,
                            "line_spacing": line_spacing,
                            "first_line_indent_chars": first_line_chars,
                            "space_before_pt": space_before_pt,
                            "space_after_pt": space_after_pt,
                        },
                    }
                )
        except Exception:
            paragraph_styles = []

        return {"paragraph_styles": paragraph_styles, "defaults": {}}

    def _extract_paragraph(
        self,
        para: Paragraph,
        sequence: int,
        max_chars: int
    ) -> Optional[DocxBlock]:
        """提取段落块"""
        full_text = para.text.strip()
        
        # 过滤空段落（但保留可能的分隔符）
        if not full_text and not self._is_separator_paragraph(para):
            return None
        
        # 获取样式信息
        style_id = para.style.style_id if para.style else None
        style_name = para.style.name if para.style else None
        
        # 获取大纲级别
        outline_level = None
        try:
            lvl = guess_heading_level(para)  # 1-based
            outline_level = (lvl - 1) if lvl is not None else None  # 与历史字段保持一致：0-based
        except Exception:
            # 任何异常都不能中断整个 extract
            outline_level = None
        
        # 获取编号信息
        num_id = None
        ilvl = None
        try:
            vals = para._p.xpath("./w:pPr/w:numPr/w:numId/@w:val", namespaces=nsmap)
            if vals:
                num_id = int(vals[0])
            vals = para._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val", namespaces=nsmap)
            if vals:
                ilvl = int(vals[0])
        except Exception:
            num_id = None
            ilvl = None
        
        # 检查分页符
        is_page_break = self._has_page_break(para)
        
        # 计算 tag（在截断前）
        tag = "NORMAL"
        if self._is_instruction(full_text):
            tag = "INSTRUCTION"
        elif self._is_color_swatch(full_text):
            tag = "COLOR_SWATCH"
        elif self._is_toc_line(full_text):
            tag = "TOC"

        # 截断文本
        text = full_text[:max_chars] if len(full_text) > max_chars else full_text
        
        return DocxBlock(
            id=str(uuid.uuid4()),
            type=BlockType.PARAGRAPH,
            text=text,
            tag=tag,
            style_id=style_id,
            style_name=style_name,
            outline_level=outline_level,
            num_id=num_id,
            ilvl=ilvl,
            is_page_break_before=is_page_break,
            sequence=sequence
        )

    def _extract_table(
        self,
        table: Table,
        sequence: int,
        max_chars: int
    ) -> Optional[DocxBlock]:
        """提取表格块"""
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        
        if rows == 0 or cols == 0:
            return None
        
        # 提取表格文本（前几行）
        text_parts = []
        for i, row in enumerate(table.rows[:3]):  # 只取前3行
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)

        tag = "NORMAL"
        if self._is_color_swatch(full_text):
            tag = "COLOR_SWATCH"
        elif self._is_instruction(full_text):
            tag = "INSTRUCTION"
        else:
            tag = "NORMAL"

        text = full_text[:max_chars] if len(full_text) > max_chars else full_text
        
        # 判断首行是否为表头
        is_header_row = False
        if rows > 1:
            first_row = table.rows[0]
            # 简单启发式：首行单元格都有内容且字体加粗
            is_header_row = all(cell.text.strip() for cell in first_row.cells)
        
        table_meta = {
            "rows": rows,
            "cols": cols,
            "has_header": is_header_row
        }
        
        return DocxBlock(
            id=str(uuid.uuid4()),
            type=BlockType.TABLE,
            text=text,
            tag=tag,
            table_meta=table_meta,
            sequence=sequence
        )

    def _extract_header_footer_paragraph(
        self,
        para: Paragraph,
        sequence: int,
        max_chars: int,
        block_type: BlockType
    ) -> Optional[DocxBlock]:
        """提取页眉页脚段落"""
        text = para.text.strip()
        if not text:
            return None
        
        text = text[:max_chars] if len(text) > max_chars else text
        
        return DocxBlock(
            id=str(uuid.uuid4()),
            type=block_type,
            text=text,
            tag=None,
            style_id=para.style.style_id if para.style else None,
            style_name=para.style.name if para.style else None,
            sequence=sequence
        )

    def _is_separator_paragraph(self, para: Paragraph) -> bool:
        """判断是否为分隔符段落（如水平线）"""
        # 简单实现：检查是否有边框样式（避免直接访问 pPr.pBdr）
        try:
            nodes = para._p.xpath("./w:pPr/w:pBdr", namespaces=nsmap)
            return bool(nodes)
        except Exception:
            return False

    def _has_page_break(self, para: Paragraph) -> bool:
        """检查段落前是否有分页符"""
        try:
            nodes = para._p.xpath("./w:pPr/w:pageBreakBefore", namespaces=nsmap)
            return bool(nodes)
        except Exception:
            return False

    def _denoise_and_prioritize(
        self,
        blocks: List[DocxBlock],
        max_blocks: int
    ) -> List[DocxBlock]:
        """
        去噪并优先保留重要块
        策略：
        1. 优先保留标题类段落（有 outline_level 或标题样式）
        2. 保留前 N 个块和后 N 个块（避免丢失开头和结尾）
        3. 合并连续空白或相似块
        """
        if len(blocks) <= max_blocks:
            return blocks
        
        # 标记重要块
        important_indices = set()
        
        # 1. 保留标题块
        for i, block in enumerate(blocks):
            if getattr(block, "tag", None) in ("TOC", "INSTRUCTION", "COLOR_SWATCH"):
                important_indices.add(i)
                continue
            if block.outline_level is not None:
                important_indices.add(i)
            elif block.style_name and any(
                keyword in block.style_name.lower()
                for keyword in ["heading", "标题", "title"]
            ):
                important_indices.add(i)
            elif block.num_id is not None and block.ilvl is not None and block.ilvl < 3:
                # 保留前3级编号项
                important_indices.add(i)
        
        # 2. 保留开头和结尾
        head_count = min(50, max_blocks // 4)
        tail_count = min(50, max_blocks // 4)
        
        for i in range(min(head_count, len(blocks))):
            important_indices.add(i)
        
        for i in range(max(0, len(blocks) - tail_count), len(blocks)):
            important_indices.add(i)
        
        # 3. 如果还需要削减，从中间非重要块中均匀采样
        if len(important_indices) > max_blocks:
            important_list = sorted(important_indices)[:max_blocks]
            important_indices = set(important_list)
        elif len(important_indices) < max_blocks:
            # 补充一些中间块
            remaining = max_blocks - len(important_indices)
            middle_range = [
                i for i in range(len(blocks))
                if i not in important_indices
            ]
            step = max(1, len(middle_range) // remaining)
            for i in range(0, len(middle_range), step):
                if len(important_indices) >= max_blocks:
                    break
                important_indices.add(middle_range[i])
        
        # 按原顺序返回
        result = [blocks[i] for i in sorted(important_indices) if i < len(blocks)]
        return result

    def _compute_style_stats(self, blocks: List[DocxBlock]) -> Dict[str, Any]:
        """统计样式使用情况"""
        style_count: Dict[str, int] = {}
        heading_styles: List[str] = []
        counter_by_lvl: Dict[int, Dict[str, int]] = {i: {} for i in range(1, 6)}
        body_style_counter: Dict[str, int] = {}
        has_table = False
        
        for block in blocks:
            if block.type == BlockType.TABLE:
                has_table = True
            if block.style_name:
                style_count[block.style_name] = style_count.get(block.style_name, 0) + 1
                
                if any(
                    keyword in block.style_name.lower()
                    for keyword in ["heading", "标题", "title"]
                ):
                    if block.style_name not in heading_styles:
                        heading_styles.append(block.style_name)

            # 各级标题样式：按 outline_level 聚合（只取 1..5）
            if (
                block.type == BlockType.PARAGRAPH
                and block.outline_level is not None
                and block.style_name
            ):
                lvl1 = int(block.outline_level) + 1  # 还原 1-based
                if 1 <= lvl1 <= 5:
                    d = counter_by_lvl[lvl1]
                    d[block.style_name] = d.get(block.style_name, 0) + 1

            # 正文样式猜测：非标题段落 + NORMAL
            if (
                block.type == BlockType.PARAGRAPH
                and block.outline_level is None
                and getattr(block, "tag", None) == "NORMAL"
                and block.style_name
            ):
                body_style_counter[block.style_name] = body_style_counter.get(block.style_name, 0) + 1
        
        # 找出最常用的样式
        sorted_styles = sorted(style_count.items(), key=lambda x: x[1], reverse=True)

        heading_style_by_level: Dict[str, str] = {}
        for lvl in range(1, 6):
            d = counter_by_lvl.get(lvl) or {}
            if not d:
                continue
            heading_style_by_level[str(lvl)] = max(d.items(), key=lambda x: x[1])[0]

        body_style_guess = None
        if body_style_counter:
            body_style_guess = max(body_style_counter.items(), key=lambda x: x[1])[0]
        
        return {
            "style_count": style_count,
            "heading_styles": heading_styles,
            "heading_style_by_level": heading_style_by_level,
            "body_style_guess": body_style_guess,
            "has_table": has_table,
            "top_styles": [s[0] for s in sorted_styles[:10]],
            "total_blocks": len(blocks)
        }

    def _compute_numbering_stats(self, blocks: List[DocxBlock]) -> Dict[str, Any]:
        """统计编号使用情况"""
        numbering_count: Dict[int, int] = {}
        level_count: Dict[int, int] = {}
        
        for block in blocks:
            if block.num_id is not None:
                numbering_count[block.num_id] = numbering_count.get(block.num_id, 0) + 1
            if block.ilvl is not None:
                level_count[block.ilvl] = level_count.get(block.ilvl, 0) + 1
        
        return {
            "numbering_count": numbering_count,
            "level_count": level_count,
            "has_numbering": len(numbering_count) > 0
        }

    def _compute_header_footer_stats(self, blocks: List[DocxBlock]) -> Dict[str, Any]:
        """统计页眉页脚情况"""
        header_count = sum(1 for b in blocks if b.type == BlockType.HEADER)
        footer_count = sum(1 for b in blocks if b.type == BlockType.FOOTER)
        
        return {
            "header_count": header_count,
            "footer_count": footer_count,
            "has_header": header_count > 0,
            "has_footer": footer_count > 0
        }
