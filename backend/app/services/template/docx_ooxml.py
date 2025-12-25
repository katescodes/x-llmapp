"""
DOCX OOXML 操作工具
提供按 body 顺序遍历、删除、插入段落、查找锚点等功能
"""
from __future__ import annotations
from typing import Iterator, Tuple, Optional, Any
import logging

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

CONTENT_MARKER = "[[CONTENT]]"


def iter_block_items(doc: Document) -> Iterator[Tuple[str, Any]]:
    """
    按 body 顺序遍历段落和表格
    
    Returns:
        ("p", Paragraph) 或 ("t", Table)
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "t", Table(child, doc)


def remove_block(block_obj: Any):
    """删除一个块（段落或表格）"""
    el = block_obj._element
    el.getparent().remove(el)


def insert_paragraph_after(
    p: Paragraph, 
    text: str = "", 
    style_name: str | None = None
) -> Paragraph:
    """
    在指定段落后插入新段落
    
    Args:
        p: 参考段落
        text: 段落文本
        style_name: 样式名称
        
    Returns:
        新创建的段落
    """
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    if text:
        para.add_run(text)
    if style_name:
        try:
            para.style = style_name
        except Exception as e:
            logger.warning(f"设置样式失败: {style_name}, error: {e}")
    return para


def paragraph_text(p: Paragraph) -> str:
    """获取段落文本"""
    return p.text or ""


def find_anchor(doc: Document) -> Paragraph:
    """
    查找内容插入锚点（更稳版）
    
    优先级：
    1. 文本标记 [[CONTENT]] 的段落
    2. 第一处 +标题1 段落，但跳过"封面/目录/模板说明"等
    3. fallback：文档最后一个段落
    
    Args:
        doc: Document对象
        
    Returns:
        锚点段落
    """
    # 1) 查找 [[CONTENT]] marker
    for kind, blk in iter_block_items(doc):
        if kind == "p":
            txt = paragraph_text(blk) or ""
            if CONTENT_MARKER in txt:
                # 清理 marker 文本
                blk.text = txt.replace(CONTENT_MARKER, "").strip()
                logger.info(f"找到 [[CONTENT]] 标记锚点")
                return blk
    
    def is_bad_anchor_text(t: str) -> bool:
        """判断是否为不适合作为锚点的文本（封面/目录/说明等）"""
        t = (t or "").strip()
        if not t:
            return True
        bad_keywords = ["目录", "封面", "模板使用说明", "样式说明", "色卡", "横版A4", "横版A3"]
        return any(k in t for k in bad_keywords)
    
    # 2) 查找第一个 +标题1 (skip bad keywords)
    for kind, blk in iter_block_items(doc):
        if kind == "p":
            try:
                if blk.style and blk.style.name in ["+标题1", "标题 1", "Heading 1"]:
                    txt = paragraph_text(blk)
                    if not is_bad_anchor_text(txt):
                        logger.info(f"使用第一个标题1作为锚点: {txt[:50]}")
                        return blk
                    else:
                        logger.debug(f"跳过不适合的锚点候选: {txt[:50]}")
            except Exception:
                continue
    
    # 3) fallback: 最后一个段落
    last_p = None
    for kind, blk in iter_block_items(doc):
        if kind == "p":
            last_p = blk
    
    if last_p is None:
        logger.warning("文档没有段落，创建新段落作为锚点")
        last_p = doc.add_paragraph("")
    else:
        logger.info("使用最后一个段落作为锚点（fallback）")
    
    return last_p


def find_anchor_by_sdt_tag(doc: Document, tag: str = "CONTENT_BODY") -> Paragraph | None:
    """
    通过 SDT (Structured Document Tag) 查找锚点（可选实现）
    
    Args:
        doc: Document对象
        tag: SDT标签名
        
    Returns:
        锚点段落或None
    """
    # TODO: 实现 SDT 查找逻辑（MVP可跳过）
    return None


def prune_after_anchor(doc: Document, anchor: Paragraph):
    """
    默认清理策略：删除 anchor 后面的所有 block（段落+表格）
    
    Args:
        doc: Document对象
        anchor: 锚点段落
    """
    seen_anchor = False
    to_delete = []
    
    for kind, blk in iter_block_items(doc):
        if kind == "p" and blk._element is anchor._element:
            seen_anchor = True
            continue
        if seen_anchor:
            to_delete.append(blk)
    
    logger.info(f"清理策略：删除锚点后 {len(to_delete)} 个块")
    for blk in to_delete:
        remove_block(blk)


def prune_by_keep_plan(doc: Document, keep_block_ids: list[str], delete_block_ids: list[str]):
    """
    按 LLM 的 keepPlan 精细化删除（可选，未来实现）
    
    Args:
        doc: Document对象
        keep_block_ids: 保留的块ID列表
        delete_block_ids: 删除的块ID列表
    """
    # TODO: 实现精细化删除逻辑
    logger.warning("精细化删除尚未实现，使用默认清理策略")


# === TOC helpers ===

TOC_MARKER = "[[TOC]]"


def find_marker_paragraph(doc: Document, marker: str) -> Paragraph | None:
    """
    查找包含特定标记的段落
    
    Args:
        doc: Document对象
        marker: 要查找的标记文本
        
    Returns:
        包含标记的段落（标记文本会被清除），如果未找到返回 None
    """
    for kind, blk in iter_block_items(doc):
        if kind != "p":
            continue
        txt = paragraph_text(blk) or ""
        if marker in txt:
            blk.text = txt.replace(marker, "").strip()
            logger.info(f"找到标记: {marker}")
            return blk
    return None


def find_toc_anchor(doc: Document) -> Paragraph | None:
    """
    查找目录锚点
    
    优先级：
    1. [[TOC]] 标记
    2. 文本为"目录"的段落
    
    Args:
        doc: Document对象
        
    Returns:
        目录锚点段落，如果未找到返回 None
    """
    # 1) 优先查找 [[TOC]] 标记
    p = find_marker_paragraph(doc, TOC_MARKER)
    if p:
        logger.info("使用 [[TOC]] 标记作为目录锚点")
        return p
    
    # 2) 其次查找文本为"目录"的段落
    for kind, blk in iter_block_items(doc):
        if kind == "p":
            t = (paragraph_text(blk) or "").strip()
            if t == "目录":
                logger.info("使用'目录'段落作为目录锚点")
                return blk
    
    logger.warning("未找到目录锚点")
    return None


def remove_blocks_between(doc: Document, start_p: Paragraph, end_p: Paragraph):
    """
    删除 start_p 之后、end_p 之前的所有段落/表格
    （不删除边界本身）
    
    Args:
        doc: Document对象
        start_p: 起始段落（不会被删除）
        end_p: 结束段落（不会被删除）
    """
    if not start_p or not end_p:
        logger.warning("起始或结束段落为空，跳过区间删除")
        return
    
    started = False
    to_delete = []
    
    for kind, blk in iter_block_items(doc):
        if kind == "p" and blk._element is start_p._element:
            started = True
            continue
        if started:
            if kind == "p" and blk._element is end_p._element:
                break
            to_delete.append(blk)
    
    logger.info(f"删除目录区间内的 {len(to_delete)} 个块")
    for blk in to_delete:
        remove_block(blk)


# === TOC 内容控件替换工具 ===

def is_toc_sdt(el) -> bool:
    """
    判断一个 w:sdt 是否为 Word 目录（TOC）控件
    
    经验判据：
    - 含 instrText 'TOC'
    - 或 docPartGallery / docPartObj（Table of Contents）
    
    Args:
        el: OOXML 元素
        
    Returns:
        是否为 TOC 控件
    """
    try:
        xml = el.xml if hasattr(el, "xml") else ""
        if "w:instrText" in xml and "TOC" in xml:
            return True
        if "docPartGallery" in xml and ("Table of Contents" in xml or "目录" in xml):
            return True
        if "w:docPartObj" in xml:
            return True
    except Exception:
        pass
    return False


def insert_paragraph_before_element(
    doc: Document, 
    el, 
    text: str = "", 
    style_name: str | None = None
) -> Paragraph:
    """
    在任意 OOXML 元素（如 w:sdt）之前插入段落
    
    Args:
        doc: Document对象
        el: OOXML 元素
        text: 段落文本
        style_name: 样式名称
        
    Returns:
        新创建的段落
    """
    new_p = OxmlElement("w:p")
    el.addprevious(new_p)
    p = Paragraph(new_p, doc._body)
    if text:
        p.add_run(text)
    if style_name:
        try:
            p.style = style_name
        except Exception as e:
            logger.warning(f"设置样式失败: {style_name}, error: {e}")
    return p


def replace_toc_sdt_with_plain_toc(
    doc: Document, 
    outline_nodes: list[dict], 
    role_mapping: dict
) -> bool:
    """
    若文档存在 TOC 的 w:sdt：删除它，并在原位置插入"纯目录行（无页码）"
    
    Args:
        doc: Document对象
        outline_nodes: 目录节点列表
        role_mapping: 样式角色映射
        
    Returns:
        是否替换成功
    """
    # 先把 outline 扁平化（兼容树形结构和扁平结构）
    def _flat(nodes):
        # 兼容扁平结构：若本身就是扁平则直接返回
        if nodes and isinstance(nodes[0], dict) and "children" not in nodes[0]:
            return nodes
        out = []
        def dfs(n):
            out.append(n)
            for c in n.get("children", []) or []:
                dfs(c)
        for n in nodes:
            dfs(n)
        return out

    flat = _flat(outline_nodes)

    # 找到第一个 TOC sdt
    body = doc.element.body
    toc_sdt = None
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sdt") and is_toc_sdt(child):
            toc_sdt = child
            logger.info("找到 TOC 内容控件，准备替换")
            break
    
    if toc_sdt is None:
        logger.info("未找到 TOC 内容控件，跳过替换")
        return False

    # 目录行样式：尽量用模板的 TOC1..TOC5（没有就 Normal）
    def toc_style(level: int) -> str:
        level = max(1, min(5, level))
        s = role_mapping.get(f"toc{level}")
        if s:
            return s
        for name in [f"TOC {level}", f"TOC{level}", f"toc {level}", f"toc {level}".lower()]:
            try:
                _ = doc.styles[name]
                return name
            except Exception:
                pass
        return role_mapping.get("body") or "Normal"

    # 在 sdt 之前插入一个"目录"标题（避免原模板目录标题藏在 sdt 里导致没了）
    cur = insert_paragraph_before_element(
        doc, 
        toc_sdt, 
        "目录", 
        style_name=role_mapping.get("h1") or role_mapping.get("body") or "Normal"
    )
    logger.info("插入目录标题段落")

    # 插入纯目录行（无页码/无点线）
    for n in flat:
        lvl = int(n.get("level") or 1)
        title = (n.get("title") or "").strip()
        numbering = (n.get("numbering") or "").strip()
        if not title:
            continue
        line = f"{numbering} {title}".strip() if numbering else title
        cur = insert_paragraph_before_element(doc, toc_sdt, line, style_name=toc_style(lvl))
    
    logger.info(f"插入 {len(flat)} 行纯文本目录（无页码）")

    # 删除原 TOC sdt
    toc_sdt.getparent().remove(toc_sdt)
    logger.info("✓ 已删除 TOC 内容控件，替换为纯文本目录")
    
    return True


# === 强化版 TOC 替换（更硬核，必命中）===

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}


def _make_p(doc: Document, text: str, style_name: str | None = None):
    """
    创建一个段落的 OOXML 元素（带样式）
    
    Args:
        doc: Document对象
        text: 段落文本
        style_name: 样式名称
        
    Returns:
        段落 OOXML 元素
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    if style_name:
        try:
            style = doc.styles[style_name]
            style_id = getattr(style, "style_id", None)
            if style_id:
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), style_id)
                pPr.append(pStyle)
        except Exception:
            pass

    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    # 保留空格
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def replace_all_toc_sdt_with_plain_toc(
    doc: Document, 
    flat_nodes: list[dict], 
    role_mapping: dict
) -> dict:
    """
    强制替换 Word TOC（w:sdt + instrText 'TOC'）为纯目录行（无页码）
    
    这是更硬核的版本，确保：
    1. 找到所有 TOC sdt（通过 instrText 'TOC' 判断）
    2. 在第一个 TOC 位置插入"目录"标题和所有目录行
    3. 删除所有 TOC sdt（防止模板里有多个 TOC）
    
    Args:
        doc: Document对象
        flat_nodes: 扁平化的目录节点列表
        role_mapping: 样式角色映射
        
    Returns:
        诊断信息：{toc_sdts_found, toc_sdts_removed, toc_lines_written}
    """
    body = doc.element.body  # <w:body/>
    children = list(body.iterchildren())

    # 1) 找所有 TOC sdt（instrText 含 TOC）
    toc_sdts = []
    for el in children:
        if not el.tag.endswith("}sdt"):
            continue
        instrs = el.xpath(".//w:instrText", namespaces=_NS)
        if any("TOC" in ((x.text or "")) for x in instrs):
            toc_sdts.append(el)

    diag = {
        "toc_sdts_found": len(toc_sdts),
        "toc_sdts_removed": 0,
        "toc_lines_written": 0,
    }

    if not toc_sdts:
        logger.info("未找到 TOC sdt（通过 instrText 'TOC' 判断）")
        return diag

    logger.info(f"找到 {len(toc_sdts)} 个 TOC sdt，准备强制替换")

    # 2) 目录行样式：toc1..toc5 -> 常见 TOC 样式 -> body/Normal
    def toc_style(level: int) -> str:
        level = max(1, min(5, level))
        s = role_mapping.get(f"toc{level}")
        if s:
            return s
        for name in (f"TOC {level}", f"TOC{level}", f"toc {level}", f"toc{level}"):
            try:
                _ = doc.styles[name]
                return name
            except Exception:
                pass
        return role_mapping.get("body") or "Normal"

    # 3) 以第一个 TOC sdt 的位置作为插入点：在它之前插入"目录 + 行"，再删除 sdt
    insert_before = toc_sdts[0]

    # 目录标题（你如果有专门目录标题样式，可放 role_mapping["toc_title"]）
    title_style = role_mapping.get("toc_title") or role_mapping.get("h1") or (role_mapping.get("body") or "Normal")
    insert_before.addprevious(_make_p(doc, "目录", title_style))
    logger.info(f"插入目录标题段落（样式: {title_style}）")

    # 插入目录行（纯标题，不包含编号，符合目录规范）
    import re
    last_line_text = ""
    
    for n in flat_nodes:
        lvl = int(n.get("level") or 1)
        title = (n.get("title") or "").strip()
        
        if not title:
            continue
        
        # 清理 title 中自带的编号前缀（目录中只显示标题名称）
        title_clean = re.sub(r"^\s*\d+(\.\d+)*\s*", "", title).strip()
        title_clean = re.sub(r"^\s*[一二三四五六七八九十]+\s*[、.．]\s*", "", title_clean).strip()
        
        # ✅ 修复：目录行不包含编号，只显示标题名称
        line = title_clean
        last_line_text = line
        
        insert_before.addprevious(_make_p(doc, line, toc_style(lvl)))
        diag["toc_lines_written"] += 1

    logger.info(f"插入 {diag['toc_lines_written']} 行纯文本目录（无页码）")
    
    # 记录最后一行文本（用于后续定位锚点）
    diag["toc_last_line_text"] = last_line_text if last_line_text else ""

    # 删除所有 TOC sdt（防止模板里多个 TOC）
    for el in toc_sdts:
        el.getparent().remove(el)
        diag["toc_sdts_removed"] += 1

    logger.info(f"✓ 已删除 {diag['toc_sdts_removed']} 个 TOC sdt")
    
    return diag

