"""
DOCX 结构提取器
从 Word 文档中提取 blocks（段落、表格），用于 LLM 分析
"""
from __future__ import annotations

import logging
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def extract_doc_blocks(docx_path: str, max_text_length: int = 300) -> List[Dict[str, Any]]:
    """
    提取文档的 blocks（段落、表格）列表
    
    每个 block 包含：
    - blockId: 唯一标识（如 "b0", "b1"）
    - type: "paragraph" 或 "table"
    - text: 文本内容（截断到 max_text_length）
    - styleName: 段落样式名称
    - isHeadingCandidate: 是否可能是标题
    - markerFlags: 特殊标记（如 hasContentMarker）
    - docPos: 文档位置信息（用于后续删除）
    
    Args:
        docx_path: Word 文档路径
        max_text_length: 每个 block 的最大文本长度
        
    Returns:
        blocks 列表
    """
    logger.info(f"提取文档结构: {docx_path}")
    
    doc = Document(docx_path)
    blocks: List[Dict[str, Any]] = []
    
    # 混排序号：paragraph 和 table 按文档顺序
    para_idx = 0
    table_idx = 0
    
    # 遍历文档的 body 元素（按顺序）
    body = doc._element.body
    
    for element in body:
        tag = element.tag
        
        # 段落
        if tag == qn('w:p'):
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                block = _extract_paragraph_block(para, len(blocks), para_idx, max_text_length)
                if block:  # 过滤空段落
                    blocks.append(block)
                para_idx += 1
        
        # 表格
        elif tag == qn('w:tbl'):
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                block = _extract_table_block(table, len(blocks), table_idx, max_text_length)
                if block:  # 过滤空表格
                    blocks.append(block)
                table_idx += 1
    
    logger.info(f"提取了 {len(blocks)} 个 blocks")
    
    return blocks


def _extract_paragraph_block(
    para,
    block_idx: int,
    para_idx: int,
    max_text_length: int
) -> Optional[Dict[str, Any]]:
    """提取段落 block"""
    text = para.text.strip()
    
    # 跳过空段落
    if not text:
        return None
    
    # 获取样式名称
    style_name = para.style.name if para.style else "Normal"
    
    # 判断是否是标题候选
    is_heading_candidate = _is_heading_candidate(para, style_name)
    
    # 检查特殊标记
    marker_flags = _check_markers(text)
    
    # 截断文本
    text_snippet = text[:max_text_length]
    if len(text) > max_text_length:
        text_snippet += "..."
    
    return {
        "blockId": f"b{block_idx}",
        "type": "paragraph",
        "text": text_snippet,
        "styleName": style_name,
        "isHeadingCandidate": is_heading_candidate,
        "markerFlags": marker_flags,
        "docPos": {
            "kind": "p",
            "index": para_idx
        }
    }


def _extract_table_block(
    table,
    block_idx: int,
    table_idx: int,
    max_text_length: int
) -> Optional[Dict[str, Any]]:
    """提取表格 block"""
    # 提取前两行作为摘要
    rows_text = []
    for i, row in enumerate(table.rows):
        if i >= 2:
            break
        row_text = " | ".join(cell.text.strip() for cell in row.cells)
        if row_text:
            rows_text.append(row_text)
    
    if not rows_text:
        return None
    
    text = "\n".join(rows_text)
    
    # 截断
    text_snippet = text[:max_text_length]
    if len(text) > max_text_length:
        text_snippet += "..."
    
    return {
        "blockId": f"b{block_idx}",
        "type": "table",
        "text": text_snippet,
        "styleName": "Table",
        "isHeadingCandidate": False,
        "markerFlags": {},
        "docPos": {
            "kind": "t",
            "index": table_idx
        }
    }


def _is_heading_candidate(para, style_name: str) -> bool:
    """
    判断段落是否可能是标题
    
    判断依据：
    1. 样式名包含 "Heading" 或 "标题"
    2. 有 outlineLvl 属性
    3. 字体较大（>14pt）
    """
    # 检查样式名
    if "heading" in style_name.lower() or "标题" in style_name:
        return True
    
    # 检查 outline level
    try:
        pPr = para._element.pPr
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                return True
    except Exception:
        pass
    
    # 检查字体大小（简化判断）
    try:
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.size and first_run.font.size.pt > 14:
                return True
    except Exception:
        pass
    
    return False


def _check_markers(text: str) -> Dict[str, bool]:
    """
    检查文本中的特殊标记
    
    Returns:
        标记字典，如 {"hasContentMarker": True}
    """
    markers = {}
    
    # 检查内容标记
    content_markers = [
        r'\[\[CONTENT\]\]',
        r'\[\[正文\]\]',
        r'\[\[BODY\]\]',
        r'<CONTENT>',
    ]
    
    for pattern in content_markers:
        if re.search(pattern, text, re.IGNORECASE):
            markers["hasContentMarker"] = True
            break
    
    # 检查是否包含填写说明
    instruction_patterns = [
        r'填写说明',
        r'注意事项',
        r'示例',
        r'样例',
        r'参考',
        r'\[请.*填写\]',
        r'【.*填.*】',
        r'（填）',
    ]
    
    for pattern in instruction_patterns:
        if re.search(pattern, text):
            markers["hasInstructionText"] = True
            break
    
    return markers


def get_block_by_id(blocks: List[Dict[str, Any]], block_id: str) -> Optional[Dict[str, Any]]:
    """根据 blockId 查找 block"""
    for block in blocks:
        if block["blockId"] == block_id:
            return block
    return None


def get_blocks_summary(blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
    """获取 blocks 的统计摘要"""
    total = len(blocks)
    paragraphs = sum(1 for b in blocks if b["type"] == "paragraph")
    tables = sum(1 for b in blocks if b["type"] == "table")
    heading_candidates = sum(1 for b in blocks if b.get("isHeadingCandidate", False))
    has_content_marker = any(b.get("markerFlags", {}).get("hasContentMarker", False) for b in blocks)
    
    return {
        "total": total,
        "paragraphs": paragraphs,
        "tables": tables,
        "heading_candidates": heading_candidates,
        "has_content_marker": has_content_marker,
    }

