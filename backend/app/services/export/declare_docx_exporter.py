"""
申报书 DOCX 导出器
"""
import logging
import os
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入自动生成功能
from app.services.export.docx_exporter import (
    AutoWriteCfg,
    build_project_context_string,
    _is_empty_or_placeholder,
    generate_section_text_by_title,
)

logger = logging.getLogger(__name__)


class DeclareDocxExporter:
    """申报书 DOCX 导出器"""
    
    def __init__(self, dao: Any):
        self.dao = dao
    
    async def export(
        self,
        project_id: str,
        output_dir: Optional[str] = None,
        auto_generate_content: bool = True,
        model_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        导出申报书为 DOCX
        
        Args:
            project_id: 项目ID
            output_dir: 输出目录（可选，默认使用环境变量）
            auto_generate_content: 是否自动生成空内容（默认 True）
            model_id: LLM模型ID（可选）
        
        Returns:
            {
                "document_id": "...",
                "storage_path": "...",
                "filename": "...",
                "file_size": 12345
            }
        """
        # 获取项目信息
        project = self.dao.get_project(project_id)
        if not project:
            raise ValueError(f"Project not found: {project_id}")
        
        project_name = project.get("name", "申报书")
        
        # 获取申报要求
        requirements = self.dao.get_requirements(project_id)
        
        # 获取目录节点
        nodes = self.dao.get_active_directory_nodes(project_id)
        if not nodes:
            raise ValueError("No active directory nodes found")
        
        # 获取章节内容
        sections = self.dao.get_active_sections(project_id)
        sections_by_node_id = {s.get("node_id"): s for s in sections}
        
        # 准备自动生成配置
        cfg = None
        project_context = ""
        content_cache = {}
        
        if auto_generate_content:
            cfg = AutoWriteCfg(
                min_words_h1=1200,
                min_words_h2=800,
                min_words_h3=500,
                min_words_h4=300,
                max_tokens=1600,
                multi_round=True,
            )
            
            # 自动构建项目上下文
            project_context = build_project_context_string(project)
            logger.info(f"[DeclareDocxExporter] 自动生成已启用，项目上下文: {len(project_context)} 字符")
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(project_name, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加摘要（如果有申报要求）
        if requirements:
            data_json = requirements.get("data_json", {})
            summary = data_json.get("summary")
            if summary:
                doc.add_heading("摘要", level=1)
                doc.add_paragraph(summary)
                doc.add_page_break()
        
        # 添加目录（Word 不支持自动目录生成，这里简化处理）
        doc.add_heading("目录", level=1)
        for node in nodes:
            level = node.get("level", 1)
            title = node.get("title", "")
            numbering = node.get("numbering", "")
            indent = "  " * (level - 1)
            doc.add_paragraph(f"{indent}{numbering} {title}")
        doc.add_page_break()
        
        # 添加章节内容
        for node in nodes:
            node_id = node.get("id")
            title = node.get("title", "")
            level = node.get("level", 1)
            numbering = node.get("numbering", "")
            
            # 添加章节标题
            heading_text = f"{numbering} {title}" if numbering else title
            doc.add_heading(heading_text, level=min(level, 3))
            
            # 添加章节内容
            section = sections_by_node_id.get(node_id)
            content_md = section.get("content_md", "") if section else ""
            
            # 判断是否需要自动生成
            if auto_generate_content and _is_empty_or_placeholder(content_md):
                try:
                    logger.info(f"[DeclareDocxExporter] 自动生成内容: title={title}, level={level}")
                    
                    # 调用自动生成函数
                    generated_text = await generate_section_text_by_title(
                        title=title,
                        level=level,
                        project_context=project_context,
                        cfg=cfg,
                        cache=content_cache,
                        model_id=model_id,
                    )
                    
                    # 按空行分段写入 docx
                    paragraphs = [
                        p.strip() 
                        for p in re.split(r"\n{2,}|\r\n{2,}", generated_text) 
                        if p.strip()
                    ]
                    
                    for para in paragraphs:
                        doc.add_paragraph(para)
                    
                    logger.info(f"[DeclareDocxExporter] 自动生成完成: {len(paragraphs)} 个段落, {len(generated_text)} 字符")
                    
                except Exception as e:
                    logger.error(f"[DeclareDocxExporter] 自动生成失败: title={title}, error={e}", exc_info=True)
                    doc.add_paragraph(f"【自动生成内容失败：{str(e)}】")
            
            # 已有内容直接写入（不覆盖）
            elif content_md and not _is_empty_or_placeholder(content_md):
                # 处理 Markdown 内容，支持图片占位符
                self._render_markdown_with_images(doc, content_md, project_id)
            
            # 完全没有内容时的占位符
            else:
                if not auto_generate_content:
                    doc.add_paragraph("（待补充）")
            
            # 添加空行
            doc.add_paragraph("")
        
        # 保存文档
        if not output_dir:
            output_dir = os.getenv("DECLARE_STORAGE_DIR", "./data/declare/documents")
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"{project_name}_申报书.docx"
        storage_path = os.path.join(output_dir, f"{project_id}_{filename}")
        
        doc.save(storage_path)
        file_size = os.path.getsize(storage_path)
        
        # 创建文档记录
        document_id = self.dao.create_document(
            project_id=project_id,
            filename=filename,
            storage_path=storage_path,
            file_size=file_size,
            format="docx",
        )
        
        logger.info(f"[DeclareDocxExporter] Exported document: {storage_path} size={file_size}")
        
        return {
            "document_id": document_id,
            "storage_path": storage_path,
            "filename": filename,
            "file_size": file_size,
        }
    
    def _render_markdown_with_images(self, doc: Document, content_md: str, project_id: str):
        """
        渲染 Markdown 内容到 DOCX，支持图片占位符 {image:filename}
        
        Args:
            doc: DOCX Document 对象
            content_md: Markdown 内容
            project_id: 项目ID（用于查找图片）
        """
        # 按行处理，检测图片占位符
        lines = content_md.split("\n")
        current_para = []
        
        image_pattern = re.compile(r'\{image:([^}]+)\}')
        
        for line in lines:
            stripped = line.strip()
            
            # 检测图片占位符
            match = image_pattern.search(stripped)
            if match:
                # 先写入当前段落（如果有）
                if current_para:
                    doc.add_paragraph("\n".join(current_para).strip())
                    current_para = []
                
                # 插入图片
                image_filename = match.group(1).strip()
                self._insert_image(doc, project_id, image_filename)
                
                # 图片占位符行的其他文本（如果有）
                remaining_text = image_pattern.sub('', stripped).strip()
                if remaining_text:
                    current_para.append(remaining_text)
                
                continue
            
            # 空行分段
            if not stripped:
                if current_para:
                    doc.add_paragraph("\n".join(current_para).strip())
                    current_para = []
                continue
            
            # 普通文本行
            current_para.append(line)
        
        # 写入最后的段落
        if current_para:
            doc.add_paragraph("\n".join(current_para).strip())
    
    def _insert_image(self, doc: Document, project_id: str, filename: str):
        """
        插入图片到 DOCX
        
        Args:
            doc: DOCX Document 对象
            project_id: 项目ID
            filename: 图片文件名
        """
        try:
            # 查找图片资产
            assets = self.dao.list_assets(project_id, kind="image")
            image_asset = None
            
            for asset in assets:
                if asset.get("filename") == filename:
                    image_asset = asset
                    break
            
            if not image_asset:
                logger.warning(f"[DeclareDocxExporter] Image not found: {filename}")
                doc.add_paragraph(f"【图片未找到：{filename}】")
                return
            
            storage_path = image_asset.get("storage_path")
            if not storage_path or not os.path.exists(storage_path):
                logger.warning(f"[DeclareDocxExporter] Image file not exists: {storage_path}")
                doc.add_paragraph(f"【图片文件不存在：{filename}】")
                return
            
            # 插入图片
            doc.add_picture(storage_path, width=Inches(5))
            
            # 添加图片说明（如果有）
            meta_json = image_asset.get("meta_json", {})
            description = meta_json.get("description", "")
            if description:
                caption_para = doc.add_paragraph(description)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption_para.style = 'Caption'
            
            logger.info(f"[DeclareDocxExporter] Inserted image: {filename}")
            
        except Exception as e:
            logger.error(f"[DeclareDocxExporter] Failed to insert image {filename}: {e}", exc_info=True)
            doc.add_paragraph(f"【图片插入失败：{filename}】")

