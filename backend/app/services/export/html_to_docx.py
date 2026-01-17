"""
HtmlToDocxInserter - HTML转Word文档插入器
用于将用户编辑的HTML内容插入到Word文档中
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from html.parser import HTMLParser


class HtmlToDocxInserter:
    """HTML转Word文档插入器"""
    
    @staticmethod
    def insert(dest_doc: Document, html_content: str):
        """
        将HTML内容插入到目标文档
        
        Args:
            dest_doc: 目标Document对象
            html_content: HTML内容字符串
        """
        if not html_content:
            return
        
        # 简单的HTML解析
        # 支持: <p>, <br>, <b>/<strong>, <i>/<em>, <ul>/<ol>/<li>, <table>
        
        # 移除HTML标签外的多余空白
        html_content = html_content.strip()
        
        # ✨ 优先处理表格（完整提取）
        table_pattern = re.compile(r'<table[^>]*>.*?</table>', re.DOTALL | re.IGNORECASE)
        tables = []
        
        def table_placeholder(match):
            """用占位符替换表格，并保存表格内容"""
            tables.append(match.group(0))
            return f'__TABLE_{len(tables)-1}__'
        
        # 将表格替换为占位符
        html_without_tables = table_pattern.sub(table_placeholder, html_content)
        
        # 分割成段落（按 <p> 标签）
        paragraphs = re.split(r'<p[^>]*>|</p>', html_without_tables)
        
        for para_html in paragraphs:
            para_html = para_html.strip()
            if not para_html:
                continue
            
            # ✅ 检查段落中是否包含表格占位符（使用正则表达式精确匹配）
            table_placeholders = re.findall(r'__TABLE_(\d+)__', para_html)
            
            if table_placeholders:
                # 如果段落只包含表格占位符（可能有多个，用换行分隔）
                # 处理每个表格占位符
                for table_idx_str in table_placeholders:
                    table_idx = int(table_idx_str)
                    if 0 <= table_idx < len(tables):
                        HtmlToDocxInserter._insert_table(dest_doc, tables[table_idx])
                
                # 如果段落中还有其他内容，移除表格占位符后插入
                remaining_text = re.sub(r'__TABLE_\d+__', '', para_html).strip()
                if remaining_text:
                    HtmlToDocxInserter._insert_paragraph(dest_doc, remaining_text)
            # 检查是否是列表项
            elif '<li>' in para_html or '<ul>' in para_html or '<ol>' in para_html:
                HtmlToDocxInserter._insert_list(dest_doc, para_html)
            else:
                HtmlToDocxInserter._insert_paragraph(dest_doc, para_html)
    
    @staticmethod
    def _insert_paragraph(dest_doc: Document, para_html: str):
        """插入一个段落"""
        para = dest_doc.add_paragraph()
        
        # 处理换行符 <br>
        segments = re.split(r'<br\s*/?>', para_html)
        
        for i, segment in enumerate(segments):
            if i > 0:
                # 添加换行
                para.add_run('\n')
            
            HtmlToDocxInserter._insert_runs(para, segment)
    
    @staticmethod
    def _insert_runs(para, html_segment: str):
        """解析HTML片段并插入runs"""
        # 简单状态机解析
        text = ""
        bold = False
        italic = False
        underline = False
        
        i = 0
        while i < len(html_segment):
            # 检查标签
            if html_segment[i] == '<':
                # 保存当前文本
                if text:
                    run = para.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    text = ""
                
                # 查找标签结束
                tag_end = html_segment.find('>', i)
                if tag_end == -1:
                    break
                
                tag = html_segment[i:tag_end+1].lower()
                
                # 解析标签
                if tag in ('<b>', '<strong>'):
                    bold = True
                elif tag in ('</b>', '</strong>'):
                    bold = False
                elif tag in ('<i>', '<em>'):
                    italic = True
                elif tag in ('</i>', '</em>'):
                    italic = False
                elif tag in ('<u>',):
                    underline = True
                elif tag in ('</u>',):
                    underline = False
                
                i = tag_end + 1
            else:
                # 普通字符
                text += html_segment[i]
                i += 1
        
        # 插入剩余文本
        if text:
            # 解码HTML实体
            text = HtmlToDocxInserter._decode_html_entities(text)
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
    
    @staticmethod
    def _insert_list(dest_doc: Document, list_html: str):
        """插入列表"""
        # 提取列表项
        items = re.findall(r'<li[^>]*>(.*?)</li>', list_html, re.DOTALL)
        
        # 判断是有序还是无序
        is_ordered = '<ol' in list_html.lower()
        
        for i, item_html in enumerate(items):
            item_html = item_html.strip()
            if not item_html:
                continue
            
            # 移除内层HTML标签（简化处理）
            item_text = re.sub(r'<[^>]+>', '', item_html)
            item_text = HtmlToDocxInserter._decode_html_entities(item_text)
            
            # 添加前缀
            if is_ordered:
                prefix = f"{i+1}. "
            else:
                prefix = "• "
            
            para = dest_doc.add_paragraph()
            para.add_run(prefix + item_text)
    
    @staticmethod
    def _insert_table(dest_doc: Document, table_html: str):
        """
        插入HTML表格到Word文档
        
        Args:
            dest_doc: 目标Document对象
            table_html: 表格HTML字符串
        """
        # 提取所有行
        rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
        
        if not rows_html:
            return
        
        # 解析每一行，获取最大列数
        parsed_rows = []
        max_cols = 0
        
        for row_html in rows_html:
            # 提取单元格（th或td）
            cells = re.findall(r'<(th|td)[^>]*>(.*?)</\1>', row_html, re.DOTALL | re.IGNORECASE)
            row_data = []
            
            for tag, cell_content in cells:
                # 清理cell内容
                cell_text = re.sub(r'<br\s*/?>', '\n', cell_content)  # 保留换行
                cell_text = re.sub(r'<[^>]+>', '', cell_text)  # 移除其他标签
                cell_text = HtmlToDocxInserter._decode_html_entities(cell_text)
                cell_text = cell_text.strip()
                
                row_data.append({
                    'text': cell_text,
                    'is_header': tag.lower() == 'th'
                })
            
            if row_data:
                parsed_rows.append(row_data)
                max_cols = max(max_cols, len(row_data))
        
        if not parsed_rows or max_cols == 0:
            return
        
        # 创建Word表格
        table = dest_doc.add_table(rows=len(parsed_rows), cols=max_cols)
        table.style = 'Table Grid'  # 使用带边框的表格样式
        
        # 填充数据
        for row_idx, row_data in enumerate(parsed_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < max_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_data['text']
                    
                    # 表头加粗
                    if cell_data['is_header']:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # 设置单元格边框
                    HtmlToDocxInserter._set_cell_border(cell)
    
    @staticmethod
    def _set_cell_border(cell, **kwargs):
        """
        为表格单元格设置边框
        
        Args:
            cell: docx表格单元格对象
            kwargs: 边框设置参数
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # 创建边框元素
        tcBorders = OxmlElement('w:tcBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 边框宽度
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 黑色边框
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    @staticmethod
    def _decode_html_entities(text: str) -> str:
        """解码HTML实体"""
        replacements = {
            '&nbsp;': ' ',
            '&lt;': '<',
            '&gt;': '>',
            '&amp;': '&',
            '&quot;': '"',
            '&#39;': "'",
        }
        for entity, char in replacements.items():
            text = text.replace(entity, char)
        return text
