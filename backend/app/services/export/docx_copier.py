"""
DocxBodyElementCopier - Word 文档元素拷贝工具
用于在导出时保真拷贝源文档中的段落和表格
"""
from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxBodyElementCopier:
    """Word 文档元素拷贝器"""
    
    @staticmethod
    def copy_range(
        src_docx_path: str,
        start_index: int,
        end_index: int,
        dest_doc: Document
    ):
        """
        从源docx文件拷贝指定范围的元素到目标文档
        
        Args:
            src_docx_path: 源docx文件路径
            start_index: 起始索引（包含）
            end_index: 结束索引（包含）
            dest_doc: 目标Document对象
        """
        # 读取源文档
        src_doc = Document(src_docx_path)
        src_elements = list(src_doc.element.body)
        
        # 验证索引范围
        if start_index < 0 or end_index >= len(src_elements) or start_index > end_index:
            return
        
        # 拷贝元素
        for i in range(start_index, end_index + 1):
            element = src_elements[i]
            
            if isinstance(element, CT_P):
                # 拷贝段落
                src_para = Paragraph(element, src_doc)
                DocxBodyElementCopier._copy_paragraph(src_para, dest_doc)
            
            elif isinstance(element, CT_Tbl):
                # 拷贝表格
                src_table = Table(element, src_doc)
                DocxBodyElementCopier._copy_table(src_table, dest_doc)
    
    @staticmethod
    def _copy_paragraph(src_para: Paragraph, dest_doc: Document):
        """拷贝段落（尽可能保留格式）"""
        dest_para = dest_doc.add_paragraph()
        
        # 拷贝段落样式
        if src_para.style:
            try:
                dest_para.style = src_para.style.name
            except:
                pass
        
        # 拷贝对齐方式
        if src_para.alignment:
            dest_para.alignment = src_para.alignment
        
        # 拷贝runs
        for src_run in src_para.runs:
            dest_run = dest_para.add_run(src_run.text)
            
            # 拷贝字体属性
            try:
                dest_run.bold = src_run.bold
                dest_run.italic = src_run.italic
                dest_run.underline = src_run.underline
                
                if src_run.font.size:
                    dest_run.font.size = src_run.font.size
                
                if src_run.font.name:
                    dest_run.font.name = src_run.font.name
                
                if src_run.font.color and src_run.font.color.rgb:
                    dest_run.font.color.rgb = src_run.font.color.rgb
            except:
                pass
    
    @staticmethod
    def _copy_table(src_table: Table, dest_doc: Document):
        """拷贝表格（尽可能保留格式）"""
        rows = len(src_table.rows)
        cols = len(src_table.columns) if src_table.rows else 0
        
        if rows == 0 or cols == 0:
            return
        
        # 创建表格
        dest_table = dest_doc.add_table(rows=rows, cols=cols)
        
        # 尝试应用表格样式
        try:
            if src_table.style:
                dest_table.style = src_table.style.name
        except:
            pass
        
        # 拷贝单元格内容
        for i, src_row in enumerate(src_table.rows):
            for j, src_cell in enumerate(src_row.cells):
                try:
                    dest_cell = dest_table.rows[i].cells[j]
                    
                    # 清空目标单元格
                    for para in dest_cell.paragraphs:
                        para.clear()
                    
                    # 拷贝源单元格的所有段落
                    for k, src_para in enumerate(src_cell.paragraphs):
                        if k == 0:
                            # 使用已有的第一个段落
                            dest_para = dest_cell.paragraphs[0]
                        else:
                            # 添加新段落
                            dest_para = dest_cell.add_paragraph()
                        
                        # 拷贝段落内容
                        for src_run in src_para.runs:
                            dest_run = dest_para.add_run(src_run.text)
                            try:
                                dest_run.bold = src_run.bold
                                dest_run.italic = src_run.italic
                                if src_run.font.size:
                                    dest_run.font.size = src_run.font.size
                            except:
                                pass
                except:
                    continue
