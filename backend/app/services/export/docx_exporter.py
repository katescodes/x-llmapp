"""
Word 文档导出器
使用模板母版生成包含目录树的 Word 文档
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from lxml import etree

from .tree_builder import DirNode
from .docx_template_loader import PageVariant, SectPrototype

logger = logging.getLogger(__name__)


# ============================================================================
# 自动生成申报书内容的辅助函数
# ============================================================================

@dataclass
class AutoWriteCfg:
    """自动写作配置：字数越多越好，按标题层级给更高下限"""
    min_words_h1: int = 1200
    min_words_h2: int = 800
    min_words_h3: int = 500
    min_words_h4: int = 300
    max_tokens: int = 1600  # 单次调用上限（避免超时）
    multi_round: bool = True  # 是否分多次生成（字数越多越好的工程化做法）


def _is_empty_or_placeholder(content: Optional[str]) -> bool:
    """
    判断内容是否为空或占位符
    
    Args:
        content: 内容字符串
        
    Returns:
        True 如果为空或占位符
    """
    if not content:
        return True
    
    content = content.strip()
    
    # 检查是否为常见占位符
    placeholders = [
        "【填写】",
        "【待补】",
        "【待填写】",
        "[填写]",
        "[待补]",
        "待填写",
        "待补充",
        "TODO",
        "TBD",
    ]
    
    return content in placeholders or len(content) < 5


def build_project_context_string(project_data: Optional[Dict] = None) -> str:
    """
    构建项目上下文字符串（用于 LLM 生成时提供背景信息）
    
    从项目数据中提取关键信息，拼接成上下文字符串。
    
    Args:
        project_data: 项目数据字典，可能包含：
            - name: 项目名称
            - company: 企业名称
            - summary: 项目摘要
            - patents: 专利清单
            - devices: 设备清单
            - achievements: 成果清单
            - meta_json: 其他元数据
            
    Returns:
        上下文字符串
    """
    if not project_data:
        return ""
    
    lines = []
    
    # 1. 项目名称
    if project_data.get("name"):
        lines.append(f"【项目名称】{project_data['name']}")
    
    # 2. 企业信息
    if project_data.get("company"):
        lines.append(f"【企业名称】{project_data['company']}")
    
    # 3. 项目摘要
    if project_data.get("summary"):
        lines.append(f"【项目摘要】{project_data['summary']}")
    
    # 4. 从 meta_json 中提取信息
    meta = project_data.get("meta_json", {})
    
    if meta.get("industry"):
        lines.append(f"【所属行业】{meta['industry']}")
    
    if meta.get("budget"):
        lines.append(f"【预算金额】{meta['budget']}")
    
    if meta.get("duration"):
        lines.append(f"【建设周期】{meta['duration']}")
    
    # 5. 专利清单（如果有）
    patents = project_data.get("patents", [])
    if patents:
        lines.append(f"【专利数量】{len(patents)} 项")
        # 只列举前 3 个专利名称
        for i, patent in enumerate(patents[:3], 1):
            if isinstance(patent, dict) and patent.get("name"):
                lines.append(f"  {i}. {patent['name']}")
    
    # 6. 设备清单（如果有）
    devices = project_data.get("devices", [])
    if devices:
        lines.append(f"【主要设备】{len(devices)} 项")
    
    # 7. 成果清单（如果有）
    achievements = project_data.get("achievements", [])
    if achievements:
        lines.append(f"【已有成果】{len(achievements)} 项")
    
    # 如果没有任何信息，返回提示
    if not lines:
        return "（暂无项目上下文信息）"
    
    return "\n".join(lines)


def _target_min_words(level: int, cfg: AutoWriteCfg) -> int:
    """根据标题层级返回目标最小字数"""
    if level <= 1:
        return cfg.min_words_h1
    if level == 2:
        return cfg.min_words_h2
    if level == 3:
        return cfg.min_words_h3
    return cfg.min_words_h4


def _infer_section_style(title: str) -> str:
    """根据标题关键词给写作侧重点，提升'像申报书'的命中率。"""
    t = title.strip()
    rules = [
        (r"概况|背景|意义|必要性|总体情况", "写背景现状、政策依据、问题痛点、建设必要性与总体目标。"),
        (r"目标|指标|成效|效益|预期", "写可量化目标指标（效率/质量/成本/能耗/交付/良率等）+对标+预期成效。"),
        (r"建设内容|建设方案|技术方案|总体架构|架构", "写总体架构（业务/数据/应用/安全）+关键系统/集成关系+技术路线。"),
        (r"场景|应用|业务流程|流程", "写典型场景：现状→改造→系统支撑→数据闭环→指标提升。至少5-8个小点。"),
        (r"组织|保障|管理|制度|运维|安全", "写组织架构、制度流程、数据治理、安全与运维保障、风险与应对。"),
        (r"投资|预算|资金|费用", "写投资构成（软硬件/服务/集成/运维）+测算口径+资金来源与使用计划。"),
        (r"进度|计划|里程碑|实施", "写阶段划分、里程碑、交付物、验收与持续改进机制。"),
    ]
    for pattern, hint in rules:
        if re.search(pattern, t):
            return hint
    return "写成标准政府项目申报口径：现状→问题→方案→系统与数据→实施→成效与指标→保障。"


async def generate_section_text_by_title(
    *,
    title: str,
    level: int,
    project_context: str,
    cfg: AutoWriteCfg,
    cache: dict,
    model_id: Optional[str] = None,
) -> str:
    """
    生成申报书某一标题下正文：字数尽量多、分段清晰、口径像申报材料。
    必须：不编造企业具体数字/资质/专利编号（没有就用【待补】占位）。
    
    采用多轮生成策略：避免单次超时，同时实现"字数越多越好"。
    
    Args:
        title: 章节标题
        level: 标题层级
        project_context: 项目上下文（来自已解析/已填充的数据）
        cfg: 自动写作配置
        cache: 缓存字典（避免重复生成）
        model_id: LLM模型ID（可选）
    
    Returns:
        生成的正文内容
    """
    key = f"L{level}:{title.strip()}"
    if key in cache:
        return cache[key]

    min_words = _target_min_words(level, cfg)
    style_hint = _infer_section_style(title)

    try:
        # 使用项目现有的 LLM 调用方式
        from app.services.llm_model_store import get_llm_store
        
        store = get_llm_store()
        model = None
        
        # 如果指定了 model_id，使用指定的模型
        if model_id:
            try:
                model = store.get_model(model_id)
            except Exception as e:
                logger.warning(f"无法获取指定模型 {model_id}，使用默认模型: {e}")
        
        # 否则使用默认模型
        if not model:
            model = store.get_default_model()
        
        if not model:
            logger.error("未找到可用的 LLM 模型")
            return "【错误：未配置 LLM 模型，无法生成内容】"
        
        from app.services.llm_client import generate_answer_with_model
        
        # 根据配置决定是否采用多轮生成
        if cfg.multi_round and min_words >= 800:
            # 多轮生成：分 2-3 次，避免单次超时
            text = await _generate_multi_round(
                title=title,
                level=level,
                project_context=project_context,
                style_hint=style_hint,
                min_words=min_words,
                cfg=cfg,
                model=model,
            )
        else:
            # 单次生成
            text = await _generate_single_round(
                title=title,
                level=level,
                project_context=project_context,
                style_hint=style_hint,
                min_words=min_words,
                cfg=cfg,
                model=model,
            )
        
        text = text.strip()
        
        # 缓存结果
        cache[key] = text
        return text
        
    except Exception as e:
        logger.error(f"生成章节内容失败: title={title}, error={e}", exc_info=True)
        return f"【生成内容失败：{str(e)}】"


async def _generate_single_round(
    *,
    title: str,
    level: int,
    project_context: str,
    style_hint: str,
    min_words: int,
    cfg: AutoWriteCfg,
    model,
) -> str:
    """单次生成"""
    from app.services.llm_client import generate_answer_with_model
    
    system = (
        "你是省级项目/工厂申报材料撰写专家，擅长把标题扩写成规范、正式、可落地的申报书正文。"
        "写作要求：只基于给定上下文，不得虚构企业名称、金额、日期、专利号、检测报告号等；"
        "若缺信息，用【待补：xxx】占位；文风正式、条理清晰，尽量多写。"
    )

    user = f"""
【章节标题】{title}
【标题层级】H{level}

【写作侧重点】
{style_hint}

【可用上下文（来自已解析/已填充的数据，可能为空）】
{project_context}

【输出要求】
1) 直接输出正文，不要再写标题。
2) 至少 {min_words} 字（中文），分为 6~12 段，每段 2~5 句；可使用（1）（2）（3）等条目。
3) 若需要数字指标但上下文未给出，用【待补：指标口径/数值】占位。
4) 禁止输出"作为AI/无法"等元话术。
"""

    text = await generate_answer_with_model(
        system_prompt=system,
        user_message=user,
        history=[],
        model=model,
        overrides={
            "max_tokens": cfg.max_tokens,
            "temperature": 0.7,
        }
    )
    
    return text


async def _generate_multi_round(
    *,
    title: str,
    level: int,
    project_context: str,
    style_hint: str,
    min_words: int,
    cfg: AutoWriteCfg,
    model,
) -> str:
    """
    多轮生成：分 2-3 次调用 LLM，最后拼接
    这样可以实现"字数越多越好"，同时避免单次超时
    """
    from app.services.llm_client import generate_answer_with_model
    
    system = (
        "你是省级项目/工厂申报材料撰写专家，擅长把标题扩写成规范、正式、可落地的申报书正文。"
        "写作要求：只基于给定上下文，不得虚构企业名称、金额、日期、专利号、检测报告号等；"
        "若缺信息，用【待补：xxx】占位；文风正式、条理清晰。"
    )
    
    parts = []
    
    # 第 1 轮：主体正文
    user1 = f"""
【章节标题】{title}
【标题层级】H{level}

【写作侧重点】
{style_hint}

【可用上下文】
{project_context}

【本轮任务】
写该章节的**主体正文**（背景、现状、问题、必要性等核心内容），至少 {min_words // 2} 字，分 4~6 段。
直接输出正文，不要标题，不要"作为AI"等元话术。
"""
    
    try:
        text1 = await generate_answer_with_model(
            system_prompt=system,
            user_message=user1,
            history=[],
            model=model,
            overrides={
                "max_tokens": cfg.max_tokens,
                "temperature": 0.7,
            }
        )
        parts.append(text1.strip())
        logger.debug(f"第 1 轮生成完成: {len(text1)} 字符")
    except Exception as e:
        logger.error(f"第 1 轮生成失败: {e}", exc_info=True)
        parts.append(f"【第 1 轮生成失败】")
    
    # 第 2 轮：实施路径、保障措施
    user2 = f"""
【章节标题】{title}
【标题层级】H{level}

【写作侧重点】
{style_hint}

【可用上下文】
{project_context}

【本轮任务】
补充该章节的**实施路径、保障措施、组织管理**等内容，至少 {min_words // 3} 字，分 3~5 段。
直接输出正文，不要标题，不要重复前面的内容。
"""
    
    try:
        text2 = await generate_answer_with_model(
            system_prompt=system,
            user_message=user2,
            history=[],
            model=model,
            overrides={
                "max_tokens": cfg.max_tokens,
                "temperature": 0.7,
            }
        )
        parts.append(text2.strip())
        logger.debug(f"第 2 轮生成完成: {len(text2)} 字符")
    except Exception as e:
        logger.error(f"第 2 轮生成失败: {e}", exc_info=True)
        parts.append(f"【第 2 轮生成失败】")
    
    # 第 3 轮：指标对标、创新点（仅对重要章节）
    if level <= 2 and min_words >= 1000:
        user3 = f"""
【章节标题】{title}
【标题层级】H{level}

【写作侧重点】
{style_hint}

【可用上下文】
{project_context}

【本轮任务】
补充该章节的**目标指标、对标分析、创新点、预期成效**等内容，至少 {min_words // 4} 字，分 2~4 段。
直接输出正文，不要标题，不要重复前面的内容。
若上下文缺具体指标，用【待补：指标名称/数值】占位。
"""
        
        try:
            text3 = await generate_answer_with_model(
                system_prompt=system,
                user_message=user3,
                history=[],
                model=model,
                overrides={
                    "max_tokens": cfg.max_tokens,
                    "temperature": 0.7,
                }
            )
            parts.append(text3.strip())
            logger.debug(f"第 3 轮生成完成: {len(text3)} 字符")
        except Exception as e:
            logger.error(f"第 3 轮生成失败: {e}", exc_info=True)
            parts.append(f"【第 3 轮生成失败】")
    
    # 拼接所有部分（用双换行分隔）
    full_text = "\n\n".join(p for p in parts if p and not p.startswith("【") and not p.endswith("失败】"))
    
    logger.info(f"多轮生成完成: title={title}, 总字符数={len(full_text)}, 轮数={len(parts)}")
    
    return full_text


def is_toc_title(text: str) -> bool:
    """
    判断文本是否为目录标题
    
    Args:
        text: 待判断的文本
        
    Returns:
        True 如果是目录标题
    """
    if not text:
        return False
    
    # 去除空白和全角空格
    normalized = text.strip().replace(" ", "").replace("\u3000", "").lower()
    
    # 匹配常见的目录标题
    toc_keywords = ["目录", "目录结构", "toc", "投标文件目录", "tableofcontents"]
    return normalized in toc_keywords


def _get_toc_title_style(
    heading_style_map: Optional[Dict[int, str]],
    doc: Document
) -> str:
    """
    获取目录标题的样式名称（确保不是带编号的 Heading 样式）
    
    优先级：
    1. heading_style_map 中的 "toc_title" 或 0 级
    2. 文档中存在的 "TOC Heading" / "TOC Title" / "目录标题"
    3. 回退到 "Normal"
    
    Args:
        heading_style_map: 样式映射（可能包含 toc_title 或 0 级）
        doc: Document 对象，用于检查样式是否存在
        
    Returns:
        样式名称
    """
    # 1. 优先使用 heading_style_map 中的 toc_title 或 0 级
    if heading_style_map:
        if "toc_title" in heading_style_map:
            return heading_style_map["toc_title"]
        if 0 in heading_style_map:
            return heading_style_map[0]
    
    # 2. 尝试常见的 TOC 专用样式
    candidates = ["TOC Heading", "TOC Title", "目录标题", "目录"]
    for candidate in candidates:
        try:
            _ = doc.styles[candidate]
            logger.info(f"使用目录标题样式: {candidate}")
            return candidate
        except KeyError:
            continue
    
    # 3. 回退到 Normal（确保不会是带编号的 Heading）
    logger.warning("未找到目录标题专用样式，使用 Normal")
    return "Normal"


def _get_heading_style_name(
    level: int,
    heading_style_map: Optional[Dict[int, str]] = None
) -> str:
    """
    获取指定层级的标题样式名称
    
    优先使用 heading_style_map 中的样式，如果没有则回退到默认的 Heading {level}
    
    Args:
        level: 层级（1~9）
        heading_style_map: 标题样式映射（level -> style_name），来自模板配置
        
    Returns:
        样式名称
    """
    lv = max(1, min(level, 9))
    
    # 优先使用自定义样式映射
    if heading_style_map and lv in heading_style_map:
        return heading_style_map[lv]
    
    # 回退到默认 Heading 1~9
    return f"Heading {lv}"


def _clear_body_keep_sectpr(doc: Document, apply_assets: Optional[Dict[str, Any]] = None) -> None:
    """
    清空文档正文内容，智能保留模板底板（封面、声明页等）
    
    优先级：
    1. 如果有 applyAssets.keepPlan：根据 LLM 分析结果精细化保留/删除
    2. 如果有 [[CONTENT]] 标记：保留标记前的内容
    3. 否则：智能识别内容开始位置（关键词匹配）
    4. 兜底：清空所有内容
    
    Args:
        doc: Document 对象（会原地修改）
        apply_assets: LLM 分析的保留/删除计划，包含：
            - anchors: 内容插入点列表
            - keepPlan.keepBlockIds: 应保留的块ID列表
            - keepPlan.deleteBlockIds: 应删除的块ID列表
    """
    body = doc._element.body
    
    # 保存最后的 sectPr
    last_sectpr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            last_sectpr = etree.fromstring(etree.tostring(child))
    
    # 方案1：使用 LLM 分析的 applyAssets（最智能）
    if apply_assets and isinstance(apply_assets, dict):
        keep_plan = apply_assets.get("keepPlan", {})
        keep_block_ids = keep_plan.get("keepBlockIds", []) if keep_plan else []
        delete_block_ids = keep_plan.get("deleteBlockIds", []) if keep_plan else []
        anchors = apply_assets.get("anchors", [])
        
        # 如果有 keepPlan，使用精细化的保留/删除策略
        if keep_block_ids or delete_block_ids:
            logger.info(f"使用 LLM keepPlan: keep={len(keep_block_ids)}, delete={len(delete_block_ids)}")
            
            # 将块ID转换为索引（假设块ID为 "b0", "b1", "b2" ...）
            keep_indices = set()
            delete_indices = set()
            
            for block_id in keep_block_ids:
                if isinstance(block_id, str) and block_id.startswith("b"):
                    try:
                        idx = int(block_id[1:])
                        keep_indices.add(idx)
                    except ValueError:
                        pass
            
            for block_id in delete_block_ids:
                if isinstance(block_id, str) and block_id.startswith("b"):
                    try:
                        idx = int(block_id[1:])
                        delete_indices.add(idx)
                    except ValueError:
                        pass
            
            logger.info(f"keep_indices 范围: {min(keep_indices) if keep_indices else 'N/A'} - {max(keep_indices) if keep_indices else 'N/A'}")
            logger.info(f"delete_indices 范围: {min(delete_indices) if delete_indices else 'N/A'} - {max(delete_indices) if delete_indices else 'N/A'}")
            
            # 遍历body元素，删除应该删除的块
            children_list = list(body)
            block_idx = 0  # 当前块索引（只计算段落和表格）
            
            for child in children_list:
                # 只处理段落和表格（跳过 sectPr 等）
                if child.tag in [qn("w:p"), qn("w:tbl")]:
                    # 如果这个块在 delete_indices 中，删除它
                    if block_idx in delete_indices:
                        logger.debug(f"删除块 b{block_idx}")
                        body.remove(child)
                    else:
                        logger.debug(f"保留块 b{block_idx}")
                    block_idx += 1
            
            # 删除所有TOC域（Word目录域）
            _remove_all_toc_fields(doc)
            
            # ✅ 额外清理：删除所有使用TOC样式的段落（这些是模板的旧目录条目）
            _remove_toc_style_paragraphs(doc)
            
            # 恢复 sectPr
            if last_sectpr is not None:
                body.append(last_sectpr)
            
            logger.info(f"根据 keepPlan 完成清理，保留了 {len(keep_indices)} 个块")
            return
        
        # 如果有锚点但没有 keepPlan，尝试使用锚点逻辑（备用）
        if anchors:
            first_anchor = anchors[0] if isinstance(anchors, list) else anchors
            anchor_block_id = first_anchor.get("blockId") if isinstance(first_anchor, dict) else None
            
            if anchor_block_id:
                logger.info(f"使用 LLM 分析的锚点: blockId={anchor_block_id}")
                # 找到锚点对应的元素索引（简化版：根据顺序估计）
                anchor_reason = first_anchor.get("reason", "") if isinstance(first_anchor, dict) else ""
                logger.info(f"锚点原因: {anchor_reason}")
                
                # 尝试从 reason 中提取关键词
                for idx, child in enumerate(list(body)):
                    if child.tag == qn("w:p"):
                        t_elements = child.findall('.//{%s}' % qn('w:t'))
                        paragraph_text = ''.join([t.text for t in t_elements if t.text])
                        
                        # 检查是否包含 [[CONTENT]] 标记
                        if "[[CONTENT]]" in paragraph_text:
                            logger.info(f"通过锚点找到 [[CONTENT]] 标记，位于第 {idx} 个元素")
                            children_to_remove = list(body)[idx:]
                            for child in children_to_remove:
                                body.remove(child)
                            if last_sectpr is not None:
                                body.append(last_sectpr)
                            return
    
    # 方案2：查找 [[CONTENT]] 标记（明确标记）
    content_marker_found = False
    content_marker_index = -1
    
    for idx, child in enumerate(list(body)):
        if child.tag == qn("w:p"):  # 段落
            # 获取段落文本
            t_elements = child.findall('.//{%s}' % qn('w:t'))
            paragraph_text = ''.join([t.text for t in t_elements if t.text])
            
            if "[[CONTENT]]" in paragraph_text:
                content_marker_found = True
                content_marker_index = idx
                logger.info(f"在第 {idx} 个元素中找到 [[CONTENT]] 标记")
                break
    
    if content_marker_found:
        # 删除 [[CONTENT]] 标记及其后的所有内容
        logger.info(f"保留前 {content_marker_index} 个元素（封面等），删除 [[CONTENT]] 及其后的内容")
        children_to_remove = list(body)[content_marker_index:]
        for child in children_to_remove:
            body.remove(child)
        if last_sectpr is not None:
            body.append(last_sectpr)
        return
    
    # 方案3：智能识别内容开始位置（关键词匹配）
    content_start_keywords = ["投标函", "目录", "第一章", "第1章", "一、", "1.", "1、"]
    
    for idx, child in enumerate(list(body)):
        if child.tag == qn("w:p"):  # 段落
            if idx <= 5:  # 至少保留前5个段落（封面信息）
                continue
            
            # 获取段落文本和样式
            t_elements = child.findall('.//{%s}' % qn('w:t'))
            paragraph_text = ''.join([t.text for t in t_elements if t.text])
            
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'))
                    # 检查是否是标题样式
                    if style_val and ('标题' in style_val or 'Heading' in style_val):
                        # 检查是否包含内容开始关键词
                        for keyword in content_start_keywords:
                            if keyword in paragraph_text:
                                content_marker_found = True
                                content_marker_index = idx
                                logger.info(f"智能识别内容开始位置：第 {idx} 个元素，标题为「{paragraph_text[:30]}」")
                                break
            if content_marker_found:
                break
    
    if content_marker_found:
        # 删除识别出的内容开始位置及其后的所有内容
        logger.info(f"保留前 {content_marker_index} 个元素（封面等），删除智能识别的内容区")
        children_to_remove = list(body)[content_marker_index:]
        for child in children_to_remove:
            body.remove(child)
        if last_sectpr is not None:
            body.append(last_sectpr)
        return
    
    # 方案4：兜底 - 清空所有内容
    logger.warning("未找到任何内容标记或锚点，清空所有正文内容（保留 sectPr）")
    for child in list(body):
        body.remove(child)
    
    # 恢复 sectPr
    if last_sectpr is not None:
        body.append(last_sectpr)


def _remove_all_toc_fields(doc: Document) -> int:
    """
    删除文档中所有的TOC域（Word目录域）和SDT中的旧目录
    
    Args:
        doc: Document 对象
        
    Returns:
        删除的TOC域数量
    """
    removed_count = 0
    body = doc._element.body
    
    # ✅ 新增：删除所有SDT（结构化文档标签）中的旧目录
    # 模板的旧目录通常在SDT中，python-docx看不到但Word会渲染
    sdts = body.findall(qn('w:sdt'))
    for sdt in sdts:
        # 检查SDT中的内容是否包含旧目录
        sdt_texts = []
        for elem in sdt.iter():
            if elem.tag == qn('w:t'):
                sdt_texts.append(elem.text or '')
        
        sdt_content = ''.join(sdt_texts)
        # 如果包含"一、"、"二、"等，认为是旧目录
        if '一、' in sdt_content or '二、' in sdt_content or 'TOC' in sdt_content:
            logger.info(f"删除SDT中的旧目录，包含 {len(sdt_texts)} 个文本节点")
            body.remove(sdt)
            removed_count += 1
    
    # 查找所有段落中的 fldSimple 元素（TOC域）
    for para in body.findall(qn('w:p')):
        # 查找 w:fldSimple 元素
        fld_simples = para.findall(qn('w:fldSimple'))
        for fld in fld_simples:
            instr = fld.get(qn('w:instr'), '')
            # 检查是否是TOC域
            if 'TOC' in instr.upper():
                logger.info(f"删除TOC域: {instr}")
                para.remove(fld)
                removed_count += 1
        
        # 查找 w:fldChar（复杂域的开始/结束标记）
        # TOC域也可能以 fldChar 的形式存在
        fld_chars = para.findall(qn('w:fldChar'))
        if fld_chars:
            # 检查是否包含TOC指令
            for run in para.findall(qn('w:r')):
                instr_text = run.find(qn('w:instrText'))
                if instr_text is not None and 'TOC' in (instr_text.text or '').upper():
                    # 删除整个段落（因为TOC域可能跨越多个run）
                    logger.info(f"删除包含TOC域的段落")
                    body.remove(para)
                    removed_count += 1
                    break
    
    if removed_count > 0:
        logger.info(f"共删除 {removed_count} 个TOC域")
    else:
        logger.info("未发现TOC域")
    
    return removed_count


def _remove_toc_style_paragraphs(doc: Document) -> int:
    """
    删除文档中所有使用TOC样式的段落（模板的旧目录条目）
    
    这些段落通常是模板原有的目录列表，应该被删除，
    系统会重新生成新的目录。
    
    Args:
        doc: Document 对象
        
    Returns:
        删除的段落数量
    """
    removed_count = 0
    paras_to_remove = []
    
    # 遍历所有段落，找出使用TOC样式的段落
    for para in doc.paragraphs:
        if para.style and 'toc' in para.style.name.lower():
            # 记录要删除的段落
            paras_to_remove.append(para)
            logger.debug(f"标记删除TOC样式段落: [{para.style.name}] {para.text[:50]}")
    
    # 删除段落（使用XML API）
    for para in paras_to_remove:
        para._element.getparent().remove(para._element)
        removed_count += 1
    
    if removed_count > 0:
        logger.info(f"共删除 {removed_count} 个TOC样式段落（模板旧目录）")
    else:
        logger.info("未发现TOC样式段落")
    
    return removed_count


def _remove_template_sample_headings(doc: Document) -> int:
    """
    删除模板中的示例标题（如"投标函"、"法定代表人授权书"等）
    
    这些标题通常是模板文件中的示例章节标题，应该被删除，
    系统会插入实际项目的章节。
    
    判断标准：
    1. 使用标题样式（Heading或标题）
    2. 出现在"目录"标题之前
    3. 不是"目录"本身
    
    Args:
        doc: Document 对象
        
    Returns:
        删除的段落数量
    """
    removed_count = 0
    paras_to_remove = []
    
    # 找到"目录"标题的位置
    toc_index = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ["目录", "目 录", "TOC", "TABLE OF CONTENTS"]:
            toc_index = i
            logger.info(f"⚠️ 找到目录标题位置: 段落{i}")
            break
    
    if toc_index == -1:
        logger.warning("⚠️ 未找到目录标题，跳过删除模板示例标题")
        return 0
    
    # 删除"目录"之前的所有标题段落（这些是模板的示例标题）
    for i, para in enumerate(doc.paragraphs):
        # 只处理"目录"之前的段落
        if i >= toc_index:
            break
        
        # 检查是否是标题样式
        if para.style and ('heading' in para.style.name.lower() or '标题' in para.style.name):
            text = para.text.strip()
            # 排除封面部分的短文本
            if len(text) > 1:
                paras_to_remove.append(para)
                logger.info(f"⚠️ 标记删除模板示例标题 [{i}]: [{para.style.name}] {text[:50]}")
    
    # 删除段落
    for para in paras_to_remove:
        para._element.getparent().remove(para._element)
        removed_count += 1
    
    if removed_count > 0:
        logger.info(f"共删除 {removed_count} 个模板示例标题")
    else:
        logger.info("未发现模板示例标题")
    
    return removed_count


def _ensure_cover_single_page(doc: Document) -> int:
    """
    确保封面内容只占一页
    
    删除封面之后的空段落、分页符、多余内容，确保目录能从第2页开始
    
    策略：
    1. 保留前5个段落（封面核心信息）
    2. 删除第6个到最后一个段落（模板的多余内容）
    3. 确保封面简洁，为目录留出空间
    
    Args:
        doc: Document 对象
        
    Returns:
        删除的段落数量
    """
    removed_count = 0
    
    # 只保留前5个段落（封面标题、项目名称、日期等）
    paras = list(doc.paragraphs)
    
    if len(paras) <= 5:
        logger.info(f"封面段落数量 <= 5，无需清理")
        return 0
    
    # 删除第6个之后的所有段落
    for i in range(5, len(paras)):
        para = paras[i]
        try:
            para._element.getparent().remove(para._element)
            removed_count += 1
        except Exception as e:
            logger.warning(f"删除段落 {i} 失败: {e}")
    
    if removed_count > 0:
        logger.info(f"✅ 清理封面多余内容：删除了 {removed_count} 个段落，确保封面只占一页")
    
    return removed_count


def _add_toc_field(doc: Document, levels: str = "1-5", include_numbering: bool = False) -> None:
    """
    添加目录域（TOC field）
    
    注意：LibreOffice 可能不会自动更新目录，需要在 Word 中打开按 F9 更新
    
    Args:
        doc: Document 对象
        levels: 目录层级范围（如 "1-5"）
        include_numbering: 是否在目录中包含标题编号（默认 False）
    """
    p = doc.add_paragraph()
    
    # TOC 开关说明：
    # \o "1-5" - 使用大纲级别 1-5
    # \h - 超链接
    # \z - 隐藏页码和制表符
    # \u - 使用段落大纲级别
    # \n - 不包含页码前的编号（新增，用于隐藏标题编号）
    toc_switches = f'TOC \\o "{levels}" \\h \\z \\u'
    if not include_numbering:
        toc_switches += ' \\n'  # 添加 \n 开关，隐藏标题前的编号
    
    # 创建 fldSimple 元素
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), toc_switches)
    
    # 添加占位文本
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "[目录将在 Word 中更新]"
    r.append(t)
    fld.append(r)
    
    p._p.append(fld)


def _add_plain_text_toc(
    doc: Document,
    roots: List[DirNode],
    heading_style_map: Optional[Dict[int, str]],
    prefix_numbering: bool = False
) -> int:
    """
    生成纯文本目录（立即可见，无需在 Word 中更新）
    
    Args:
        doc: Document 对象
        roots: 根节点列表
        heading_style_map: 标题样式映射
        prefix_numbering: 是否在目录中包含编号
        
    Returns:
        生成的目录行数
    """
    from docx.shared import Pt, Inches
    
    toc_lines = 0
    
    def _get_toc_style_name(level: int) -> str:
        """获取目录样式名称"""
        # 尝试使用标准的目录样式（按优先级）
        toc_style_names = [
            f"TOC {level}",
            f"toc {level}",
            f"TOC{level}",
            f"toc{level}",
            f"目录 {level}",
            f"目录{level}",
        ]
        
        # 检查文档中是否存在这些样式
        for style_name in toc_style_names:
            try:
                # 尝试直接访问样式
                style = doc.styles[style_name]
                if style:
                    logger.debug(f"找到目录样式: {style_name}")
                    return style_name
            except KeyError:
                # 样式不存在
                continue
            except Exception as e:
                logger.debug(f"检查样式 {style_name} 时出错: {e}")
                continue
        
        # 如果没有专用的目录样式，返回 None（使用默认样式）
        logger.debug(f"未找到level {level}的目录样式，将使用默认样式")
        return None
    
    def _add_toc_entry(node: DirNode, depth: int = 0):
        """递归添加目录条目"""
        nonlocal toc_lines
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        
        # 跳过"目录"节点本身
        if is_toc_title(node.title):
            # 但要处理其子节点
            for child in node.children:
                _add_toc_entry(child, depth)
            return
        
        # 跳过标记为目录的节点
        if node.meta_json.get("is_toc") or node.meta_json.get("type") == "toc":
            for child in node.children:
                _add_toc_entry(child, depth)
            return
        
        # 生成目录文本（带编号）
        # ✅ TOC条目始终显示编号（一、二、三、）
        if node.numbering:
            toc_text = f"{node.numbering}、{node.title}"
        else:
            toc_text = node.title
        
        # ✅ 添加Tab符和页码占位符
        # 格式：一、章节标题\t[页码占位符]
        page_num_placeholder = "X"  # 占位符，提示用户更新
        toc_text_with_tab = f"{toc_text}\t{page_num_placeholder}"
        
        # 获取目录样式
        toc_style = _get_toc_style_name(min(node.level, 5))
        
        # 添加目录行
        try:
            if toc_style:
                para = doc.add_paragraph(toc_text_with_tab, style=toc_style)
                logger.debug(f"使用目录样式 {toc_style} 添加: {toc_text[:30]}")
            else:
                # 如果没有专用样式，使用默认样式并手动设置格式
                para = doc.add_paragraph(toc_text_with_tab)
                
                # 设置缩进（根据层级）
                para.paragraph_format.left_indent = Inches(depth * 0.25)
                
                # 设置字体大小（根据层级递减）
                if para.runs:
                    font_size = max(10, 12 - depth)
                    para.runs[0].font.size = Pt(font_size)
                    
                logger.debug(f"使用默认样式添加: {toc_text[:30]} (depth={depth})")
            
            # ✅ 增强：添加Tab stop with dot leader（点号连线到页码）
            # 这样TOC条目看起来像：章节标题........页码
            tab_stops = para.paragraph_format.tab_stops
            
            # 清除现有的tab stops
            for ts in list(tab_stops):
                tab_stops.clear_all()
            
            # 添加右对齐的tab stop，位于页面右侧，使用点号引导符
            # 16cm是A4纸宽度减去左右边距后的合理位置
            tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            logger.debug(f"为TOC条目添加Tab stop: {toc_text[:30]}")
            
        except Exception as e:
            logger.warning(f"添加目录行失败: {e}")
            para = doc.add_paragraph(toc_text)
        
        toc_lines += 1
        
        # 递归处理子节点
        for child in node.children:
            _add_toc_entry(child, depth + 1)
    
    # 遍历所有根节点
    for root in roots:
        _add_toc_entry(root, 0)
    
    logger.info(f"生成纯文本目录: {toc_lines} 行")
    return toc_lines


def _add_section_break(doc: Document, sectpr_xml: etree._Element) -> None:
    """
    添加分节符（section break）
    
    Args:
        doc: Document 对象
        sectpr_xml: sectPr XML 元素
    """
    # 添加一个空段落，并在其 pPr 中插入 sectPr
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    
    # 深拷贝 sectPr（避免多次引用同一个对象）
    sectpr_copy = etree.fromstring(etree.tostring(sectpr_xml))
    ppr.append(sectpr_copy)


def _maybe_prefix_numbering(text: str, numbering: Optional[str], enabled: bool) -> str:
    """
    如果启用了编号前缀，将编号添加到标题前面
    
    Args:
        text: 原始标题文本
        numbering: 编号（如 "1.2.3"）
        enabled: 是否启用
        
    Returns:
        处理后的标题文本
    """
    if not enabled or not numbering:
        return text
    
    # 避免重复：如果 text 本来就以编号开头，则不再添加
    import re
    if re.match(r"^\s*\d+(\.\d+)*\s+", text):
        return text
    
    return f"{numbering} {text}"


def _extract_template_images(template_path: str) -> Dict[str, List[Any]]:
    """
    从模板中提取图片及其关联的标题
    
    返回格式：
    {
        "投标人营业执照": [图片段落对象列表],
        "投标人资质证书": [图片段落对象列表],
        ...
    }
    """
    from docx import Document as TempDocument
    
    template_images = {}
    current_heading = None
    
    try:
        template_doc = TempDocument(template_path)
        
        for idx, para in enumerate(template_doc.paragraphs):
            # 检查是否是标题
            if para.style and ('标题' in para.style.name or 'heading' in para.style.name.lower()):
                current_heading = para.text.strip()
                logger.debug(f"发现标题: {current_heading}")
            
            # 检查段落是否包含图片
            has_image = False
            for run in para.runs:
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                if drawings:
                    has_image = True
                    break
            
            # 如果包含图片且有当前标题，保存图片段落
            if has_image and current_heading:
                # 将整个段落的XML保存下来，以便复制图片
                if current_heading not in template_images:
                    template_images[current_heading] = []
                template_images[current_heading].append(para)
                logger.info(f"提取图片: {current_heading} (段落{idx})")
        
        logger.info(f"从模板中提取了 {len(template_images)} 个标题的图片")
        for heading, paras in template_images.items():
            logger.info(f"  - {heading}: {len(paras)}张图片")
        
    except Exception as e:
        logger.error(f"提取模板图片失败: {e}", exc_info=True)
    
    return template_images


def _insert_images_to_paragraph(source_para, target_doc):
    """
    将源段落中的图片复制到目标文档
    
    Args:
        source_para: 源段落（来自模板）
        target_doc: 目标文档
    """
    try:
        # 创建一个新段落用于存放图片
        target_para = target_doc.add_paragraph()
        
        # 复制图片runs
        for run in source_para.runs:
            # 检查run中是否有图片
            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawings:
                # 创建新run并复制图片元素
                new_run = target_para.add_run()
                for drawing in drawings:
                    # 直接复制drawing元素到新run
                    new_run._element.append(drawing)
                logger.debug("复制了一张图片")
        
        return target_para
    except Exception as e:
        logger.error(f"插入图片失败: {e}", exc_info=True)
        return None


async def render_directory_tree_to_docx(
    template_path: str,
    output_path: str,
    roots: List[DirNode],
    section_prototypes: Dict[PageVariant, SectPrototype],
    *,
    include_toc: bool = True,
    prefix_numbering_in_text: bool = False,
    heading_style_map: Optional[Dict[int, str]] = None,
    normal_style_name: Optional[str] = None,
    apply_assets: Optional[Dict[str, Any]] = None,
    insert_section_body: Optional[callable] = None,
    auto_generate_content: bool = False,
    auto_write_cfg: Optional[AutoWriteCfg] = None,
    project_context: str = "",
    model_id: Optional[str] = None,
    parallel_generation: bool = True,
    max_concurrent: int = 5,
) -> None:
    """
    将目录树渲染为 Word 文档（使用模板母版）
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        roots: 根节点列表
        section_prototypes: 页面布局原型映射
        include_toc: 是否包含目录
        prefix_numbering_in_text: 是否在标题前添加编号
        heading_style_map: 标题样式映射（level -> style_name），来自模板配置
        normal_style_name: 正文样式名称，用于 summary 段落
        apply_assets: LLM分析的保留/删除计划（anchors, keepPlan等）
        insert_section_body: 插入节正文的回调函数（node: DirNode, doc: Document）
        auto_generate_content: 是否自动生成内容（当 summary 为空时）
        auto_write_cfg: 自动写作配置（如果未提供则使用默认值）
        project_context: 项目上下文信息（用于自动生成内容）
        model_id: LLM模型ID（用于自动生成内容）
        parallel_generation: 是否并行生成内容（默认True）
        max_concurrent: 最大并发数（默认5）
    """
    logger.info(f"开始渲染文档: template={template_path}, output={output_path}, auto_generate={auto_generate_content}, parallel={parallel_generation}")
    
    # ✅ 提取模板中的图片（资质、证书类）
    template_images = _extract_template_images(template_path)
    
    # 初始化自动写作配置和缓存
    if auto_generate_content and auto_write_cfg is None:
        auto_write_cfg = AutoWriteCfg()
    
    content_cache = {} if auto_generate_content else None
    
    # ✅ 如果启用并行生成，预先收集所有需要生成的节点并并行生成
    if auto_generate_content and parallel_generation:
        logger.info("启用并行内容生成")
        await _parallel_generate_all_content(
            roots=roots,
            project_context=project_context,
            cfg=auto_write_cfg,
            cache=content_cache,
            model_id=model_id,
            max_concurrent=max_concurrent,
        )
    
    # 1. 加载模板（保留页眉页脚）
    doc = Document(template_path)
    
    # 2. 清空正文内容（保留最后的 sectPr，使用 LLM 分析的保留计划）
    _clear_body_keep_sectpr(doc, apply_assets=apply_assets)
    
    # 2.5 ✅ 确保封面只占一页：删除封面后的所有空段落和多余内容
    # 这样可以保证目录从第2页开始
    _ensure_cover_single_page(doc)
    
    # 3. 添加目录页（可选）
    if include_toc:
        # ✅ 在目录前添加分页符，确保目录从第2页开始
        doc.add_page_break()
        logger.info("在目录前添加分页符")
        
        # ✅ 使用专用的 TOC 标题样式，避免使用带编号的 Heading 1
        toc_style = _get_toc_title_style(heading_style_map, doc)
        logger.info(f"使用目录标题样式: {toc_style}")
        
        try:
            toc_title = doc.add_paragraph("目录", style=toc_style)
            logger.info(f"目录标题创建成功，样式: {toc_title.style.name}")
        except Exception as e:
            logger.error(f"使用样式 {toc_style} 失败: {e}，回退到 Normal")
            toc_title = doc.add_paragraph("目录")
            try:
                toc_title.style = "Normal"
            except:
                pass
        
        toc_title.alignment = 1  # 居中
        
        # 3.5. 删除模板中保留的示例标题（在插入"目录"标题后，生成TOC条目前）
        # 这些标题是模板的示例结构，会与系统生成的目录冲突
        logger.info("⚠️ 准备删除模板示例标题...")
        removed = _remove_template_sample_headings(doc)
        logger.info(f"⚠️ 删除模板示例标题完成: 删除了 {removed} 个段落")
        
        # ✅ 生成纯文本目录（立即可见，无需在 Word 中更新）
        _add_plain_text_toc(doc, roots, heading_style_map, prefix_numbering_in_text)
        
        # ✅ 额外清理：删除模板的示例标题（如"投标函"、"法定代表人授权书"等）
        # 必须在插入"目录"之后调用，因为需要依赖"目录"的位置来判断哪些是模板示例标题
        _remove_template_sample_headings(doc)
        
        doc.add_page_break()
    
    # 4. DFS 遍历目录树，写入内容
    async def emit_node(node: DirNode, depth: int = 0):
        """递归输出节点"""
        # 4.0 过滤"目录"节点：避免把目录标题当作章节插入正文
        if is_toc_title(node.title):
            logger.info(f"跳过目录节点: {node.title}")
            # 递归处理子节点
            for child in node.children:
                await emit_node(child, depth + 1)
            return
        
        # 4.0.5 过滤明确标记为目录的节点
        if node.meta_json.get("is_toc") or node.meta_json.get("type") == "toc":
            logger.info(f"跳过标记为目录的节点: {node.title}")
            # 递归处理子节点
            for child in node.children:
                await emit_node(child, depth + 1)
            return
        
        # 4.1 检查是否需要插入分节符（横版/特殊布局）
        page_variant = node.meta_json.get("page_variant")
        if page_variant and page_variant in section_prototypes:
            logger.debug(f"插入分节符: {page_variant} for node {node.title}")
            _add_section_break(doc, section_prototypes[page_variant].sectPr_xml)
        
        # 4.2 添加标题
        title_text = _maybe_prefix_numbering(node.title, node.numbering, prefix_numbering_in_text)
        style_name = _get_heading_style_name(node.level, heading_style_map)
        
        try:
            para = doc.add_paragraph(title_text, style=style_name)
        except Exception as e:
            # 样式不存在时回退到默认 Heading
            logger.warning(f"样式 {style_name} 不存在，使用默认 Heading: {e}")
            h_level = min(max(node.level, 1), 9)
            para = doc.add_heading(title_text, level=h_level)
        
        # ✅ 保留标题的自动编号（显示"一、二、三、"）
        # 模板的+标题1样式配置了自动编号(numId=1)，这正是用户需要的
        logger.debug(f"添加标题（保留编号）: {title_text[:30]}")
        
        # 4.3 添加 summary（正文内容）
        content_to_add = node.summary
        
        # 如果启用了自动生成且当前节点没有实质内容（空或占位符），则自动生成
        if auto_generate_content and _is_empty_or_placeholder(content_to_add):
            # 如果启用了并行生成，从缓存中获取（已预先生成）
            if parallel_generation:
                key = f"L{node.level}:{node.title.strip()}"
                generated_text = content_cache.get(key, "")
                if not generated_text:
                    logger.warning(f"并行生成的内容未找到: {key}")
                    generated_text = "【内容生成失败】"
            else:
                # 串行生成（原有逻辑）
            try:
                logger.info(f"自动生成内容: title={node.title}, level={node.level}")
                generated_text = await generate_section_text_by_title(
                    title=node.title,
                    level=node.level,
                    project_context=project_context,
                    cfg=auto_write_cfg,
                    cache=content_cache,
                    model_id=model_id,
                )
                except Exception as e:
                    logger.error(f"自动生成内容失败: title={node.title}, error={e}", exc_info=True)
                    generated_text = f"【自动生成内容失败：{str(e)}】"
                
                # 将生成的文本按空行分段，写入多个段落
                # 这样保持了 docx 的段落结构，更美观
                paragraphs = [
                    p.strip() 
                    for p in re.split(r"\n{2,}|\r\n{2,}", generated_text) 
                    if p.strip()
                ]
                
                for para in paragraphs:
                    if normal_style_name:
                        try:
                            doc.add_paragraph(para, style=normal_style_name)
                        except Exception as e:
                            logger.warning(f"样式 {normal_style_name} 不存在，使用默认段落: {e}")
                            doc.add_paragraph(para)
                    else:
                        doc.add_paragraph(para)
                
            logger.info(f"添加生成内容: {len(paragraphs)} 个段落, 总字符数 {len(generated_text)}")
        
        # 否则，添加原有内容（如果有的话）
        elif content_to_add and not _is_empty_or_placeholder(content_to_add):
            # 优先使用 normal_style_name，如果未指定则使用默认段落
            if normal_style_name:
                try:
                    doc.add_paragraph(content_to_add, style=normal_style_name)
                except Exception as e:
                    logger.warning(f"样式 {normal_style_name} 不存在，使用默认段落: {e}")
                    doc.add_paragraph(content_to_add)
            else:
                doc.add_paragraph(content_to_add)
        
        # 4.4 插入节正文内容（如果提供了回调）
        if insert_section_body:
            try:
                insert_section_body(node, doc)
            except Exception as e:
                logger.error(f"插入节正文失败: node={node.id}, error={e}", exc_info=True)
                doc.add_paragraph(f"[正文内容加载失败: {str(e)}]")
        
        # ✅ 4.5 插入模板中的图片（资质、证书类）
        # 根据当前节点的标题匹配模板中的图片
        if template_images:
            # 尝试精确匹配
            if node.title in template_images:
                logger.info(f"为章节'{node.title}'插入 {len(template_images[node.title])} 张图片")
                for img_para in template_images[node.title]:
                    _insert_images_to_paragraph(img_para, doc)
            # 尝试模糊匹配（关键词）
            else:
                for heading, img_paras in template_images.items():
                    # 检查标题中是否包含关键词
                    keywords = ['营业执照', '资质证书', '授权书', '许可证', '证明']
                    if any(kw in node.title and kw in heading for kw in keywords):
                        logger.info(f"为章节'{node.title}'插入图片（模糊匹配: {heading}）")
                        for img_para in img_paras:
                            _insert_images_to_paragraph(img_para, doc)
                        break
        
        # 4.6 递归处理子节点
        for child in node.children:
            await emit_node(child, depth + 1)
    
    # 遍历所有根节点
    for root in roots:
        await emit_node(root)
    
    # 5. 最后切回默认布局（通常是 A4 竖版）
    back_variant = PageVariant.A4_PORTRAIT if PageVariant.A4_PORTRAIT in section_prototypes else PageVariant.DEFAULT
    if back_variant in section_prototypes:
        _add_section_break(doc, section_prototypes[back_variant].sectPr_xml)
    
    # 6. 保存文档
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)
    logger.info(f"文档渲染完成: {output_path}")


async def _parallel_generate_all_content(
    roots: List[DirNode],
    project_context: str,
    cfg: AutoWriteCfg,
    cache: dict,
    model_id: Optional[str],
    max_concurrent: int,
) -> None:
    """
    并行生成所有需要生成内容的节点
    
    Args:
        roots: 根节点列表
        project_context: 项目上下文
        cfg: 自动写作配置
        cache: 内容缓存字典
        model_id: LLM模型ID
        max_concurrent: 最大并发数
    """
    import asyncio
    
    # 收集所有需要生成的节点
    nodes_to_generate = []
    
    def collect_nodes(node: DirNode):
        """递归收集节点"""
        # 过滤目录节点
        if is_toc_title(node.title):
            for child in node.children:
                collect_nodes(child)
            return
        
        if node.meta_json.get("is_toc") or node.meta_json.get("type") == "toc":
            for child in node.children:
                collect_nodes(child)
            return
        
        # 如果节点需要生成内容
        if _is_empty_or_placeholder(node.summary):
            nodes_to_generate.append(node)
        
        # 递归子节点
        for child in node.children:
            collect_nodes(child)
    
    for root in roots:
        collect_nodes(root)
    
    total = len(nodes_to_generate)
    if total == 0:
        logger.info("没有需要生成内容的节点")
        return
    
    logger.info(f"开始并行生成 {total} 个节点的内容，最大并发数: {max_concurrent}")
    
    # 创建信号量控制并发
    semaphore = asyncio.Semaphore(max_concurrent)
    completed = 0
    
    async def generate_one(node: DirNode, index: int):
        """生成单个节点的内容"""
        nonlocal completed
        
        async with semaphore:
            try:
                logger.info(f"[{index+1}/{total}] 开始生成: {node.title}")
                
                # 生成内容（会自动使用缓存）
                await generate_section_text_by_title(
                    title=node.title,
                    level=node.level,
                    project_context=project_context,
                    cfg=cfg,
                    cache=cache,
                    model_id=model_id,
                )
                
                completed += 1
                logger.info(f"[{index+1}/{total}] 生成成功: {node.title}")
                
            except Exception as e:
                logger.error(f"[{index+1}/{total}] 生成失败: {node.title} - {e}", exc_info=True)
    
    # 并行执行所有生成任务
    tasks = [generate_one(node, i) for i, node in enumerate(nodes_to_generate)]
    await asyncio.gather(*tasks, return_exceptions=True)
    
    logger.info(f"并行生成完成: 成功 {completed}/{total} 个节点")


def render_simple_outline_to_docx(
    output_path: str,
    roots: List[DirNode],
    *,
    include_toc: bool = True,
    prefix_numbering_in_text: bool = False,
) -> None:
    """
    渲染简单的目录文档（不使用模板）
    
    Args:
        output_path: 输出文件路径
        roots: 根节点列表
        include_toc: 是否包含目录
        prefix_numbering_in_text: 是否在标题前添加编号
    """
    logger.info(f"开始渲染简单文档: output={output_path}")
    
    # 1. 创建新文档
    doc = Document()
    
    # 2. 添加目录页（可选）
    if include_toc:
        doc.add_heading("目录", level=1)
        _add_toc_field(doc, "1-5")
        doc.add_page_break()
    
    # 3. DFS 遍历目录树
    def emit_node(node: DirNode):
        title_text = _maybe_prefix_numbering(node.title, node.numbering, prefix_numbering_in_text)
        h_level = min(max(node.level, 1), 9)
        doc.add_heading(title_text, level=h_level)
        
        if node.summary:
            doc.add_paragraph(node.summary)
        
        for child in node.children:
            emit_node(child)
    
    for root in roots:
        emit_node(root)
    
    # 4. 保存
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)
    logger.info(f"简单文档渲染完成: {output_path}")

