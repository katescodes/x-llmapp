"""
OutlineSampleAttacher - 目录范本挂载器
将抽取的范本片段自动挂载到目录节点的正文槽位
"""
from typing import Any, Dict, List, Optional
import logging
import os

from app.services.dao.tender_dao import TenderDAO
from app.services.fragment.fragment_matcher import FragmentTitleMatcher
from app.services.fragment.fragment_type import FragmentType
from app.services.fragment.semantic_matcher import (
    search_best_match_in_pdf,
    batch_search_for_directory,
    batch_search_for_directory_async,  # ✅ 新增异步版本
)
from app.services.fragment.pdf_layout_extractor import extract_pdf_items
from app.services.fragment.pdf_content_extractor import extract_fragment_content

logger = logging.getLogger(__name__)


class OutlineSampleAttacher:
    """目录范本挂载器（增强版：支持LLM兜底）"""
    
    def __init__(self, dao: TenderDAO, llm_client=None):
        """
        初始化
        
        Args:
            dao: TenderDAO实例
            llm_client: LLM客户端（可选，用于LLM兜底）
        """
        self.dao = dao
        self.matcher = FragmentTitleMatcher()
        self.llm_matcher = None
        self.llm_client = llm_client  # ✅ 保存llm_client引用
        
        # 如果提供了llm_client，初始化LLM匹配器
        if llm_client:
            try:
                from app.services.fragment.llm_matcher import LLMFragmentMatcher
                self.llm_matcher = LLMFragmentMatcher(llm_client)
                logger.info("[OutlineSampleAttacher] LLM matcher enabled")
            except Exception as e:
                logger.warning(f"[OutlineSampleAttacher] Failed to init LLM matcher: {e}")
    
    def attach(self, project_id: str, outline_nodes: List[Dict[str, Any]], use_llm: bool = True) -> int:
        """
        将范本片段挂载到目录节点（同步版本）
        
        Args:
            project_id: 项目ID
            outline_nodes: 目录节点列表（扁平结构）
            use_llm: 是否使用LLM兜底（默认True）
            
        Returns:
            成功挂载的节点数量
        """
        # 调用异步版本
        from app.platform.utils.async_runner import run_async
        return run_async(self.attach_async(project_id, outline_nodes, use_llm))
    
    async def attach_async(
        self,
        project_id: str,
        outline_nodes: List[Dict[str, Any]],
        use_llm: bool = True
    ) -> int:
        """
        将范本片段挂载到目录节点（异步版本，增强版：支持LLM兜底）
        
        Args:
            project_id: 项目ID
            outline_nodes: 目录节点列表（扁平结构）
            use_llm: 是否使用LLM兜底（默认True）
            
        Returns:
            成功挂载的节点数量
        """
        attached_count = 0
        # 1. 获取项目的所有范本片段
        fragments = self.dao.list_fragments("PROJECT", project_id)

        def _is_valid_fragment(frag: Dict[str, Any]) -> bool:
            try:
                src = (frag.get("source_file_key") or "").strip()
                if not src or not str(src).lower().endswith(".docx"):
                    return False
                s = frag.get("start_body_index")
                e = frag.get("end_body_index")
                if s is None or e is None:
                    return False
                s = int(s)
                e = int(e)
                if s < 0 or e < 0 or s > e:
                    return False
                return True
            except Exception:
                return False
        
        # 按类型组织片段（仅保留合法片段）
        fragments_by_type: Dict[str, List[Dict[str, Any]]] = {}
        for frag in fragments:
            if not _is_valid_fragment(frag):
                # 避免“挂载空”：没有合法 start/end/source 的片段不参与匹配
                continue
            ftype = frag.get("fragment_type")
            if ftype not in fragments_by_type:
                fragments_by_type[ftype] = []
            fragments_by_type[ftype].append(frag)

        # 如果没有任何可用片段：清理历史的 TEMPLATE_SAMPLE 标记，避免“挂载但没有内容”
        if not fragments_by_type:
            for node in outline_nodes:
                node_id = node.get("id")
                if not node_id:
                    continue
                existing_body = self.dao.get_section_body(project_id, node_id)
                if not existing_body:
                    continue
                if existing_body.get("source") != "TEMPLATE_SAMPLE":
                    continue
                # 不覆盖 USER/AI；这里仅清理 TEMPLATE_SAMPLE
                self.dao.upsert_section_body(
                    project_id=project_id,
                    node_id=node_id,
                    source="EMPTY",
                    fragment_id=None,
                    content_html=None,
                )
            return 0
        
        # 2. 遍历目录节点，匹配并挂载
        for node in outline_nodes:
            node_id = node.get("id")
            if not node_id:
                continue
            
            # attach 策略：
            # - 仅当该 node 当前 section_body 不是 USER（或 USER 为空）时才覆盖挂载
            existing_body = self.dao.get_section_body(project_id, node_id)
            if existing_body:
                # 如果已有用户编辑内容，跳过
                if existing_body.get("source") == "USER":
                    if existing_body.get("content_html") or existing_body.get("content_json"):
                        continue
                # 如果已有 AI 生成内容，跳过
                if existing_body.get("source") == "AI":
                    continue
                # 如果已有 TEMPLATE_SAMPLE 且有内容，跳过（避免重复覆盖）
                if existing_body.get("source") == "TEMPLATE_SAMPLE":
                    if existing_body.get("content_json") or existing_body.get("content_html"):
                        continue
            
            # 归一化节点标题
            node_title = node.get("title", "")
            node_title_norm = self.matcher.normalize(node_title)
            
            if not node_title_norm:
                continue
            
            # ✨ 增强匹配逻辑：规则 + LLM混合策略
            best_fragment = await self._match_fragment_hybrid(
                node=node,
                node_title_norm=node_title_norm,
                fragments=fragments,
                fragments_by_type=fragments_by_type,
                use_llm=use_llm
            )
            
            if not best_fragment:
                continue
            
            # 二次校验（防止 candidates 中混入不合法）
            if not _is_valid_fragment(best_fragment):
                continue
            
            # 挂载片段并写入content（从fragment的content字段获取）
            # ✅ 优先使用fragment中已提取的content_html和content_text
            content_html = best_fragment.get("content_html")
            content_json = None
            
            if not content_html:
                # 兜底：使用旧的blocks_json方式
                content_json = self._extract_fragment_blocks(best_fragment)
            
            self.dao.upsert_section_body(
                project_id=project_id,
                node_id=node_id,
                source="TEMPLATE_SAMPLE",
                fragment_id=best_fragment["id"],
                content_html=content_html,  # ✅ 直接使用fragment的HTML内容
                content_json=content_json,  # 兜底使用blocks
            )
            attached_count += 1
            
            logger.info(
                f"[OutlineSampleAttacher] Attached fragment '{best_fragment.get('title')}' to node '{node_title}' "
                f"(content_type: {best_fragment.get('content_type', 'unknown')}, "
                f"content_len: {len(content_html or content_json or '')})"
            )

        return attached_count
    
    async def _match_fragment_hybrid(
        self,
        node: Dict[str, Any],
        node_title_norm: str,
        fragments: List[Dict[str, Any]],
        fragments_by_type: Dict[str, List[Dict[str, Any]]],
        use_llm: bool
    ) -> Optional[Dict[str, Any]]:
        """
        混合匹配策略：规则优先，LLM兜底
        
        Args:
            node: 目录节点
            node_title_norm: 归一化后的节点标题
            fragments: 所有fragments
            fragments_by_type: 按类型组织的fragments
            use_llm: 是否使用LLM兜底
            
        Returns:
            匹配到的fragment，如果没有匹配则返回 None
        """
        # Phase 1: 规则匹配（带置信度）
        ftype, confidence = self.matcher.match_type_with_confidence(node_title_norm)
        
        best_fragment = None
        
        if ftype and confidence >= 0.9:
            # 高置信度（≥0.9），直接使用规则匹配
            best_fragment = self._find_best_fragment(
                node_title_norm,
                fragments_by_type.get(str(ftype), [])
            )
            if best_fragment:
                logger.info(
                    f"[OutlineSampleAttacher] Node '{node.get('title')}' matched by rules "
                    f"(confidence: {confidence:.2f}, type: {ftype})"
                )
                return best_fragment
        
        elif ftype and confidence >= 0.6:
            # 中等置信度（0.6-0.9），先尝试规则
            best_fragment = self._find_best_fragment(
                node_title_norm,
                fragments_by_type.get(str(ftype), [])
            )
            
            if best_fragment:
                logger.info(
                    f"[OutlineSampleAttacher] Node '{node.get('title')}' matched by rules "
                    f"(medium confidence: {confidence:.2f}, type: {ftype})"
                )
                return best_fragment
        
        # Phase 2: LLM兜底（低置信度或规则无匹配）
        if use_llm and self.llm_matcher:
            if not best_fragment or (ftype and confidence < 0.9):
                logger.info(
                    f"[OutlineSampleAttacher] Node '{node.get('title')}' using LLM fallback "
                    f"(rule confidence: {confidence:.2f if ftype else 0.0})"
                )
                
                try:
                    llm_fragment = await self.llm_matcher.match_async(node, fragments)
                    if llm_fragment:
                        logger.info(
                            f"[OutlineSampleAttacher] Node '{node.get('title')}' matched by LLM"
                        )
                        return llm_fragment
                except Exception as e:
                    logger.error(f"[OutlineSampleAttacher] LLM match failed: {type(e).__name__}: {e}")
        
        # Phase 3: 返回规则匹配结果（如果有）
        if best_fragment:
            logger.info(
                f"[OutlineSampleAttacher] Node '{node.get('title')}' using rule match "
                f"as fallback (confidence: {confidence:.2f if ftype else 0.0})"
            )
            return best_fragment
        
        # 无匹配
        logger.debug(f"[OutlineSampleAttacher] Node '{node.get('title')}' no match found")
        return None
    
    def _find_best_fragment(
        self,
        node_title_norm: str,
        candidates: List[Dict[str, Any]]
    ) -> Optional[Dict[str, Any]]:
        """
        从候选片段中找到最佳匹配
        
        匹配规则：
        1. 优先标题完全相等
        2. 其次标题包含关系
        3. 最后按置信度和范围长度排序
        """
        if not candidates:
            return None
        
        # 精确匹配
        for frag in candidates:
            frag_title_norm = frag.get("title_norm", "")
            if frag_title_norm == node_title_norm:
                return frag
        
        # 包含匹配
        matches_with_score = []
        for frag in candidates:
            frag_title_norm = frag.get("title_norm", "")
            
            if node_title_norm in frag_title_norm or frag_title_norm in node_title_norm:
                # 计算得分：置信度 + 范围长度
                confidence = frag.get("confidence") or 0.0
                range_len = frag.get("end_body_index", 0) - frag.get("start_body_index", 0)
                score = confidence * 1000 + range_len
                matches_with_score.append((score, frag))
        
        if matches_with_score:
            matches_with_score.sort(key=lambda x: x[0], reverse=True)
            return matches_with_score[0][1]
        
        # 如果没有包含匹配，返回第一个（已按置信度排序）
        return candidates[0] if candidates else None
    
    def _extract_fragment_blocks(self, fragment: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        从源文件中提取片段的 blocks 内容
        
        Returns:
            blocks 列表，格式：[{type:'paragraph'|'table', text/tableData/...}, ...]
        """
        try:
            from docx import Document
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            
            source_file = fragment.get("source_file_key")
            if not source_file or not os.path.exists(source_file):
                logger.warning(f"[attacher] source_file not found: {source_file}")
                return []
            
            start_idx = int(fragment.get("start_body_index", 0))
            end_idx = int(fragment.get("end_body_index", 0))
            
            if start_idx < 0 or end_idx < 0 or start_idx > end_idx:
                logger.warning(f"[attacher] invalid indices: start={start_idx}, end={end_idx}")
                return []
            
            doc = Document(source_file)
            
            # 使用与 fragment_extractor 相同的逻辑：按 body 顺序提取
            para_map = {p._p: p for p in doc.paragraphs}
            tbl_map = {t._tbl: t for t in doc.tables}
            
            blocks = []
            idx = 0
            
            for child in doc.element.body.iterchildren():
                # 只提取指定范围的 blocks
                if idx < start_idx:
                    idx += 1
                    continue
                if idx > end_idx:
                    break
                
                tag = child.tag.split("}")[-1]
                if tag == "p":
                    p = para_map.get(child)
                    if p:
                        text = (p.text or "").strip()
                        blocks.append({
                            "type": "paragraph",
                            "text": text,
                            "styleName": getattr(getattr(p, "style", None), "name", None),
                        })
                elif tag == "tbl":
                    t = tbl_map.get(child)
                    if t:
                        rows = []
                        for r in t.rows:
                            rows.append([(c.text or "").strip() for c in r.cells])
                        blocks.append({
                            "type": "table",
                            "tableData": rows,
                            "styleName": getattr(getattr(t, "style", None), "name", None),
                        })
                
                idx += 1
            
            logger.info(f"[attacher] extracted {len(blocks)} blocks from fragment {fragment.get('id')}")
            return blocks
            
        except Exception as e:
            logger.error(f"[attacher] _extract_fragment_blocks failed: {type(e).__name__}: {e}")
            return []
    
    async def attach_from_pdf_semantic_async(
        self,
        project_id: str,
        outline_nodes: List[Dict[str, Any]],
        min_confidence: float = 0.4,  # ✅ 降低阈值到0.4
        use_llm: bool = True,  # ✅ 支持LLM验证
    ) -> int:
        """
        ✅ 新策略：从目录标题出发，语义搜索PDF中最匹配的表格内容（异步版本）
        
        策略：
        - 方案A：混合策略（关键词 + LLM语义判断）
        - 方案C：上下文增强搜索（表格前后段落）
        
        Args:
            project_id: 项目ID
            outline_nodes: 目录节点列表（扁平结构）
            min_confidence: 最低置信度阈值（默认0.4）
            use_llm: 是否使用LLM验证（默认True）
            
        Returns:
            成功挂载的节点数量
        """
        attached_count = 0
        
        try:
            # 1. 获取项目的招标文件PDF路径
            tender_assets = self.dao.list_assets(project_id)  # ✅ 修正方法名
            pdf_path = None
            
            for asset in tender_assets:
                # 检查是否为招标文件PDF
                kind = asset.get("kind", "")
                storage_path = asset.get("storage_path", "")
                filename = asset.get("filename", "")
                
                if kind == "tender" and storage_path.lower().endswith(".pdf"):
                    if os.path.exists(storage_path):
                        pdf_path = storage_path
                        break
            
            if not pdf_path:
                logger.warning(f"[attach_from_pdf_semantic] No PDF tender file found for project {project_id}")
                return 0
            
            logger.info(f"[attach_from_pdf_semantic] Using PDF: {pdf_path}")
            
            # 2. 解析PDF
            pdf_items, diag = extract_pdf_items(pdf_path, max_pages=500)
            logger.info(f"[attach_from_pdf_semantic] Extracted {len(pdf_items)} items from PDF ({diag.get('pages')} pages)")
            
            # 3. 批量搜索匹配（✅ 使用异步版本 + LLM）
            llm_client = self.llm_client  # ✅ 直接使用保存的llm_client
            
            matches = await batch_search_for_directory_async(
                outline_nodes, 
                pdf_items, 
                llm_client,
                min_confidence,
                use_llm=use_llm
            )
            
            logger.info(f"[attach_from_pdf_semantic] Found {len(matches)} matches (LLM={'ON' if use_llm and llm_client else 'OFF'})")
            
            # 4. 为每个匹配结果提取内容并填充
            for node_id, match in matches.items():
                try:
                    item = match["item"]
                    item_index = match["item_index"]
                    confidence = match["confidence"]
                    llm_verified = match.get("llm_verified", False)
                    
                    # 提取内容
                    content = extract_fragment_content(pdf_items, item_index, item_index)
                    
                    # 获取节点标题（用于日志）
                    node = next((n for n in outline_nodes if n.get("id") == node_id), None)
                    node_title = node.get("title", "Unknown") if node else "Unknown"
                    
                    logger.info(
                        f"[attach_from_pdf_semantic] Filling '{node_title}' "
                        f"(confidence={confidence:.2f}, llm={llm_verified}, type={content['type']})"
                    )
                    
                    # 保存到数据库（只传递upsert_section_body接受的参数）
                    self.dao.upsert_section_body(
                        project_id=project_id,
                        node_id=node_id,
                        source="PDF_SEMANTIC_MATCH",
                        fragment_id=None,
                        content_html=content["html"],
                        content_json=content["items"],  # 使用content_json存储结构化内容
                    )
                    
                    attached_count += 1
                    
                except Exception as e:
                    logger.error(f"[attach_from_pdf_semantic] Failed to fill node {node_id}: {e}")
                    continue
            
            logger.info(f"[attach_from_pdf_semantic] Successfully attached {attached_count}/{len(outline_nodes)} nodes")
            return attached_count
            
        except Exception as e:
            logger.error(f"[attach_from_pdf_semantic] Failed: {type(e).__name__}: {e}", exc_info=True)
            return 0
    
    def attach_from_pdf_semantic(
        self,
        project_id: str,
        outline_nodes: List[Dict[str, Any]],
        min_confidence: float = 0.4,
        use_llm: bool = True,
    ) -> int:
        """
        同步版本（包装异步方法）
        """
        from app.platform.utils.async_runner import run_async
        return run_async(self.attach_from_pdf_semantic_async(
            project_id, outline_nodes, min_confidence, use_llm
        ))
