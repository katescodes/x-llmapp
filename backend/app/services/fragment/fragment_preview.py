"""
fragment_preview - 将 docx 中的范本片段渲染为最小可用 HTML 预览

要求：
- 遍历顺序与 DocxBodyElementCopier 一致：list(Document(...).element.body)
- 支持段落(w:p)与表格(w:tbl)的最小渲染（忽略复杂样式）
- 控制预览体积：max_elems / 表格行列截断 / 总字符数截断
"""

from __future__ import annotations

import html
from typing import List, Tuple

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table


def _escape_text(s: str) -> str:
    # 先 escape，再把换行转为 <br/>，保证表格/段落里的多行文本可读
    return html.escape(s or "").replace("\n", "<br/>")


# -------- 新版（带截断控制） --------

_MAX_PREVIEW_TEXT_CHARS = 12000
_MAX_TABLE_ROWS = 15
_MAX_TABLE_COLS = 8


def build_fragment_preview_meta(
    docx_path: str,
    start: int,
    end: int,
    max_elems: int = 60,
) -> Tuple[str, int, List[str]]:
    """
    构建范本片段预览（简化 HTML）。

    - 遍历 doc.element.body 的 CT_P / CT_Tbl，按 index 取 [start..end]（包含）
    - 段落：<p> + escape(text)
    - 表格：<table><tr><td>...</td></tr>...</table>
      - 只取前 15 行、每行前 8 个单元格，避免太大
    - 总字符超过阈值（默认 12000）则截断，并追加 “…已截断”

    Returns:
        (html, text_len, warnings)
    """
    warnings: List[str] = []

    try:
        doc = Document(docx_path)
    except Exception as e:
        return (
            f"<div style='color:#b00020'>[范本预览渲染失败：无法打开 docx：{html.escape(str(e))}]</div>",
            0,
            [f"open_docx_failed: {type(e).__name__}: {str(e)}"],
        )

    elements = list(doc.element.body)

    if start < 0 or end >= len(elements) or start > end:
        return (
            "<div style='color:#b00020'>[范本预览范围无效]</div>",
            0,
            [f"invalid_range: start={start}, end={end}, n_body={len(elements)}"],
        )

    # 只计数 CT_P/CT_Tbl（其它元素忽略），避免 max_elems 被空元素消耗
    rendered_elems = 0
    text_len = 0
    out: List[str] = []

    def _append_html(piece_html: str, piece_text_len: int):
        nonlocal text_len
        if not piece_html:
            return
        out.append(piece_html)
        text_len += int(piece_text_len or 0)

    for i in range(start, end + 1):
        el = elements[i]

        if isinstance(el, CT_P):
            if rendered_elems >= max_elems:
                warnings.append(f"max_elems_reached: {max_elems}")
                break
            p = Paragraph(el, doc)
            txt = (p.text or "").strip()
            _append_html(f"<p>{_escape_text(txt)}</p>", len(txt))
            rendered_elems += 1

        elif isinstance(el, CT_Tbl):
            if rendered_elems >= max_elems:
                warnings.append(f"max_elems_reached: {max_elems}")
                break
            t = Table(el, doc)
            rows_html: List[str] = []
            row_count = 0
            col_truncated = False
            for row in t.rows[:_MAX_TABLE_ROWS]:
                row_count += 1
                tds: List[str] = []
                if len(row.cells) > _MAX_TABLE_COLS:
                    col_truncated = True
                for cell in row.cells[:_MAX_TABLE_COLS]:
                    cell_txt = (cell.text or "").strip()
                    # 单元格文本也做轻量控制，避免超大表格导致预览暴涨
                    if len(cell_txt) > 800:
                        cell_txt = cell_txt[:800] + "…"
                        warnings.append("cell_text_truncated")
                    tds.append(f"<td>{_escape_text(cell_txt)}</td>")
                    text_len += len(cell_txt)
                rows_html.append("<tr>" + "".join(tds) + "</tr>")
            if len(t.rows) > _MAX_TABLE_ROWS:
                warnings.append(f"table_rows_truncated: {len(t.rows)}->({_MAX_TABLE_ROWS})")
            if col_truncated:
                warnings.append(f"table_cols_truncated: >{_MAX_TABLE_COLS}")
            out.append(
                "<table border='1' cellspacing='0' cellpadding='6' style='border-collapse:collapse;max-width:100%'>"
                + "".join(rows_html)
                + "</table>"
            )
            rendered_elems += 1

        else:
            # 其它元素类型：忽略（例如 sectPr）
            continue

        if text_len >= _MAX_PREVIEW_TEXT_CHARS:
            warnings.append(f"preview_truncated_at_{_MAX_PREVIEW_TEXT_CHARS}_chars")
            break

    html_out = "<div class='template-sample-preview'>" + "\n".join(out) + "</div>"
    if text_len >= _MAX_PREVIEW_TEXT_CHARS:
        html_out += "<div style='color:#64748b;font-size:12px;margin-top:8px'>…已截断</div>"
    return html_out, int(text_len), warnings


def build_fragment_preview_html(docx_path: str, start: int, end: int, max_elems: int = 60) -> str:
    """只返回 HTML（给外部快速使用的最简接口）。"""
    html_out, _, _ = build_fragment_preview_meta(docx_path=docx_path, start=start, end=end, max_elems=max_elems)
    return html_out


# -------- 兼容旧接口（章节正文预览使用） --------

def render_fragment_html(source_docx_path: str, start_idx: int, end_idx: int) -> str:
    """
    将 docx 中 [start_idx, end_idx]（包含）的 body elements 渲染成 HTML。

    index 规则必须与 DocxBodyElementCopier.copy_range() 完全一致：
    src_elements = list(Document(...).element.body)
    """
    # 章节预览允许稍多一些元素，避免“预览太短看不到表格主体”
    return build_fragment_preview_html(
        docx_path=source_docx_path,
        start=int(start_idx),
        end=int(end_idx),
        max_elems=120,
    )


