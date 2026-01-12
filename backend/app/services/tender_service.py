"""
æ‹›æŠ•æ ‡åº”ç”¨ - ä¸šåŠ¡é€»è¾‘å±‚ (Service)
åŒ…å« LLM è°ƒç”¨ã€æ–‡ä»¶è§£æã€è§„åˆ™æŠ½å–ã€å®¡æ ¸å åŠ ç­‰æ ¸å¿ƒé€»è¾‘
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import uuid
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document  # type: ignore[import-untyped]
from fastapi import UploadFile

from app.config import get_settings, get_feature_flags
from app.schemas.project_delete import ProjectDeletePlanResponse, ProjectDeleteRequest
from app.services.dao.tender_dao import TenderDAO
from app.services.project_delete import ProjectDeletionOrchestrator
from app.services.template.docx_extractor import DocxBlockExtractor
from app.services.template.llm_analyzer import TemplateLlmAnalyzer, get_analysis_cache
from app.services.template.template_spec import TemplateSpec, create_minimal_spec, BasePolicyMode
from app.services.template.outline_merger import OutlineMerger
from app.services.template.template_parse_preview import DocxTemplateDeterministicParser, TemplatePreviewGenerator

logger = logging.getLogger(__name__)


# ==================== å·¥å…·å‡½æ•° ====================

def _safe_mkdir(p: str):
    """å®‰å…¨åˆ›å»ºç›®å½•"""
    os.makedirs(p, exist_ok=True)


def _sha256(b: bytes) -> str:
    """è®¡ç®—SHA256å“ˆå¸Œ"""
    return hashlib.sha256(b).hexdigest()


def _extract_json(text: str) -> Any:
    """
    ä» LLM è¾“å‡ºä¸­å®¹é”™æå– JSON
    æ”¯æŒ markdown code fence åŒ…è£¹çš„ JSON
    """
    if not text:
        raise ValueError("empty llm output")
    
    # å°è¯•æå– markdown code fence ä¸­çš„å†…å®¹
    m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text, re.IGNORECASE)
    if m:
        text = m.group(1).strip()
    
    # å°è¯•æå–ç¬¬ä¸€ä¸ª JSON å¯¹è±¡æˆ–æ•°ç»„
    m2 = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", text.strip())
    if m2:
        text = m2.group(1)
    
    return json.loads(text)


def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    """
    å°†æ–‡æœ¬åˆ†å—
    
    Args:
        text: åŸå§‹æ–‡æœ¬
        max_chars: æ¯å—æœ€å¤§å­—ç¬¦æ•°
        overlap: é‡å å­—ç¬¦æ•°
    
    Returns:
        æ–‡æœ¬å—åˆ—è¡¨
    """
    text = (text or "").strip()
    if not text:
        return []
    
    out = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + max_chars)
        out.append(text[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)
    return out


def _read_text_from_file_bytes(filename: str, data: bytes) -> str:
    """
    ä»æ–‡ä»¶å­—èŠ‚ä¸­è¯»å–æ–‡æœ¬
    æ”¯æŒ txt/md/pdf/docx æ ¼å¼
    """
    name = (filename or "").lower()
    
    # TXT/MD æ–‡ä»¶
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return data.decode(errors="ignore")
    
    # PDF æ–‡ä»¶
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(data))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n".join(parts)
        except Exception:
            return ""
    
    # DOCX æ–‡ä»¶
    if name.endswith(".docx"):
        try:
            import io
            from docx import Document as Doc
            d = Doc(io.BytesIO(data))
            return "\n".join([p.text for p in d.paragraphs if p.text])
        except Exception:
            return ""
    
    # å…œåº•ï¼šå°è¯• UTF-8 è§£ç 
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _build_marked_context(chunks: List[Dict[str, Any]]) -> str:
    """
    å°† chunks æ„å»ºæˆå¸¦æ ‡è®°çš„ä¸Šä¸‹æ–‡
    ç”¨äº LLM èƒ½å¤Ÿå¼•ç”¨ chunk_id
    """
    parts = []
    for c in chunks:
        parts.append(f"[DOC {c.get('doc_id')} CHUNK {c.get('chunk_id')} POS {c.get('position')}]")
        parts.append(c.get("content") or "")
        parts.append("")  # ç©ºè¡Œ
    return "\n".join(parts).strip()


# ==================== LLM è°ƒç”¨æ•°æ®ç»“æ„ ====================

@dataclass
class LLMCall:
    """LLM è°ƒç”¨å‚æ•°"""
    model_id: Optional[str]
    messages: List[Dict[str, str]]


# ==================== Service ä¸»ç±» ====================

class TenderService:
    """æ‹›æŠ•æ ‡ä¸šåŠ¡é€»è¾‘æœåŠ¡"""

    def __init__(self, dao: TenderDAO, llm_orchestrator: Any, jobs_service: Any = None):
        """
        åˆå§‹åŒ– Service
        
        Args:
            dao: TenderDAO å®ä¾‹
            llm_orchestrator: LLM è°ƒåº¦å™¨ï¼ˆduck typingï¼‰
            jobs_service: å¹³å°ä»»åŠ¡æœåŠ¡ï¼ˆå¯é€‰ï¼Œç”¨äºæ—è·¯åŒå†™ï¼‰
        """
        self.dao = dao
        self.llm = llm_orchestrator
        self.jobs_service = jobs_service
        self.settings = get_settings()
        self.feature_flags = get_feature_flags()
        self._docx_extractor: Optional[DocxBlockExtractor] = None
        self._llm_analyzer: Optional[TemplateLlmAnalyzer] = None
        self._deletion_orchestrator: Optional[ProjectDeletionOrchestrator] = None

    @property
    def docx_extractor(self) -> DocxBlockExtractor:
        """å»¶è¿Ÿåˆå§‹åŒ– DocxBlockExtractor"""
        if self._docx_extractor is None:
            self._docx_extractor = DocxBlockExtractor()
        return self._docx_extractor

    @property
    def llm_analyzer(self) -> TemplateLlmAnalyzer:
        """å»¶è¿Ÿåˆå§‹åŒ– TemplateLlmAnalyzer"""
        if self._llm_analyzer is None:
            self._llm_analyzer = TemplateLlmAnalyzer()
        return self._llm_analyzer
    
    @property
    def deletion_orchestrator(self) -> ProjectDeletionOrchestrator:
        """å»¶è¿Ÿåˆå§‹åŒ– ProjectDeletionOrchestrator"""
        if self._deletion_orchestrator is None:
            self._deletion_orchestrator = ProjectDeletionOrchestrator(self.dao.pool)
        return self._deletion_orchestrator

    # ==================== LLM è°ƒç”¨ï¼ˆDuck Typingï¼‰ ====================

    def _llm_text(self, call: LLMCall) -> str:
        """
        è°ƒç”¨ LLM å¹¶è¿”å›æ–‡æœ¬
        ä½¿ç”¨ duck typing å…¼å®¹å¤šç§ orchestrator æ¥å£
        """
        if not self.llm:
            raise RuntimeError("LLM orchestrator not available")

        # å°è¯•å¸¸è§çš„æ–¹æ³•å
        last_error = None
        for method_name in ("chat", "complete", "generate", "run", "ask"):
            fn = getattr(self.llm, method_name, None)
            if not fn:
                continue
            
            try:
                # å°è¯• (messages, model_id) ç­¾å
                res = fn(messages=call.messages, model_id=call.model_id)
                
                # å¤„ç†è¿”å›å€¼
                if isinstance(res, str):
                    return res
                if isinstance(res, dict):
                    # å°è¯•å¸¸è§çš„é”®
                    for k in ("content", "text", "output"):
                        if k in res and isinstance(res[k], str):
                            return res[k]
                    # OpenAI-like æ ¼å¼
                    if "choices" in res and res["choices"]:
                        ch = res["choices"][0]
                        if isinstance(ch, dict):
                            msg = ch.get("message") or {}
                            if isinstance(msg, dict) and isinstance(msg.get("content"), str):
                                return msg["content"]
                # å…œåº•
                return str(res)
            
            except TypeError as e:
                # å°è¯• (prompt, model_id) ç­¾å
                try:
                    prompt = "\n".join([f"{m['role']}: {m['content']}" for m in call.messages])
                    res = fn(prompt=prompt, model_id=call.model_id)
                    return res if isinstance(res, str) else str(res)
                except Exception as inner_e:
                    last_error = inner_e
                    continue
            except Exception as e:
                # ä¿å­˜é”™è¯¯å¹¶ç»§ç»­å°è¯•å…¶ä»–æ–¹æ³•
                last_error = e
                continue

        # å¦‚æœæ‰€æœ‰æ–¹æ³•éƒ½å¤±è´¥äº†ï¼ŒæŠ›å‡ºæœ€åä¸€ä¸ªé”™è¯¯
        if last_error:
            raise RuntimeError(f"LLM call failed: {str(last_error)}") from last_error
        else:
            raise RuntimeError("No compatible LLM method found on orchestrator")

    # ==================== LLM Prompts ====================

    PROJECT_INFO_PROMPT = """
ä½ æ˜¯æ‹›æŠ•æ ‡åŠ©æ‰‹ã€‚è¯·ä»"æ‹›æ ‡æ–‡ä»¶åŸæ–‡ç‰‡æ®µ"ä¸­æŠ½å–é¡¹ç›®ä¿¡æ¯ï¼Œå¹¶è¾“å‡ºä¸¥æ ¼ JSONï¼š
{
  "data": {
    "projectName": "é¡¹ç›®åç§°",
    "ownerName": "æ‹›æ ‡äºº/ä¸šä¸»",
    "agencyName": "ä»£ç†æœºæ„",
    "bidDeadline": "æŠ•æ ‡æˆªæ­¢æ—¶é—´",
    "bidOpeningTime": "å¼€æ ‡æ—¶é—´",
    "budget": "é¢„ç®—é‡‘é¢",
    "maxPrice": "æœ€é«˜é™ä»·",
    "bidBond": "æŠ•æ ‡ä¿è¯é‡‘",
    "schedule": "å·¥æœŸè¦æ±‚",
    "quality": "è´¨é‡è¦æ±‚",
    "location": "é¡¹ç›®åœ°ç‚¹/äº¤ä»˜åœ°ç‚¹",
    "contact": "è”ç³»äººä¸ç”µè¯",

    "technicalParameters": [
      {
        "category": "åŠŸèƒ½/æŠ€æœ¯è¦æ±‚/è®¾å¤‡å‚æ•°/æ€§èƒ½æŒ‡æ ‡/æ¥å£åè®® ç­‰åˆ†ç±»ï¼ˆå¯é€‰ï¼‰",
        "item": "æ¡ç›®æ ‡é¢˜æˆ–åŠŸèƒ½ç‚¹",
        "requirement": "è¦æ±‚æè¿°ï¼ˆå¯åŒ…å«å‹å·ã€æ•°é‡ã€èŒƒå›´ç­‰ï¼‰",
        "parameters": [
          {"name": "å‚æ•°å", "value": "å‚æ•°å€¼/æŒ‡æ ‡", "unit": "å•ä½ï¼ˆå¯ç©ºï¼‰", "remark": "å¤‡æ³¨ï¼ˆå¯ç©ºï¼‰"}
        ],
        "evidence_chunk_ids": ["CHUNK_xxx"]
      }
    ],

    "businessTerms": [
      {
        "term": "æ¡æ¬¾åç§°ï¼ˆä»˜æ¬¾/éªŒæ”¶/è´¨ä¿/äº¤ä»˜/è¿çº¦/å‘ç¥¨/ç¨è´¹/æœåŠ¡/åŸ¹è®­/å”®åç­‰ï¼‰",
        "requirement": "æ¡æ¬¾å†…å®¹ä¸è¦æ±‚ï¼ˆå°½é‡ç»“æ„åŒ–æè¿°ï¼‰",
        "evidence_chunk_ids": ["CHUNK_xxx"]
      }
    ],

    "scoringCriteria": {
      "evaluationMethod": "è¯„æ ‡åŠæ³•/è¯„åˆ†åŠæ³•ï¼ˆå¦‚ç»¼åˆè¯„åˆ†æ³•ã€æœ€ä½è¯„æ ‡ä»·æ³•ç­‰ï¼Œæ²¡æœ‰åˆ™ç©ºå­—ç¬¦ä¸²ï¼‰",
      "items": [
        {
          "category": "è¯„åˆ†å¤§é¡¹ï¼ˆå•†åŠ¡/æŠ€æœ¯/ä»·æ ¼/èµ„ä¿¡/æœåŠ¡ç­‰ï¼‰",
          "item": "è¯„åˆ†ç»†åˆ™/å­é¡¹",
          "score": "åˆ†å€¼ï¼ˆæ•°å­—æˆ–åŸæ–‡ï¼‰",
          "rule": "å¾—åˆ†è§„åˆ™/æ‰£åˆ†æ¡ä»¶/åŠ åˆ†æ¡ä»¶",
          "evidence_chunk_ids": ["CHUNK_xxx"]
        }
      ]
    }
  },
  "evidence_chunk_ids": ["CHUNK_xxx", "CHUNK_yyy"]
}

è¦æ±‚ï¼š
- data é‡Œçš„å­—æ®µå¯ä»¥ä¸ºç©ºå­—ç¬¦ä¸²ï¼›æ•°ç»„å­—æ®µå¯ä¸ºç©ºæ•°ç»„ï¼›scoringCriteria å¯ä»¥ä¸º {} ä½†å¿…é¡»å­˜åœ¨
- technicalParameters/businessTerms/scoringCriteria.items å¦‚æœæ–‡ä¸­æ‰¾ä¸åˆ°å°±è¾“å‡º []
- evidence_chunk_ids å¿…é¡»æ¥è‡ªä¸Šä¸‹æ–‡æ ‡è®°ä¸­çš„ CHUNK id
- ä¸è¦è¾“å‡ºé™¤ JSON ä»¥å¤–çš„ä»»ä½•æ–‡å­—
"""

    RISK_PROMPT = """
ä½ æ˜¯æ‹›æŠ•æ ‡åŠ©æ‰‹ã€‚è¯·ä»"æ‹›æ ‡æ–‡ä»¶åŸæ–‡ç‰‡æ®µ"ä¸­æå–æ‹›æ ‡è¦æ±‚ä¸æ³¨æ„äº‹é¡¹ï¼Œè¾“å‡ºä¸¥æ ¼ JSON æ•°ç»„ï¼š
[
  {
    "risk_type": "mustReject",  // æˆ– "other"
    "title": "é£é™©æ ‡é¢˜",
    "description": "è¯¦ç»†æè¿°",
    "suggestion": "å»ºè®®æªæ–½",
    "severity": "critical",  // low, medium, high, critical
    "tags": ["èµ„æ ¼", "ä¿è¯é‡‘"],
    "evidence_chunk_ids": ["chunk_xxx"]
  }
]

è¦æ±‚ï¼š
- mustRejectï¼šç¼ºå…³é”®èµ„è´¨/æœªæŒ‰è¦æ±‚ç­¾ç« /ä¿è¯é‡‘/æ ¼å¼æ€§åºŸæ ‡ç­‰"å¿…åºŸæ ‡"ç‚¹
- otherï¼šæ˜“é”™ç‚¹ã€æ‰£åˆ†ç‚¹ã€æ—¶é—´èŠ‚ç‚¹ã€è£…è®¢/ä»½æ•°/å¯†å°ç­‰æ³¨æ„äº‹é¡¹
- evidence_chunk_ids å¿…é¡»æ¥è‡ªä¸Šä¸‹æ–‡ CHUNK id
- ä¸è¦è¾“å‡ºé™¤ JSON ä»¥å¤–çš„ä»»ä½•æ–‡å­—
"""

    DIRECTORY_PROMPT = """
ä½ æ˜¯æ‹›æŠ•æ ‡åŠ©æ‰‹ã€‚ä½ è¦ç”Ÿæˆçš„æ˜¯ã€æŠ•æ ‡æ–‡ä»¶/å“åº”æ–‡ä»¶ã€‘çš„ç›®å½•ç»“æ„ï¼Œè€Œä¸æ˜¯æ‹›æ ‡æ–‡ä»¶æœ¬èº«çš„ç›®å½•ã€‚

ä½ ä¼šæ”¶åˆ°â€œæ‹›æ ‡æ–‡ä»¶åŸæ–‡ç‰‡æ®µâ€ï¼ˆå¸¦ CHUNK idï¼‰ã€‚ä½ å¿…é¡»ï¼š
- åªä»æ‹›æ ‡æ–‡ä»¶ä¸­ä¸â€œæŠ•æ ‡æ–‡ä»¶ç»„æˆ/å“åº”æ–‡ä»¶ç»„æˆ/æŠ•æ ‡æ–‡ä»¶æ ¼å¼/æäº¤èµ„æ–™/é™„ä»¶è¡¨æ ¼/æŠ•æ ‡æ–‡ä»¶åº”åŒ…æ‹¬â€ç›¸å…³çš„æ®µè½ä¸­æŠ½å–è¦æ±‚ï¼›
- ç”ŸæˆæŠ•æ ‡æ–‡ä»¶ç›®å½•ï¼ˆæŒ‰æŠ•æ ‡æ–‡ä»¶åº”æäº¤çš„ææ–™æ¥ç»„ç»‡ï¼‰ï¼Œä¸å¾—å¤åˆ»æ‹›æ ‡æ–‡ä»¶ç›®å½•ç« èŠ‚ï¼ˆå¦‚ï¼šæ‹›æ ‡å…¬å‘Šã€æŠ•æ ‡äººé¡»çŸ¥ã€è¯„æ ‡åŠæ³•ã€åˆåŒæ¡æ¬¾ã€æŠ€æœ¯è§„èŒƒæ­£æ–‡ç­‰ï¼‰ã€‚

ã€ä¸¥é‡ç¦æ­¢ã€‘å¦‚æœä½ è¾“å‡ºåŒ…å«ä»¥ä¸‹ä»»ä½•â€œæ‹›æ ‡æ–‡ä»¶ç›®å½•å¼ç« èŠ‚â€ï¼ŒåŸºæœ¬å¯åˆ¤ä¸ºé”™è¯¯ï¼ˆé™¤éæ‹›æ ‡æ˜ç¡®è¦æ±‚æŠ•æ ‡æ–‡ä»¶ä¹Ÿè¦æŒ‰è¿™äº›ç« èŠ‚æäº¤å“åº”ï¼‰ï¼š
- æ‹›æ ‡å…¬å‘Š / æŠ•æ ‡äººé¡»çŸ¥ / è¯„æ ‡åŠæ³• / åˆåŒæ¡æ¬¾ / æŠ€æœ¯è§„èŒƒ / å·¥ç¨‹é‡æ¸…å•è¯´æ˜ / å¼€æ ‡è¯„æ ‡åŠæ³• / èµ„æ ¼é¢„å®¡æ–‡ä»¶ ç­‰ã€‚

è¾“å‡ºä¸¥æ ¼ JSON æ•°ç»„ï¼ˆæŒ‰ numbering ä»å°åˆ°å¤§ï¼‰ï¼š
[
  {
    "numbering": "1",
    "level": 1,
    "title": "æŠ•æ ‡å‡½åŠæŠ•æ ‡å‡½é™„å½•",
    "required": true,
    "notes": "å¯é€‰å¤‡æ³¨ï¼ˆå¦‚ï¼šæŒ‰æ‹›æ ‡æ–‡ä»¶é™„ä»¶æ ¼å¼ï¼‰",
    "evidence_chunk_ids": ["<æ¥è‡ªä¸Šä¸‹æ–‡æ ‡è®°çš„ CHUNK id>"]
  }
]

ç”Ÿæˆè§„åˆ™ï¼š
1) ç›®å½•åº”è¯¥ä½“ç°â€œæŠ•æ ‡äººè¦äº¤å“ªäº›æ–‡ä»¶â€ï¼Œè€Œä¸æ˜¯â€œæ‹›æ ‡æ–‡ä»¶åœ¨è®²ä»€ä¹ˆâ€ã€‚
2) ä¼˜å…ˆä»æ‹›æ ‡æ–‡ä»¶æ˜ç¡®è¦æ±‚ä¸­æŠ½å–ï¼šå¿…é¡»æäº¤/åº”æäº¤/é¡»æä¾›/æŒ‰é™„ä»¶æ ¼å¼/è¡¨X/é™„å½•Xã€‚
3) è‹¥æ‹›æ ‡æ–‡ä»¶æœªæ˜ç¡®åˆ—å‡ºå®Œæ•´ç›®å½•ï¼Œä½ å¯ä»¥æŒ‰è¡Œä¸šæƒ¯ä¾‹è¡¥é½ä¸€ä¸ªåˆç†çš„æŠ•æ ‡æ–‡ä»¶ç»“æ„ï¼ˆrequired=falseï¼Œevidence_chunk_ids=[]ï¼Œnotes å†™â€œè¡Œä¸šæƒ¯ä¾‹è¡¥é½â€ï¼‰ã€‚
4) å¸¸è§ï¼ˆå¯ä½œä¸ºå…œåº•ï¼‰ä¸€çº§ç»“æ„å»ºè®®ï¼ˆæŒ‰ä½ åˆ¤æ–­å¯åˆå¹¶/åˆ å‡ï¼‰ï¼š
   - æŠ•æ ‡å‡½åŠé™„å½•ï¼ˆæŠ•æ ‡å‡½ã€æŠ¥ä»·å‡½/å“åº”å‡½ã€æ³•å®šä»£è¡¨äººèº«ä»½è¯æ˜ã€æˆæƒå§”æ‰˜ä¹¦ï¼‰
   - èµ„æ ¼å®¡æŸ¥æ–‡ä»¶ï¼ˆè¥ä¸šæ‰§ç…§ã€èµ„è´¨ã€ä¸šç»©ã€è´¢åŠ¡ã€ä¿¡èª‰ã€æ‰¿è¯ºç­‰ï¼‰
   - å•†åŠ¡å“åº”æ–‡ä»¶ï¼ˆå•†åŠ¡æ¡æ¬¾å“åº”/åç¦»è¡¨ã€æœåŠ¡æ‰¿è¯ºã€é¡¹ç›®ç®¡ç†/è¿›åº¦è®¡åˆ’ç­‰ï¼‰
   - æŠ€æœ¯å“åº”æ–‡ä»¶ï¼ˆæŠ€æœ¯æ–¹æ¡ˆã€æŠ€æœ¯å‚æ•°å“åº”/åç¦»è¡¨ã€è®¾å¤‡/ç³»ç»Ÿè¯´æ˜ç­‰ï¼‰
   - æŠ¥ä»·æ–‡ä»¶ï¼ˆå¼€æ ‡ä¸€è§ˆè¡¨/æŠ¥ä»·æ±‡æ€»è¡¨ã€åˆ†é¡¹æŠ¥ä»·è¡¨ã€æ¸…å•/æ˜ç»†ï¼‰
   - å…¶ä»–èµ„æ–™ä¸é™„ä»¶ï¼ˆæ‹›æ ‡è¦æ±‚çš„å…¶ä»–è¡¨æ ¼ã€å£°æ˜ã€è¯æ˜ææ–™ï¼‰
5) numbering å¿…é¡»ç”¨ 1/1.1/1.1.1 å½¢å¼ï¼Œlevel ä¸ numbering å¯¹åº”ã€‚
6) evidence_chunk_ids å¿…é¡»æ¥è‡ªä¸Šä¸‹æ–‡æ ‡è®°ä¸­çš„ CHUNK idï¼›æ‰¾ä¸åˆ°è¯æ®å°±ç”¨ []ã€‚
7) ä¸è¦è¾“å‡ºé™¤ JSON ä»¥å¤–çš„ä»»ä½•æ–‡å­—ã€‚
"""

    CUSTOM_RULE_PROMPT = """
ä½ æ˜¯"ä¼ä¸šå†…éƒ¨æ‹›æŠ•æ ‡å®¡æ ¸è§„åˆ™æŠ½å–å™¨"ã€‚è¯·ä»"è§„åˆ™æ–‡ä»¶åŸæ–‡ç‰‡æ®µ"ä¸­æŠ½å–ç»“æ„åŒ–è§„åˆ™ï¼Œè¾“å‡ºä¸¥æ ¼ JSON æ•°ç»„ï¼š
[
  {
    "dimension": "èµ„æ ¼å®¡æŸ¥",  // èµ„æ ¼å®¡æŸ¥|æŠ¥ä»·å®¡æŸ¥|æŠ€æœ¯å®¡æŸ¥|å•†åŠ¡å®¡æŸ¥|å·¥æœŸä¸è´¨é‡|æ–‡æ¡£ç»“æ„|å…¶ä»–
    "title": "è§„åˆ™æ ‡é¢˜",
    "check": "å¯æ‰§è¡Œçš„æ£€æŸ¥æè¿°ï¼ˆæ¸…æ™°ã€å…·ä½“ï¼‰",
    "rigid": true,  // trueè¡¨ç¤ºä¸æ»¡è¶³å°±åº”åˆ¤ fail
    "severity": "high",  // low, medium, high, critical
    "tags": ["èµ„è´¨", "ä¸šç»©"],
    "evidence_chunk_ids": ["chunk_xxx"]
  }
]

è¦æ±‚ï¼š
- rigid=true è¡¨ç¤ºåˆšæ€§è¦æ±‚ï¼Œä¸æ»¡è¶³å°±åº”åˆ¤ fail æˆ– mustReject
- evidence_chunk_ids å¿…é¡»æ¥è‡ªä¸Šä¸‹æ–‡ CHUNK id
- ä¸è¦è¾“å‡ºé™¤ JSON ä»¥å¤–çš„ä»»ä½•æ–‡å­—
"""

    # æ³¨æ„ï¼šCUSTOM_RULE_PROMPT æš‚æ—¶ä¿ç•™ï¼Œç”¨äºæœªæ¥å¯èƒ½çš„è§„åˆ™æ–‡ä»¶è§£æåŠŸèƒ½
    # å½“å‰ç‰ˆæœ¬ä¸­ï¼Œè§„åˆ™æ–‡ä»¶ç›´æ¥ä½œä¸ºåŸæ–‡ç‰‡æ®µå åŠ ï¼Œä¸å†å•ç‹¬æŠ½å–ä¸º JSON è§„åˆ™é›†

    REVIEW_PROMPT = """
ä½ æ˜¯æ‹›æŠ•æ ‡"æŠ•æ ‡æ–‡ä»¶å®¡æ ¸å‘˜"ã€‚ä½ ä¼šæ”¶åˆ°ï¼š
1) æ‹›æ ‡æ–‡ä»¶åŸæ–‡ç‰‡æ®µï¼ˆå¸¦ CHUNK idï¼‰
2) æŠ•æ ‡æ–‡ä»¶åŸæ–‡ç‰‡æ®µï¼ˆå¸¦ CHUNK idï¼‰
3) å¯é€‰ï¼šè‡ªå®šä¹‰å®¡æ ¸è§„åˆ™æ–‡ä»¶åŸæ–‡ç‰‡æ®µï¼ˆå¸¦ CHUNK idï¼Œå¯ä¸ºç©ºï¼‰

è¯·è¾“å‡ºä¸¥æ ¼ JSON æ•°ç»„ï¼š
[
  {
    "dimension": "èµ„æ ¼å®¡æŸ¥",  // èµ„æ ¼å®¡æŸ¥|æŠ¥ä»·å®¡æŸ¥|æŠ€æœ¯å®¡æŸ¥|å•†åŠ¡å®¡æŸ¥|å·¥æœŸä¸è´¨é‡|æ–‡æ¡£ç»“æ„|å…¶ä»–
    "requirement_text": "æ‹›æ ‡è¦æ±‚ï¼ˆæ‘˜è¦ï¼‰",
    "response_text": "æŠ•æ ‡å“åº”ï¼ˆæ‘˜è¦ï¼‰",
    "result": "pass",  // pass, risk, fail
    "remark": "åŸå› /å»ºè®®/ç¼ºå¤±ç‚¹/å†²çªç‚¹",
    "rigid": false,  // æ˜¯å¦åˆšæ€§è¦æ±‚
    "tender_evidence_chunk_ids": ["chunk_xxx"],
    "bid_evidence_chunk_ids": ["chunk_yyy"]
  }
]

è§„åˆ™ï¼š
- ç»“æœå«ä¹‰ï¼špass=æ˜ç¡®ç¬¦åˆï¼›fail=æ˜ç¡®ä¸ç¬¦åˆï¼›risk=ä¸ç¡®å®š/ç¼ºææ–™/å†²çª/éœ€è¦äººå·¥ç¡®è®¤
- è‡ªå®šä¹‰è§„åˆ™æ–‡ä»¶ï¼ˆå¦‚æœ‰ï¼‰ä¸æ‹›æ ‡è¦æ±‚"å åŠ "ï¼šä¹Ÿè¦äº§å‡ºå¯¹åº”çš„å®¡æ ¸é¡¹ï¼ˆå¯åˆå¹¶åˆ°åŒç»´åº¦ï¼‰
- evidence_chunk_ids å¿…é¡»æ¥è‡ªä¸Šä¸‹æ–‡ CHUNK id
- ä¸è¦è¾“å‡ºé™¤ JSON ä»¥å¤–çš„ä»»ä½•æ–‡å­—
"""

    # ==================== æ–‡ä»¶å…¥åº“ ====================

    def _ingest_to_kb(
        self,
        kb_id: str,
        filename: str,
        kind: str,
        bidder_name: Optional[str],
        data: bytes,
    ) -> str:
        """
        REMOVED: Legacy KB ingest path deleted.
        
        This method previously wrote to kb_documents/kb_chunks (deprecated tables).
        All ingest must now go through platform/ingest/v2_service.py (DocStore).
        
        If you see this error, you are using OLD/SHADOW mode which is no longer supported.
        Set INGEST_MODE=NEW_ONLY in your environment.
        """
        raise RuntimeError(
            f"[REMOVED] Legacy tender pipeline (_ingest_to_kb) has been deleted. "
            f"NEW_ONLY mode is required. Use platform/ingest/v2_service.py (DocStore). "
            f"File: {filename}, kind: {kind}"
        )

    def _load_context_by_assets(
        self,
        project_id: str,
        kinds: List[str],
        bidder_name: Optional[str],
        bid_asset_ids: List[str],
        limit: int,
    ) -> Tuple[List[Dict[str, Any]], List[str]]:
        """
        æ ¹æ®èµ„äº§æ¡ä»¶åŠ è½½ä¸Šä¸‹æ–‡ chunks
        
        Args:
            kinds: èµ„äº§ç±»å‹åˆ—è¡¨ï¼ˆå¦‚ ["tender"] æˆ– ["bid"]ï¼‰
            bidder_name: æŠ•æ ‡äººåç§°ï¼ˆç”¨äºè¿‡æ»¤ bidï¼‰
            bid_asset_ids: æŠ•æ ‡èµ„äº§IDåˆ—è¡¨ï¼ˆç²¾ç¡®æŒ‡å®šï¼‰
            limit: æœ€å¤šåŠ è½½å¤šå°‘ä¸ª chunks
        
        Returns:
            (chunks, doc_ids)
        """
        # è·å–æ‰€æœ‰èµ„äº§
        assets = self.dao.list_assets(project_id)
        
        # è¿‡æ»¤èµ„äº§
        filtered = []
        for a in assets:
            if a.get("kind") not in kinds:
                continue
            
            # ç‰¹æ®Šå¤„ç† bid èµ„äº§
            if a.get("kind") == "bid":
                if bid_asset_ids:
                    # ç²¾ç¡®æŒ‡å®š
                    if a.get("id") not in bid_asset_ids:
                        continue
                elif bidder_name:
                    # æŒ‰æŠ•æ ‡äººåç§°è¿‡æ»¤
                    if (a.get("bidder_name") or "") != bidder_name:
                        continue
            
            filtered.append(a)

        # æå– doc_ids å¹¶åŠ è½½ chunks
        doc_ids = [a.get("kb_doc_id") for a in filtered if a.get("kb_doc_id")]
        chunks = self.dao.load_chunks_by_doc_ids(doc_ids, limit=limit)
        return chunks, doc_ids

    # ==================== å…¬å¼€ API ====================

    async def import_assets(
        self,
        project_id: str,
        kind: str,
        files: List[UploadFile],
        bidder_name: Optional[str],
    ) -> List[Dict[str, Any]]:
        """
        é¡¹ç›®å†…ä¸Šä¼ æ–‡ä»¶å¹¶è‡ªåŠ¨ç»‘å®š
        
        Args:
            project_id: é¡¹ç›®ID
            kind: tender | bid | template | custom_rule | company_profile | tech_doc | case_study | finance_doc | cert_doc
            files: ä¸Šä¼ çš„æ–‡ä»¶åˆ—è¡¨
            bidder_name: æŠ•æ ‡äººåç§°ï¼ˆkind=bid æ—¶å¿…å¡«ï¼‰
        
        Returns:
            åˆ›å»ºçš„èµ„äº§åˆ—è¡¨
        """
        # è·å–é¡¹ç›®ä¿¡æ¯
        proj = self.dao.get_project(project_id)
        if not proj:
            raise ValueError("project not found")

        kb_id = proj["kb_id"]
        assets_out = []

        # åˆ›å»ºå­˜å‚¨ç›®å½•
        base_dir = os.path.join("data", "tender_assets", project_id)
        _safe_mkdir(base_dir)

        for f in files:
            b = await f.read()  # å¼‚æ­¥è¯»å–æ–‡ä»¶
            filename = f.filename or "file"
            mime = getattr(f, "content_type", None)
            size = len(b)

            kb_doc_id = None
            storage_path: Optional[str] = None
            doc_version_id = None  # æ–°å¢ï¼šDocStore ç‰ˆæœ¬ID
            tpl_meta = {}  # åˆå§‹åŒ– meta_json
            
            # âœ… æ£€æµ‹æ˜¯å¦ä¸ºå›¾ç‰‡æ–‡ä»¶
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            image_exts = {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff', 'svg'}
            is_image = file_ext in image_exts
            
            # âœ… å¦‚æœæ˜¯å›¾ç‰‡æ–‡ä»¶ï¼Œåªä¿å­˜ä¸å…¥åº“å‘é‡
            if is_image:
                storage_path = os.path.join(base_dir, f"{kind}_{uuid.uuid4().hex}_{filename}")
                with open(storage_path, "wb") as w:
                    w.write(b)
                
                # åˆ›å»ºèµ„äº§è®°å½•ï¼ˆä¸å…¥åº“å‘é‡ï¼‰
                asset = self.dao.create_asset(
                    project_id=project_id,
                    kind=kind,
                    filename=filename,
                    storage_path=storage_path,
                    file_size=size,
                    mime_type=mime,
                    kb_doc_id=None,
                    bidder_name=bidder_name,
                    meta_json={"asset_type": "image", "skip_ingest": True}
                )
                assets_out.append(asset)
                continue  # è·³è¿‡åç»­çš„å…¥åº“æµç¨‹

            if kind == "template":
                # æ¨¡æ¿æ–‡ä»¶ï¼šä¿å­˜åˆ°ç£ç›˜
                storage_path = os.path.join(base_dir, f"{kind}_{uuid.uuid4().hex}_{filename}")
                with open(storage_path, "wb") as w:
                    w.write(b)
                
                # æ–°å¢ï¼šè§£ææ¨¡æ¿ç›®å½•/æ ·å¼æ‘˜è¦ï¼Œå†™å…¥ meta_json
                tpl_meta = self._parse_template_meta(storage_path)
            
            # Step 4: æ–°å…¥åº“é€»è¾‘ï¼ˆcutover æ§åˆ¶ï¼‰
            # åªæ”¯æŒ NEW_ONLY æ¨¡å¼ï¼Œåˆ é™¤OLD/SHADOW/PREFER_NEWåˆ†æ”¯
            ingest_v2_result = None
            
            # âœ… æ‰©å±•ï¼štemplate å’Œä¼ä¸šèµ„æ–™ä¹Ÿéœ€è¦å…¥åº“åˆ°çŸ¥è¯†åº“
            if kind in ("tender", "bid", "custom_rule", "template", "company_profile", "tech_doc", "case_study", "finance_doc", "cert_doc"):
                from app.core.cutover import get_cutover_config
                cutover = get_cutover_config()
                ingest_mode = cutover.get_mode("ingest", project_id)
                tpl_meta["ingest_mode_used"] = ingest_mode.value
                
                # å¼ºåˆ¶è¦æ±‚ NEW_ONLY
                if ingest_mode.value != "NEW_ONLY":
                    raise RuntimeError(
                        f"[REMOVED] Legacy tender pipeline deleted. "
                        f"INGEST_MODE={ingest_mode.value} is no longer supported. "
                        f"Set INGEST_MODE=NEW_ONLY. File: {filename}"
                    )
                
                # åªèµ° NEW_ONLY è·¯å¾„
                from app.platform.ingest.v2_service import IngestV2Service
                from app.services.db.postgres import _get_pool
                pool = _get_pool()
                ingest_v2 = IngestV2Service(pool)
                
                # ç¡®ä¿ storage_path å­˜åœ¨
                if not storage_path:
                    storage_path = os.path.join(base_dir, f"{kind}_{uuid.uuid4().hex}_{filename}")
                    with open(storage_path, "wb") as w:
                        w.write(b)
                
                temp_asset_id = f"temp_{uuid.uuid4().hex}"
                
                # æ˜ å°„æ–‡æ¡£ç±»å‹åˆ°çŸ¥è¯†åº“åˆ†ç±»
                from app.utils.doc_type_mapper import map_doc_type_to_kb_category
                kb_category = map_doc_type_to_kb_category(kind)
                
                ingest_v2_result = await ingest_v2.ingest_asset_v2(
                    project_id=project_id,
                    asset_id=temp_asset_id,
                    file_bytes=b,
                    filename=filename,
                    doc_type=kb_category,  # ä½¿ç”¨æ˜ å°„åçš„çŸ¥è¯†åº“åˆ†ç±»
                    owner_id=proj.get("owner_id"),
                    storage_path=storage_path,
                    kb_id=kb_id,  # âœ… ä¼ é€’ kb_id
                )
                
                # æ–°å…¥åº“æˆåŠŸ
                tpl_meta["doc_version_id"] = ingest_v2_result.doc_version_id
                tpl_meta["ingest_v2_status"] = "success"
                tpl_meta["ingest_v2_segments"] = ingest_v2_result.segment_count
                
                # é‡è¦ï¼šä» doc_version_id è·å– document_id ä½œä¸º kb_doc_id
                try:
                    with self.dao.pool.connection() as conn:
                        with conn.cursor() as cur:
                            cur.execute("""
                                SELECT document_id 
                                FROM document_versions 
                                WHERE id = %s
                            """, (ingest_v2_result.doc_version_id,))
                            row = cur.fetchone()
                            if row:
                                # pool ä½¿ç”¨ dict_row factoryï¼Œæ‰€ä»¥ row æ˜¯ dict
                                kb_doc_id = row['document_id']
                                logger.info(f"âœ“ IngestV2: Got document_id={kb_doc_id} from doc_version_id={ingest_v2_result.doc_version_id}")
                            else:
                                logger.error(f"âŒ IngestV2: document_versions table has no record for doc_version_id={ingest_v2_result.doc_version_id}")
                except Exception as e:
                    logger.error(f"âŒ Failed to query document_id from document_versions: {e}", exc_info=True)
                
                logger.info(
                    f"IngestV2 NEW_ONLY success: "
                    f"doc_version_id={ingest_v2_result.doc_version_id} "
                    f"document_id={kb_doc_id} "
                    f"segments={ingest_v2_result.segment_count}"
                )
            
            # REMOVED: SHADOW mode ingest deleted (lines 612-659)
            # NEW_ONLY is now the only supported mode
            
            # æ—§åŒå†™é€»è¾‘ï¼ˆå…¼å®¹ Step 2ï¼Œå¦‚æœ DOCSTORE_DUALWRITE=true ä¸”æœªè¢« v2 è¦†ç›–ï¼‰
            if self.feature_flags.DOCSTORE_DUALWRITE and "doc_version_id" not in tpl_meta:
                try:
                    from app.platform.docstore.service import DocStoreService
                    from app.services.db.postgres import _get_pool
                    pool = _get_pool()
                    docstore = DocStoreService(pool)
                    
                    document_id = docstore.create_document(
                        namespace="tender",
                        doc_type=kind,
                        owner_id=proj.get("owner_id")
                    )
                    
                    doc_version_id = docstore.create_document_version(
                        document_id=document_id,
                        filename=filename,
                        file_content=b,
                        storage_path=storage_path
                    )
                    
                    tpl_meta["doc_version_id"] = doc_version_id
                    
                except Exception as e:
                    logger.error(f"DocStore dual-write failed: {e}", exc_info=True)
            
            # æ—è·¯è§£æï¼šRuleSetï¼ˆå¦‚æœå¯ç”¨ä¸” kind=custom_ruleï¼‰
            if kind == "custom_rule" and self.feature_flags.RULESET_PARSE_ENABLED:
                try:
                    from app.services.platform.ruleset_service import RuleSetService
                    from app.services.db.postgres import _get_pool
                    pool = _get_pool()
                    ruleset_service = RuleSetService(pool)
                    
                    # 1. è§£ç æ–‡ä»¶å†…å®¹ä¸ºæ–‡æœ¬
                    try:
                        content_text = b.decode('utf-8')
                    except UnicodeDecodeError:
                        # å°è¯•å…¶ä»–ç¼–ç 
                        try:
                            content_text = b.decode('gbk')
                        except Exception:
                            content_text = b.decode('latin-1', errors='ignore')
                    
                    # 2. è§£æå¹¶æ ¡éªŒ
                    is_valid, message, parsed_data = ruleset_service.parse_and_validate(content_text)
                    
                    # 3. åˆ›å»º rule_setï¼ˆå¦‚æœä¸å­˜åœ¨ï¼‰
                    rule_set_id = None
                    existing_rule_sets = ruleset_service.list_rule_sets_by_project(project_id, limit=1)
                    if existing_rule_sets:
                        rule_set_id = existing_rule_sets[0]["id"]
                    else:
                        rule_set_id = ruleset_service.create_rule_set(
                            namespace="tender",
                            scope="project",
                            name=f"è§„åˆ™é›†-{project_id}",
                            project_id=project_id
                        )
                    
                    # 4. åˆ›å»º rule_set_version
                    validate_status = "valid" if is_valid else "invalid"
                    rule_set_version_id = ruleset_service.create_version(
                        rule_set_id=rule_set_id,
                        content_yaml=content_text,
                        validate_status=validate_status,
                        validate_message=message
                    )
                    
                    # 5. å°† rule_set_version_id è®°å½•åˆ° meta_json
                    tpl_meta["rule_set_version_id"] = rule_set_version_id
                    tpl_meta["validate_status"] = validate_status
                    tpl_meta["validate_message"] = message
                    
                    print(f"[INFO] RuleSet parsed: version_id={rule_set_version_id}, status={validate_status}")
                    
                except Exception as e:
                    # é™çº§ï¼šRuleSet è§£æå¤±è´¥ä¸å½±å“ä¸»æµç¨‹
                    print(f"[WARN] Failed to parse RuleSet: {e}")
                    tpl_meta["validate_status"] = "error"
                    tpl_meta["validate_message"] = f"Parsing error: {str(e)}"

            # åˆ›å»ºèµ„äº§è®°å½•
            asset = self.dao.create_asset(
                project_id=project_id,
                kind=kind,
                filename=filename,
                mime_type=mime,
                size_bytes=size,
                kb_doc_id=kb_doc_id,
                storage_path=storage_path,
                bidder_name=bidder_name,
                meta_json=tpl_meta,
            )
            
            # âœ… åŒæ­¥åˆ›å»º kb_documents è®°å½•ï¼ˆè®©æ–‡æ¡£åœ¨çŸ¥è¯†åº“ä¸­å¯è§ï¼‰
            # è¯Šæ–­æ—¥å¿—ï¼šè®°å½•æ˜¯å¦æ»¡è¶³æ¡ä»¶
            logger.info(f"Check kb_document creation: kb_doc_id={kb_doc_id}, kind={kind}, kb_id={kb_id}")
            
            if not kb_doc_id:
                logger.warning(f"âš ï¸ kb_doc_id is None for {filename}, skipping kb_documents creation. doc_version_id={tpl_meta.get('doc_version_id')}")
            elif kind not in ("tender", "bid", "custom_rule", "template"):  # âœ… æ·»åŠ  template
                logger.warning(f"âš ï¸ kind '{kind}' not in allowed list for {filename}, skipping kb_documents creation")
            else:
                try:
                    from app.services.dao import kb_dao
                    
                    # è®¡ç®—æ–‡ä»¶å“ˆå¸Œ
                    import hashlib
                    content_hash = hashlib.sha256(b).hexdigest()
                    
                    # æ˜ å°„æ–‡æ¡£åˆ†ç±»
                    if kind == "tender":
                        kb_category = "tender_doc"
                    elif kind == "bid":
                        kb_category = "bid_doc"
                    elif kind == "custom_rule":
                        kb_category = "custom_rule"
                    elif kind == "template":
                        kb_category = "template_doc"
                    else:
                        kb_category = "general_doc"
                    
                    logger.info(f"Creating kb_document: kb_id={kb_id}, doc_id={kb_doc_id}, filename={filename}, category={kb_category}")
                    
                    # åˆ›å»º kb_documents è®°å½•ï¼ˆä½¿ç”¨å·²æœ‰çš„ document_idï¼‰
                    kb_dao.create_kb_document_with_id(
                        kb_id=kb_id,
                        doc_id=kb_doc_id,  # ä½¿ç”¨ documents.id
                        filename=filename,
                        source="tender_upload",
                        content_hash=content_hash,
                        status="ready",
                        kb_category=kb_category,
                        meta={
                            "project_id": project_id,
                            "asset_id": asset["id"],
                            "kind": kind,
                            "bidder_name": bidder_name,
                            "doc_version_id": tpl_meta.get("doc_version_id"),
                            "size": size,
                        }
                    )
                    logger.info(f"âœ“ Successfully created kb_document: kb_id={kb_id}, doc_id={kb_doc_id}, filename={filename}")
                except Exception as e:
                    logger.error(f"âŒ Failed to create kb_document for {filename}: {e}", exc_info=True)
                    # ä¸å½±å“ä¸»æµç¨‹ï¼Œç»§ç»­æ‰§è¡Œ
            
            assets_out.append(asset)

        return assets_out

    def list_assets(self, project_id: str) -> List[Dict[str, Any]]:
        """åˆ—å‡ºé¡¹ç›®çš„æ‰€æœ‰èµ„äº§"""
        return self.dao.list_assets(project_id)

    def delete_asset(self, project_id: str, asset_id: str):
        """
        åˆ é™¤èµ„äº§
        - åˆ é™¤çŸ¥è¯†åº“æ–‡æ¡£åŠå…¶chunksï¼ˆå¦‚æœæœ‰ï¼‰
        - åˆ é™¤ç£ç›˜æ–‡ä»¶ï¼ˆå¦‚æœæ˜¯æ¨¡æ¿æ–‡ä»¶ï¼‰
        - åˆ é™¤é¡¹ç›®æ–‡æ¡£ç»‘å®šè®°å½•ï¼ˆå…¼å®¹æ—§APIï¼‰
        - åˆ é™¤æ•°æ®åº“assetè®°å½•
        - å¦‚æœåˆ é™¤çš„æ–‡æ¡£è¢«é¡¹ç›®ä¿¡æ¯ã€é£é™©ã€ç›®å½•ã€å®¡æ ¸ç­‰å¼•ç”¨ï¼Œç›¸å…³çš„evidence_chunk_idsä¼šè‡ªåŠ¨å¤±æ•ˆ
        """
        import os
        import logging
        logger = logging.getLogger(__name__)
        
        # è·å–èµ„äº§ä¿¡æ¯
        asset = self.dao.get_asset_by_id(asset_id)
        if not asset:
            raise ValueError("Asset not found")
        
        if asset["project_id"] != project_id:
            raise ValueError("Asset does not belong to this project")
        
        # è·å–é¡¹ç›®ä¿¡æ¯ï¼ˆç”¨äºè·å– kb_idï¼‰
        proj = self.dao.get_project(project_id)
        if not proj:
            raise ValueError("Project not found")
        
        kb_id = proj.get("kb_id")
        kb_doc_id = asset.get("kb_doc_id")
        
        # 1. åˆ é™¤çŸ¥è¯†åº“æ–‡æ¡£åŠå…¶chunks
        if kb_doc_id and kb_id:
            try:
                from app.services import kb_service
                # ä¼ å…¥ skip_asset_cleanup=Trueï¼Œé¿å…å¾ªç¯è°ƒç”¨
                # kb_service.delete_document ä¼šåˆ é™¤ï¼š
                # - kb_documents è®°å½•
                # - kb_chunks è®°å½•
                # - milvus å‘é‡
                # ä½†ä¸ä¼šå†æ¬¡åˆ é™¤ assetï¼ˆé¿å…å¾ªç¯ï¼‰
                kb_service.delete_document(kb_id, kb_doc_id, skip_asset_cleanup=True)
                logger.info(f"Deleted KB document {kb_doc_id} from knowledge base {kb_id}")
            except Exception as e:
                # è®°å½•é”™è¯¯ä½†ç»§ç»­åˆ é™¤ï¼Œé¿å…å­¤å„¿æ•°æ®
                logger.warning(f"Failed to delete KB document {kb_doc_id}: {e}")
        
        # 2. åˆ é™¤ç£ç›˜æ–‡ä»¶ï¼ˆæ¨¡æ¿æ–‡ä»¶ï¼‰
        if asset.get("storage_path"):
            try:
                storage_path = asset["storage_path"]
                if os.path.exists(storage_path):
                    os.remove(storage_path)
                    logger.info(f"Deleted file {storage_path}")
            except Exception as e:
                logger.warning(f"Failed to delete file {storage_path}: {e}")
        
        # 3. åˆ é™¤é¡¹ç›®æ–‡æ¡£ç»‘å®šè®°å½•ï¼ˆå…¼å®¹æ—§APIï¼‰
        if kb_doc_id:
            try:
                self.dao._execute(
                    "DELETE FROM tender_project_documents WHERE project_id=%s AND kb_doc_id=%s",
                    (project_id, kb_doc_id)
                )
                logger.info(f"Deleted project document binding for doc {kb_doc_id}")
            except Exception as e:
                logger.warning(f"Failed to delete project document binding: {e}")
        
        # 4. åˆ é™¤assetè®°å½•
        self.dao.delete_asset(asset_id)
        logger.info(f"Deleted asset {asset_id} from project {project_id}")
        
        # æ³¨æ„ï¼šä¸éœ€è¦æ˜¾å¼åˆ é™¤ tender_project_info, tender_directory_nodes, 
        # tender_review_items ä¸­å¼•ç”¨è¯¥æ–‡æ¡£ chunks çš„æ•°æ®ï¼Œå› ä¸ºï¼š
        # - è¿™äº›è¡¨ä¸­çš„ evidence_chunk_ids æ˜¯æ•°ç»„ç±»å‹ï¼Œåˆ é™¤chunkåè¿™äº›IDä¼šè‡ªç„¶å¤±æ•ˆ
        # - ä¿ç•™è¿™äº›è®°å½•å¯ä»¥è®©ç”¨æˆ·çŸ¥é“æ›¾ç»æœ‰å“ªäº›åˆ†æç»“æœï¼Œåªæ˜¯è¯æ®é“¾æ–­äº†
        # - å¦‚æœéœ€è¦é‡æ–°åˆ†æï¼Œç”¨æˆ·å¯ä»¥é‡æ–°ä¸Šä¼ æ–‡æ¡£å¹¶è¿è¡Œç›¸åº”çš„æå–ä»»åŠ¡

    def extract_project_info(
        self,
        project_id: str,
        model_id: Optional[str],
        run_id: Optional[str] = None,
        owner_id: Optional[str] = None,
    ):
        """
        æŠ½å–é¡¹ç›®ä¿¡æ¯
        
        REMOVED: Only NEW_ONLY mode supported.
        OLD/SHADOW/PREFER_NEW modes have been deleted.
        """
        # å¼ºåˆ¶æ£€æŸ¥æ¨¡å¼
        from app.core.cutover import get_cutover_config
        cutover = get_cutover_config()
        extract_mode = cutover.get_mode("extract", project_id)
        
        if extract_mode.value != "NEW_ONLY":
            raise RuntimeError(
                f"[REMOVED] Legacy tender extraction deleted. "
                f"EXTRACT_MODE={extract_mode.value} is no longer supported. "
                f"Set EXTRACT_MODE=NEW_ONLY. Method: extract_project_info"
            )
        
        # æ—è·¯åŒå†™ï¼šåˆ›å»º platform jobï¼ˆå¦‚æœå¯ç”¨ï¼‰
        job_id = None
        if self.feature_flags.PLATFORM_JOBS_ENABLED and self.jobs_service and run_id:
            try:
                job_id = self.jobs_service.create_job(
                    namespace="tender",
                    biz_type="extract_project_info",
                    biz_id=project_id,
                    owner_id=owner_id,
                    initial_status="running",
                    initial_message="æ­£åœ¨æå–é¡¹ç›®ä¿¡æ¯..."
                )
            except Exception as e:
                # é™çº§ï¼šjob åˆ›å»ºå¤±è´¥ä¸å½±å“ä¸»æµç¨‹
                print(f"[WARN] Failed to create platform job: {e}")
        
        try:
            # åªæ‰§è¡Œ NEW_ONLY è·¯å¾„
            import asyncio
            from app.works.tender.extract_v2_service import ExtractV2Service
            from app.services.db.postgres import _get_pool
            
            logger.info(f"NEW_ONLY extract_project_info: using v2 for project={project_id}")
            pool = _get_pool()
            extract_v2 = ExtractV2Service(pool, self.llm)
            
            # v2 æŠ½å–
            v2_result = asyncio.run(extract_v2.extract_project_info_v2(
                project_id=project_id,
                model_id=model_id,
                run_id=run_id
            ))
            
            # âœ… v2_result å·²ç»æ˜¯å®Œæ•´çš„ V3 ç»“æ„ï¼ŒåŒ…å« schema_version å’Œå…­å¤§ç±»
            # æå–è¯æ®å’Œè¿½è¸ªä¿¡æ¯
            eids = v2_result.get("evidence_chunk_ids") or []
            trace = v2_result.get("retrieval_trace") or {}
            
            # âœ… æ„å»ºè¦ä¿å­˜çš„æ•°æ®ï¼šåªä¿ç•™æ ¸å¿ƒå­—æ®µï¼ˆå…­å¤§ç±» + schema_versionï¼‰
            data_to_save = {
                "schema_version": v2_result.get("schema_version", "tender_info_v3"),
                "project_overview": v2_result.get("project_overview", {}),
                "bidder_qualification": v2_result.get("bidder_qualification", {}),
                "evaluation_and_scoring": v2_result.get("evaluation_and_scoring", {}),
                "business_terms": v2_result.get("business_terms", {}),
                "technical_requirements": v2_result.get("technical_requirements", {}),
                "document_preparation": v2_result.get("document_preparation", {}),
            }
            
            obj = {"data_json": data_to_save, "evidence_chunk_ids": eids}
            
            # âœ… æ•°æ®å·²ç»åœ¨_extract_project_info_stagedä¸­ä¿å­˜è¿‡äº†ï¼Œè¿™é‡Œä¸éœ€è¦é‡å¤ä¿å­˜
            # åªæ›´æ–°runçŠ¶æ€å³å¯
            logger.info(f"é¡¹ç›®ä¿¡æ¯æå–å®Œæˆï¼Œå‡†å¤‡æ›´æ–°runçŠ¶æ€: project={project_id}")
            
            # æ›´æ–°è¿è¡ŒçŠ¶æ€
            if run_id:
                # æ„å»ºresult_jsonï¼Œç¡®ä¿traceæ˜¯dictæ ¼å¼
                result_json_data = {
                    **obj,
                    "extract_v2_status": "ok",
                    "extract_mode_used": "NEW_ONLY",
                }
                
                # å¤„ç†traceï¼šå¦‚æœæ˜¯liståˆ™åŒ…è£…æˆdictï¼Œå¦‚æœæ˜¯dictåˆ™å±•å¼€
                if isinstance(trace, dict):
                    result_json_data.update(trace)
                elif isinstance(trace, list):
                    result_json_data["retrieval_trace"] = trace
                else:
                    result_json_data["retrieval_trace"] = trace
                
                self.dao.update_run(
                    run_id, "success", progress=1.0, 
                    message="é¡¹ç›®ä¿¡æ¯æå–å®Œæˆ", 
                    result_json=result_json_data
                )
            
            logger.info(f"NEW_ONLY extract_project_info: v2 succeeded for project={project_id}")
            
            # æ—è·¯åŒå†™ï¼šæ›´æ–° job æˆåŠŸï¼ˆå¦‚æœå¯ç”¨ï¼‰
            if job_id and self.jobs_service:
                try:
                    self.jobs_service.finish_job_success(
                        job_id=job_id,
                        result={"summary": "é¡¹ç›®ä¿¡æ¯æå–å®Œæˆ"},
                        message="æˆåŠŸ"
                    )
                except Exception as e:
                    print(f"[WARN] Failed to update platform job: {e}")
        
        except Exception as e:
            logger.error(f"é¡¹ç›®ä¿¡æ¯æå–å¤±è´¥: {e}", exc_info=True)
            
            # æ›´æ–°runçŠ¶æ€ä¸ºå¤±è´¥
            if run_id:
                self.dao.update_run(
                    run_id, "failed", progress=0, 
                    message=f"æå–å¤±è´¥: {str(e)[:200]}"
                )
            
            # æ›´æ–° job å¤±è´¥çŠ¶æ€ï¼ˆå¦‚æœå¯ç”¨ï¼‰
            if job_id and self.jobs_service:
                try:
                    self.jobs_service.finish_job_fail(job_id=job_id, error=str(e))
                except Exception as je:
                    print(f"[WARN] Failed to update platform job on error: {je}")
            # é‡æ–°æŠ›å‡ºåŸå§‹å¼‚å¸¸
            raise

    # extract_risks å·²åˆ é™¤
    # è¯·ä½¿ç”¨ POST /api/apps/tender/projects/{project_id}/extract/risks (å®é™…è°ƒç”¨requirements_v1)
    # risksæ¨¡å—å·²åºŸå¼ƒï¼Œç»Ÿä¸€ä½¿ç”¨requirementsæ¨¡å—

    def _filter_chunks_for_bid_directory(self, chunks: List[Dict[str, Any]], limit: int = 80) -> List[Dict[str, Any]]:
        keywords = [
            "æŠ•æ ‡æ–‡ä»¶", "å“åº”æ–‡ä»¶", "æ–‡ä»¶ç»„æˆ", "åº”åŒ…æ‹¬", "é¡»æä¾›", "æäº¤", "é™„ä»¶", "æ ¼å¼", "è¡¨", "é™„å½•",
            "æŠ•æ ‡å‡½", "æˆæƒå§”æ‰˜ä¹¦", "æ³•å®šä»£è¡¨äºº", "å¼€æ ‡ä¸€è§ˆè¡¨", "æŠ¥ä»·", "åˆ†é¡¹æŠ¥ä»·", "æ¸…å•",
            "èµ„æ ¼", "èµ„è´¨", "ä¸šç»©", "è´¢åŠ¡", "ä¿¡èª‰", "æ‰¿è¯º", "åç¦»è¡¨", "æŠ€æœ¯å“åº”", "å•†åŠ¡å“åº”"
        ]
        scored = []
        for c in chunks:
            t = (c.get("content") or "")
            s = 0
            for k in keywords:
                if k in t:
                    s += 2
            # è½»å¾®åŠ åˆ†ï¼šè¶ŠçŸ­è¶Šå¯èƒ½æ˜¯æ¡æ¬¾/æ¸…å•ç±»
            if 0 < len(t) < 1200:
                s += 1
            scored.append((s, c))
        scored.sort(key=lambda x: x[0], reverse=True)
        # è‡³å°‘ä¿ç•™ä¸€äº›å¤´éƒ¨ï¼Œé¿å…å…¨æ˜¯0åˆ†
        top = [c for s, c in scored if s > 0][:limit]
        if len(top) < 20:
            top = [c for _, c in scored[:limit]]
        return top

    def _looks_like_tender_toc(self, nodes: List[Dict[str, Any]]) -> bool:
        tender_like = ["æ‹›æ ‡å…¬å‘Š", "æŠ•æ ‡äººé¡»çŸ¥", "è¯„æ ‡åŠæ³•", "åˆåŒæ¡æ¬¾", "æŠ€æœ¯è§„èŒƒ", "å·¥ç¨‹é‡æ¸…å•", "å¼€æ ‡", "èµ„æ ¼é¢„å®¡"]
        titles = " ".join([(n.get("title") or "") for n in nodes[:20] if isinstance(n, dict)])
        hit = sum(1 for k in tender_like if k in titles)
        return hit >= 2  # å‘½ä¸­ä¸¤ä¸ªä»¥ä¸ŠåŸºæœ¬å¯åˆ¤ä¸ºè·‘å

    def generate_directory(
        self,
        project_id: str,
        model_id: Optional[str],
        run_id: Optional[str] = None,
    ):
        """ç”Ÿæˆç›®å½• - ä½¿ç”¨ V2 å¼•æ“"""
        
        # ğŸ” DEBUG: å†™å…¥è°ƒè¯•æ—¥å¿—
        debug_log = open("/app/tender_service_debug.log", "a")
        debug_log.write(f"\n=== TenderService.generate_directory START ===\n")
        debug_log.write(f"project_id: {project_id}\n")
        debug_log.write(f"model_id: {model_id}\n")
        debug_log.write(f"run_id: {run_id}\n")
        debug_log.flush()
        
        # 1. æ£€æŸ¥æ¨¡å¼
        from app.core.cutover import get_cutover_config
        cutover = get_cutover_config()
        extract_mode = cutover.get_mode("extract", project_id)
        debug_log.write(f"extract_mode: {extract_mode.value}\n")
        debug_log.flush()
        
        if extract_mode.value != "NEW_ONLY":
            debug_log.write(f"æ¨¡å¼æ£€æŸ¥å¤±è´¥ï¼Œé€€å‡º\n")
            debug_log.close()
            raise RuntimeError("Legacy directory generation deleted. Set EXTRACT_MODE=NEW_ONLY")
        
        # 2. åˆ›å»º platform job (å¯é€‰)
        if self.jobs_service:
            try:
                job_id = self.jobs_service.create_job(
                    job_type="extract",
                    project_id=project_id,
                    run_id=run_id
                )
                logger.info(f"[generate_directory] platform job created: {job_id}")
            except Exception as e:
                logger.warning(f"[generate_directory] Failed to create platform job: {e}")
        
        # 3. è°ƒç”¨ V2 æŠ½å–æœåŠ¡
        from app.works.tender.extract_v2_service import ExtractV2Service
        from app.services.db.postgres import _get_pool
        from app.platform.utils.async_runner import run_async
        
        debug_log.write(f"å‡†å¤‡è°ƒç”¨ ExtractV2Service\n")
        debug_log.flush()
        
        pool = _get_pool()
        extract_v2 = ExtractV2Service(pool, self.llm)
        
        debug_log.write(f"å¼€å§‹è°ƒç”¨ generate_directory_v2...\n")
        debug_log.flush()
        
        v2_result = run_async(extract_v2.generate_directory_v2(
            project_id=project_id,
            model_id=model_id,
            run_id=run_id
        ))
        
        debug_log.write(f"generate_directory_v2 è¿”å›: keys={list(v2_result.keys())}\n")
        debug_log.flush()
        
        # 4. æå– nodes å’Œç”Ÿæˆæ¨¡å¼ä¿¡æ¯
        nodes = v2_result.get("data", {}).get("nodes", [])
        if not nodes:
            raise ValueError("Directory nodes empty")
        
        # ä¿å­˜ç”Ÿæˆæ¨¡å¼ä¿¡æ¯
        generation_mode = v2_result.get("generation_mode", "llm")
        fast_stats = v2_result.get("fast_stats", {})
        
        logger.info(f"[generate_directory] V2 extracted {len(nodes)} nodes, mode={generation_mode}")
        
        # âœ… 4.1 é€šç”¨ç›®å½•è§„èŒƒåŒ–ï¼ˆæ–°å¢ï¼šwrapperæŠ˜å  + ä¸‰åˆ†å†Œä¸€çº§ + è¯­ä¹‰çº åï¼‰
        nodes = self._normalize_directory_nodes(nodes)
        logger.info(f"[generate_directory] normalized nodes -> {len(nodes)}")
        
        # 5. åå¤„ç†: æ’åº + æ„å»ºæ ‘ + ç”Ÿæˆ numbering
        # ğŸ”¥ å¦‚æœæ˜¯ä»æ‹›æ ‡ä¹¦åŸæ–‡æå–çš„ï¼Œè·³è¿‡æ ‘æ„å»ºï¼ˆå› ä¸ºå·²ç»æœ‰æ­£ç¡®çš„parent_idå’Œorder_noï¼‰
        if generation_mode == "extracted_from_tender":
            logger.info(f"[generate_directory] è·³è¿‡æ ‘æ„å»ºï¼ˆå·²ä»æ‹›æ ‡ä¹¦æå–å®Œæ•´ç»“æ„ï¼‰")
            nodes_with_tree = nodes  # ç›´æ¥ä½¿ç”¨å·²æœ‰ç»“æ„
        else:
            nodes_sorted = self._sort_directory_nodes_for_tree(nodes)
            nodes_with_tree = self._build_directory_tree(nodes_sorted)
        
        # 6. ä¿å­˜ï¼ˆä½¿ç”¨replace_directoryï¼‰
        # âŒ ç¦ç”¨ï¼šèŠ‚ç‚¹å·²åœ¨extract_v2_service.pyçš„_save_nodes_to_dbä¸­ä¿å­˜
        # é‡å¤ä¿å­˜ä¼šå¯¼è‡´sourceè¢«è¦†ç›–ä¸º"tender"
        logger.info(f"[generate_directory] è·³è¿‡replace_directoryï¼ˆèŠ‚ç‚¹å·²åœ¨_save_nodes_to_dbä¸­ä¿å­˜ï¼‰")
        
        # âœ¨ 7. è‡ªåŠ¨å¡«å……èŒƒæœ¬ï¼ˆé›†æˆï¼šä¸€é”®å®Œæˆç›®å½•ç”Ÿæˆ+èŒƒæœ¬å¡«å……ï¼‰
        try:
            logger.info(f"[generate_directory] Starting auto_fill_samples for project {project_id}")
            diag = self.auto_fill_samples(project_id)
            
            # è®°å½•å¡«å……ç»“æœ
            attached = diag.get("attached_sections", 0)
            extracted = diag.get("tender_fragments_upserted", 0)
            
            if diag.get("ok"):
                logger.info(
                    f"[generate_directory] auto_fill_samples success: "
                    f"extracted {extracted} fragments, attached {attached} sections"
                )
            else:
                warnings = diag.get("warnings", [])
                logger.warning(
                    f"[generate_directory] auto_fill_samples partial success: "
                    f"attached {attached} sections, warnings: {warnings}"
                )
        except Exception as e:
            logger.error(f"[generate_directory] auto_fill_samples failed: {type(e).__name__}: {e}")
        
        # 8. æ›´æ–°çŠ¶æ€
        if run_id:
            self.dao.update_run(
                run_id,
                "success",
                progress=1.0,
                message="Directory generated with auto-filled samples",
                result_json=v2_result
            )
    
    # ==================== ç›®å½•è§„èŒƒåŒ–æ–¹æ³•ï¼ˆé€šç”¨ç‰ˆï¼‰ ====================
    
    def _bucket_by_title(self, title: str) -> str:
        """æ ¹æ®æ ‡é¢˜å†…å®¹åˆ¤æ–­æ‰€å±åˆ†æ¡¶"""
        import re
        _BUCKET_PRICE = re.compile(r"(æŠ¥ä»·|ä»·æ ¼|æ˜ç»†|æ±‡æ€»|æ€»ä»·|åˆ†é¡¹|æŠ¥ä»·è¡¨|æŠ¥ä»·å•|æŠ•æ ‡æŠ¥ä»·|ç£‹å•†æŠ¥ä»·|æŠ¥ä»·å“åº”)", re.I)
        _BUCKET_TECH  = re.compile(r"(æŠ€æœ¯|æ–¹æ¡ˆ|è§„æ ¼|å‚æ•°|åç¦»|æ ·æœ¬|æ‰‹å†Œ|å®æ–½|ç»„ç»‡|æ¶æ„|æµ‹è¯•|é…ç½®|é€‰å‹|æŠ€æœ¯è§„æ ¼)", re.I)
        _BUCKET_BIZ   = re.compile(r"(è¥ä¸šæ‰§ç…§|èµ„è´¨|è¯ä¹¦|ç¤¾ä¿|ä¿¡ç”¨|æˆæƒ|å§”æ‰˜|æ‰¿è¯º|å£°æ˜|åŸºæœ¬æƒ…å†µ|ä¿¡èª‰|è‡ªè¯„|è¯æ˜|å»ºè®®|ä¸è½¬åŒ…|åˆ†åŒ…)", re.I)
        
        t = (title or "").strip()
        if not t:
            return "unknown"
        if _BUCKET_PRICE.search(t):
            return "price"
        if _BUCKET_TECH.search(t):
            return "tech"
        if _BUCKET_BIZ.search(t):
            return "biz"
        return "unknown"
    
    def _infer_parent_index_by_level(self, nodes: list) -> list:
        """æ ¹æ® level æ¨æ–­çˆ¶èŠ‚ç‚¹ç´¢å¼•"""
        parent = [-1] * len(nodes)
        stack = []  # [(level, index)]
        for i, n in enumerate(nodes):
            lv = int(n.get("level") or 1)
            while stack and stack[-1][0] >= lv:
                stack.pop()
            parent[i] = stack[-1][1] if stack else -1
            stack.append((lv, i))
        return parent
    
    def _find_section_titles(self, nodes: list) -> dict:
        """æŸ¥æ‰¾ä¸‰åˆ†å†Œå’Œ wrapper æ ‡é¢˜"""
        import re
        _WRAPPER_RE = re.compile(r"(æŠ•æ ‡æ–‡ä»¶|å“åº”æ–‡ä»¶|ç£‹å•†å“åº”æ–‡ä»¶|æŠ•æ ‡å“åº”æ–‡ä»¶|å“åº”æ–‡ä»¶ç›®å½•|æŠ•æ ‡æ–‡ä»¶ç›®å½•)", re.I)
        _SECTION_BIZ_RE = re.compile(r"(èµ„ä¿¡|å•†åŠ¡|èµ„æ ¼)", re.I)
        _SECTION_TECH_RE = re.compile(r"(æŠ€æœ¯)", re.I)
        _SECTION_PRICE_RE = re.compile(r"(æŠ¥ä»·|ä»·æ ¼|ç£‹å•†æŠ¥ä»·|æŠ¥ä»·å“åº”)", re.I)
        
        biz = tech = price = wrapper = None
        for n in nodes:
            title = (n.get("title") or "").strip()
            if not title:
                continue
            if wrapper is None and _WRAPPER_RE.search(title):
                wrapper = title
            if biz is None and _SECTION_BIZ_RE.search(title):
                biz = title
            if tech is None and _SECTION_TECH_RE.search(title):
                tech = title
            if price is None and _SECTION_PRICE_RE.search(title):
                price = title
        return {"biz": biz, "tech": tech, "price": price, "wrapper": wrapper}
    
    def _collapse_wrapper(self, nodes: list) -> list:
        """æŠ˜å  wrapper èŠ‚ç‚¹ï¼ˆæŠ•æ ‡æ–‡ä»¶/å“åº”æ–‡ä»¶ç­‰æ€»æ ‡é¢˜ï¼‰"""
        from collections import deque
        
        if not nodes:
            return nodes
        sec = self._find_section_titles(nodes)
        wrapper = sec["wrapper"]
        if not wrapper:
            return nodes
        if not (sec["biz"] and sec["tech"] and sec["price"]):
            return nodes

        title_to_first_idx = {}
        for i, n in enumerate(nodes):
            t = (n.get("title") or "").strip()
            if t and t not in title_to_first_idx:
                title_to_first_idx[t] = i

        w_idx = title_to_first_idx.get(wrapper)
        if w_idx is None:
            return nodes

        parent = self._infer_parent_index_by_level(nodes)
        children = [[] for _ in nodes]
        for i, p in enumerate(parent):
            if p >= 0:
                children[p].append(i)

        sub = set()
        q = deque([w_idx])
        while q:
            x = q.popleft()
            sub.add(x)
            for c in children[x]:
                q.append(c)

        # ä¸‰åˆ†å†Œå¿…é¡»éƒ½åœ¨ wrapper å­æ ‘é‡Œæ‰æŠ˜å ï¼ˆé¿å…è¯¯ä¼¤ï¼‰
        b = title_to_first_idx.get(sec["biz"])
        t = title_to_first_idx.get(sec["tech"])
        p = title_to_first_idx.get(sec["price"])
        if not (b in sub and t in sub and p in sub):
            return nodes

        new_nodes = []
        for i, n in enumerate(nodes):
            if i == w_idx:
                continue  # remove wrapper
            nn = dict(n)
            if i in sub:
                lv = int(nn.get("level") or 1)
                nn["level"] = max(1, lv - 1)
                title = (nn.get("title") or "").strip()
                if title in (sec["biz"], sec["tech"], sec["price"]):
                    nn["parent_ref"] = ""
            new_nodes.append(nn)

        return new_nodes
    
    def _ensure_sections_are_level1(self, nodes: list) -> list:
        """ç¡®ä¿ä¸‰åˆ†å†Œä¸ºä¸€çº§æ ‡é¢˜"""
        if not nodes:
            return nodes
        sec = self._find_section_titles(nodes)
        if not (sec["biz"] and sec["tech"] and sec["price"]):
            return nodes

        title_to_first_idx = {}
        for i, n in enumerate(nodes):
            t = (n.get("title") or "").strip()
            if t and t not in title_to_first_idx:
                title_to_first_idx[t] = i

        idxs = [title_to_first_idx.get(sec["biz"]), title_to_first_idx.get(sec["tech"]), title_to_first_idx.get(sec["price"])]
        if any(i is None for i in idxs):
            return nodes

        new_nodes = [dict(n) for n in nodes]
        for i in idxs:
            new_nodes[i]["level"] = 1
            new_nodes[i]["parent_ref"] = ""

        # é¡¶å±‚åˆ†å†Œ order_no æŒ‰å‡ºç°é¡ºåºé‡æ’
        top_seq = 1
        for i in sorted(idxs):
            new_nodes[i]["order_no"] = top_seq
            top_seq += 1

        return new_nodes
    
    def _rebucket_to_sections(self, nodes: list) -> list:
        """è¯­ä¹‰åˆ†æ¡¶çº åï¼šæŠŠæ¡ç›®æŒ‚åˆ°æ­£ç¡®åˆ†å†Œ"""
        from collections import defaultdict
        
        if not nodes:
            return nodes
        sec = self._find_section_titles(nodes)
        if not (sec["biz"] and sec["tech"] and sec["price"]):
            return nodes

        biz_title, tech_title, price_title = sec["biz"], sec["tech"], sec["price"]
        section_titles = {biz_title, tech_title, price_title}

        new_nodes = [dict(n) for n in nodes]

        # åˆ¤æ–­æ˜¯å¦"å…¨æŒ‚æŠ¥ä»·"çš„å…¸å‹é”™æŒ‚ï¼Œè§¦å‘ aggressive
        cnt = defaultdict(int)
        for n in new_nodes:
            pr = (n.get("parent_ref") or "").strip()
            if pr in section_titles:
                cnt[pr] += 1
        aggressive = (cnt.get(biz_title, 0) == 0 and cnt.get(tech_title, 0) == 0 and cnt.get(price_title, 0) >= 6)

        for n in new_nodes:
            title = (n.get("title") or "").strip()
            if not title:
                continue
            if title in section_titles:
                continue
            import re
            _WRAPPER_RE = re.compile(r"(æŠ•æ ‡æ–‡ä»¶|å“åº”æ–‡ä»¶|ç£‹å•†å“åº”æ–‡ä»¶|æŠ•æ ‡å“åº”æ–‡ä»¶|å“åº”æ–‡ä»¶ç›®å½•|æŠ•æ ‡æ–‡ä»¶ç›®å½•)", re.I)
            if _WRAPPER_RE.search(title):
                continue  # å…œåº•ï¼šwrapperæ®‹ç•™ä¸å¤„ç†

            bucket = self._bucket_by_title(title)
            if bucket == "unknown":
                continue

            target_parent = {"biz": biz_title, "tech": tech_title, "price": price_title}[bucket]
            cur_pr = (n.get("parent_ref") or "").strip()

            # aggressive æˆ–è€…æ˜æ˜¾é”™æŒ‚/æ— æŒ‚è½½ -> çº å
            if aggressive or cur_pr in ("", price_title) or cur_pr not in section_titles:
                n["parent_ref"] = target_parent
                n["level"] = 2  # å‹ç¼©åˆ°åˆ†å†Œä¸‹äºŒçº§ï¼Œä¿è¯ç¨³å®šå¯ç”¨

        # åˆ†å†Œä¸‹äºŒçº§èŠ‚ç‚¹ order_no ç¨³å®šé‡æ’ï¼ˆæŒ‰åŸå‡ºç°é¡ºåºï¼‰
        bucket_items = defaultdict(list)
        for idx, n in enumerate(new_nodes):
            if int(n.get("level") or 1) == 2:
                pr = (n.get("parent_ref") or "").strip()
                if pr in section_titles:
                    bucket_items[pr].append((idx, n))

        for pr, items in bucket_items.items():
            items.sort(key=lambda x: x[0])
            seq = 1
            for _, n in items:
                n["order_no"] = seq
                seq += 1

        # å…œåº•ï¼šä»ä¸ºç©ºçš„äºŒçº§èŠ‚ç‚¹ï¼Œé»˜è®¤å½’å…¥èµ„ä¿¡åŠå•†åŠ¡ï¼ˆæ¯”å…¨æŒ‚æŠ¥ä»·æ›´åˆç†ï¼‰
        for n in new_nodes:
            if int(n.get("level") or 1) > 1 and not (n.get("parent_ref") or "").strip():
                n["parent_ref"] = biz_title
                n["level"] = 2

        return new_nodes
    
    def _normalize_directory_nodes(self, nodes: list) -> list:
        """é€šç”¨ç›®å½•è§„èŒƒåŒ–ï¼šwrapperæŠ˜å  + ä¸‰åˆ†å†Œä¸€çº§ + è¯­ä¹‰çº å"""
        # âŒ ç¦ç”¨æ‰€æœ‰è§„èŒƒåŒ–é€»è¾‘ï¼ˆè‡ªåŠ¨åˆ›é€ ä¸‰åˆ†å†Œçš„ç½ªé­ç¥¸é¦–ï¼‰
        # ç›´æ¥è¿”å›åŸå§‹nodesï¼Œä¸åšä»»ä½•ä¿®æ”¹
        return nodes or []
    
    def _sort_directory_nodes_for_tree(self, nodes: list) -> list:
        """
        å…³é”®ï¼šä¸èƒ½æŒ‰ (level, order_no) å…¨å±€æ’åºï¼Œå¦åˆ™æ‰€æœ‰ level=2 ä¼šå †åˆ°æœ€åï¼Œ
        _build_directory_tree ç”¨æ ˆæ¨ parent æ—¶ä¼šå…¨éƒ¨æŒ‚åˆ°æœ€åä¸€ä¸ª level=1ï¼ˆé€šå¸¸æ˜¯æŠ¥ä»·æ–‡ä»¶ï¼‰ã€‚
        è¿™é‡ŒæŒ‰ "æ‰€å±åˆ†å†Œ(root)" åˆ†ç»„åå†æ’åºï¼Œä¿è¯æ¯ä¸ªåˆ†å†Œçš„å­èŠ‚ç‚¹ç´§è·Ÿå…¶åã€‚
        """
        nodes = nodes or []
        # æ‰¾ä¸€çº§åˆ†å†Œï¼ˆæŒ‰å‡ºç°é¡ºåºï¼‰
        top = [n for n in nodes if int(n.get("level") or 1) == 1]
        top = sorted(top, key=lambda n: n.get("order_no", 999))

        section_titles = [ (n.get("title") or "").strip() for n in top if (n.get("title") or "").strip() ]
        section_order = {t: i for i, t in enumerate(section_titles)}

        def root_key(n: dict) -> int:
            title = (n.get("title") or "").strip()
            lv = int(n.get("level") or 1)
            if lv == 1 and title in section_order:
                return section_order[title]
            pr = (n.get("parent_ref") or "").strip()
            if pr in section_order:
                return section_order[pr]
            # fallbackï¼šæœªçŸ¥çš„å…ˆæ”¾æœ€å
            return 999

        return sorted(
            nodes,
            key=lambda n: (
                root_key(n),
                int(n.get("level") or 99),
                int(n.get("order_no") or 999),
            )
        )
    
    def _build_directory_tree(self, nodes: List[Dict]) -> List[Dict]:
        """æ„å»ºç›®å½•æ ‘: ä¸¤éæ³• - å…ˆç”Ÿæˆidå’Œtitleæ˜ å°„ï¼Œå†ç»Ÿä¸€åˆ†é…parent_id"""
        import uuid
        
        # 0) é¢„å¤„ç†ï¼šç¡®ä¿æ¯ä¸ªèŠ‚ç‚¹æœ‰ id
        out = []
        for n in nodes:
            nn = dict(n)
            nn.setdefault("id", f"node_{uuid.uuid4().hex[:16]}")
            nn["title"] = (nn.get("title") or "").strip()
            nn["parent_ref"] = (nn.get("parent_ref") or "").strip()
            nn["level"] = int(nn.get("level") or 1)
            nn["order_no"] = int(nn.get("order_no") or 0)
            out.append(nn)

        # 1) title -> first idï¼ˆç”¨äº parent_ref è§£æï¼‰
        title_to_first_id = {}
        for n in out:
            if n["title"] and n["title"] not in title_to_first_id:
                title_to_first_id[n["title"]] = n["id"]

        # 2) åˆ†é… parent_idï¼šä¼˜å…ˆ parent_refï¼Œå…¶æ¬¡ fallback level æ ˆ
        stack = []  # (level, id)
        for n in out:
            if n["parent_ref"] and n["parent_ref"] in title_to_first_id:
                n["parent_id"] = title_to_first_id[n["parent_ref"]]
            else:
                # fallbackï¼šæ ˆæ¨æ–­ï¼ˆåªå¯¹ç¡®å®ç¼º parent_ref çš„æƒ…å†µï¼‰
                while stack and stack[-1][0] >= n["level"]:
                    stack.pop()
                n["parent_id"] = stack[-1][1] if stack else None

            stack.append((n["level"], n["id"]))

        # 3) æ„ children æ˜ å°„
        children = {}
        by_id = {n["id"]: n for n in out}
        for n in out:
            children.setdefault(n["id"], [])
        roots = []
        for n in out:
            pid = n.get("parent_id")
            if not pid:
                roots.append(n)
            else:
                children.setdefault(pid, []).append(n)

        # 4) æ¯ä¸ªçˆ¶èŠ‚ç‚¹ä¸‹æŒ‰ order_no æ’åºï¼Œé€’å½’ç”Ÿæˆ numbering + æ‰å¹³åŒ–è¾“å‡º
        def sort_k(x: dict):
            return (int(x.get("order_no") or 0), x.get("title") or "")

        for pid in list(children.keys()):
            children[pid].sort(key=sort_k)

        flat = []
        def walk(node: dict, prefix: str):
            # ğŸ”¥ ä¿ç•™åŸå§‹numberingï¼Œä¸è¦é‡æ–°ç”Ÿæˆ
            flat.append(node)
            kids = children.get(node["id"], [])
            for idx, c in enumerate(kids, start=1):
                # ğŸ”¥ åªæœ‰å½“numberingä¸ºç©ºæˆ–ä¸å­˜åœ¨æ—¶ï¼Œæ‰è‡ªåŠ¨ç”Ÿæˆ
                if not c.get("numbering"):
                    c["numbering"] = f"{prefix}.{idx}" if prefix else str(idx)
                walk(c, c["numbering"])

        # roots å¿…é¡»ç¨³å®šæ’åºï¼ˆä¼˜å…ˆæŒ‰level=1ï¼Œç„¶åæŒ‰order_noï¼‰
        def root_sort_key(r):
            # ä¼˜å…ˆæŒ‰numberingä¸­çš„æ•°å­—æ’åºï¼ˆæå–ï¼ˆ1ï¼‰ã€ï¼ˆ2ï¼‰ä¸­çš„æ•°å­—ï¼‰
            import re
            numbering = r.get("numbering", "")
            match = re.search(r'[ï¼ˆ\(](\d+)[ï¼‰\)]', numbering)
            if match:
                return (0, int(match.group(1)))  # L1èŠ‚ç‚¹ï¼ŒæŒ‰ï¼ˆ1ï¼‰ã€ï¼ˆ2ï¼‰æ’åº
            return (1, int(r.get("order_no") or 999))  # å…¶ä»–èŠ‚ç‚¹æŒ‰order_noæ’åº
        
        roots.sort(key=root_sort_key)
        
        for idx, r in enumerate(roots, start=1):
            # ğŸ”¥ åªæœ‰å½“numberingä¸ºç©ºæˆ–ä¸å­˜åœ¨æ—¶ï¼Œæ‰è‡ªåŠ¨ç”Ÿæˆ
            if not r.get("numbering"):
                r["numbering"] = str(idx)
            walk(r, r["numbering"])

        # ğŸ”¥ é‡æ–°åˆ†é…order_noï¼ŒæŒ‰ç…§æ·±åº¦ä¼˜å…ˆéå†çš„é¡ºåº
        for i, node in enumerate(flat, start=1):
            node["order_no"] = i

        return flat
    
    def _pick_latest_asset(self, assets: List[Dict[str, Any]], require_storage_path: bool = False) -> Optional[Dict[str, Any]]:
        """
        ä»èµ„äº§åˆ—è¡¨ä¸­é€‰æ‹©â€œæœ€æ–°â€çš„ä¸€æ¡ï¼š
        - ä¼˜å…ˆæŒ‰ created_at DESC
        - created_at ç¼ºå¤±åˆ™ä¿æŒåŸé¡ºåºï¼Œå–æœ€åä¸€ä¸ªï¼ˆæ›´æ¥è¿‘â€œæœ€æ–°ä¸Šä¼ â€ï¼‰
        - require_storage_path=True æ—¶ï¼Œä¼šä¼˜å…ˆé€‰æ‹© storage_path éç©ºçš„è®°å½•
        """
        if not assets:
            return None

        cands = assets
        if require_storage_path:
            with_path = [a for a in assets if (a.get("storage_path") or "").strip()]
            cands = with_path or assets

        def _key(a: Dict[str, Any]):
            # psycopg/RealDictCursor é€šå¸¸ä¼šè¿”å› datetimeï¼›ç¼ºå¤±åˆ™ç”¨ç©ºä¸²å…œåº•
            return a.get("created_at") or ""

        try:
            cands_sorted = sorted(cands, key=_key, reverse=True)
            return cands_sorted[0] if cands_sorted else None
        except Exception:
            return cands[-1] if cands else None

    def _auto_extract_and_attach_samples(self, project_id: str):
        """è‡ªåŠ¨æŠ½å–æ‹›æ ‡ä¹¦èŒƒæœ¬å¹¶æŒ‚è½½åˆ°ç›®å½•èŠ‚ç‚¹"""
        import logging
        logger = logging.getLogger(__name__)

        # 1. è·å–æ‹›æ ‡ä¹¦æ–‡ä»¶
        assets = self.dao.list_assets(project_id)
        tender_assets = [a for a in assets if a.get("kind") == "tender"]
        # é€‰æ‹©â€œæœ€æ–°ä¸”å¯ç”¨â€çš„ tenderï¼ˆä¼˜å…ˆ storage_path å­˜åœ¨ï¼‰
        tender_asset = self._pick_latest_asset(tender_assets, require_storage_path=True)
        
        if not tender_asset:
            return
        
        # 2. è·å–æ–‡ä»¶è·¯å¾„
        storage_path = tender_asset.get("storage_path")
        if not storage_path:
            logger.warning(f"[samples] tender asset has no storage_path, skip. project_id={project_id}")
            return

        # æ”¯æŒ DOCX å’Œ PDF æ ¼å¼
        ext = os.path.splitext(storage_path.lower())[1]
        if ext not in [".docx", ".pdf"]:
            logger.warning(f"[samples] tender asset is not docx/pdf, skip. project_id={project_id}, storage_path={storage_path}")
            return

        if not os.path.exists(storage_path):
            logger.warning(f"[samples] tender file not found on disk, skip. project_id={project_id}, storage_path={storage_path}")
            return
        
        # 3. æŠ½å–èŒƒæœ¬ï¼ˆæ”¯æŒDOCXå’ŒPDFï¼‰
        from app.services.fragment.fragment_extractor import TenderSampleFragmentExtractor
        logger.info(f"[samples] extracting fragments from tender file ({ext}). project_id={project_id}, path={storage_path}")
        extractor = TenderSampleFragmentExtractor(self.dao)
        extractor.extract_and_upsert(
            project_id=project_id,
            tender_docx_path=storage_path,
            file_key=storage_path,
        )

        try:
            fragments_count = len(self.dao.list_fragments("PROJECT", project_id))
        except Exception:
            fragments_count = -1
        logger.info(f"[samples] extracted fragments. project_id={project_id}, count={fragments_count}")
        
        # 4. æŒ‚è½½åˆ°ç›®å½•èŠ‚ç‚¹
        from app.services.fragment.outline_attacher import OutlineSampleAttacher
        nodes = self.dao.list_directory(project_id)
        attacher = OutlineSampleAttacher(self.dao)
        attached_count = attacher.attach(project_id, nodes)
        logger.info(f"[samples] attached fragments to outline nodes. project_id={project_id}, attached_count={attached_count}")

    def save_directory(self, project_id: str, nodes: List[Dict[str, Any]]):
        """ä¿å­˜ç›®å½•ï¼ˆç”¨æˆ·ç¼–è¾‘åï¼‰"""
        self.dao.replace_directory(project_id, nodes)
    
    def get_directory_with_body_meta(self, project_id: str) -> List[Dict]:
        """
        è·å–ç›®å½•ï¼ˆå¸¦æ­£æ–‡å…ƒä¿¡æ¯ï¼‰
        ä¸ºæ¯ä¸ªèŠ‚ç‚¹é™„åŠ  bodyMeta ä¿¡æ¯
        """
        nodes = self.dao.list_directory(project_id)
        bodies = self.dao.list_section_bodies(project_id)
        
        # æ„å»º node_id -> body æ˜ å°„
        body_map = {b["node_id"]: b for b in bodies}
        
        # ä¸ºæ¯ä¸ªèŠ‚ç‚¹æ·»åŠ  bodyMeta
        for node in nodes:
            node_id = node.get("id")
            body = body_map.get(node_id)
            
            if body:
                node["bodyMeta"] = {
                    "source": body.get("source"),
                    "fragmentId": body.get("fragment_id"),
                    "hasContent": bool(body.get("content_html")),
                }
            else:
                node["bodyMeta"] = {
                    "source": "EMPTY",
                    "fragmentId": None,
                    "hasContent": False,
                }
        
        return nodes
    
    def _render_snippet_blocks_to_html(self, blocks: List[Dict]) -> str:
        """
        å°†æ ¼å¼èŒƒæ–‡çš„blocksæ¸²æŸ“ä¸ºHTML
        
        Args:
            blocks: doc_blocksæ ¼å¼çš„blocksåˆ—è¡¨
        
        Returns:
            HTMLå­—ç¬¦ä¸²
        """
        html_parts = []
        
        for block in blocks:
            block_type = block.get("type", "")
            
            # æ®µè½
            if block_type == "p":
                text = block.get("text", "").strip()
                if text:
                    # ç®€å•å¤„ç†æ¢è¡Œå’Œç©ºæ ¼
                    text_html = text.replace("\n", "<br>")
                    html_parts.append(f"<p>{text_html}</p>")
            
            # è¡¨æ ¼
            elif block_type == "table":
                rows = block.get("rows", [])
                if not rows:
                    continue
                
                # ä½¿ç”¨æ›´çœŸå®çš„è¡¨æ ¼æ ·å¼
                html_parts.append('''<table style="
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 16px 0;
                    border: 1px solid #d0d0d0;
                    font-size: 14px;
                ">''')
                
                # è¡¨å¤´
                if len(rows) > 0:
                    html_parts.append("<thead><tr>")
                    for cell in rows[0]:
                        cell_text = str(cell).replace("\n", "<br>")
                        html_parts.append(f'''<th style="
                            background-color: #f5f5f5; 
                            font-weight: 600; 
                            text-align: center; 
                            padding: 10px 8px;
                            border: 1px solid #d0d0d0;
                            color: #333;
                        ">{cell_text}</th>''')
                    html_parts.append("</tr></thead>")
                
                # è¡¨ä½“
                if len(rows) > 1:
                    html_parts.append("<tbody>")
                    for row in rows[1:]:
                        html_parts.append("<tr>")
                        for cell in row:
                            cell_text = str(cell).replace("\n", "<br>")
                            html_parts.append(f'''<td style="
                                padding: 10px 8px;
                                border: 1px solid #d0d0d0;
                                vertical-align: top;
                            ">{cell_text}</td>''')
                        html_parts.append("</tr>")
                    html_parts.append("</tbody>")
                
                html_parts.append("</table>")
        
        return "\n".join(html_parts)
    
    def get_section_body_content(self, project_id: str, node_id: str) -> Optional[Dict]:
        """
        è·å–ç« èŠ‚æ­£æ–‡å†…å®¹
        - ä¼˜å…ˆæ£€æŸ¥èŠ‚ç‚¹çš„ meta_json.snippet_blocksï¼ˆæ ¼å¼èŒƒæ–‡ï¼‰
        - å¦‚æœæ˜¯ç”¨æˆ·ç¼–è¾‘å†…å®¹ï¼Œè¿”å› HTML
        - å¦‚æœæ˜¯èŒƒæœ¬æŒ‚è½½ï¼Œè¿”å›ç®€åŒ–çš„é¢„è§ˆHTMLï¼ˆä»æºæ–‡æ¡£ç”Ÿæˆï¼‰
        - å¦‚æœæ˜¯PDFè¯­ä¹‰åŒ¹é…ï¼Œç›´æ¥è¿”å›content_html
        """
        # é¦–å…ˆæ£€æŸ¥èŠ‚ç‚¹æ˜¯å¦æœ‰æ ¼å¼èŒƒæ–‡
        with self.dao.pool.connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT meta_json
                    FROM tender_directory_nodes
                    WHERE project_id = %s AND id = %s
                """, (project_id, node_id))
                node_row = cur.fetchone()
                
                if node_row:
                    meta_json = node_row.get("meta_json") or {}
                    snippet_blocks = meta_json.get("snippet_blocks")
                    snippet_id = meta_json.get("snippet_id")
                    
                    if snippet_blocks and isinstance(snippet_blocks, list) and len(snippet_blocks) > 0:
                        # æœ‰æ ¼å¼èŒƒæ–‡ï¼Œæ¸²æŸ“snippet_blocksä¸ºHTML
                        html = self._render_snippet_blocks_to_html(snippet_blocks)
                        return {
                            "source": "SNIPPET",
                            "contentHtml": html,
                            "fragmentId": None,
                            "snippetId": snippet_id
                        }
        
        # æ²¡æœ‰æ ¼å¼èŒƒæ–‡ï¼Œæ£€æŸ¥ project_section_body
        body = self.dao.get_section_body(project_id, node_id)
        if not body:
            return None
        
        source = body.get("source")
        
        # ç”¨æˆ·ç¼–è¾‘å†…å®¹
        if source == "USER" and body.get("content_html"):
            return {
                "source": source,
                "contentHtml": body["content_html"],
                "fragmentId": body.get("fragment_id"),
            }
        
        # PDFè¯­ä¹‰åŒ¹é… - ç›´æ¥è¿”å›content_htmlï¼ˆå·²åœ¨æŒ‚è½½æ—¶æå–ï¼‰
        if source == "PDF_SEMANTIC_MATCH" and body.get("content_html"):
            return {
                "source": source,
                "contentHtml": body["content_html"],
                "fragmentId": body.get("fragment_id"),
            }
        
        # èŒƒæœ¬æŒ‚è½½ - ç”Ÿæˆé¢„è§ˆHTMLï¼ˆå³ä½¿å¤±è´¥ä¹Ÿå¿…é¡»è¿”å›éç©º contentHtmlï¼‰
        if source == "TEMPLATE_SAMPLE":
            import logging
            logger = logging.getLogger(__name__)

            frag_id = body.get("fragment_id")
            if not frag_id:
                return {
                    "source": source,
                    "contentHtml": "<div class='template-sample-preview' style='color:#92400e'>[è¯¥ç« èŠ‚å·²æŒ‚è½½èŒƒæœ¬ï¼Œä½† fragment_id ä¸ºç©º]</div>",
                    "fragmentId": None,
                }

            fragment = self.dao.get_fragment_by_id(frag_id)
            if not fragment:
                return {
                    "source": source,
                    "contentHtml": "<div class='template-sample-preview' style='color:#92400e'>[è¯¥ç« èŠ‚å·²æŒ‚è½½èŒƒæœ¬ï¼Œä½†ç‰‡æ®µä¸å­˜åœ¨æˆ–å·²è¢«åˆ é™¤]</div>",
                    "fragmentId": frag_id,
                }

            src_docx_path = fragment.get("source_file_key")
            start_idx = fragment.get("start_body_index")
            end_idx = fragment.get("end_body_index")

            if not src_docx_path:
                logger.warning(f"[samples] fragment has no source_file_key. project_id={project_id}, node_id={node_id}, fragment_id={fragment.get('id')}")
                return {
                    "source": source,
                    "contentHtml": "<div class='template-sample-preview' style='color:#92400e'>[å·²æŒ‚è½½èŒƒæœ¬ï¼šæºæ–‡ä»¶è·¯å¾„ç¼ºå¤±ï¼Œå¯¼å‡ºæ—¶å¯èƒ½æ— æ³•æ‹·è´]</div>",
                    "fragmentId": fragment.get("id"),
                }
            if not str(src_docx_path).lower().endswith(".docx"):
                logger.warning(f"[samples] fragment source is not docx. project_id={project_id}, node_id={node_id}, fragment_id={fragment.get('id')}, source={src_docx_path}")
                return {
                    "source": source,
                    "contentHtml": "<div class='template-sample-preview' style='color:#92400e'>[å·²æŒ‚è½½èŒƒæœ¬ï¼šæºæ–‡ä»¶ä¸æ˜¯ docxï¼Œæ— æ³•é¢„è§ˆ]</div>",
                    "fragmentId": fragment.get("id"),
                }
            if start_idx is None or end_idx is None:
                logger.warning(f"[samples] fragment indices missing. project_id={project_id}, node_id={node_id}, fragment_id={fragment.get('id')}")
                return {
                    "source": source,
                    "contentHtml": "<div class='template-sample-preview' style='color:#92400e'>[å·²æŒ‚è½½èŒƒæœ¬ï¼šç‰‡æ®µèŒƒå›´ç¼ºå¤±ï¼Œæ— æ³•é¢„è§ˆ]</div>",
                    "fragmentId": fragment.get("id"),
                }

            from app.services.fragment.fragment_preview import render_fragment_html
            try:
                html = render_fragment_html(str(src_docx_path), int(start_idx), int(end_idx))
                if not (html or "").strip():
                    html = "<div class='template-sample-preview' style='color:#92400e'>[å·²æŒ‚è½½èŒƒæœ¬ï¼šé¢„è§ˆä¸ºç©º]</div>"
                return {
                    "source": source,
                    "contentHtml": html,
                    "fragmentId": fragment.get("id"),
                }
            except Exception as e:
                logger.warning(f"[samples] render_fragment_html failed: {e}. project_id={project_id}, node_id={node_id}, fragment_id={fragment.get('id')}")
                return {
                    "source": source,
                    "contentHtml": f"<div class='template-sample-preview' style='color:#b00020'>[èŒƒæœ¬é¢„è§ˆæ¸²æŸ“å¤±è´¥: {str(e)}]</div>",
                    "fragmentId": fragment.get("id"),
                }
        
        return {
            "source": source or "EMPTY",
            "contentHtml": "",
            "fragmentId": None,
        }
    
    def update_section_body(self, project_id: str, node_id: str, content_html: str):
        """æ›´æ–°ç« èŠ‚æ­£æ–‡ï¼ˆç”¨æˆ·ç¼–è¾‘ï¼‰"""
        # ä¿ç•™åŸæœ‰çš„ fragment_idï¼ˆç”¨äºæ¢å¤èŒƒæœ¬ï¼‰
        existing = self.dao.get_section_body(project_id, node_id)
        fragment_id = existing.get("fragment_id") if existing else None
        
        self.dao.upsert_section_body(
            project_id=project_id,
            node_id=node_id,
            source="USER",
            fragment_id=fragment_id,  # ä¿ç•™ä»¥ä¾¿æ¢å¤
            content_html=content_html,
        )
    
    def restore_sample_for_section(self, project_id: str, node_id: str):
        """æ¢å¤ç« èŠ‚çš„èŒƒæœ¬å†…å®¹"""
        # æŸ¥æ‰¾è¯¥èŠ‚ç‚¹çš„ç›®å½•ä¿¡æ¯
        nodes = self.dao.list_directory(project_id)
        node = next((n for n in nodes if n.get("id") == node_id), None)
        
        if not node:
            return
        
        # å°è¯•é‡æ–°åŒ¹é…èŒƒæœ¬
        from app.services.fragment.fragment_matcher import FragmentTitleMatcher
        matcher = FragmentTitleMatcher()
        
        node_title = node.get("title", "")
        node_title_norm = matcher.normalize(node_title)
        
        ftype = matcher.match_type(node_title_norm)
        if not ftype:
            # æ— æ³•åŒ¹é…ï¼Œæ¸…ç©ºæ­£æ–‡
            self.dao.upsert_section_body(
                project_id=project_id,
                node_id=node_id,
                source="EMPTY",
                fragment_id=None,
                content_html=None,
            )
            return
        
        # æŸ¥æ‰¾æœ€ä½³ç‰‡æ®µ
        candidates = self.dao.find_fragments_by_type("PROJECT", project_id, str(ftype))
        if not candidates:
            # æ²¡æœ‰åŒ¹é…ç‰‡æ®µ
            self.dao.upsert_section_body(
                project_id=project_id,
                node_id=node_id,
                source="EMPTY",
                fragment_id=None,
                content_html=None,
            )
            return
        
        # ä½¿ç”¨ç¬¬ä¸€ä¸ªå€™é€‰ï¼ˆå·²æŒ‰ç½®ä¿¡åº¦æ’åºï¼‰
        best = candidates[0]
        self.dao.upsert_section_body(
            project_id=project_id,
            node_id=node_id,
            source="TEMPLATE_SAMPLE",
            fragment_id=best["id"],
            content_html=None,
        )

    # ==================== èŒƒæœ¬ç‰‡æ®µï¼šåˆ—è¡¨ + é¢„è§ˆï¼ˆç›®å½•é¡µä¾§è¾¹æ ï¼‰ ====================

    def list_sample_fragments(self, project_id: str) -> List[Dict[str, Any]]:
        """
        åˆ—å‡ºæœ¬é¡¹ç›®ä¸‹æŠ½å–åˆ°çš„èŒƒæœ¬ç‰‡æ®µï¼ˆè½»é‡åˆ—è¡¨ï¼Œä¸å«å¤§æ­£æ–‡ï¼‰ã€‚
        """
        rows = self.dao.list_sample_fragments(project_id)
        out: List[Dict[str, Any]] = []
        for r in rows or []:
            out.append(
                {
                    "id": r.get("id"),
                    "title": r.get("title") or "",
                    "fragment_type": r.get("fragment_type") or "",
                    "confidence": r.get("confidence"),
                }
            )
        return out

    def get_sample_fragment_preview(self, project_id: str, fragment_id: str, max_elems: int = 60) -> Dict[str, Any]:
        """
        è·å–å•æ¡èŒƒæœ¬ç‰‡æ®µçš„é¢„è§ˆï¼ˆæ‡’åŠ è½½ï¼‰ã€‚

        - é˜²è¶Šæƒï¼šowner_id å¿…é¡»ç­‰äº project_idï¼ˆä¸” owner_type=PROJECTï¼‰
        - ä» source_file_key æ‰“å¼€ docxï¼ŒæŒ‰ [start_body_index..end_body_index] æŠ½å–å†…å®¹æ¸²æŸ“ä¸ºç®€æ˜“ HTML
        """
        fragment = self.dao.get_fragment_by_id(fragment_id)
        if not fragment:
            raise ValueError("fragment not found")

        if str(fragment.get("owner_type") or "") != "PROJECT" or str(fragment.get("owner_id") or "") != str(project_id):
            raise PermissionError("fragment not in project")

        title = fragment.get("title") or ""
        ftype = fragment.get("fragment_type") or ""
        warnings: List[str] = []

        src = (fragment.get("source_file_key") or "").strip()
        start_idx = fragment.get("start_body_index")
        end_idx = fragment.get("end_body_index")

        if not src:
            warnings.append("source_file_key_missing")
            return {
                "id": fragment_id,
                "title": title,
                "fragment_type": ftype,
                "preview_html": "<div style='color:#92400e'>[æºæ–‡ä»¶ç¼ºå¤±ï¼Œæ— æ³•é¢„è§ˆ]</div>",
                "text_len": 0,
                "warnings": warnings,
            }

        if not str(src).lower().endswith(".docx"):
            warnings.append("source_not_docx")
            return {
                "id": fragment_id,
                "title": title,
                "fragment_type": ftype,
                "preview_html": "<div style='color:#92400e'>[æºæ–‡ä»¶ä¸æ˜¯ docxï¼Œæ— æ³•é¢„è§ˆ]</div>",
                "text_len": 0,
                "warnings": warnings,
            }

        if not os.path.exists(src):
            warnings.append("source_docx_not_found")
            return {
                "id": fragment_id,
                "title": title,
                "fragment_type": ftype,
                "preview_html": "<div style='color:#92400e'>[æº docx ä¸å­˜åœ¨/æœªè½ç›˜ï¼Œæ— æ³•é¢„è§ˆ]</div>",
                "text_len": 0,
                "warnings": warnings,
            }

        if start_idx is None or end_idx is None:
            warnings.append("range_missing")
            return {
                "id": fragment_id,
                "title": title,
                "fragment_type": ftype,
                "preview_html": "<div style='color:#92400e'>[ç‰‡æ®µèŒƒå›´ç¼ºå¤±ï¼Œæ— æ³•é¢„è§ˆ]</div>",
                "text_len": 0,
                "warnings": warnings,
            }

        from app.services.fragment.fragment_preview import build_fragment_preview_meta

        html_out, text_len, w = build_fragment_preview_meta(
            docx_path=str(src),
            start=int(start_idx),
            end=int(end_idx),
            max_elems=int(max_elems or 60),
        )
        warnings.extend(list(w or []))

        return {
            "id": fragment_id,
            "title": title,
            "fragment_type": ftype,
            "preview_html": html_out,
            "text_len": int(text_len or 0),
            "warnings": warnings,
        }
    
    def auto_fill_samples(self, project_id: str) -> Dict[str, Any]:
        """
        è‡ªåŠ¨å¡«å……æ‰€æœ‰ç« èŠ‚çš„èŒƒæœ¬ï¼ˆæ°¸ä¸æŠ›å¼‚å¸¸ï¼Œå¿…é¡»è¿”å›è¯Šæ–­ dictï¼Œç¦æ­¢ return Noneï¼‰

        å›ºå®šè¿”å›ç»“æ„ï¼š
        {
          "ok": bool,
          "project_id": str,
          "tender_asset_id": str|None,
          "tender_filename": str|None,
          "tender_storage_path": str|None,
          "storage_path_exists": bool,
          "needs_reupload": bool,
          "tender_fragments_upserted": int,   # æœ¬æ¬¡ä»æ‹›æ ‡ä¹¦æ–°æŠ½å‡ºæ¥å¹¶å†™åº“çš„æ•°é‡
          "tender_fragments_total": int,      # å½“å‰é¡¹ç›®åº“é‡Œ fragment æ€»æ•°ï¼ˆowner=PROJECTï¼‰
          "attached_sections_template_sample": int, # æœ¬æ¬¡æŒ‚è½½ source=TEMPLATE_SAMPLE çš„ç« èŠ‚æ•°
          "attached_sections_builtin": int,         # æœ¬æ¬¡æŒ‚è½½ source=BUILTIN_SAMPLE çš„ç« èŠ‚æ•°
          # å…¼å®¹å­—æ®µï¼ˆæ—§å‰ç«¯/è„šæœ¬ï¼‰
          "extracted_fragments": int,
          "attached_sections": int,
          "warnings": [str, ...],
          "nodes": [...å¸¦ bodyMeta...],
        }
        """
        import logging

        logger = logging.getLogger(__name__)

        # å›ºå®šç»“æ„åˆå§‹åŒ–ï¼ˆç¡®ä¿ä»»ä½•åˆ†æ”¯éƒ½ä¸ä¼šç¼ºå­—æ®µï¼‰
        diag: Dict[str, Any] = {
            "ok": False,
            "project_id": project_id,
            "tender_asset_id": None,
            "tender_filename": None,
            "tender_storage_path": None,
            "storage_path_exists": False,
            "needs_reupload": False,
            "tender_fragments_upserted": 0,
            "tender_fragments_total": 0,
            "attached_sections_template_sample": 0,
            "attached_sections_builtin": 0,
            # å…¼å®¹å­—æ®µ
            "extracted_fragments": 0,
            "attached_sections": 0,
            "warnings": [],
            "nodes": [],
            # è¯Šæ–­å¢å¼ºå­—æ®µ
            "tender_storage_path_ext": "",
            "body_items_count": 0,
            "fragments_detected_by_rules": 0,
            "llm_used": False,
        }

        try:
            warnings: List[str] = []

            # A) é€‰æ‹© tender èµ„äº§ï¼ˆä¼˜å…ˆæœ€æ–°ï¼Œä¸¥æ ¼è¿‡æ»¤å¥—ç”¨æ ¼å¼äº§ç‰©ï¼‰
            assets = self.dao.list_assets(project_id)
            
            def _asset_text(a):
                """è·å–èµ„äº§æ–‡æœ¬ï¼ˆç”¨äºå…³é”®è¯è¿‡æ»¤ï¼‰"""
                return f"{a.get('filename','')} {a.get('storage_path','')}".lower()
            
            # æ’é™¤å…³é”®è¯ï¼šå¥—ç”¨æ ¼å¼ã€å¯¼å‡ºã€æŠ•æ ‡æ–‡ä»¶äº§ç‰©
            deny_kw = [
                "å¥—ç”¨æ ¼å¼", "render_", "template_renders", "export_", "å¯¼å‡º",
                "æŠ•æ ‡æ–‡ä»¶", "bid_", "skeleton", "éª¨æ¶"
            ]
            
            tenders = []
            for a in (assets or []):
                if (a or {}).get("kind") != "tender":
                    continue
                sp = (a.get("storage_path") or "").lower()
                if not (sp.endswith(".docx") or sp.endswith(".pdf")):
                    continue
                if any(k in _asset_text(a) for k in deny_kw):
                    continue
                tenders.append(a)
            
            if not tenders:
                diag["needs_reupload"] = True
                warnings.append("æœªæ‰¾åˆ°å¯ç”¨çš„æ‹›æ ‡ä¹¦(tender)èµ„äº§ï¼šè¯·ä¸Šä¼ æ‹›æ ‡æ–‡ä»¶ï¼ˆdocx/pdfï¼‰ï¼Œç³»ç»Ÿå·²æ’é™¤å¥—ç”¨æ ¼å¼/å¯¼å‡º/æŠ•æ ‡æ–‡ä»¶äº§ç‰©")
                diag["warnings"] = warnings
                diag["nodes"] = self.get_directory_with_body_meta(project_id)
                return diag

            def _sort_key(a: Dict[str, Any]) -> str:
                v = a.get("created_at")
                return str(v or "")

            try:
                tenders_sorted = sorted(tenders, key=_sort_key, reverse=True)
                tender_asset = tenders_sorted[0]
            except Exception:
                tender_asset = tenders[-1]

            diag["tender_asset_id"] = tender_asset.get("id")
            diag["tender_filename"] = tender_asset.get("filename")

            # B) åªç”¨ storage_pathï¼Œä¸”å¿…é¡»å­˜åœ¨ä¸”æ˜¯ .docx/.pdfï¼›è‹¥ç¼ºå¤±åˆ™å°è¯•æ¢å¤/å…œåº•
            path = str((tender_asset.get("storage_path") or "")).strip() or None
            diag["tender_storage_path"] = path
            diag["storage_path_exists"] = bool(path and os.path.exists(path))
            diag["tender_storage_path_ext"] = os.path.splitext(path or "")[1].lower()

            if not path or not diag["storage_path_exists"]:
                # C-2) æ—§é¡¹ç›®ï¼šå°è¯•ä»ç£ç›˜/å­˜å‚¨æ¢å¤ï¼ˆè‹¥ç³»ç»Ÿæ²¡æœ‰åŸå§‹ bytesï¼Œåˆ™é€šå¸¸åªèƒ½èµ° fallbackï¼‰
                restored_path = self._try_restore_tender_docx_from_disk(project_id, tender_asset)
                # ä¿®å¤ï¼šç¡®ä¿ restored_ext ä¸ä¸ºç©ºï¼Œé¿å… NoneType é”™è¯¯
                restored_ext = os.path.splitext(restored_path or "")[1].lower() if restored_path else ""
                
                if restored_path and os.path.exists(restored_path) and restored_ext in (".docx", ".pdf"):
                    path = restored_path
                    diag["tender_storage_path"] = restored_path
                    diag["storage_path_exists"] = True
                else:
                    # C-3) æ— æ³•æ¢å¤ï¼šèµ°å†…ç½®èŒƒæœ¬åº“ fallbackï¼ˆåŠŸèƒ½å¯ç”¨ï¼‰ï¼Œå¹¶æç¤ºéœ€è¦é‡ä¼ ä¿çœŸæŠ½å–
                    diag["needs_reupload"] = True
                    warnings.append(
                        "æ— æ³•ä»æ‹›æ ‡ä¹¦æŠ½å–èŒƒæœ¬ï¼ˆç¼ºå°‘åŸä»¶/æœªè½ç›˜ï¼‰ï¼Œå·²ä½¿ç”¨å†…ç½®èŒƒæœ¬åº“ï¼›å¦‚éœ€ä¿çœŸæŠ½å–è¯·é‡æ–°ä¸Šä¼ æ‹›æ ‡ä¹¦ï¼ˆdocx/pdfï¼‰"
                    )
                    attached = self._auto_fill_samples_builtin(project_id, warnings)
                    diag["attached_sections_builtin"] = int(attached or 0)
                    diag["attached_sections_template_sample"] = 0
                    diag["tender_fragments_upserted"] = 0
                    try:
                        diag["tender_fragments_total"] = len(self.dao.list_fragments("PROJECT", project_id))
                    except Exception:
                        diag["tender_fragments_total"] = 0
                    # å…¼å®¹å­—æ®µ
                    diag["extracted_fragments"] = int(diag["tender_fragments_upserted"])
                    diag["attached_sections"] = int(diag["attached_sections_template_sample"]) + int(diag["attached_sections_builtin"])
                    try:
                        diag["nodes"] = self.get_directory_with_body_meta(project_id)
                    except Exception as e:
                        warnings.append(f"get_directory_with_body_meta exception: {type(e).__name__}: {str(e)}")
                        diag["nodes"] = []
                    diag["ok"] = bool(diag["attached_sections"] and int(diag["attached_sections"]) > 0)
                    diag["warnings"] = warnings
                    return diag

            # C) å…è®¸ PDF / DOCX éƒ½èµ° extractorï¼›åªæœ‰ extractor çœŸçš„æŠ½ä¸åˆ°å†…å®¹æ‰ fallback
            use_path = path
            
            # âœ… PDFæ–‡ä»¶è·³è¿‡fragmentsæŠ½å–ï¼Œç›´æ¥ä½¿ç”¨è¯­ä¹‰æœç´¢
            pdf_ext = os.path.splitext(use_path)[1].lower()
            if pdf_ext == ".pdf":
                logger.info(f"[auto_fill_samples] PDF detected, skipping fragment extraction, using semantic search")
                diag["tender_storage_path_ext"] = ".pdf"
                diag["body_items_count"] = 0
                diag["tender_fragments_upserted"] = 0
                diag["fragments_detected_by_rules"] = 0
                diag["llm_used"] = False
                # ç›´æ¥è·³åˆ°è¯­ä¹‰æœç´¢éƒ¨åˆ†
            else:
                # DOCXæ–‡ä»¶èµ°ä¼ ç»Ÿçš„fragmentsæŠ½å–æµç¨‹
                try:
                    from app.services.fragment.fragment_extractor import TenderSampleFragmentExtractor
                    
                    extractor = TenderSampleFragmentExtractor(self.dao)
                    summary = extractor.extract_and_upsert_summary(
                        project_id=project_id,
                        tender_docx_path=use_path,
                        file_key=path,
                    )
                    
                    diag["body_items_count"] = int(summary.get("body_elements") or 0)
                    diag["tender_fragments_upserted"] = int(summary.get("upserted_fragments") or 0)
                    diag["fragments_detected_by_rules"] = int(summary.get("fragments_detected") or 0)
                    diag["llm_used"] = bool((summary.get("llm_spans_raw") or 0) > 0)
                    diag["tender_storage_path_ext"] = str(summary.get("input_ext") or os.path.splitext(use_path)[1].lower())

                    # å¦‚æœDOCXæŠ½å–ç»“æœä¸º 0ï¼Œfallback builtin
                    if diag["tender_fragments_upserted"] <= 0:
                        ext_name = diag.get("tender_storage_path_ext", "").upper().replace(".", "")
                        warnings.append(f"{ext_name} æœªèƒ½æŠ½å–åˆ°èŒƒæœ¬ç‰‡æ®µï¼ˆå¯èƒ½ä¸ºæ‰«æä»¶ã€æ— èŒƒæœ¬åŒºåŸŸæˆ–æ ¼å¼ä¸è§„èŒƒï¼‰ï¼Œå·²ä½¿ç”¨å†…ç½®èŒƒæœ¬åº“ï¼›å»ºè®®ä¸Šä¼ æ ‡å‡†æ ¼å¼æ‹›æ ‡æ–‡ä»¶")
                        attached = self._auto_fill_samples_builtin(project_id, warnings)
                        diag["attached_sections_builtin"] = attached
                        diag["warnings"] = warnings
                        diag["nodes"] = self.get_directory_with_body_meta(project_id)
                        diag["ok"] = True
                        return diag
                except Exception as e:
                    warnings.append(f"extractor exception: {type(e).__name__}: {str(e)}")
                    logger.exception(f"[samples] extractor exception: project_id={project_id}")
                    diag["tender_fragments_upserted"] = 0
                    diag["fragments_detected_by_rules"] = 0

            # å½“å‰åº“å†… fragments æ€»æ•°ï¼ˆç”¨äºåŒºåˆ†â€œæœ¬æ¬¡æŠ½å–â€ vs â€œå†å²å·²å­˜åœ¨â€ï¼‰
            try:
                diag["tender_fragments_total"] = len(self.dao.list_fragments("PROJECT", project_id))
            except Exception:
                diag["tender_fragments_total"] = 0

            try:
                from app.services.fragment.outline_attacher import OutlineSampleAttacher

                nodes = self.dao.list_directory(project_id)
                attacher = OutlineSampleAttacher(self.dao, llm_client=self.llm)
                
                # âœ… æ–°ç­–ç•¥ï¼šä»ç›®å½•æ ‡é¢˜å‡ºå‘ï¼Œè¯­ä¹‰æœç´¢PDFå†…å®¹ï¼ˆé€‚ç”¨äºPDFï¼‰
                # æ—§ç­–ç•¥ï¼šä»fragmentsè¡¨æŸ¥æ‰¾ï¼ˆé€‚ç”¨äºDOCXï¼‰
                if diag["tender_storage_path_ext"] == ".pdf":
                    logger.info(f"[auto_fill_samples] Using PDF semantic search (keyword-only, LLM disabled) for project {project_id}")
                    # âœ… æš‚æ—¶ç¦ç”¨LLMéªŒè¯ï¼Œåªä½¿ç”¨å…³é”®è¯åŒ¹é…
                    attached_count = int(attacher.attach_from_pdf_semantic(
                        project_id, nodes, min_confidence=0.4, use_llm=False  # âœ… ç¦ç”¨LLM
                    ) or 0)
                else:
                    logger.info(f"[auto_fill_samples] Using traditional fragment matching for project {project_id}")
                    attached_count = int(attacher.attach(project_id, nodes, use_llm=True) or 0)
                
                diag["attached_sections_template_sample"] = attached_count
                diag["attached_write_mode"] = "content_json"  # è¯Šæ–­ä¿¡æ¯ï¼šå†™å…¥æ¨¡å¼
            except Exception as e:
                warnings.append(f"attacher exception: {type(e).__name__}: {str(e)}")
                diag["attached_sections_template_sample"] = 0

            # æœ¬åˆ†æ”¯æœªèµ° builtin
            diag["attached_sections_builtin"] = 0

            # å…¼å®¹å­—æ®µï¼ˆæ—§å‰ç«¯/è„šæœ¬ï¼‰
            diag["extracted_fragments"] = int(diag["tender_fragments_upserted"])
            diag["attached_sections"] = int(diag["attached_sections_template_sample"]) + int(diag["attached_sections_builtin"])

            # nodesï¼ˆå¸¦ bodyMetaï¼‰
            try:
                diag["nodes"] = self.get_directory_with_body_meta(project_id)
            except Exception as e:
                warnings.append(f"get_directory_with_body_meta exception: {type(e).__name__}: {str(e)}")
                diag["nodes"] = []

            diag["ok"] = bool(diag["attached_sections"] and int(diag["attached_sections"]) > 0)
            diag["warnings"] = warnings
            return diag
        except Exception as e:
            # æœ€åä¸€é“é˜²çº¿ï¼šç»ä¸æŠ›å‡º
            logger.exception("auto_fill_samples unexpected exception project_id=%s", project_id)
            w = diag.get("warnings")
            if not isinstance(w, list):
                w = []
            w.append(f"auto_fill_samples unexpected: {type(e).__name__}: {str(e)}")
            diag["warnings"] = w
            try:
                diag["nodes"] = self.get_directory_with_body_meta(project_id)
            except Exception:
                diag["nodes"] = []
            return diag

    def _auto_fill_samples_builtin(self, project_id: str, warnings: List[str]) -> int:
        """
        D. å…œåº•ï¼šå†…ç½®èŒƒæœ¬åº“ï¼ˆè‡³å°‘æŠ•æ ‡å‡½/æˆæƒä¹¦/æŠ¥ä»·è¡¨ä¸‰ç±»ï¼‰
        å†™å…¥ section_body.source=BUILTIN_SAMPLEï¼ˆå‰ç«¯å¯ç›´æ¥å±•ç¤ºï¼‰ï¼Œå¹¶æç¤º warningsã€‚
        """
        from app.services.fragment.fragment_matcher import FragmentTitleMatcher
        from app.services.fragment.builtin_samples import BUILTIN_SAMPLE_HTML_BY_TYPE

        warnings.append("ä½¿ç”¨å†…ç½®èŒƒæœ¬åº“å¡«å……ï¼ˆéä»æ‹›æ ‡ä¹¦docxæŠ½å–ï¼‰")

        matcher = FragmentTitleMatcher()
        nodes = self.dao.list_directory(project_id)

        attached = 0
        for node in nodes:
            node_id = node.get("id")
            if not node_id:
                continue
            title_norm = matcher.normalize(node.get("title") or "")
            ftype = matcher.match_type(title_norm)
            if not ftype:
                continue
            html = BUILTIN_SAMPLE_HTML_BY_TYPE.get(ftype)
            if not html:
                continue

            existing = self.dao.get_section_body(project_id, node_id)
            # ä¸è¦†ç›–ç”¨æˆ·å·²æœ‰å†…å®¹
            if existing and existing.get("source") == "USER" and (existing.get("content_html") or "").strip():
                continue
            # ä¸è¦†ç›–å·²æœ‰ AI å†…å®¹
            if existing and existing.get("source") == "AI":
                continue

            self.dao.upsert_section_body(
                project_id=project_id,
                node_id=node_id,
                source="BUILTIN_SAMPLE",
                fragment_id=None,
                content_html=html,
            )
            attached += 1

        return attached

    def _try_restore_tender_docx_from_disk(self, project_id: str, tender_asset: Dict[str, Any]) -> Optional[str]:
        """
        æ—§é¡¹ç›®æ¢å¤å°è¯•ï¼ˆç£ç›˜ä¾§ï¼‰ï¼šå¦‚æœé¡¹ç›®ç›®å½•ä¸‹å·²å­˜åœ¨ docxï¼Œä½† DB storage_path ä¸¢å¤±/ä¸ºç©ºï¼Œ
        åˆ™é€‰æ‹©ä¸€ä¸ª docx å›å†™ storage_path å¹¶è¿”å›ã€‚

        æ³¨æ„ï¼šå½“å‰ç³»ç»Ÿ KB ä»…ä¿å­˜è§£æåçš„æ–‡æœ¬ chunksï¼Œé€šå¸¸æ— æ³•ä» KB è¿˜åŸåŸå§‹ docx bytesï¼›
        å› æ­¤è¿™é‡Œä¼˜å…ˆåšâ€œç£ç›˜å­˜åœ¨ä½† DB æ²¡å†™â€çš„è‡ªæ„ˆã€‚
        """
        import glob
        import logging

        logger = logging.getLogger(__name__)
        base_dir = os.path.join("data", "tender_assets", project_id)
        try:
            cands = sorted(glob.glob(os.path.join(base_dir, "*.docx")), reverse=True)
        except Exception:
            cands = []

        if not cands:
            return None

        # ä¼˜å…ˆåŒ¹é…æ–‡ä»¶ååŒ…å« tender_asset.filename
        filename = str((tender_asset or {}).get("filename") or "").strip()
        pick = None
        if filename:
            for p in cands:
                if filename in os.path.basename(p):
                    pick = p
                    break
        if not pick:
            pick = cands[0]

        asset_id = (tender_asset or {}).get("id")
        if asset_id:
            try:
                self.dao.update_asset_storage_path(str(asset_id), pick)
            except Exception as e:
                logger.warning("update_asset_storage_path failed asset_id=%s: %s", asset_id, e)

        logger.info("[samples] restored tender docx from disk. project_id=%s pick=%s", project_id, pick)
        return pick

    # extract_rule_set æ–¹æ³•å·²åˆ é™¤ï¼Œè§„åˆ™æ–‡ä»¶ç°åœ¨ç›´æ¥ä½œä¸ºå®¡æ ¸ä¸Šä¸‹æ–‡å åŠ 

    def run_review(
        self,
        project_id: str,
        model_id: Optional[str],
        custom_rule_asset_ids: List[str],
        bidder_name: Optional[str],
        bid_asset_ids: List[str],
        custom_rule_pack_ids: Optional[List[str]] = None,
        use_llm_semantic: bool = True,  # âœ… é»˜è®¤å¯ç”¨QAéªŒè¯
        run_id: Optional[str] = None,
        owner_id: Optional[str] = None,
    ):
        """
        è¿è¡Œå®¡æ ¸ï¼ˆæ‹›æ ‡è§„åˆ™ + è‡ªå®šä¹‰è§„åˆ™æ–‡ä»¶å åŠ ï¼‰
        
        Args:
            custom_rule_asset_ids: è‡ªå®šä¹‰è§„åˆ™æ–‡ä»¶èµ„äº§IDåˆ—è¡¨ï¼ˆç›´æ¥å åŠ åŸæ–‡ï¼‰
            custom_rule_pack_ids: è‡ªå®šä¹‰è§„åˆ™åŒ…IDåˆ—è¡¨ï¼ˆåº”ç”¨è§„åˆ™åŒ…ä¸­çš„è§„åˆ™ï¼‰
            bidder_name: æŠ•æ ‡äººåç§°ï¼ˆé€‰æ‹©æŠ•æ ‡äººï¼‰
            bid_asset_ids: æŠ•æ ‡èµ„äº§IDåˆ—è¡¨ï¼ˆç²¾ç¡®æŒ‡å®šæ–‡ä»¶ï¼‰
            use_llm_semantic: æ˜¯å¦ä½¿ç”¨LLMè¯­ä¹‰å®¡æ ¸ï¼ˆQAéªŒè¯ï¼Œé»˜è®¤Trueï¼‰
            owner_id: ä»»åŠ¡æ‰€æœ‰è€…IDï¼ˆå¯é€‰ï¼‰
        """
        # æ—è·¯åŒå†™ï¼šåˆ›å»º platform jobï¼ˆå¦‚æœå¯ç”¨ï¼‰
        job_id = None
        if self.feature_flags.PLATFORM_JOBS_ENABLED and self.jobs_service and run_id:
            try:
                job_id = self.jobs_service.create_job(
                    namespace="tender",
                    biz_type="review_run",
                    biz_id=project_id,
                    owner_id=owner_id,
                    initial_status="running",
                    initial_message="æ­£åœ¨è¿è¡Œå®¡æ ¸..."
                )
            except Exception as e:
                print(f"[WARN] Failed to create platform job: {e}")
        
        try:
            # ä½¿ç”¨ ReviewV3Service (æ”¯æŒè§„åˆ™å¼•æ“å’Œè‡ªå®šä¹‰è§„åˆ™åŒ…)
            
            # å¦‚æœæ²¡æœ‰æŒ‡å®šmodel_idï¼Œä½¿ç”¨é»˜è®¤æ¨¡å‹
            if not model_id:
                from app.services.llm_model_store import get_llm_store
                store = get_llm_store()
                default_model = store.get_default_model()
                if default_model:
                    model_id = default_model.id
                    logger.info(f"No model_id provided, using default: {model_id}")
                else:
                    logger.warning("No model_id provided and no default model configured")
            
            logger.info(f"ReviewV3 mode: project={project_id}, model_id={model_id}, bidder={bidder_name}, custom_rule_pack_ids={custom_rule_pack_ids}")
            
            arr = []  # åˆå§‹åŒ–å®¡æ ¸é¡¹åˆ—è¡¨
            try:
                import asyncio
                from app.works.tender.review_v3_service import ReviewV3Service
                from app.services.db.postgres import _get_pool
                
                pool = _get_pool()
                review_v3 = ReviewV3Service(pool, self.llm)
                logger.info("Created ReviewV3Service")
                
                # è¿è¡Œ v3 å®¡æ ¸
                logger.info(f"Calling run_review_v3 with use_llm_semantic={use_llm_semantic}...")
                v3_results = asyncio.run(review_v3.run_review_v3(
                    project_id=project_id,
                    bidder_name=bidder_name,
                    model_id=model_id,
                    custom_rule_pack_ids=custom_rule_pack_ids,
                    use_llm_semantic=use_llm_semantic,
                    run_id=run_id
                ))
                logger.info(f"run_review_v3 completed")
                
                # v3 æˆåŠŸï¼šä½¿ç”¨ v3 ç»“æœ
                arr = v3_results.get("items", [])  # æå– items åˆ—è¡¨
                logger.info(f"Extracted {len(arr)} review items")
                
                # âš ï¸ æ³¨æ„ï¼šReviewPipelineV3å·²ç»ç›´æ¥ä¿å­˜åˆ°tender_review_itemsè¡¨
                # ä¸éœ€è¦å†è°ƒç”¨replace_review_itemsï¼Œå¦åˆ™ä¼šåˆ é™¤å·²ä¿å­˜çš„V3æ•°æ®ï¼
                # self.dao.replace_review_items(project_id, arr)  # å·²åºŸå¼ƒ
                
                # æ›´æ–°è¿è¡ŒçŠ¶æ€
                if run_id:
                    self.dao.update_run(
                        run_id, "success", progress=1.0,
                        message="ok",
                        result_json={
                            "count": len(arr),
                            "review_v3_status": "ok",
                            "review_mode": v3_results.get("review_mode", "UNKNOWN"),
                            "pass_count": v3_results.get("pass_count", 0),
                            "fail_count": v3_results.get("fail_count", 0),
                            "warn_count": v3_results.get("warn_count", 0),
                            "pending_count": v3_results.get("pending_count", 0)
                        }
                    )
                
                logger.info(f"Review completed successfully for project={project_id}, count={len(arr)}")
                
            except Exception as e:
                # å®¡æ ¸å¤±è´¥ï¼šè®°å½•å¹¶æŠ›é”™
                error_msg = f"Review failed: {str(e)}"
                logger.error(
                    f"Review failed for project={project_id}: {e}",
                    exc_info=True
                )
                
                # æ›´æ–°è¿è¡ŒçŠ¶æ€ä¸ºå¤±è´¥
                if run_id:
                    self.dao.update_run(
                        run_id, "failed", progress=0.0,
                        message=error_msg,
                        result_json={
                            "review_v3_status": "failed",
                            "review_v3_error": str(e)
                        }
                    )
                
                # æŠ›å‡ºé”™è¯¯
                raise ValueError(error_msg) from e
            
            # æ—è·¯åŒå†™ï¼šReviewCaseï¼ˆå¦‚æœå¯ç”¨ï¼‰
            case_id = None
            review_run_id = None
            if self.feature_flags.REVIEWCASE_DUALWRITE:
                try:
                    from app.services.platform.reviewcase_service import ReviewCaseService
                    from app.services.db.postgres import _get_pool
                    pool = _get_pool()
                    reviewcase_service = ReviewCaseService(pool)
                    
                    # 1. æ”¶é›†æ–‡æ¡£ç‰ˆæœ¬IDï¼ˆä» assets çš„ meta_json ä¸­æå–ï¼‰
                    tender_doc_version_ids = []
                    bid_doc_version_ids = []
                    
                    # è·å–æ‹›æ ‡æ–‡ä»¶çš„ doc_version_id
                    tender_assets = self.dao.list_assets(project_id)
                    for asset in tender_assets:
                        if asset.get("kind") == "tender":
                            meta_json = asset.get("meta_json") or {}
                            doc_version_id = meta_json.get("doc_version_id")
                            if doc_version_id:
                                tender_doc_version_ids.append(doc_version_id)
                    
                    # è·å–æŠ•æ ‡æ–‡ä»¶çš„ doc_version_id
                    for asset in tender_assets:
                        if asset.get("kind") == "bid" and asset.get("bidder_name") == bidder_name:
                            meta_json = asset.get("meta_json") or {}
                            doc_version_id = meta_json.get("doc_version_id")
                            if doc_version_id:
                                bid_doc_version_ids.append(doc_version_id)
                    
                    # è·å– custom_rule çš„ rule_set_version_idï¼ˆå¦‚æœå­˜åœ¨ï¼‰
                    rule_set_version_id = None
                    for asset in tender_assets:
                        if asset.get("kind") == "custom_rule":
                            meta_json = asset.get("meta_json") or {}
                            rule_set_version_id = meta_json.get("rule_set_version_id")
                            if rule_set_version_id:
                                break  # åªå–ç¬¬ä¸€ä¸ª
                    
                    # 2. åˆ›å»º review_case
                    case_id = reviewcase_service.create_case(
                        namespace="tender",
                        project_id=project_id,
                        tender_doc_version_ids=tender_doc_version_ids,
                        bid_doc_version_ids=bid_doc_version_ids
                    )
                    
                    # 3. åˆ›å»º review_runï¼ˆè®°å½• rule_set_version_idï¼‰
                    review_run_id = reviewcase_service.create_run(
                        case_id=case_id,
                        model_id=model_id,
                        rule_set_version_id=rule_set_version_id,
                        status="running"
                    )
                    
                    # 4. æ‰¹é‡åˆ›å»º review_findingsï¼ˆä» arr è½¬æ¢ï¼‰
                    findings = []
                    for item in arr:
                        source = item.get("source", "compare")
                        finding = {
                            "source": source,
                            "dimension": item.get("dimension", "å…¶ä»–"),
                            "requirement_text": item.get("requirement_text") or item.get("title", ""),
                            "response_text": item.get("response_text", ""),
                            "result": item.get("result", "risk"),
                            "rigid": item.get("rigid", False),
                            "remark": item.get("remark", ""),
                            "evidence_jsonb": {
                                "tender_chunk_ids": item.get("tender_evidence_chunk_ids", []),
                                "bid_chunk_ids": item.get("bid_evidence_chunk_ids", []),
                                "evidence_chunk_ids": item.get("evidence_chunk_ids", []),
                                "rule_id": item.get("rule_id") if source == "rule" else None
                            }
                        }
                        findings.append(finding)
                    
                    reviewcase_service.batch_create_findings(review_run_id, findings)
                    
                    # 5. æ›´æ–° review_run çŠ¶æ€
                    reviewcase_service.update_run_status(
                        review_run_id,
                        status="succeeded",
                        result_json={
                            "total_findings": len(arr),
                            "compare_findings": len(arr)
                        }
                    )
                    
                    print(f"[INFO] ReviewCase dual-write succeeded: case_id={case_id}, review_run_id={review_run_id}")
                    
                except Exception as e:
                    # é™çº§ï¼šReviewCase åŒå†™å¤±è´¥ä¸å½±å“ä¸»æµç¨‹
                    print(f"[WARN] Failed to write to ReviewCase: {e}")
            
            # æ—è·¯åŒå†™ï¼šæ›´æ–° job æˆåŠŸï¼ˆå¦‚æœå¯ç”¨ï¼‰
            if job_id and self.jobs_service:
                try:
                    self.jobs_service.finish_job_success(
                        job_id=job_id,
                        result={"review_items_count": len(arr)},
                        message="æˆåŠŸ"
                    )
                except Exception as e:
                    print(f"[WARN] Failed to update platform job: {e}")
        
        except Exception as e:
            # æ›´æ–° job å¤±è´¥çŠ¶æ€ï¼ˆå¦‚æœå¯ç”¨ï¼‰
            if job_id and self.jobs_service:
                try:
                    self.jobs_service.finish_job_fail(job_id=job_id, error=str(e))
                except Exception as je:
                    print(f"[WARN] Failed to update platform job on error: {je}")
            # é‡æ–°æŠ›å‡ºåŸå§‹å¼‚å¸¸
            raise

    def generate_docx(
        self,
        project_id: str,
        template_asset_id: Optional[str],
    ) -> bytes:
        """
        ç”Ÿæˆ Word æ–‡æ¡£
        
        Args:
            template_asset_id: æ¨¡æ¿èµ„äº§IDï¼ˆå¯é€‰ï¼‰
        
        Returns:
            Word æ–‡æ¡£å­—èŠ‚
        """
        # åŠ è½½ç›®å½•èŠ‚ç‚¹
        nodes = self.dao.list_directory(project_id)

        # åŠ è½½æ¨¡æ¿ï¼ˆå¦‚æœæŒ‡å®šï¼‰
        tpl_doc = None
        template_spec: Optional[TemplateSpec] = None
        
        if template_asset_id:
            assets = self.dao.get_assets_by_ids(project_id, [template_asset_id])
            if assets:
                asset = assets[0]
                path = asset.get("storage_path")
                
                # å°è¯•åŠ è½½ TemplateSpec
                meta = asset.get("meta_json") or {}
                if isinstance(meta, str):
                    try:
                        meta = json.loads(meta)
                    except Exception:
                        meta = {}
                
                # å¦‚æœæ˜¯æ ¼å¼æ¨¡æ¿ï¼Œå°è¯•åŠ è½½ spec
                template_db_id = meta.get("format_template_id")
                if template_db_id:
                    template_spec = self.get_format_template_spec(template_db_id)
                
                # åŠ è½½æ¨¡æ¿æ–‡æ¡£
                if path and os.path.exists(path):
                    try:
                        tpl_doc = Document(path)
                    except Exception:
                        pass

        # å¦‚æœæœ‰æ¨¡æ¿ä¸”æœ‰ specï¼Œä½¿ç”¨ spec æŒ‡å¯¼æ ·å¼åº”ç”¨
        if tpl_doc and template_spec:
            return self._generate_docx_with_spec(nodes, tpl_doc, template_spec, project_id)
        elif tpl_doc:
            # ä¼ ç»Ÿæ¨¡å¼ï¼šç®€å•å¤åˆ¶æ¨¡æ¿ç»“æ„
            doc = tpl_doc
        else:
            # æ— æ¨¡æ¿ï¼šåˆ›å»ºç©ºæ–‡æ¡£
            doc = Document()

        # æ ¹æ®ç›®å½•ç”Ÿæˆéª¨æ¶ï¼ˆå¹¶æ’å…¥æ­£æ–‡å†…å®¹ï¼‰
        for n in nodes:
            title = n.get("title") or ""
            level = int(n.get("level") or 1)
            # docx heading level 1..9
            h = min(max(level, 1), 9)
            doc.add_heading(title, level=h)
            
            # æ’å…¥æ­£æ–‡å†…å®¹ï¼ˆèŒƒæœ¬æˆ–ç”¨æˆ·ç¼–è¾‘ï¼‰
            self._insert_section_body(doc, project_id, n)
            
            # æ·»åŠ å¤‡æ³¨ï¼ˆå¦‚æœæœ‰ä¸”æœªè¢«æ­£æ–‡è¦†ç›–ï¼‰
            notes = n.get("notes") or ""
            if notes:
                doc.add_paragraph(notes)

        # ä¿å­˜åˆ°å†…å­˜
        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_docx_v2(self, project_id: str, format_template_id: Optional[str] = None) -> bytes:
        """
        æ¨èå¯¼å‡ºæ¥å£ï¼š
        - æ”¯æŒ format_template_idï¼ˆæ¥è‡ª format_templates è¡¨ï¼‰
        - è‹¥æœªä¼ ï¼Œåˆ™å°è¯•ä»ç›®å½•èŠ‚ç‚¹ meta_json æ¨æ–­å·²å¥—ç”¨æ¨¡æ¿
        """
        # 1) åŠ è½½ç›®å½•
        nodes = self.dao.list_directory(project_id)

        # 2) æ¨æ–­å·²å¥—ç”¨æ¨¡æ¿ï¼ˆå¦‚æœæœªæŒ‡å®šï¼‰
        if not format_template_id:
            for n in nodes:
                meta = n.get("meta_json") or {}
                if isinstance(meta, str):
                    try:
                        meta = json.loads(meta)
                    except Exception:
                        meta = {}
                if isinstance(meta, dict) and meta.get("format_template_id"):
                    format_template_id = str(meta.get("format_template_id"))
                    break

        # 3) æ— æ¨¡æ¿ï¼šç›´æ¥å¯¼å‡ºï¼ˆæ— åº•æ¿ï¼‰
        if not format_template_id:
            return self.generate_docx(project_id, template_asset_id=None)

        # 4) æœ‰æ¨¡æ¿ï¼šåŠ è½½ spec + åº•æ¿ docx
        spec = self.get_format_template_spec(format_template_id)
        tpl_doc = self._load_format_template_doc(format_template_id)
        if not spec:
            return self.generate_docx(project_id, template_asset_id=None)

        if tpl_doc:
            return self._generate_docx_with_spec(nodes, tpl_doc, spec, project_id)

        # å…œåº•ï¼šæ²¡æ‰¾åˆ°åº•æ¿ docxï¼ˆæ—§æ•°æ®/æ–‡ä»¶ä¸¢å¤±ï¼‰ï¼Œç”¨ REBUILD é¿å… style ä¸å­˜åœ¨
        try:
            spec.base_policy.mode = BasePolicyMode.REBUILD
        except Exception:
            pass
        return self._generate_docx_with_spec(nodes, Document(), spec, project_id)

    def _insert_section_body(self, doc: Document, project_id: str, node: Dict):
        """
        æ’å…¥ç« èŠ‚æ­£æ–‡å†…å®¹
        - å¦‚æœæœ‰ç”¨æˆ·ç¼–è¾‘å†…å®¹ï¼Œæ’å…¥HTMLè½¬æ¢çš„å†…å®¹
        - å¦åˆ™å¦‚æœæœ‰æŒ‚è½½çš„èŒƒæœ¬ï¼Œæ‹·è´æºdocxçš„å†…å®¹
        - å¦åˆ™ä¸æ’å…¥å†…å®¹ï¼ˆä¿æŒç©ºï¼‰
        """
        from app.services.export.docx_copier import DocxBodyElementCopier
        from app.services.export.html_to_docx import HtmlToDocxInserter
        
        node_id = node.get("id")
        if not node_id:
            return
        
        # æŸ¥è¯¢ç« èŠ‚æ­£æ–‡
        body = self.dao.get_section_body(project_id, node_id)
        if not body:
            return
        
        source = body.get("source")
        
        # ç”¨æˆ·ç¼–è¾‘å†…å®¹ä¼˜å…ˆ
        if source == "USER" and body.get("content_html"):
            HtmlToDocxInserter.insert(doc, body["content_html"])
            return
        
        # èŒƒæœ¬æŒ‚è½½
        if source == "TEMPLATE_SAMPLE" and body.get("fragment_id"):
            fragment = self.dao.get_fragment_by_id(body["fragment_id"])
            if fragment:
                # è·å–æºæ–‡ä»¶è·¯å¾„
                source_file_key = fragment.get("source_file_key")
                start_idx = fragment.get("start_body_index")
                end_idx = fragment.get("end_body_index")
                
                if source_file_key and start_idx is not None and end_idx is not None:
                    try:
                        # æ‹·è´æºæ–‡æ¡£å†…å®¹
                        DocxBodyElementCopier.copy_range(
                            source_file_key,
                            start_idx,
                            end_idx,
                            doc
                        )
                    except Exception as e:
                        # æ‹·è´å¤±è´¥ï¼Œæ·»åŠ é”™è¯¯æç¤º
                        doc.add_paragraph(f"[èŒƒæœ¬å†…å®¹æ‹·è´å¤±è´¥: {str(e)}]")
    
    def _generate_docx_with_spec(
        self,
        nodes: List[Dict],
        tpl_doc: Document,
        spec: TemplateSpec,
        project_id: str,
    ) -> bytes:
        """
        ä½¿ç”¨ TemplateSpec ç”Ÿæˆæ–‡æ¡£
        
        ç­–ç•¥ï¼š
        1. æ ¹æ® base_policy æ„é€ åº•æ¿
        2. ä½¿ç”¨ style_hints åº”ç”¨æ ·å¼
        """
        import io
        
        # 1. æ ¹æ® base_policy å¤„ç†åº•æ¿
        if spec.base_policy.mode.value == "REBUILD":
            # å®Œå…¨é‡å»ºï¼šåˆ›å»ºç©ºæ–‡æ¡£
            doc = Document()
        else:
            # KEEP_ALL æˆ– KEEP_RANGEï¼šä¿ç•™æ¨¡æ¿
            doc = tpl_doc
            if spec.base_policy.mode.value == "KEEP_RANGE":
                self._apply_keep_range(doc, spec)

        # 1.5 ç¡®ä¿æ ·å¼å¯ç”¨ï¼šå¦‚æœæ¨¡æ¿ä¸­ä¸å­˜åœ¨ spec.style_hints æŒ‡å‘çš„æ ·å¼ï¼Œåˆ™æ ¹æ® spec.style_rules è‡ªåŠ¨åˆ›å»º AI_* æ ·å¼
        try:
            from app.services.export.style_applier import ensure_styles_from_spec
            ensure_styles_from_spec(doc, spec)
        except Exception:
            pass
        
        # 2. è¿½åŠ ç›®å½•èŠ‚ç‚¹ï¼Œä½¿ç”¨ style_hints
        for n in nodes:
            title = n.get("title") or ""
            level = int(n.get("level") or 1)
            
            # è·å–æ ·å¼æç¤º
            style_name = OutlineMerger.get_style_hint_for_level(level, spec)
            
            if style_name:
                # å°è¯•ä½¿ç”¨æŒ‡å®šæ ·å¼
                try:
                    para = doc.add_paragraph(title, style=style_name)
                except Exception:
                    # æ ·å¼ä¸å­˜åœ¨ï¼Œå›é€€åˆ° heading
                    h = min(max(level, 1), 9)
                    doc.add_heading(title, level=h)
            else:
                # ä½¿ç”¨é»˜è®¤ heading
                h = min(max(level, 1), 9)
                doc.add_heading(title, level=h)
            
            # æ’å…¥æ­£æ–‡å†…å®¹ï¼ˆèŒƒæœ¬æˆ–ç”¨æˆ·ç¼–è¾‘ï¼‰
            self._insert_section_body(doc, project_id, n)
            
            # æ·»åŠ å¤‡æ³¨ï¼ˆå¦‚æœæœ‰ä¸”æœªè¢«æ­£æ–‡è¦†ç›–ï¼‰
            notes = n.get("notes") or ""
            if notes:
                doc.add_paragraph(notes)
        
        # ä¿å­˜åˆ°å†…å­˜
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _apply_keep_range(self, doc: Document, spec: TemplateSpec) -> None:
        """
        KEEP_RANGEï¼šå°½é‡æŒ‰ range_anchor.start_text/end_text è£å‰ªæ¨¡æ¿æ­£æ–‡èŒƒå›´ï¼Œä¿ç•™é¡µçœ‰é¡µè„š/é¡µè¾¹è·ç­‰ section å±æ€§ã€‚
        è¯´æ˜ï¼šè¿™æ˜¯â€œæœ€å°å¯ç”¨å®ç°â€ï¼Œä¸ä¾èµ– block_idï¼ˆpython-docx ä¸ä¿ç•™ extractor çš„ block_idï¼‰ã€‚
        """
        try:
            ra = getattr(spec.base_policy, "range_anchor", None)
            if not ra:
                return
            start_text = (getattr(ra, "start_text", None) or "").strip()
            end_text = (getattr(ra, "end_text", None) or "").strip()
            if not start_text or not end_text:
                return

            body_elms = list(doc.element.body)
            idx_map: List[int] = []
            for i, el in enumerate(body_elms):
                # sectPr ä¸èƒ½åˆ 
                if getattr(el, "tag", "").endswith("}sectPr"):
                    continue
                idx_map.append(i)

            def _elm_text(el) -> str:
                try:
                    from docx.text.paragraph import Paragraph
                    p = Paragraph(el, doc)  # type: ignore[arg-type]
                    return p.text or ""
                except Exception:
                    return ""

            start_i = None
            end_i = None
            for i in idx_map:
                t = _elm_text(body_elms[i])
                if start_i is None and start_text in t:
                    start_i = i
                if start_i is not None and end_text in t:
                    end_i = i
                    break

            if start_i is None or end_i is None or start_i > end_i:
                return

            for i in reversed(idx_map):
                if i < start_i or i > end_i:
                    try:
                        doc.element.body.remove(body_elms[i])
                    except Exception:
                        continue
        except Exception:
            return

    # list_rule_sets æ–¹æ³•å·²åˆ é™¤

    def lookup_chunks(self, chunk_ids: List[str]) -> List[Dict[str, Any]]:
        """æŸ¥è¯¢ chunksï¼ˆè¯æ®å›æº¯ï¼‰"""
        return self.dao.lookup_chunks(chunk_ids)

    def _parse_template_meta(self, docx_path: str) -> Dict[str, Any]:
        """
        è§£ææ¨¡æ¿ï¼š
        - outline_nodes: [{numbering, level, title, notes?}]  (ä» Heading / å¤§çº²æ ·å¼æ¨å¯¼)
        - style_hint:    ç®€è¦ä¿¡æ¯ï¼ˆæ˜¯å¦æœ‰é¡µçœ‰é¡µè„š/å›¾ç‰‡ç­‰ï¼‰
        """
        try:
            from docx import Document as Doc
        except Exception:
            return {}

        doc = Doc(docx_path)

        # style_hintï¼ˆMVPï¼šåªåšå­˜åœ¨æ€§ï¼‰
        has_header = False
        has_footer = False
        has_images = False
        try:
            for s in doc.sections:
                if s.header and s.header.paragraphs:
                    has_header = True
                if s.footer and s.footer.paragraphs:
                    has_footer = True
        except Exception:
            pass
        try:
            # æœ‰æ— å›¾ç‰‡ï¼ˆç²—ç•¥ï¼‰
            has_images = bool(doc.part._package.image_parts)
        except Exception:
            pass

        # outlineï¼šæŒ‰ paragraph style.name åŒ…å« "Heading" æ¨ level
        outline = []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
            sname = (getattr(p.style, "name", "") or "")
            if "Heading" in sname:
                # Heading 1/2/3...
                level = 1
                m = re.search(r"Heading\s+(\d+)", sname)
                if m:
                    level = int(m.group(1))
                outline.append({"title": txt, "level": level})

        # ç»™ outline ç”Ÿæˆ numberingï¼ˆMVPï¼šæŒ‰ level è‡ªåŠ¨ç¼–å·ï¼‰
        numbering_stack = []
        out_nodes = []
        for node in outline:
            lvl = max(1, min(9, int(node["level"])))
            while len(numbering_stack) < lvl:
                numbering_stack.append(0)
            while len(numbering_stack) > lvl:
                numbering_stack.pop()
            numbering_stack[lvl-1] += 1
            for i in range(lvl, len(numbering_stack)):
                numbering_stack[i] = 0
            nums = [str(n) for n in numbering_stack if n > 0]
            out_nodes.append({
                "numbering": ".".join(nums),
                "level": lvl,
                "title": node["title"],
                "source": "template",
                "is_required": False,
                "evidence_chunk_ids": [],
                "meta_json": {},
            })

        return {
            "template_outline_nodes": out_nodes,
            "style_hint": {
                "has_header": has_header,
                "has_footer": has_footer,
                "has_images": has_images,
            },
        }

    async def analyze_template_with_llm(
        self,
        docx_bytes: bytes,
        template_sha256: str,
        force: bool = False
    ) -> TemplateSpec:
        """
        ä½¿ç”¨ LLM åˆ†ææ¨¡æ¿ç»“æ„
        
        Args:
            docx_bytes: Word æ–‡æ¡£å­—èŠ‚å†…å®¹
            template_sha256: æ¨¡æ¿æ–‡ä»¶ SHA256 å“ˆå¸Œ
            force: æ˜¯å¦å¼ºåˆ¶é‡æ–°åˆ†æï¼ˆå¿½ç•¥ç¼“å­˜ï¼‰
            
        Returns:
            TemplateSpec: æ¨¡æ¿è§„æ ¼
        """
        # æ£€æŸ¥ feature flag
        if not self.settings.TEMPLATE_LLM_ANALYSIS_ENABLED:
            return create_minimal_spec(confidence=0.0, error_msg="LLM analysis disabled")
        
        # æ£€æŸ¥ç¼“å­˜
        if not force and self.settings.TEMPLATE_LLM_ANALYSIS_CACHE_BY_SHA256:
            cache = get_analysis_cache()
            cached_spec_json = cache.get(
                template_sha256,
                self.settings.TEMPLATE_LLM_ANALYSIS_VERSION,
                self.settings.TEMPLATE_LLM_ANALYSIS_MODEL
            )
            if cached_spec_json:
                try:
                    return TemplateSpec.from_json(cached_spec_json)
                except Exception:
                    pass  # ç¼“å­˜å¤±æ•ˆï¼Œç»§ç»­åˆ†æ
        
        try:
            # 1. ç¡®å®šæ€§ç»“æ„åŒ–æå–
            extract_result = self.docx_extractor.extract(
                docx_bytes,
                max_blocks=self.settings.TEMPLATE_LLM_ANALYSIS_MAX_BLOCKS,
                max_chars_per_block=self.settings.TEMPLATE_LLM_ANALYSIS_MAX_CHARS_PER_BLOCK
            )
            
            # 2. LLM åˆ†æ
            spec = await self.llm_analyzer.analyze(extract_result)
            
            # 3. å†™å…¥ç¼“å­˜
            if self.settings.TEMPLATE_LLM_ANALYSIS_CACHE_BY_SHA256 and spec.diagnostics.confidence > 0:
                cache = get_analysis_cache()
                cache.set(
                    template_sha256,
                    self.settings.TEMPLATE_LLM_ANALYSIS_VERSION,
                    self.settings.TEMPLATE_LLM_ANALYSIS_MODEL,
                    spec.to_json()
                )
            
            return spec
            
        except Exception as e:
            # Fallback: è¿”å›æœ€å°è§„æ ¼
            error_msg = f"Template analysis failed: {type(e).__name__}: {str(e)}"
            return create_minimal_spec(confidence=0.0, error_msg=error_msg)

    async def import_format_template_with_analysis(
        self,
        name: str,
        docx_bytes: bytes,
        description: Optional[str] = None,
        owner_id: Optional[str] = None,
        is_public: bool = False,
        force_analyze: bool = False
    ) -> Dict[str, Any]:
        """
        å¯¼å…¥æ ¼å¼æ¨¡æ¿å¹¶è¿›è¡Œ LLM åˆ†æ
        
        Args:
            name: æ¨¡æ¿åç§°
            docx_bytes: Word æ–‡æ¡£å­—èŠ‚å†…å®¹
            description: æ¨¡æ¿æè¿°
            owner_id: æ‰€æœ‰è€… ID
            is_public: æ˜¯å¦å…¬å¼€
            force_analyze: æ˜¯å¦å¼ºåˆ¶é‡æ–°åˆ†æ
            
        Returns:
            æ¨¡æ¿è®°å½•ï¼ˆåŒ…å« spec åˆ†æç»“æœï¼‰
        """
        # 1. è®¡ç®— SHA256
        template_sha256 = _sha256(docx_bytes)
        
        # 2. æ£€æŸ¥æ˜¯å¦å·²æœ‰ç›¸åŒ hash çš„æ¨¡æ¿ï¼ˆç¼“å­˜å¤ç”¨ï¼‰
        cached_template = None
        if not force_analyze and self.settings.TEMPLATE_LLM_ANALYSIS_CACHE_BY_SHA256:
            cached_template = self.dao.get_format_template_by_sha256(template_sha256)
        
        # 3. åˆ›å»ºæ¨¡æ¿è®°å½•
        template = self.dao.create_format_template(
            name=name,
            description=description,
            style_config={},  # å¯åç»­æ‰©å±•
            owner_id=owner_id,
            is_public=is_public
        )
        
        template_id = template["id"]
        
        # 3.1 è½ç›˜ä¿å­˜æ¨¡æ¿ docxï¼ˆç”¨äºå¯¼å‡º KEEP_ALL/KEEP_RANGE ä¿ç•™é¡µçœ‰é¡µè„š/é¡µè¾¹è·/åº•æ¿ï¼‰
        self._persist_format_template_docx(template_id=template_id, docx_bytes=docx_bytes)
        
        # 4. LLM åˆ†æï¼ˆå¼‚æ­¥ï¼‰
        spec: Optional[TemplateSpec] = None
        
        if cached_template and cached_template.get("template_spec_json"):
            # å¤ç”¨ç¼“å­˜çš„ spec
            try:
                spec = TemplateSpec.from_json(cached_template["template_spec_json"])
                spec_json = cached_template["template_spec_json"]
                spec_version = cached_template["template_spec_version"]
                diagnostics_json = cached_template.get("template_spec_diagnostics_json")
            except Exception:
                # ç¼“å­˜å¤±è´¥ï¼Œé‡æ–°åˆ†æ
                spec = await self.analyze_template_with_llm(docx_bytes, template_sha256, force=False)
                spec_json = spec.to_json()
                spec_version = spec.version
                diagnostics_json = json.dumps(spec.diagnostics.to_dict() if hasattr(spec.diagnostics, 'to_dict') else {})
        else:
            # æ–°åˆ†æ
            spec = await self.analyze_template_with_llm(docx_bytes, template_sha256, force=force_analyze)
            spec_json = spec.to_json()
            spec_version = spec.version
            diagnostics_data = {
                "confidence": spec.diagnostics.confidence,
                "warnings": spec.diagnostics.warnings,
                "ignored_as_instructions_block_ids": spec.diagnostics.ignored_as_instructions_block_ids,
                "analysis_duration_ms": spec.diagnostics.analysis_duration_ms,
                "llm_model": spec.diagnostics.llm_model
            }
            diagnostics_json = json.dumps(diagnostics_data)
        
        # 5. æ›´æ–°æ¨¡æ¿è®°å½•
        self.dao.update_format_template_spec(
            template_id=template_id,
            template_sha256=template_sha256,
            template_spec_json=spec_json,
            template_spec_version=spec_version,
            template_spec_diagnostics_json=diagnostics_json
        )
        
        # 6. è¿”å›å®Œæ•´æ¨¡æ¿è®°å½•
        updated_template = self.dao.get_format_template(template_id)
        result = updated_template or template
        
        # ç¡®ä¿æ—¥æœŸå­—æ®µæ˜¯å­—ç¬¦ä¸²ï¼ˆé˜²æ­¢æœªè½¬æ¢çš„ datetime å¯¹è±¡ï¼‰
        if result:
            for k in ("created_at", "updated_at", "template_spec_analyzed_at"):
                if result.get(k) and hasattr(result[k], "isoformat"):
                    result[k] = result[k].isoformat()
        
        return result

    async def reanalyze_format_template(
        self,
        template_id: str,
        docx_bytes: bytes,
        force: bool = True,
    ) -> Dict[str, Any]:
        """
        å¼ºåˆ¶é‡æ–°åˆ†ææ ¼å¼æ¨¡æ¿
        
        Args:
            template_id: æ¨¡æ¿ ID
            docx_bytes: Word æ–‡æ¡£å­—èŠ‚å†…å®¹
            
        Returns:
            æ›´æ–°åçš„æ¨¡æ¿è®°å½•
        """
        # 1. è®¡ç®— SHA256
        template_sha256 = _sha256(docx_bytes)

        # 1.1 æ›´æ–°æ¨¡æ¿ docx æ–‡ä»¶ï¼ˆç”¨äºå¯¼å‡º KEEP_ALL/KEEP_RANGEï¼‰
        self._persist_format_template_docx(template_id=template_id, docx_bytes=docx_bytes)
        
        # 2. å¼ºåˆ¶é‡æ–°åˆ†æ
        spec = await self.analyze_template_with_llm(docx_bytes, template_sha256, force=force)
        
        # 3. æ›´æ–°æ¨¡æ¿è®°å½•
        spec_json = spec.to_json()
        diagnostics_data = {
            "confidence": spec.diagnostics.confidence,
            "warnings": spec.diagnostics.warnings,
            "ignored_as_instructions_block_ids": spec.diagnostics.ignored_as_instructions_block_ids,
            "analysis_duration_ms": spec.diagnostics.analysis_duration_ms,
            "llm_model": spec.diagnostics.llm_model
        }
        diagnostics_json = json.dumps(diagnostics_data)
        
        self.dao.update_format_template_spec(
            template_id=template_id,
            template_sha256=template_sha256,
            template_spec_json=spec_json,
            template_spec_version=spec.version,
            template_spec_diagnostics_json=diagnostics_json
        )
        
        # 4. è¿”å›æ›´æ–°åçš„æ¨¡æ¿
        return self.dao.get_format_template(template_id) or {}

    # ==================== æ ¼å¼æ¨¡æ¿ docx å­˜å‚¨/åŠ è½½ ====================

    def _format_templates_storage_dir(self) -> str:
        # ä½¿ç”¨ APP_DATA_DIRï¼ˆdocker-compose å·²æŒ‚è½½ ./data -> /app/dataï¼‰ï¼Œä¿è¯å®¹å™¨é‡å¯å¯å¤ç”¨
        base = os.path.join(self.settings.APP_DATA_DIR, "format_templates")
        _safe_mkdir(base)
        return base

    def _format_template_docx_path(self, template_id: str) -> str:
        return os.path.join(self._format_templates_storage_dir(), f"{template_id}.docx")

    def _format_template_work_dir(self, template_id: str) -> str:
        """æ¨¡æ¿è§£æ/é¢„è§ˆå·¥ä½œç›®å½•ï¼ˆAPP_DATA_DIR ä¸‹ï¼Œå¯æŒä¹…åŒ–ï¼‰"""
        base = os.path.join(self._format_templates_storage_dir(), template_id)
        _safe_mkdir(base)
        return base

    def _format_template_assets_dir(self, template_id: str) -> str:
        d = os.path.join(self._format_template_work_dir(template_id), "assets")
        _safe_mkdir(d)
        return d

    def _format_template_preview_dir(self, template_id: str) -> str:
        d = os.path.join(self._format_template_work_dir(template_id), "preview")
        _safe_mkdir(d)
        return d

    def _persist_format_template_docx(self, template_id: str, docx_bytes: bytes) -> Optional[str]:
        """æŠŠæ ¼å¼æ¨¡æ¿çš„ docx æ–‡ä»¶ä¿å­˜åˆ°ç£ç›˜ï¼Œå¹¶å†™å…¥ format_templates.template_storage_pathï¼ˆå¦‚åˆ—å­˜åœ¨ï¼‰"""
        try:
            path = self._format_template_docx_path(template_id)
            with open(path, "wb") as w:
                w.write(docx_bytes)
            try:
                self.dao.update_format_template_storage_path(template_id, path)
            except Exception:
                # å…¼å®¹ï¼šè‹¥æ•°æ®åº“å°šæœªæ‰§è¡Œæ–°å¢åˆ— migrationï¼Œä¸é˜»æ–­ä¸»è¦æµç¨‹
                pass
            # best-effortï¼šå†™å…¥ SOURCE_DOCX èµ„äº§è®°å½•ï¼ˆè‹¥å·²æ‰§è¡Œ 014 migrationï¼‰
            try:
                self.dao.delete_format_template_assets(template_id, asset_types=["SOURCE_DOCX"])
            except Exception:
                pass
            try:
                self.dao.create_format_template_asset(
                    template_id=template_id,
                    asset_type="SOURCE_DOCX",
                    variant="DEFAULT",
                    storage_path=path,
                    file_name=os.path.basename(path),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception:
                pass
            return path
        except Exception:
            return None

    # ==================== æ ¼å¼æ¨¡æ¿ç¡®å®šæ€§è§£æ / è§£æé¢„è§ˆ ====================

    def parse_format_template(self, template_id: str, force: bool = True) -> Dict[str, Any]:
        """
        ç¡®å®šæ€§è§£ææ¨¡æ¿ï¼ˆheader/footer å›¾ç‰‡ã€section å‚æ•°ã€heading æ ·å¼å­˜åœ¨æ€§ç­‰ï¼‰ï¼Œå¹¶è½åº“ï¼š
        - format_templates.parse_status/parse_result_json/parse_error
        - format_template_assetsï¼šHEADER_IMG/FOOTER_IMG
        """
        tpl = self.dao.get_format_template(template_id)
        if not tpl:
            raise ValueError("Template not found")

        path = str((tpl.get("template_storage_path") or "")).strip()
        if not path or not os.path.exists(path):
            raise ValueError("Template docx not found on disk (template_storage_path missing)")

        # é‡æ–°è§£æï¼šæ¸…ç†æ—§èµ„äº§ï¼ˆDB è®°å½•ï¼‰+ é¢„è§ˆè·¯å¾„
        if force:
            try:
                self.dao.delete_format_template_assets(
                    template_id,
                    asset_types=["HEADER_IMG", "FOOTER_IMG", "PREVIEW_DOCX", "PREVIEW_PDF"],
                )
            except Exception:
                pass
            try:
                self.dao.clear_format_template_preview_paths(template_id)
            except Exception:
                pass

        assets_dir = self._format_template_assets_dir(template_id)

        try:
            doc = Document(path)
            parser = DocxTemplateDeterministicParser()
            parse_result, images = parser.parse(doc)

            # å†™å›¾ç‰‡èµ„äº§ï¼ˆåŸæ · bytesï¼‰
            for img in images:
                try:
                    out_path = os.path.join(assets_dir, img.file_name)
                    with open(out_path, "wb") as w:
                        w.write(img.blob)
                    self.dao.create_format_template_asset(
                        template_id=template_id,
                        asset_type="HEADER_IMG" if img.where == "header" else "FOOTER_IMG",
                        variant=img.variant,
                        storage_path=out_path,
                        file_name=img.file_name,
                        content_type=img.content_type,
                        width_px=img.width_px,
                        height_px=img.height_px,
                    )
                except Exception:
                    # å•å¼ å›¾ç‰‡å¤±è´¥ä¸é˜»æ–­æ•´ä½“è§£æ
                    continue

            # å†™ parse_resultï¼ˆbest-effortï¼Œå…¼å®¹æœªæ‰§è¡Œ migration çš„ç¯å¢ƒï¼‰
            try:
                self.dao.update_format_template_parse_result(
                    template_id=template_id,
                    parse_status="SUCCESS",
                    parse_result_json=parse_result,
                    parse_error=None,
                    preview_docx_path=None,
                    preview_pdf_path=None,
                )
            except Exception:
                pass

            return {"template_id": template_id, "parse_status": "SUCCESS", "parse_result": parse_result}
        except Exception as e:
            err = f"{type(e).__name__}: {str(e)}"
            try:
                self.dao.update_format_template_parse_result(
                    template_id=template_id,
                    parse_status="FAILED",
                    parse_result_json={},
                    parse_error=err,
                    preview_docx_path=None,
                    preview_pdf_path=None,
                )
            except Exception:
                pass
            return {"template_id": template_id, "parse_status": "FAILED", "parse_error": err}

    def get_format_template_parse_summary(self, template_id: str) -> Dict[str, Any]:
        tpl = self.dao.get_format_template(template_id)
        if not tpl:
            raise ValueError("Template not found")

        assets: List[Dict[str, Any]] = []
        try:
            assets = self.dao.list_format_template_assets(template_id)
        except Exception:
            assets = []

        # èµ„äº§ç»Ÿè®¡ï¼švariant -> counts
        by_variant: Dict[str, Dict[str, int]] = {}
        for a in assets:
            v = str((a.get("variant") or "DEFAULT")).strip() or "DEFAULT"
            t = str((a.get("asset_type") or "")).strip()
            by_variant.setdefault(v, {"HEADER_IMG": 0, "FOOTER_IMG": 0, "PREVIEW_DOCX": 0, "PREVIEW_PDF": 0})
            if t in by_variant[v]:
                by_variant[v][t] += 1

        return {
            "template_id": template_id,
            "parse_status": tpl.get("parse_status") or "PENDING",
            "parse_error": tpl.get("parse_error"),
            "parse_updated_at": tpl.get("parse_updated_at"),
            "parse_result": tpl.get("parse_result_json") or {},
            "assets_by_variant": by_variant,
        }

    def generate_format_template_preview(self, template_id: str, fmt: str = "pdf") -> Dict[str, Any]:
        """
        ç”Ÿæˆå¹¶ç¼“å­˜ç¤ºèŒƒé¢„è§ˆæ–‡ä»¶ï¼š
        - æ€»æ˜¯å…ˆç¡®ä¿æœ‰ parse_resultï¼ˆå¿…è¦æ—¶è§¦å‘ parseï¼‰
        - ç”Ÿæˆ sample.docxï¼ˆPREVIEW_DOCXï¼‰
        - è‹¥ fmt=pdf ä¸” LibreOffice å¯ç”¨ï¼Œåˆ™ç”Ÿæˆ sample.pdfï¼ˆPREVIEW_PDFï¼‰
        """
        tpl = self.dao.get_format_template(template_id)
        if not tpl:
            raise ValueError("Template not found")

        # ensure parse
        status = str((tpl.get("parse_status") or "PENDING")).strip()
        if status != "SUCCESS" or not (tpl.get("parse_result_json") or {}):
            self.parse_format_template(template_id, force=True)
            tpl = self.dao.get_format_template(template_id) or tpl

        fmt = (fmt or "pdf").lower().strip()
        if fmt not in ("pdf", "docx"):
            fmt = "pdf"

        # cache hit: return existing paths if available
        if fmt == "pdf":
            p = str((tpl.get("preview_pdf_path") or "")).strip()
            if p and os.path.exists(p):
                return {"format": "pdf", "path": p}
        else:
            p = str((tpl.get("preview_docx_path") or "")).strip()
            if p and os.path.exists(p):
                return {"format": "docx", "path": p}

        template_docx_path = str((tpl.get("template_storage_path") or "")).strip()
        if not template_docx_path or not os.path.exists(template_docx_path):
            raise ValueError("Template docx not found on disk")

        # collect images from assets table
        assets: List[Dict[str, Any]] = []
        try:
            assets = self.dao.list_format_template_assets(template_id)
        except Exception:
            assets = []

        images_by_variant: Dict[str, Dict[str, List[str]]] = {}
        for a in assets:
            at = str((a.get("asset_type") or "")).strip()
            if at not in ("HEADER_IMG", "FOOTER_IMG"):
                continue
            v = str((a.get("variant") or "DEFAULT")).strip() or "DEFAULT"
            images_by_variant.setdefault(v, {"header": [], "footer": []})
            sp = str((a.get("storage_path") or "")).strip()
            if not sp or not os.path.exists(sp):
                continue
            if at == "HEADER_IMG":
                images_by_variant[v]["header"].append(sp)
            else:
                images_by_variant[v]["footer"].append(sp)

        parse_result = tpl.get("parse_result_json") or {}

        # style hints from specï¼ˆå¦‚æœå­˜åœ¨ï¼‰
        style_hints: Dict[str, Any] = {}
        try:
            spec = self.get_format_template_spec(template_id)
            if spec:
                d = spec.to_dict()
                style_hints = (d.get("style_hints") or {}) if isinstance(d, dict) else {}
        except Exception:
            style_hints = {}

        preview_dir = self._format_template_preview_dir(template_id)
        gen = TemplatePreviewGenerator(work_dir=preview_dir)
        docx_path = gen.generate_sample_docx(
            template_docx_path=template_docx_path,
            parse_result=parse_result,
            images_by_variant=images_by_variant,
            spec_style_hints=style_hints,
        )

        try:
            self.dao.create_format_template_asset(
                template_id=template_id,
                asset_type="PREVIEW_DOCX",
                variant="DEFAULT",
                storage_path=docx_path,
                file_name=os.path.basename(docx_path),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception:
            pass

        pdf_path: Optional[str] = None
        if fmt == "pdf":
            pdf_path = gen.convert_to_pdf(docx_path)
            if pdf_path:
                try:
                    self.dao.create_format_template_asset(
                        template_id=template_id,
                        asset_type="PREVIEW_PDF",
                        variant="DEFAULT",
                        storage_path=pdf_path,
                        file_name=os.path.basename(pdf_path),
                        content_type="application/pdf",
                    )
                except Exception:
                    pass

        # update template preview paths (best-effort)
        try:
            self.dao.update_format_template_parse_result(
                template_id=template_id,
                parse_status=tpl.get("parse_status") or "SUCCESS",
                parse_result_json=parse_result,
                parse_error=tpl.get("parse_error"),
                preview_docx_path=docx_path,
                preview_pdf_path=pdf_path if pdf_path else None,
            )
        except Exception:
            pass

        if fmt == "pdf" and pdf_path and os.path.exists(pdf_path):
            return {"format": "pdf", "path": pdf_path}
        return {"format": "docx", "path": docx_path}

    def _load_format_template_doc(self, template_id: str) -> Optional[Document]:
        """
        åŠ è½½æ ¼å¼æ¨¡æ¿ docxï¼ˆç”¨äºå¯¼å‡ºåº•æ¿ï¼‰
        - ä¼˜å…ˆè¯»å– format_templates.template_storage_path
        - è‹¥ä¸ºç©º/ä¸¢å¤±ï¼Œåˆ™æŒ‰ template_sha256 åœ¨ data/tender_assets é‡Œå›æº¯æŸ¥æ‰¾ï¼Œå¹¶å›å¡« storage_path
        """
        tpl = self.dao.get_format_template(template_id)
        if not tpl:
            return None

        path = tpl.get("template_storage_path")
        if path and os.path.exists(path):
            try:
                return Document(path)
            except Exception:
                pass

        sha256 = (tpl.get("template_sha256") or "").strip()
        if sha256:
            found = self._find_docx_by_sha256(sha256)
            if found and os.path.exists(found):
                try:
                    self.dao.update_format_template_storage_path(template_id, found)
                except Exception:
                    pass
                try:
                    return Document(found)
                except Exception:
                    return None

        return None

    def _find_docx_by_sha256(self, target_sha256: str) -> Optional[str]:
        """
        åœ¨å¸¸è§è½ç›˜ç›®å½•ä¸‹æŸ¥æ‰¾åŒ¹é… sha256 çš„ docx æ–‡ä»¶ï¼ˆå…¼å®¹æ—§æ•°æ®ï¼‰ã€‚
        - data/tender_assets/**ï¼ˆå†å² tender é¡¹ç›®èµ„äº§ï¼‰
        - storage/attachments/**ï¼ˆé™„ä»¶ä¸Šä¼ è½ç›˜ç›®å½•ï¼‰
        """
        try:
            roots = [
                os.path.join("data", "tender_assets"),
                os.path.join("storage", "attachments"),
            ]

            candidates: List[str] = []
            for root in roots:
                if not os.path.exists(root):
                    continue
                for dirpath, _dirnames, filenames in os.walk(root):
                    for fn in filenames:
                        if not fn.lower().endswith(".docx"):
                            continue
                        # tender_assets ä¼˜å…ˆ template_ï¼›attachments ä¸åšé™åˆ¶
                        if root.endswith(os.path.join("data", "tender_assets")) and "template_" not in fn:
                            continue
                        candidates.append(os.path.join(dirpath, fn))
                    if len(candidates) > 4000:
                        break
                if len(candidates) > 4000:
                    break

            for p in candidates:
                try:
                    with open(p, "rb") as r:
                        b = r.read()
                    if _sha256(b) == target_sha256:
                        return p
                except Exception:
                    continue
        except Exception:
            return None
        return None

    def get_format_template_spec(self, template_id: str) -> Optional[TemplateSpec]:
        """
        è·å–æ ¼å¼æ¨¡æ¿çš„ TemplateSpec
        
        Args:
            template_id: æ¨¡æ¿ ID
            
        Returns:
            TemplateSpec æˆ– None
        """
        template = self.dao.get_format_template(template_id)
        if not template or not template.get("template_spec_json"):
            return None
        
        try:
            spec = TemplateSpec.from_json(template["template_spec_json"])
            self._normalize_template_spec_style_hints(template_id, spec)
            return spec
        except Exception:
            return None

    def _normalize_template_spec_style_hints(self, template_id: str, spec: TemplateSpec) -> None:
        """
        çº é”™/è§„èŒƒåŒ– style_hintsï¼ˆé¿å… LLM è¾“å‡º '+æ ‡é¢˜1' è¿™ç±» doc.styles ä¸å­˜åœ¨çš„æ ·å¼åï¼‰ï¼š
        - å»æ‰å‰ç¼€ '+' å¹¶ trim
        - è‹¥èƒ½åŠ è½½åˆ°æ¨¡æ¿ docxï¼Œåˆ™æŠŠ heading1..5 æ˜ å°„åˆ° doc.styles ä¸­çœŸå®å­˜åœ¨çš„æ ·å¼å
        """
        try:
            hints = getattr(spec, "style_hints", None)
            if not hints:
                return

            def _clean(v: Any) -> Optional[str]:
                if not isinstance(v, str):
                    return None
                s = v.strip()
                while s.startswith("+"):
                    s = s[1:].strip()
                return s or None

            # å…ˆåšçº¯å­—ç¬¦ä¸²æ¸…æ´—
            for k in (
                "heading1_style",
                "heading2_style",
                "heading3_style",
                "heading4_style",
                "heading5_style",
                "body_style",
                "table_style",
                "list_style",
            ):
                setattr(hints, k, _clean(getattr(hints, k, None)))

            tpl_doc = self._load_format_template_doc(template_id)
            if not tpl_doc:
                return

            available = set()
            try:
                for s in tpl_doc.styles:
                    try:
                        available.add(str(s.name))
                    except Exception:
                        continue
            except Exception:
                return

            def _pick_first(cands: List[str]) -> Optional[str]:
                for c in cands:
                    if c in available:
                        return c
                return None

            # heading æ˜ å°„ï¼ˆä¼˜å…ˆè‹±æ–‡ Heading Nï¼Œå…¶æ¬¡ä¸­æ–‡ æ ‡é¢˜ N/æ ‡é¢˜Nï¼‰
            level_to_attr = {
                1: "heading1_style",
                2: "heading2_style",
                3: "heading3_style",
                4: "heading4_style",
                5: "heading5_style",
            }
            for lvl, attr in level_to_attr.items():
                cur = getattr(hints, attr, None)
                if cur and cur in available:
                    continue
                fallback = _pick_first(
                    [
                        f"Heading {lvl}",
                        f"æ ‡é¢˜ {lvl}",
                        f"æ ‡é¢˜{lvl}",
                    ]
                )
                if fallback:
                    setattr(hints, attr, fallback)
        except Exception:
            return

    def get_format_template_analysis_summary(self, template_id: str) -> Dict[str, Any]:
        """
        è·å–æ¨¡æ¿åˆ†ææ‘˜è¦
        
        Args:
            template_id: æ¨¡æ¿ ID
            
        Returns:
            åˆ†ææ‘˜è¦
        """
        template = self.dao.get_format_template(template_id)
        if not template:
            return {"error": "Template not found"}
        
        if not template.get("template_spec_json"):
            return {
                "analyzed": False,
                "message": "Template not analyzed yet"
            }
        
        try:
            spec = TemplateSpec.from_json(template["template_spec_json"])
            
            # ç»Ÿè®¡ outline èŠ‚ç‚¹æ•°
            def count_nodes(nodes):
                count = len(nodes)
                for node in nodes:
                    count += count_nodes(node.children)
                return count
            
            outline_node_count = count_nodes(spec.outline)
            
            # ç»Ÿè®¡æ ·å¼æç¤º
            style_hints_count = sum(1 for key in [
                "heading1_style", "heading2_style", "heading3_style",
                "body_style", "table_style"
            ] if getattr(spec.style_hints, key, None))
            
            return {
                "analyzed": True,
                "version": spec.version,
                "confidence": spec.diagnostics.confidence,
                "warnings": spec.diagnostics.warnings,
                "exclude_block_ids_count": len(spec.base_policy.exclude_block_ids),
                "outline_node_count": outline_node_count,
                "style_hints_count": style_hints_count,
                "base_policy_mode": spec.base_policy.mode.value,
                "analyzed_at": template.get("template_spec_analyzed_at"),
                "analysis_duration_ms": spec.diagnostics.analysis_duration_ms,
                "llm_model": spec.diagnostics.llm_model
            }
        except Exception as e:
            return {
                "analyzed": True,
                "error": f"Failed to parse spec: {str(e)}"
            }

    def apply_format_template_to_directory(self, project_id: str, format_template_id: str) -> List[Dict[str, Any]]:
        """
        å°†æ ¼å¼æ¨¡æ¿åº”ç”¨åˆ°ç›®å½•ï¼ˆå†™åº“ï¼‰ï¼š
        - å¦‚æœæ¨¡æ¿å®šä¹‰ç»“æ„ï¼ˆmerge_policy.template_defines_structure=trueï¼‰ï¼ŒæŒ‰æ¨¡æ¿ç»“æ„åˆå¹¶ç›®å½•
        - æ— è®ºæ˜¯å¦æ”¹ç»“æ„ï¼Œéƒ½æŠŠ format_template_id å†™å…¥æ¯ä¸ªèŠ‚ç‚¹ meta_jsonï¼ˆä½œä¸ºâ€œå·²å¥—ç”¨æ¨¡æ¿â€æ ‡è®°ï¼‰
        - è¿”å›å¸¦ bodyMeta çš„ç›®å½•èŠ‚ç‚¹ï¼ˆç”¨äºå‰ç«¯åŸåœ°åˆ·æ–°ï¼‰
        """
        template_spec = self.get_format_template_spec(format_template_id)
        nodes = self.dao.list_directory(project_id)

        if template_spec and getattr(template_spec, "merge_policy", None) and template_spec.merge_policy.template_defines_structure:
            nodes = OutlineMerger.merge_with_template(nodes, template_spec)

        # è®°å½•æ‰€é€‰æ ¼å¼æ¨¡æ¿ï¼ˆè½åº“åœ¨ç›®å½•èŠ‚ç‚¹ meta_json ä¸­ï¼Œé¿å…é¢å¤–å¼•å…¥é¡¹ç›® settings è¡¨ï¼‰
        for n in nodes:
            meta = n.get("meta_json") or {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except Exception:
                    meta = {}
            if not isinstance(meta, dict):
                meta = {}
            meta["format_template_id"] = format_template_id
            n["meta_json"] = meta

        self.dao.replace_directory(project_id, nodes)
        return self.get_directory_with_body_meta(project_id)

    async def extract_directory_with_template_merge(
        self,
        project_id: str,
        model_id: Optional[str],
        template_id: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        æå–ç›®å½•å¹¶ä¸æ¨¡æ¿åˆå¹¶
        
        Args:
            project_id: é¡¹ç›® ID
            model_id: LLM æ¨¡å‹ ID
            template_id: æ ¼å¼æ¨¡æ¿ IDï¼ˆå¯é€‰ï¼‰
            
        Returns:
            åˆå¹¶åçš„ç›®å½•èŠ‚ç‚¹åˆ—è¡¨
        """
        # 1. å…ˆç”¨ AI æå–ç›®å½•
        ai_nodes = await self.extract_directory(project_id, model_id)
        
        # 2. å¦‚æœæŒ‡å®šäº†æ¨¡æ¿ï¼Œè¿›è¡Œåˆå¹¶
        if template_id:
            template_spec = self.get_format_template_spec(template_id)
            if template_spec and template_spec.merge_policy.template_defines_structure:
                # ä½¿ç”¨ OutlineMerger åˆå¹¶
                ai_nodes = OutlineMerger.merge_with_template(ai_nodes, template_spec)
        
        return ai_nodes

    def preview_directory_by_template(self, project_id: str, template_asset_id: str) -> Dict[str, Any]:
        """
        preview è§„åˆ™ï¼š
        - å¦‚æœæ¨¡æ¿è§£æåˆ° outlineï¼šç”¨æ¨¡æ¿ outline ä½œä¸ºé¢„è§ˆç›®å½•ï¼ˆsource=templateï¼‰
        - å¦åˆ™ï¼šè¿”å›å½“å‰ç›®å½•ï¼ˆsource=tender/manualï¼‰
        - åŒæ—¶è¿”å›æ¨¡æ¿çš„æ ·å¼æç¤ºï¼ˆstyle_hintsï¼‰
        """
        assets = self.dao.get_assets_by_ids(project_id, [template_asset_id])
        if not assets:
            # å…¼å®¹ï¼šå‰ç«¯â€œæ ¼å¼æ¨¡æ¿â€é€‰æ‹©å™¨ä¼ å…¥çš„æ˜¯ format_templates è¡¨çš„ idï¼ˆtpl_...ï¼‰
            # è¿™é‡Œåšä¸€æ¬¡ fallbackï¼šå¦‚æœæ˜¯æ ¼å¼æ¨¡æ¿IDï¼Œè¿”å›å½“å‰ç›®å½• + è¯¥æ¨¡æ¿çš„ style_hints
            tpl = self.dao.get_format_template(template_asset_id)
            if tpl:
                spec_raw = tpl.get("template_spec_json")
                spec = {}
                if isinstance(spec_raw, str) and spec_raw.strip():
                    try:
                        spec = json.loads(spec_raw)
                    except Exception:
                        spec = {}
                elif isinstance(spec_raw, dict):
                    spec = spec_raw

                style_hints = {}
                if isinstance(spec, dict):
                    style_hints = spec.get("style_hints") or {}

                # å¦‚æœæ²¡æœ‰ style_hintsï¼Œå°è¯•ä½¿ç”¨ style_configï¼ˆå­˜å‚¨åœ¨ format_templates.style_configï¼‰
                if not style_hints:
                    style_hints = tpl.get("style_config") or {}

                # å…œåº•é»˜è®¤å€¼ï¼ˆå‰ç«¯åªå…³å¿ƒè¿™äº›å­—æ®µï¼‰
                if not style_hints:
                    style_hints = {
                        "page_background": "#ffffff",
                        "font_family": "SimSun, serif",
                        "font_size": "14px",
                        "line_height": "1.6",
                        "toc_indent_1": "0px",
                        "toc_indent_2": "20px",
                        "toc_indent_3": "40px",
                        "toc_indent_4": "60px",
                    }

                # è¿™é‡Œå…ˆä¸å¼ºè¡Œç”¨æ¨¡æ¿ outlineï¼ˆæ ¼å¼æ¨¡æ¿ outline ç»“æ„ä¸ç›®å½•èŠ‚ç‚¹ç»“æ„ä¸å®Œå…¨ä¸€è‡´ï¼‰ï¼Œ
                # å…ˆä¿è¯â€œåŠ è½½æ ·å¼â€å¯ç”¨ï¼šè¿”å›å½“å‰ç›®å½• + æ¨¡æ¿ style_hintsã€‚
                return {
                    "nodes": self.dao.list_directory(project_id),
                    "style_hints": style_hints,
                }

            raise ValueError("template asset not found")
        meta = assets[0].get("meta_json") or {}
        if isinstance(meta, str):
            try:
                meta = json.loads(meta)
            except Exception:
                meta = {}
        
        # è·å–ç›®å½•èŠ‚ç‚¹
        outline = meta.get("template_outline_nodes") or []
        nodes = outline if isinstance(outline, list) and len(outline) > 0 else self.dao.list_directory(project_id)
        
        # å°è¯•è·å– template_specï¼ˆå¦‚æœå­˜åœ¨ï¼‰
        template_spec = meta.get("template_spec") or {}
        style_hints = template_spec.get("style_hints") if isinstance(template_spec, dict) else {}
        
        # å¦‚æœæ²¡æœ‰ style_hintsï¼Œä½¿ç”¨é»˜è®¤å€¼
        if not style_hints:
            style_hints = {
                "page_background": "#ffffff",
                "font_family": "SimSun, serif",
                "font_size": "14px",
                "line_height": "1.6",
                "toc_indent_1": "0px",
                "toc_indent_2": "20px",
                "toc_indent_3": "40px",
                "toc_indent_4": "60px",
            }
        
        return {
            "nodes": nodes,
            "style_hints": style_hints,
        }

    def apply_template_to_directory(self, project_id: str, template_asset_id: str) -> int:
        """å¥—ç”¨æ¨¡æ¿åˆ°ç›®å½•"""
        result = self.preview_directory_by_template(project_id, template_asset_id)
        nodes = result["nodes"]
        # è½åº“ï¼šDAO ä¼šè¡¥ parent_id/order_no
        self.dao.replace_directory(project_id, nodes)
        return len(nodes)
    
    # ==================== é¡¹ç›®ç®¡ç†ï¼ˆç¼–è¾‘ã€åˆ é™¤ï¼‰ ====================
    
    def update_project(self, project_id: str, name: Optional[str], description: Optional[str]) -> Dict[str, Any]:
        """
        æ›´æ–°é¡¹ç›®ä¿¡æ¯
        
        Args:
            project_id: é¡¹ç›®ID
            name: æ–°é¡¹ç›®åç§°ï¼ˆå¯é€‰ï¼‰
            description: æ–°é¡¹ç›®æè¿°ï¼ˆå¯é€‰ï¼‰
            
        Returns:
            æ›´æ–°åçš„é¡¹ç›®ä¿¡æ¯
        """
        if name is not None and not name.strip():
            raise ValueError("Project name cannot be empty")
        
        return self.dao.update_project(project_id, name, description)
    
    def get_project_delete_plan(self, project_id: str) -> ProjectDeletePlanResponse:
        """
        è·å–é¡¹ç›®åˆ é™¤è®¡åˆ’ï¼ˆé¢„æ£€ï¼‰
        
        Args:
            project_id: é¡¹ç›®ID
            
        Returns:
            åˆ é™¤è®¡åˆ’ï¼ˆåŒ…å«èµ„æºæ¸…å•å’Œç¡®è®¤ä»¤ç‰Œï¼‰
        """
        return self.deletion_orchestrator.build_plan(project_id)
    
    def delete_project(self, project_id: str, confirm_request: ProjectDeleteRequest) -> None:
        """
        åˆ é™¤é¡¹ç›®ï¼ˆéœ€è¦ç¡®è®¤ï¼‰
        
        Args:
            project_id: é¡¹ç›®ID
            confirm_request: åˆ é™¤ç¡®è®¤è¯·æ±‚ï¼ˆåŒ…å«ç¡®è®¤ä»¤ç‰Œï¼‰
            
        Raises:
            ValueError: ç¡®è®¤ä¿¡æ¯ä¸åŒ¹é…
        """
        # 1. è·å–å½“å‰é¡¹ç›®ä¿¡æ¯
        project = self.dao.get_project(project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")
        
        # 2. é‡æ–°ç”Ÿæˆåˆ é™¤è®¡åˆ’å¹¶éªŒè¯ä»¤ç‰Œ
        plan = self.deletion_orchestrator.build_plan(project_id)
        if confirm_request.confirm_token != plan.confirm_token:
            raise ValueError("Confirm token mismatch. Please regenerate the delete plan.")
        
        # 3. æ‰§è¡Œåˆ é™¤
        self.deletion_orchestrator.delete(project_id)
    
    # ==================== æ ¼å¼æ¨¡æ¿ç®¡ç†æ‰©å±• ====================
    
    def update_format_template_meta(
        self,
        template_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        is_public: Optional[bool] = None
    ) -> Dict[str, Any]:
        """
        æ›´æ–°æ ¼å¼æ¨¡æ¿å…ƒæ•°æ®
        
        Args:
            template_id: æ¨¡æ¿ID
            name: æ–°åç§°ï¼ˆå¯é€‰ï¼‰
            description: æ–°æè¿°ï¼ˆå¯é€‰ï¼‰
            is_public: æ˜¯å¦å…¬å¼€ï¼ˆå¯é€‰ï¼‰
            
        Returns:
            æ›´æ–°åçš„æ¨¡æ¿è®°å½•
        """
        if name is not None and not name.strip():
            raise ValueError("Template name cannot be empty")
        
        return self.dao.update_format_template_meta(template_id, name, description, is_public)
    
    def get_format_template_extract(self, template_id: str, docx_bytes: bytes) -> Dict[str, Any]:
        """
        è·å–æ ¼å¼æ¨¡æ¿çš„è§£æè¯¦æƒ…ï¼ˆblocks + exclude ä¿¡æ¯ï¼‰
        
        Args:
            template_id: æ¨¡æ¿ID
            docx_bytes: Word æ–‡æ¡£å­—èŠ‚å†…å®¹
            
        Returns:
            è§£æè¯¦æƒ…ï¼ˆåŒ…å« blocks å’Œ exclude ä¿¡æ¯ï¼‰
        """
        # 1. æå–ç»“æ„åŒ– blocks
        extract_result = self.docx_extractor.extract(docx_bytes)
        
        # 2. è·å–æ¨¡æ¿çš„ spec
        spec = self.get_format_template_spec(template_id)
        
        # 3. æ•´ç†è¾“å‡º
        blocks = []
        excluded_block_ids = set()
        
        if spec and spec.base_policy and hasattr(spec.base_policy, 'excluded_block_ids'):
            excluded_block_ids = set(spec.base_policy.excluded_block_ids or [])
        
        for block in extract_result.blocks:
            block_info = {
                "block_id": block.block_id,
                "type": block.type,
                "content_preview": (block.content or "")[:100],
                "style": block.style or "",
                "excluded": block.block_id in excluded_block_ids,
            }
            blocks.append(block_info)
        
        return {
            "blocks": blocks,
            "total_blocks": len(blocks),
            "excluded_count": len(excluded_block_ids),
            "style_stats": extract_result.style_stats,
        }

    # ==================== è¯­ä¹‰ç›®å½•ç”Ÿæˆ ====================

    def generate_semantic_outline(
        self,
        project_id: str,
        mode: str = "FAST",
        max_depth: int = 5,
    ) -> Dict[str, Any]:
        """
        ç”Ÿæˆè¯­ä¹‰ç›®å½•ï¼ˆä»è¯„åˆ†/è¦æ±‚æ¨å¯¼ï¼‰
        
        ç°åœ¨ç›´æ¥è°ƒç”¨ works/tender/outline çš„ç»Ÿä¸€å…¥å£
        
        Args:
            project_id: é¡¹ç›®ID
            mode: ç”Ÿæˆæ¨¡å¼ FAST/FULL
            max_depth: æœ€å¤§å±‚çº§
            
        Returns:
            è¯­ä¹‰ç›®å½•ç»“æœ
        """
        from app.works.tender.outline.outline_v2_service import generate_outline_v2
        
        # è·å–é¡¹ç›®ä¿¡æ¯ï¼ˆç”¨äºowner_idï¼‰
        project = self.dao.get_project(project_id)
        owner_id = project.get("owner_id") if project else None
        
        # è°ƒç”¨ç»Ÿä¸€å…¥å£
        return generate_outline_v2(
            pool=self.pool,
            project_id=project_id,
            owner_id=owner_id,
            mode=mode,
            max_depth=max_depth,
            llm_orchestrator=self.llm,
        )

    def _flatten_outline_nodes(
        self,
        nodes: List[Any],
        outline_id: str,
        project_id: str,
        parent_id: Optional[str] = None,
        order_offset: int = 0,
    ) -> List[Dict[str, Any]]:
        """
        å°†æ ‘å½¢ç›®å½•èŠ‚ç‚¹æ‰å¹³åŒ–ä¸ºæ•°æ®åº“å­˜å‚¨æ ¼å¼
        
        Args:
            nodes: æ ‘å½¢èŠ‚ç‚¹åˆ—è¡¨
            outline_id: ç›®å½•ID
            project_id: é¡¹ç›®ID
            parent_id: çˆ¶èŠ‚ç‚¹ID
            order_offset: æ’åºåç§»é‡
            
        Returns:
            æ‰å¹³åŒ–èŠ‚ç‚¹åˆ—è¡¨
        """
        result = []
        order_no = order_offset
        
        for node in nodes:
            order_no += 1
            
            flat_node = {
                "node_id": node.node_id,
                "outline_id": outline_id,
                "project_id": project_id,
                "parent_id": parent_id,
                "level": node.level,
                "order_no": order_no,
                "numbering": node.numbering,
                "title": node.title,
                "summary": node.summary,
                "tags": node.tags,
                "evidence_chunk_ids": node.evidence_chunk_ids,
                "covered_req_ids": node.covered_req_ids,
            }
            
            result.append(flat_node)
            
            # é€’å½’å¤„ç†å­èŠ‚ç‚¹
            if node.children:
                child_nodes = self._flatten_outline_nodes(
                    nodes=node.children,
                    outline_id=outline_id,
                    project_id=project_id,
                    parent_id=node.node_id,
                    order_offset=order_no,
                )
                result.extend(child_nodes)
                order_no += len(child_nodes)
        
        return result

    def get_semantic_outline(self, outline_id: str) -> Optional[Dict[str, Any]]:
        """
        è·å–è¯­ä¹‰ç›®å½•ï¼ˆåŒ…å«å®Œæ•´çš„æ ‘å½¢ç»“æ„ï¼‰
        
        Args:
            outline_id: ç›®å½•ID
            
        Returns:
            å®Œæ•´çš„è¯­ä¹‰ç›®å½•ç»“æœï¼ŒåŒ…å«æ ‘å½¢ç»“æ„
        """
        # 1. è·å–ç›®å½•åŸºæœ¬ä¿¡æ¯
        outline = self.dao.get_semantic_outline(outline_id)
        if not outline:
            return None
        
        # 2. è·å–è¦æ±‚é¡¹
        requirements = self.dao.get_requirement_items(outline_id)
        
        # 3. è·å–èŠ‚ç‚¹ï¼ˆæ‰å¹³ï¼‰å¹¶é‡å»ºæ ‘å½¢ç»“æ„
        nodes_flat = self.dao.get_semantic_outline_nodes(outline_id)
        outline_tree = self._rebuild_outline_tree(nodes_flat)
        
        # 4. ç»„è£…è¿”å›ç»“æœ
        return {
            "outline_id": outline["outline_id"],
            "project_id": outline["project_id"],
            "mode": outline["mode"],
            "max_depth": outline["max_depth"],
            "status": outline["status"],
            "coverage_rate": outline["coverage_rate"],
            "diagnostics": outline.get("diagnostics_json", {}),
            "outline": outline_tree,
            "requirements": requirements,
            "created_at": outline.get("created_at"),
        }

    def _rebuild_outline_tree(self, nodes_flat: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        ä»æ‰å¹³èŠ‚ç‚¹åˆ—è¡¨é‡å»ºæ ‘å½¢ç»“æ„
        
        Args:
            nodes_flat: æ‰å¹³èŠ‚ç‚¹åˆ—è¡¨
            
        Returns:
            æ ‘å½¢èŠ‚ç‚¹åˆ—è¡¨
        """
        if not nodes_flat:
            return []
        
        # æ„å»º node_id -> node æ˜ å°„
        node_map = {}
        for node in nodes_flat:
            node_copy = dict(node)
            node_copy["children"] = []
            node_map[node["node_id"]] = node_copy
        
        # æ„å»ºæ ‘å½¢ç»“æ„
        root_nodes = []
        for node in nodes_flat:
            node_obj = node_map[node["node_id"]]
            parent_id = node.get("parent_id")
            
            if parent_id and parent_id in node_map:
                # æ·»åŠ åˆ°çˆ¶èŠ‚ç‚¹çš„children
                node_map[parent_id]["children"].append(node_obj)
            else:
                # æ ¹èŠ‚ç‚¹
                root_nodes.append(node_obj)
        
        return root_nodes

    def get_latest_semantic_outline(self, project_id: str) -> Optional[Dict[str, Any]]:
        """
        è·å–é¡¹ç›®æœ€æ–°çš„è¯­ä¹‰ç›®å½•
        
        Args:
            project_id: é¡¹ç›®ID
            
        Returns:
            æœ€æ–°çš„è¯­ä¹‰ç›®å½•ç»“æœ
        """
        outline = self.dao.get_latest_semantic_outline(project_id)
        if not outline:
            return None
        
        return self.get_semantic_outline(outline["outline_id"])
    
    async def _build_tender_project_context(self, project_id: str) -> str:
        """æ„å»ºæ‹›æ ‡é¡¹ç›®ä¸Šä¸‹æ–‡"""
        context_parts = []
        try:
            project_info = self.dao.get_project_info(project_id)
            if project_info and project_info.get("data_json"):
                data = project_info.get("data_json", {})
                if data.get("project_name"):
                    context_parts.append(f"é¡¹ç›®åç§°ï¼š{data['project_name']}")
                if data.get("tenderee"):
                    context_parts.append(f"æ‹›æ ‡äººï¼š{data['tenderee']}")
                if data.get("budget"):
                    context_parts.append(f"é¢„ç®—é‡‘é¢ï¼š{data['budget']}")
                if data.get("project_overview"):
                    context_parts.append(f"é¡¹ç›®æ¦‚å†µï¼š{data['project_overview']}")
        except Exception as e:
            logger.warning(f"[TenderService] è·å–é¡¹ç›®ä¿¡æ¯å¤±è´¥: {e}")
        if len(context_parts) < 3:
            context_parts.append("ï¼ˆæ³¨ï¼šé¡¹ç›®ä¿¡æ¯ä¸è¶³ï¼Œè¯·æ ¹æ®ç« èŠ‚æ ‡é¢˜ç”Ÿæˆåˆç†å†…å®¹ï¼‰")
        return "\n".join(context_parts)
    
    async def _retrieve_context_for_section(
        self, 
        project_id: str, 
        section_title: str,
        top_k: int = 5
    ) -> Dict[str, Any]:
        """
        ä¸ºç« èŠ‚æ£€ç´¢ç›¸å…³ä¼ä¸šèµ„æ–™
        
        Args:
            project_id: é¡¹ç›®ID
            section_title: ç« èŠ‚æ ‡é¢˜
            top_k: è¿”å›çš„æœ€ç›¸å…³ç‰‡æ®µæ•°é‡
            
        Returns:
            æ£€ç´¢ç»“æœ: {
                "chunks": [...],  # æ£€ç´¢åˆ°çš„æ–‡æ¡£ç‰‡æ®µ
                "quality_score": float,  # æ£€ç´¢è´¨é‡è¯„åˆ† (0-1)
                "has_relevant": bool  # æ˜¯å¦æœ‰ç›¸å…³å†…å®¹
            }
        """
        from app.platform.ingest.v2_service import IngestV2Service
        from app.services.db.postgres import _get_pool
        
        result = {
            "chunks": [],
            "quality_score": 0.0,
            "has_relevant": False
        }
        
        try:
            # è·å–é¡¹ç›®çš„çŸ¥è¯†åº“ID
            proj = self.dao.get_project(project_id)
            if not proj:
                logger.warning(f"é¡¹ç›®ä¸å­˜åœ¨: {project_id}")
                return result
            
            kb_id = proj.get("kb_id")
            if not kb_id:
                logger.warning(f"é¡¹ç›®æœªç»‘å®šçŸ¥è¯†åº“: {project_id}")
                return result
            
            # æ„å»ºæ£€ç´¢query
            query = self._build_retrieval_query(section_title)
            
            # ä»Milvusæ£€ç´¢
            ingest_service = IngestV2Service(_get_pool())
            search_results = await ingest_service.search_in_kb(
                kb_id=kb_id,
                query_text=query,
                top_k=top_k,
                filters={
                    "doc_type": [
                        "qualification_doc",  # company_profile, cert_doc
                        "technical_material",  # tech_doc
                        "history_case",       # case_study
                        "financial_doc"       # finance_doc
                    ]
                }
            )
            
            if not search_results:
                return result
            
            # è¯„ä¼°æ£€ç´¢è´¨é‡
            quality_score = self._assess_retrieval_quality(search_results)
            has_relevant = quality_score > 0.4
            
            result = {
                "chunks": search_results,
                "quality_score": quality_score,
                "has_relevant": has_relevant
            }
            
            logger.info(
                f"[æ£€ç´¢] ç« èŠ‚={section_title}, "
                f"è¿”å›={len(search_results)}æ¡, "
                f"è´¨é‡={quality_score:.2f}"
            )
            
        except Exception as e:
            logger.error(f"æ£€ç´¢å¤±è´¥: {e}", exc_info=True)
        
        return result
    
    def _build_retrieval_query(self, section_title: str) -> str:
        """
        æ ¹æ®ç« èŠ‚æ ‡é¢˜æ„å»ºæ£€ç´¢query
        
        Args:
            section_title: ç« èŠ‚æ ‡é¢˜
            
        Returns:
            æ£€ç´¢queryå­—ç¬¦ä¸²
        """
        # ç« èŠ‚æ ‡é¢˜ -> æ£€ç´¢æ„å›¾æ˜ å°„
        title_lower = section_title.lower()
        
        if any(kw in title_lower for kw in ["å…¬å¸", "ä¼ä¸š", "ç®€ä»‹", "æ¦‚å†µ", "èµ„è´¨"]):
            return f"{section_title} ä¼ä¸šç®€ä»‹ èµ„è´¨è¯ä¹¦ è£èª‰å¥–é¡¹"
        elif any(kw in title_lower for kw in ["æŠ€æœ¯", "æ–¹æ¡ˆ", "å®æ–½", "è®¾è®¡"]):
            return f"{section_title} æŠ€æœ¯æ–¹æ¡ˆ å®æ–½æ–¹æ³• æŠ€æœ¯è·¯çº¿"
        elif any(kw in title_lower for kw in ["æ¡ˆä¾‹", "ä¸šç»©", "é¡¹ç›®ç»éªŒ", "æˆåŠŸæ¡ˆä¾‹"]):
            return f"{section_title} é¡¹ç›®æ¡ˆä¾‹ æˆåŠŸä¸šç»© ç±»ä¼¼é¡¹ç›®"
        elif any(kw in title_lower for kw in ["è´¢åŠ¡", "æŠ¥è¡¨", "å®¡è®¡"]):
            return f"{section_title} è´¢åŠ¡æŠ¥è¡¨ å®¡è®¡æŠ¥å‘Š"
        else:
            return section_title
    
    def _assess_retrieval_quality(self, search_results: List[Dict]) -> float:
        """
        è¯„ä¼°æ£€ç´¢è´¨é‡
        
        Args:
            search_results: æ£€ç´¢ç»“æœåˆ—è¡¨
            
        Returns:
            è´¨é‡è¯„åˆ† (0-1)
        """
        if not search_results:
            return 0.0
        
        # åŸºäºç›¸ä¼¼åº¦åˆ†æ•°å’Œæ•°é‡è¯„ä¼°
        scores = [chunk.get("score", 0.0) for chunk in search_results]
        
        # å¹³å‡ç›¸ä¼¼åº¦
        avg_score = sum(scores) / len(scores)
        
        # æœ€é«˜ç›¸ä¼¼åº¦
        max_score = max(scores)
        
        # ç»¼åˆè¯„åˆ† (æƒé‡: æœ€é«˜0.6 + å¹³å‡0.4)
        quality = max_score * 0.6 + avg_score * 0.4
        
        return min(quality, 1.0)
    
    async def _generate_section_content(
        self,
        project_id: str,
        title: str,
        level: int,
        project_context: str,
        requirements: Optional[str] = None,  # âœ… æ–°å¢ï¼šç”¨æˆ·è‡ªå®šä¹‰è¦æ±‚
        model_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        ä¸ºå•ä¸ªç« èŠ‚ç”Ÿæˆå†…å®¹ï¼ˆä½¿ç”¨ç»Ÿä¸€ç»„ä»¶ï¼‰
        
        Returns:
            {
                "content": str,  # HTMLæ ¼å¼çš„ç« èŠ‚å†…å®¹
                "evidence_chunk_ids": List[str]  # å¼•ç”¨çš„èµ„æ–™ç‰‡æ®µID
                "quality_metrics": Dict  # è´¨é‡æŒ‡æ ‡
            }
        """
        from app.services.generation import (
            DocumentRetriever,
            RetrievalContext,
            PromptBuilder,
            PromptContext,
            ContentGenerator,
            GenerationContext,
            QualityAssessor
        )
        from app.platform.ingest.v2_service import IngestV2Service
        from app.services.db.postgres import _get_pool
        
        # Step 1: è·å–é¡¹ç›®ä¿¡æ¯
        proj = self.dao.get_project(project_id)
        if not proj:
            raise ValueError(f"é¡¹ç›®ä¸å­˜åœ¨: {project_id}")
        
        kb_id = proj.get("kb_id")
        if not kb_id:
            raise ValueError(f"é¡¹ç›®æœªç»‘å®šçŸ¥è¯†åº“: {project_id}")
        
        project_info_dict = {}
        try:
            proj_info = self.dao.get_project_info(project_id)
            if proj_info and proj_info.get("data_json"):
                project_info_dict = proj_info["data_json"]
        except Exception as e:
            logger.warning(f"è·å–é¡¹ç›®ä¿¡æ¯å¤±è´¥: {e}")
        
        # Step 2: æ£€ç´¢ç›¸å…³èµ„æ–™ï¼ˆä½¿ç”¨ç»Ÿä¸€ç»„ä»¶ï¼‰
        retriever = DocumentRetriever(_get_pool())
        retrieval_context = RetrievalContext(
            kb_id=kb_id,
            section_title=title,
            section_level=level,
            document_type="tender",
            project_info=project_info_dict
        )
        retrieval_result = await retriever.retrieve(retrieval_context, top_k=5)
        
        # Step 3: æ„å»ºPromptï¼ˆä½¿ç”¨ç»Ÿä¸€ç»„ä»¶ï¼‰
        prompt_builder = PromptBuilder()
        
        # âœ… å¦‚æœæœ‰ç”¨æˆ·è‡ªå®šä¹‰è¦æ±‚ï¼Œæ„å»ºrequirementså­—å…¸
        requirements_dict = None
        if requirements:
            requirements_dict = {"custom_requirements": requirements}
        
        prompt_context = PromptContext(
            document_type="tender",
            section_title=title,
            section_level=level,
            project_info=project_info_dict,
            retrieval_result=retrieval_result,
            requirements=requirements_dict  # âœ… ä¼ é€’ç”¨æˆ·è¦æ±‚
        )
        prompt = prompt_builder.build(prompt_context)
        
        # Step 4: ç”Ÿæˆå†…å®¹ï¼ˆä½¿ç”¨ç»Ÿä¸€ç»„ä»¶ï¼‰
        generator = ContentGenerator(self.llm)
        gen_context = GenerationContext(
            document_type="tender",
            section_title=title,
            prompt=prompt,
            model_id=model_id
        )
        generation_result = await generator.generate(gen_context)
        
        # Step 5: è¯„ä¼°è´¨é‡ï¼ˆä½¿ç”¨ç»Ÿä¸€ç»„ä»¶ï¼‰
        assessor = QualityAssessor()
        quality_metrics = assessor.assess(
            generation_result,
            retrieval_result,
            level
        )
        
        # Step 6: è®°å½•è´¨é‡æŒ‡æ ‡
        logger.info(
            f"[ç”Ÿæˆ] ç« èŠ‚={title}, "
            f"å­—æ•°={generation_result.word_count}, "
            f"è¯æ®={len(retrieval_result.chunks)}æ¡, "
            f"è´¨é‡={quality_metrics.overall_score:.2f}, "
            f"ç­‰çº§={quality_metrics.get_grade()}"
        )
        
        return {
            "content": generation_result.content,
            "evidence_chunk_ids": retrieval_result.get_chunk_ids(),
            "quality_metrics": quality_metrics.to_dict()
        }
    
    async def generate_full_content(
        self,
        project_id: str,
        model_id: Optional[str] = None,
        run_id: Optional[str] = None,
        max_concurrent: int = 5,
    ) -> Dict[str, Any]:
        """
        å¹¶è¡Œç”Ÿæˆæ ‡ä¹¦æ‰€æœ‰ç« èŠ‚å†…å®¹
        
        Args:
            project_id: é¡¹ç›®ID
            model_id: LLMæ¨¡å‹ID
            run_id: è¿è¡Œè®°å½•ID
            max_concurrent: æœ€å¤§å¹¶å‘æ•°ï¼ˆé»˜è®¤5ï¼‰
        
        Returns:
            ç”Ÿæˆç»“æœç»Ÿè®¡
        """
        import asyncio
        
        try:
            logger.info(f"[TenderService] å¼€å§‹å¹¶è¡Œç”Ÿæˆæ ‡ä¹¦å†…å®¹: project_id={project_id}, max_concurrent={max_concurrent}")
            
            # æ›´æ–° run çŠ¶æ€
            if run_id:
                self.dao.update_run(run_id, "running", progress=0.0, message="å¼€å§‹ç”Ÿæˆ...")
            
            # è·å–æ‰€æœ‰ç›®å½•èŠ‚ç‚¹
            nodes = self.dao.get_directory_nodes(project_id)
            if not nodes:
                raise ValueError("æ²¡æœ‰æ‰¾åˆ°ç›®å½•èŠ‚ç‚¹")
            
            # æ„å»ºé¡¹ç›®ä¸Šä¸‹æ–‡
            project_context = await self._build_tender_project_context(project_id)
            
            # ç­›é€‰éœ€è¦ç”Ÿæˆå†…å®¹çš„èŠ‚ç‚¹ï¼ˆæ²¡æœ‰sectionæˆ–sectionä¸ºç©ºï¼‰
            nodes_to_generate = []
            nodes_with_snippet = []  # å·²æœ‰èŒƒæœ¬çš„èŠ‚ç‚¹
            for node in nodes:
                node_id = node.get("id")
                section = self.dao.get_section_body(project_id, node_id)
                
                # æ£€æŸ¥èŠ‚ç‚¹æ˜¯å¦å·²æœ‰èŒƒæ–‡å†…å®¹
                has_snippet = False
                if section:
                    meta_json = node.get("meta_json", {})
                    snippet_blocks = meta_json.get("snippet_blocks") if isinstance(meta_json, dict) else None
                    if snippet_blocks and len(snippet_blocks) > 0:
                        has_snippet = True
                        nodes_with_snippet.append(node)
                        logger.info(f"è·³è¿‡èŠ‚ç‚¹ï¼ˆå·²æœ‰èŒƒæ–‡ï¼‰: {node.get('title')}")
                
                # å¦‚æœæ²¡æœ‰èŒƒæ–‡ï¼Œä¸”æ²¡æœ‰sectionæˆ–å†…å®¹ä¸ºç©º/å ä½ç¬¦ï¼Œåˆ™éœ€è¦ç”Ÿæˆ
                if not has_snippet and (not section or self._is_empty_section(section)):
                    nodes_to_generate.append(node)
            
            total = len(nodes_to_generate)
            logger.info(f"[TenderService] éœ€è¦ç”Ÿæˆå†…å®¹çš„èŠ‚ç‚¹æ•°: {total}")
            logger.info(f"[TenderService] å·²æœ‰èŒƒæ–‡çš„èŠ‚ç‚¹æ•°: {len(nodes_with_snippet)}")
            
            if total == 0:
                if run_id:
                    self.dao.update_run(
                        run_id,
                        "success",
                        progress=1.0,
                        message=f"æ‰€æœ‰ç« èŠ‚å·²æœ‰å†…å®¹ï¼Œæ— éœ€ç”Ÿæˆï¼ˆ{len(nodes_with_snippet)}ä¸ªå·²æœ‰èŒƒæ–‡ï¼‰",
                        result_json={
                            "generated": 0, 
                            "total": len(nodes), 
                            "skipped": len(nodes),
                            "snippet_count": len(nodes_with_snippet)
                        },
                    )
                return {
                    "generated": 0, 
                    "total": len(nodes), 
                    "skipped": len(nodes),
                    "snippet_count": len(nodes_with_snippet)
                }
            
            # åˆ›å»ºä¿¡å·é‡æ§åˆ¶å¹¶å‘
            semaphore = asyncio.Semaphore(max_concurrent)
            
            # ç”Ÿæˆä»»åŠ¡
            completed = 0
            failed = 0
            
            async def generate_one(node: Dict[str, Any], index: int) -> Tuple[bool, str]:
                """ç”Ÿæˆå•ä¸ªèŠ‚ç‚¹çš„å†…å®¹"""
                nonlocal completed, failed
                
                async with semaphore:
                    node_id = node.get("id")
                    title = node.get("title", "")
                    level = node.get("level", 1)
                    
                    try:
                        logger.info(f"[{index+1}/{total}] å¼€å§‹ç”Ÿæˆ: {title}")
                        
                        # ç”Ÿæˆå†…å®¹ï¼ˆæ–°ç‰ˆæœ¬è¿”å›å­—å…¸ï¼‰
                        result = await self._generate_section_content(
                            project_id=project_id,
                            title=title,
                            level=level,
                            project_context=project_context,
                            model_id=model_id,
                        )
                        
                        # æå–å†…å®¹å’Œè¯æ®
                        content = result.get("content", "")
                        evidence_chunk_ids = result.get("evidence_chunk_ids", [])
                        
                        # ä¿å­˜åˆ°æ•°æ®åº“
                        self.dao.upsert_section_body(
                            project_id=project_id,
                            node_id=node_id,
                            source="AI",  # æ ‡è®°ä¸ºAIç”Ÿæˆ
                            fragment_id=None,  # AIç”Ÿæˆä¸å…³è”æ¨¡æ¿ç‰‡æ®µ
                            content_html=content,
                            content_json=None,
                            evidence_chunk_ids=evidence_chunk_ids,  # å­˜å‚¨è¯æ®ID
                        )
                        
                        completed += 1
                        logger.info(
                            f"[{index+1}/{total}] ç”ŸæˆæˆåŠŸ: {title}, "
                            f"å­—æ•°={len(content)}, è¯æ®={len(evidence_chunk_ids)}æ¡"
                        )
                        
                        # æ›´æ–°è¿›åº¦
                        if run_id:
                            progress = completed / total
                            self.dao.update_run(
                                run_id,
                                "running",
                                progress=progress,
                                message=f"å·²å®Œæˆ {completed}/{total} ä¸ªç« èŠ‚",
                            )
                        
                        return True, title
                        
                    except Exception as e:
                        failed += 1
                        error_msg = f"ç”Ÿæˆå¤±è´¥: {str(e)}"
                        logger.error(f"[{index+1}/{total}] {title} - {error_msg}", exc_info=True)
                        return False, title
            
            # å¹¶è¡Œæ‰§è¡Œæ‰€æœ‰ç”Ÿæˆä»»åŠ¡
            tasks = [generate_one(node, i) for i, node in enumerate(nodes_to_generate)]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # ç»Ÿè®¡ç»“æœ
            success_count = sum(1 for r in results if isinstance(r, tuple) and r[0])
            
            result_json = {
                "generated": success_count,
                "failed": failed,
                "total": total,
                "skipped": len(nodes) - total,
                "snippet_count": len(nodes_with_snippet),
            }
            
            # æ›´æ–° run çŠ¶æ€
            if run_id:
                if failed == 0:
                    self.dao.update_run(
                        run_id,
                        "success",
                        progress=1.0,
                        message=f"ç”Ÿæˆå®Œæˆï¼æˆåŠŸ {success_count} ä¸ªç« èŠ‚ï¼Œ{len(nodes_with_snippet)} ä¸ªç« èŠ‚ä½¿ç”¨èŒƒæ–‡",
                        result_json=result_json,
                    )
                else:
                    self.dao.update_run(
                        run_id,
                        "partial",
                        progress=1.0,
                        message=f"éƒ¨åˆ†å®Œæˆï¼šæˆåŠŸ {success_count}ï¼Œå¤±è´¥ {failed}",
                        result_json=result_json,
                    )
            
            logger.info(f"[TenderService] å¹¶è¡Œç”Ÿæˆå®Œæˆ: {result_json}")
            return result_json
            
        except Exception as e:
            logger.error(f"[TenderService] å¹¶è¡Œç”Ÿæˆå¤±è´¥: {e}", exc_info=True)
            if run_id:
                self.dao.update_run(
                    run_id,
                    "failed",
                    progress=0.0,
                    message=str(e),
                    result_json={"error": str(e)},
                )
            raise
    
    def _is_empty_section(self, section: Dict[str, Any]) -> bool:
        """åˆ¤æ–­sectionæ˜¯å¦ä¸ºç©ºæˆ–åªæœ‰å ä½ç¬¦"""
        content_html = (section.get("content_html") or "").strip()
        content_md = (section.get("content_md") or "").strip()
        
        # å¦‚æœéƒ½ä¸ºç©º
        if not content_html and not content_md:
            return True
        
        # æ£€æŸ¥æ˜¯å¦åªæœ‰å ä½ç¬¦
        placeholders = [
            "ã€å¡«å†™ã€‘", "ã€å¾…è¡¥ã€‘", "ã€å¾…å¡«å†™ã€‘", "[å¡«å†™]", "[å¾…è¡¥]",
            "å¾…å¡«å†™", "å¾…è¡¥å……", "TODO", "TBD", "ï¼ˆå¾…è¡¥å……ï¼‰",
        ]
        
        content = content_html or content_md
        content = content.strip()
        
        # ç§»é™¤HTMLæ ‡ç­¾æ£€æŸ¥
        import re
        text_only = re.sub(r'<[^>]+>', '', content).strip()
        
        if not text_only or text_only in placeholders:
            return True
        
        # å†…å®¹å¤ªçŸ­ä¹Ÿè®¤ä¸ºéœ€è¦ç”Ÿæˆ
        if len(text_only) < 10:
            return True
        
        return False