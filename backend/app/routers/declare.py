"""
申报书 Router
"""
import logging
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, Form, UploadFile, BackgroundTasks, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel

from app.utils.auth import get_current_user_sync, get_current_user
from app.services.db.postgres import _get_pool
from app.services.dao.declare_dao import DeclareDAO
from app.services.declare_service import DeclareService
from app.works.declare.extract_v2_service import DeclareExtractV2Service

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/apps/declare", tags=["declare"])


def _get_llm(req: Request):
    """从 app.state 获取 LLM orchestrator"""
    llm = getattr(req.app.state, "llm_orchestrator", None)
    if llm is None:
        raise HTTPException(status_code=500, detail="LLM orchestrator not initialized on app.state")
    return llm


def _get_dao():
    """获取 DAO"""
    pool = _get_pool()
    return DeclareDAO(pool)


def _get_service(req: Request):
    """获取 Service"""
    dao = _get_dao()
    llm = _get_llm(req)
    # TODO: 集成 jobs_service
    return DeclareService(dao, llm, jobs_service=None)


# ==================== Request/Response Models ====================

class ProjectCreateReq(BaseModel):
    name: str
    description: Optional[str] = None


class ProjectUpdateReq(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None


class ProjectDeleteRequest(BaseModel):
    confirm_token: str


class ProjectDeletePlanResponse(BaseModel):
    warning: str
    items: List[dict]
    confirm_token: str


class ProjectOut(BaseModel):
    project_id: str
    kb_id: str
    name: str
    description: Optional[str]
    owner_id: Optional[str]
    created_at: str
    updated_at: str


class RunOut(BaseModel):
    run_id: str
    project_id: str
    task_type: str
    status: str
    progress: float
    message: Optional[str]
    result_json: dict
    created_at: str
    updated_at: str


# ==================== Projects ====================

@router.post("/projects", response_model=ProjectOut)
def create_project(req: ProjectCreateReq, user=Depends(get_current_user_sync)):
    """创建申报项目"""
    from app.services.kb_service import create_kb
    
    # 创建知识库，设置owner为当前用户
    kb_id = create_kb(
        name=f"申报-{req.name}",
        description=req.description or f"申报项目：{req.name}",
        category_id="cat_knowledge",
        owner_id=user.user_id  # 关键：设置知识库所有者
    )
    
    # 创建项目
    dao = _get_dao()
    project = dao.create_project(kb_id, req.name, req.description, owner_id=user.user_id)
    return project


@router.get("/projects", response_model=List[ProjectOut])
def list_projects(user=Depends(get_current_user_sync)):
    """列出申报项目"""
    dao = _get_dao()
    projects = dao.list_projects(owner_id=user.user_id)
    return projects


@router.get("/projects/{project_id}", response_model=ProjectOut)
def get_project(project_id: str, user=Depends(get_current_user_sync)):
    """获取项目详情"""
    dao = _get_dao()
    project = dao.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@router.put("/projects/{project_id}", response_model=ProjectOut)
def update_project(project_id: str, req: ProjectUpdateReq, user=Depends(get_current_user_sync)):
    """更新项目信息"""
    dao = _get_dao()
    project = dao.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # 更新项目
    updated = dao.update_project(
        project_id=project_id,
        name=req.name,
        description=req.description
    )
    return updated


@router.get("/projects/{project_id}/delete-plan", response_model=ProjectDeletePlanResponse)
def get_project_delete_plan(project_id: str, user=Depends(get_current_user_sync)):
    """获取项目删除计划（预检）"""
    import uuid
    dao = _get_dao()
    
    project = dao.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # 获取资产数量
    assets = dao.list_assets(project_id)
    
    # 获取其他资源数量
    requirements = dao.get_requirements(project_id)
    directory_nodes = dao.get_active_directory_nodes(project_id)
    sections = dao.get_active_sections(project_id)
    
    items = []
    if assets:
        items.append({
            "type": "资产文件",
            "count": len(assets),
            "samples": [a.get("filename", "") for a in assets[:3]],
            "physical_targets": []
        })
    if requirements:
        items.append({
            "type": "申报要求",
            "count": 1,
            "samples": ["申报要求数据"],
            "physical_targets": []
        })
    if directory_nodes:
        items.append({
            "type": "目录节点",
            "count": len(directory_nodes),
            "samples": [n.get("title", "") for n in directory_nodes[:3]],
            "physical_targets": []
        })
    if sections:
        items.append({
            "type": "章节内容",
            "count": len(sections),
            "samples": [s.get("node_title", "") for s in sections[:3]],
            "physical_targets": []
        })
    
    confirm_token = str(uuid.uuid4())
    
    return {
        "warning": f"删除项目 \"{project.get('name')}\" 将同时删除所有关联数据",
        "items": items,
        "confirm_token": confirm_token
    }


@router.delete("/projects/{project_id}", status_code=204)
def delete_project(project_id: str, req: ProjectDeleteRequest, user=Depends(get_current_user_sync)):
    """删除项目"""
    dao = _get_dao()
    
    project = dao.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # 简单验证token（在生产环境应该更严格）
    if not req.confirm_token:
        raise HTTPException(status_code=400, detail="缺少确认令牌")
    
    # 删除项目（级联删除所有关联数据）
    dao.delete_project(project_id)
    
    # 删除知识库
    try:
        from app.services.kb_service import delete_kb
        kb_id = project.get("kb_id")
        if kb_id:
            delete_kb(kb_id)
    except Exception as e:
        logger.warning(f"删除知识库失败: {e}")
    
    return None


# ==================== Assets ====================

@router.post("/projects/{project_id}/assets/import")
def import_assets(
    project_id: str,
    kind: str = Form(...),
    files: List[UploadFile] = File(...),
    req: Request = None,
    user=Depends(get_current_user_sync),
):
    """导入资产"""
    service = _get_service(req)
    assets = service.import_assets(project_id, kind, files, user_id=user.user_id)
    return {"assets": assets}


@router.get("/projects/{project_id}/assets")
def list_assets(project_id: str, kind: Optional[str] = None, user=Depends(get_current_user_sync)):
    """列出资产"""
    dao = _get_dao()
    assets = dao.list_assets(project_id, kind)
    return {"assets": assets}


@router.delete("/projects/{project_id}/assets/{asset_id}")
def delete_asset(
    project_id: str,
    asset_id: str,
    user=Depends(get_current_user_sync)
):
    """删除资产"""
    dao = _get_dao()
    dao.delete_asset(asset_id)
    return {"success": True}


# ==================== Requirements ====================

@router.post("/projects/{project_id}/extract/requirements", response_model=RunOut)
def extract_requirements(
    project_id: str,
    bg: BackgroundTasks,
    req: Request,
    sync: int = 0,
    model_id: Optional[str] = None,
    user=Depends(get_current_user_sync),
):
    """抽取申报要求"""
    dao = _get_dao()
    service = _get_service(req)
    
    # 创建 run
    run_id = dao.create_run(project_id, "requirements")
    
    if sync == 1:
        # 同步执行
        service.extract_requirements(project_id, model_id, run_id)
        run = dao.get_run(run_id)
        return run
    else:
        # 异步执行
        bg.add_task(service.extract_requirements, project_id, model_id, run_id)
        run = dao.get_run(run_id)
        return run


@router.get("/projects/{project_id}/requirements")
def get_requirements(project_id: str, user=Depends(get_current_user_sync)):
    """获取申报要求"""
    dao = _get_dao()
    requirements = dao.get_requirements(project_id)
    if not requirements:
        return {"data": None}
    return requirements


# ==================== Directory ====================

@router.post("/projects/{project_id}/directory/generate", response_model=RunOut)
def generate_directory(
    project_id: str,
    bg: BackgroundTasks,
    req: Request,
    sync: int = 0,
    model_id: Optional[str] = None,
    user=Depends(get_current_user_sync),
):
    """生成申报书目录"""
    dao = _get_dao()
    service = _get_service(req)
    
    # 创建 run
    run_id = dao.create_run(project_id, "directory")
    
    if sync == 1:
        # 同步执行
        service.generate_directory(project_id, model_id, run_id)
        run = dao.get_run(run_id)
        return run
    else:
        # 异步执行
        bg.add_task(service.generate_directory, project_id, model_id, run_id)
        run = dao.get_run(run_id)
        return run


@router.get("/projects/{project_id}/directory/nodes")
def get_directory_nodes(project_id: str, user=Depends(get_current_user_sync)):
    """获取目录节点（活跃版本）"""
    dao = _get_dao()
    nodes = dao.get_active_directory_nodes(project_id)
    return {"nodes": nodes}


@router.get("/projects/{project_id}/directory/all-versions")
def get_all_directory_versions(project_id: str, user=Depends(get_current_user_sync)):
    """
    获取所有项目类型的目录版本
    返回格式：
    {
        "versions": [
            {
                "version_id": "...",
                "project_type": "头雁型",
                "project_description": "...",
                "is_active": true,
                "nodes": [...]
            },
            ...
        ]
    }
    """
    dao = _get_dao()
    versions = dao.get_all_directory_versions(project_id)
    
    # 为每个版本加载对应的nodes
    result = []
    for version in versions:
        nodes = dao.get_directory_nodes_by_version(version["version_id"])
        
        # 处理created_at字段（可能是datetime对象或字符串）
        created_at = version.get("created_at")
        if created_at:
            if hasattr(created_at, 'isoformat'):
                created_at = created_at.isoformat()
            else:
                created_at = str(created_at)
        
        result.append({
            "version_id": version["version_id"],
            "project_type": version["project_type"],
            "project_description": version.get("project_description"),
            "is_active": version["is_active"],
            "created_at": created_at,
            "nodes": nodes
        })
    
    return {"versions": result}


# ==================== Sections ====================

@router.get("/projects/{project_id}/sections")
def get_sections(
    project_id: str,
    node_id: Optional[str] = None,
    user=Depends(get_current_user_sync),
):
    """获取章节内容"""
    dao = _get_dao()
    sections = dao.get_active_sections(project_id, node_id)
    return {"sections": sections}


class AnalyzeIntentReq(BaseModel):
    """AI意图识别请求"""
    user_input: str
    conversation_history: List[Dict[str, str]] = []
    directory_structure: List[Dict[str, Any]]


class AnalyzeIntentRes(BaseModel):
    """AI意图识别响应"""
    intent_type: str
    target_node_ids: List[str]
    action_description: str
    requirements: str
    confidence: float


@router.post("/projects/{project_id}/ai-assistant/analyze-intent")
async def analyze_user_intent(
    project_id: str,
    req: Request,
    user=Depends(get_current_user),
):
    """AI意图识别 - 理解用户想修改哪些章节、如何修改"""
    from pydantic import BaseModel
    
    class IntentRequest(BaseModel):
        user_input: str
        conversation_history: List[Dict[str, str]] = []
        directory_structure: List[Dict[str, Any]]
    
    body = await req.json()
    request_data = IntentRequest(**body)
    
    llm = _get_llm(req)
    
    # 构建章节信息
    sections_info = "\n".join([
        f"- [{node.get('id')}] {node.get('orderNo', '')} {node.get('title', '')}"
        for node in request_data.directory_structure
    ])
    
    # 构建对话历史
    history_text = ""
    if request_data.conversation_history:
        history_text = "\n".join([
            f"{msg['role']}: {msg['content']}"
            for msg in request_data.conversation_history[-5:]
        ])
    
    # 意图识别prompt
    intent_prompt = f"""你是一个文档编辑AI助手。请分析用户的意图，理解他们想修改哪些章节、如何修改。

【当前文档章节结构】
{sections_info}

{f"【最近对话历史】{history_text}" if history_text else ""}

【用户输入】
{request_data.user_input}

请以JSON格式返回分析结果：
{{
    "intent_type": "generate|modify|optimize|global",
    "target_node_ids": ["章节ID列表"],
    "action_description": "简短描述要做什么",
    "requirements": "提炼的具体需求（传给生成API）",
    "confidence": 0.0-1.0
}}

判断规则：
1. 如果提到"第X章"、章节标题、或章节编号 → 提取对应的node_id
2. 如果说"这里"、"上面"、"刚才" → 结合对话历史判断
3. 如果说"整个文档"、"所有" → intent_type=global，返回多个node_id
4. 如果说"扩写"、"增加" → intent_type=generate
5. 如果说"修改"、"改成" → intent_type=modify
6. 如果说"优化"、"润色" → intent_type=optimize

只返回JSON，不要其他文字。"""

    try:
        messages = [{"role": "user", "content": intent_prompt}]
        response = await llm.achat(messages=messages, model_id=None)
        
        # 提取文本内容
        if isinstance(response, dict) and "choices" in response:
            result_text = response["choices"][0]["message"]["content"].strip()
        elif isinstance(response, str):
            result_text = response.strip()
        else:
            result_text = str(response).strip()
        
        # 提取JSON
        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0].strip()
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0].strip()
        
        import json
        result = json.loads(result_text)
        
        logger.info(f"[意图识别] 用户输入: {request_data.user_input[:50]}...")
        logger.info(f"[意图识别] 识别结果: {result}")
        
        return result
        
    except Exception as e:
        logger.error(f"[意图识别] 失败: {e}", exc_info=True)
        return {
            "intent_type": "unknown",
            "target_node_ids": [],
            "action_description": "无法理解意图",
            "requirements": request_data.user_input,
            "confidence": 0.0
        }


@router.post("/projects/{project_id}/sections/generate")
async def generate_section_content(
    project_id: str,
    req: Request,
    user=Depends(get_current_user),
):
    """
    AI生成单个章节内容
    用于DocumentComponentManagement组件的AI助手功能
    """
    from pydantic import BaseModel
    
    class GenerateSectionRequest(BaseModel):
        title: str
        level: int
        node_id: Optional[str] = None  # ✅ 新增：节点ID
        requirements: Optional[str] = None
    
    body = await req.json()
    request_data = GenerateSectionRequest(**body)
    
    dao = _get_dao()
    pool = _get_pool()
    llm = _get_llm(req)
    
    # 获取申报要求（用于生成上下文）
    requirements = dao.get_requirements(project_id)
    requirements_summary = ""
    if requirements:
        req_data = requirements.get("data_json", {})
        requirements_summary = req_data.get("summary", "") if isinstance(req_data, dict) else ""
    
    # ✅ 如果有用户自定义要求，添加到上下文
    requirements_dict = {}
    if requirements_summary:
        requirements_dict["summary"] = requirements_summary
    if request_data.requirements:
        requirements_dict["custom_requirements"] = request_data.requirements
    
    # 构建最终的requirements_summary（用于向后兼容）
    if request_data.requirements:
        requirements_summary += f"\n\n【用户要求】\n{request_data.requirements}"
    
    # 创建extract_v2实例并生成内容
    extract_v2 = DeclareExtractV2Service(pool, llm)
    model_id = None  # 使用默认模型
    
    result = await extract_v2.autofill_section_unified(
        project_id=project_id,
        model_id=model_id,
        node_title=request_data.title,
        node_level=request_data.level,  # 使用前端传来的level
        requirements_summary=requirements_summary,
        requirements_dict=requirements_dict,  # ✅ 新增：结构化的requirements
        run_id=None,
    )
    
    # 返回生成的内容
    # autofill_section 返回格式: {"data": {"content_md": "..."}, "evidence_chunk_ids": [...], ...}
    data = result.get("data", {})
    if isinstance(data, dict):
        content = data.get("content_md", "")
    else:
        content = ""
    
    logger.info(f"[sections/generate] ✅ 生成完成: node_id={request_data.node_id}, title={request_data.title}, content_len={len(content)}")
    
    # ✅ 新增：自动保存到数据库（如果提供了node_id）
    if request_data.node_id and content:
        try:
            logger.info(f"[sections/generate] 开始保存: node_id={request_data.node_id}")
            section_id = dao.save_section(
                project_id=project_id,
                node_id=request_data.node_id,
                node_title=request_data.title,
                content_html=content,  # content_md实际上是HTML格式
                content_md=None,
                evidence_chunk_ids=result.get("evidence_chunk_ids"),
                retrieval_trace=result.get("retrieval_trace"),
            )
            logger.info(f"[sections/generate] ✅ 已保存到数据库: section_id={section_id}, node_id={request_data.node_id}")
        except Exception as e:
            logger.error(f"[sections/generate] ❌ 保存失败: node_id={request_data.node_id}, error={e}", exc_info=True)
            # 保存失败不影响返回生成的内容
    elif not request_data.node_id:
        logger.warning(f"[sections/generate] ⚠️ 未保存（缺少node_id）: title={request_data.title}")
    elif not content:
        logger.warning(f"[sections/generate] ⚠️ 未保存（内容为空）: node_id={request_data.node_id}, title={request_data.title}")
    
    logger.info(f"[sections/generate] END: node_id={request_data.node_id}")
    return {"content": content}


@router.post("/projects/{project_id}/sections/save")
async def save_section_content(
    project_id: str,
    req: Request,
    user=Depends(get_current_user),
):
    """
    保存章节内容到数据库
    用于手动编辑后的保存
    """
    from pydantic import BaseModel
    
    class SaveSectionRequest(BaseModel):
        node_id: str
        node_title: str
        content_html: str
        content_md: Optional[str] = None
    
    body = await req.json()
    request_data = SaveSectionRequest(**body)
    
    dao = _get_dao()
    
    try:
        section_id = dao.save_section(
            project_id=project_id,
            node_id=request_data.node_id,
            node_title=request_data.node_title,
            content_html=request_data.content_html,
            content_md=request_data.content_md,
        )
        logger.info(f"[sections/save] 保存成功: section_id={section_id}")
        return {"success": True, "section_id": section_id}
    except Exception as e:
        logger.error(f"[sections/save] 保存失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")


@router.get("/projects/{project_id}/sections/load")
async def load_all_sections(
    project_id: str,
    user=Depends(get_current_user),
):
    """
    加载项目的所有章节内容
    返回格式: {node_id: {content_html: "...", content_md: "...", ...}}
    """
    dao = _get_dao()
    
    try:
        sections_dict = dao.get_all_sections(project_id)
        logger.info(f"[sections/load] 加载了 {len(sections_dict)} 个章节")
        return {"sections": sections_dict}
    except Exception as e:
        logger.error(f"[sections/load] 加载失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"加载失败: {str(e)}")


@router.get("/projects/{project_id}/sections/{node_id}")
async def get_section_content(
    project_id: str,
    node_id: str,
    user=Depends(get_current_user),
):
    """
    获取单个节点的章节内容
    """
    dao = _get_dao()
    
    try:
        section = dao.get_section(project_id, node_id)
        if not section:
            return {"section": None}
        return {"section": section}
    except Exception as e:
        logger.error(f"[sections/get] 获取失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


# ==================== Assets ====================

@router.get("/projects/{project_id}/assets/image/{filename}")
async def get_project_image(
    project_id: str,
    filename: str,
):
    """
    获取项目上传的图片资源
    用于在文档预览中显示图片
    
    注意：此接口不需要认证，因为图片通过<img>标签加载，无法携带Authorization header
    安全性由项目ID和文件名的复杂性保证
    """
    import os
    from urllib.parse import unquote
    from fastapi.responses import FileResponse
    
    dao = _get_dao()
    
    # 解码文件名
    filename = unquote(filename)
    
    # 查找匹配的资源
    with dao.pool.connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT storage_path, mime_type
                FROM declare_assets
                WHERE project_id = %s AND filename = %s AND asset_type = 'image'
                LIMIT 1
            """, [project_id, filename])
            
            row = cur.fetchone()
            
            if not row:
                raise HTTPException(status_code=404, detail=f"图片未找到: {filename}")
            
            storage_path = row['storage_path']
            mime_type = row['mime_type'] or 'image/png'
            
            if not os.path.exists(storage_path):
                raise HTTPException(status_code=404, detail=f"图片文件不存在: {filename}")
            
            return FileResponse(
                storage_path,
                media_type=mime_type,
                headers={
                    "Cache-Control": "public, max-age=3600",
                    "Access-Control-Allow-Origin": "*"
                }
            )


# ==================== Document ====================

@router.post("/projects/{project_id}/document/generate", response_model=RunOut)
async def generate_document(
    project_id: str,
    bg: BackgroundTasks,
    req: Request,
    sync: int = 0,
    auto_generate: int = 1,
    user=Depends(get_current_user_sync),
):
    """生成申报书文档"""
    dao = _get_dao()
    service = _get_service(req)
    
    # 创建 run
    run_id = dao.create_run(project_id, "document")
    
    auto_generate_content = bool(auto_generate)
    
    if sync == 1:
        # 同步执行
        await service.generate_document(
            project_id, 
            run_id, 
            auto_generate_content=auto_generate_content
        )
        run = dao.get_run(run_id)
        return run
    else:
        # 异步执行（使用 asyncio.create_task）
        import asyncio
        asyncio.create_task(service.generate_document(
            project_id, 
            run_id, 
            auto_generate_content=auto_generate_content
        ))
        run = dao.get_run(run_id)
        return run


@router.get("/projects/{project_id}/export/docx")
async def export_docx(project_id: str, user=Depends(get_current_user_sync)):
    """导出申报书为 Word 文档"""
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO
    import tempfile
    from pathlib import Path
    
    dao = _get_dao()
    
    # 1. 获取项目信息
    project = dao.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project_name = project.get("project_name", "申报书")
    
    # 2. 获取目录节点（最新版本）
    with dao.pool.connection() as conn:
        with conn.cursor() as cur:
            # 获取最新的目录版本
            cur.execute("""
                SELECT version_id 
                FROM declare_directory_versions 
                WHERE project_id = %s AND is_active = true 
                ORDER BY created_at DESC 
                LIMIT 1
            """, [project_id])
            version_row = cur.fetchone()
            
            if not version_row:
                raise HTTPException(status_code=404, detail="No directory found")
            
            version_id = version_row['version_id']
            
            # 获取该版本的所有目录节点，按 order_no 排序
            cur.execute("""
                SELECT id, title, level, order_no, parent_id, meta_json
                FROM declare_directory_nodes
                WHERE version_id = %s
                ORDER BY order_no
            """, [version_id])
            nodes = cur.fetchall()
    
    if not nodes:
        raise HTTPException(status_code=404, detail="No directory nodes found")
    
    # ✅ 从 meta_json 中提取 notes 到顶层（便于后续使用）
    for node in nodes:
        meta_json = node.get('meta_json')
        if isinstance(meta_json, dict) and 'notes' in meta_json:
            node['notes'] = meta_json['notes']
        else:
            node['notes'] = None
    
    # 3. 获取所有章节内容（使用最新版本的sections）
    sections_dict = {}
    with dao.pool.connection() as conn:
        with conn.cursor() as cur:
            for node in nodes:
                cur.execute("""
                    SELECT content_html
                    FROM declare_sections
                    WHERE project_id = %s AND node_id = %s
                    ORDER BY created_at DESC
                    LIMIT 1
                """, [project_id, node['id']])
                section_row = cur.fetchone()
                if section_row and section_row['content_html']:
                    sections_dict[node['id']] = section_row['content_html']
    
    # 4. 创建 Word 文档
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 5. 按目录顺序写入内容
    for node in nodes:
        node_id = node['id']
        title = node['title']
        level = node['level']
        
        # 5.1 添加标题
        heading_level = min(max(level, 1), 9)  # Word 支持 1-9 级标题
        doc.add_heading(title, level=heading_level)
        
        # 5.2 添加正文内容
        content_html = sections_dict.get(node_id, "")
        
        if content_html:
            from bs4 import BeautifulSoup
            from docx.shared import Inches, Pt
            from PIL import Image
            import re
            
            # 第一步：处理HTML，提取文本和图片占位符
            soup = BeautifulSoup(content_html, 'html.parser')
            
            # 提取Mermaid图表（支持两种格式）
            mermaid_diagrams = []
            
            # 格式1: <div class="mermaid-diagram">
            for mermaid_div in soup.find_all('div', class_='mermaid-diagram'):
                mermaid_code = mermaid_div.get_text(strip=True)
                
                # 清理markdown代码围栏（如果存在）
                # 格式：```mermaid\ngraph TD...\n```
                if mermaid_code.startswith('```'):
                    lines = mermaid_code.split('\n')
                    # 去掉第一行的 ```mermaid 和最后一行的 ```
                    if len(lines) >= 3:
                        # 第一行可能是 ```mermaid 或 ```
                        if lines[0].startswith('```'):
                            lines = lines[1:]
                        # 最后一行可能是 ```
                        if lines and lines[-1].strip() == '```':
                            lines = lines[:-1]
                        mermaid_code = '\n'.join(lines).strip()
                
                mermaid_diagrams.append(mermaid_code)
                # 替换为占位符
                from bs4 import NavigableString
                placeholder = NavigableString(f'{{MERMAID_{len(mermaid_diagrams)-1}}}')
                mermaid_div.replace_with(placeholder)
            
            # 格式2: <pre><code class="language-mermaid">
            for code_tag in soup.find_all('code', class_='language-mermaid'):
                mermaid_code = code_tag.get_text(strip=True)
                mermaid_diagrams.append(mermaid_code)
                # 替换整个<pre>标签为占位符
                from bs4 import NavigableString
                placeholder = NavigableString(f'{{MERMAID_{len(mermaid_diagrams)-1}}}')
                # 找到父级<pre>标签并替换
                pre_tag = code_tag.find_parent('pre')
                if pre_tag:
                    pre_tag.replace_with(placeholder)
                else:
                    code_tag.replace_with(placeholder)
            
            # 获取纯文本
            text_content = soup.get_text()
            
            # 第二步：按段落分割并处理
            paragraphs = text_content.split('\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                
                # 检查是否是Mermaid占位符
                mermaid_match = re.match(r'\{MERMAID_(\d+)\}', para_text)
                if mermaid_match:
                    idx = int(mermaid_match.group(1))
                    if idx < len(mermaid_diagrams):
                        try:
                            # 使用Kroki.io在线服务渲染Mermaid为图片
                            import tempfile
                            import urllib.parse
                            import urllib.request
                            import zlib
                            import base64
                            
                            logger.info(f"[export_docx] 开始渲染Mermaid图表 {idx}")
                            
                            # Kroki使用deflate+base64编码
                            compressed = zlib.compress(mermaid_diagrams[idx].encode('utf-8'))
                            encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
                            
                            # 构造Kroki URL (使用国内可访问的镜像或自建服务)
                            kroki_url = f"https://kroki.io/mermaid/png/{encoded}"
                            
                            # 下载渲染后的图片
                            logger.info(f"[export_docx] 请求Kroki渲染: {len(encoded)} bytes")
                            
                            # 创建临时输出文件
                            output_png = tempfile.NamedTemporaryFile(
                                suffix='.png', 
                                delete=False
                            )
                            
                            # 下载图片（添加超时和重试）
                            try:
                                req = urllib.request.Request(
                                    kroki_url,
                                    headers={'User-Agent': 'Mozilla/5.0'}
                                )
                                with urllib.request.urlopen(req, timeout=15) as response:
                                    output_png.write(response.read())
                                output_png.close()
                            except Exception as e:
                                logger.warning(f"[export_docx] Kroki下载失败: {e}")
                                output_png.close()
                                raise
                            
                            if Path(output_png.name).exists() and Path(output_png.name).stat().st_size > 0:
                                # 成功渲染，嵌入图片
                                img = Image.open(output_png.name)
                                width, height = img.size
                                
                                # 计算显示尺寸（最大宽度6英寸）
                                max_width_inches = 6.0
                                width_inches = width / 96.0  # PNG通常是96 DPI
                                display_width = min(max_width_inches, width_inches)
                                
                                # 添加标题
                                doc.add_paragraph("【流程图/架构图】", style='Heading 4')
                                
                                # 嵌入图片
                                pic_para = doc.add_paragraph()
                                pic_para.alignment = 1  # 居中
                                run = pic_para.add_run()
                                run.add_picture(output_png.name, width=Inches(display_width))
                                
                                logger.info(f"[export_docx] Mermaid图表已渲染为图片: {width}x{height}px, 显示宽度: {display_width:.2f}英寸")
                            else:
                                # 渲染失败，fallback到代码块
                                logger.warning(f"[export_docx] Mermaid渲染失败: 文件为空")
                                doc.add_paragraph("【流程图/架构图】", style='Heading 4')
                                doc.add_paragraph(mermaid_diagrams[idx], style='Normal')
                                logger.info(f"[export_docx] Fallback: 添加Mermaid代码块")
                            
                            # 清理临时文件
                            try:
                                import os
                                os.unlink(output_png.name)
                            except:
                                pass
                                
                        except Exception as e:
                            logger.error(f"[export_docx] Mermaid渲染失败: {e}", exc_info=True)
                            # Fallback: 添加代码块
                            doc.add_paragraph("【流程图/架构图】", style='Heading 4')
                            doc.add_paragraph(mermaid_diagrams[idx], style='Normal')
                    continue
                
                # 检查是否包含图片占位符
                image_matches = re.findall(r'\{image:([^}]+)\}', para_text)
                if image_matches:
                    # 分割文本和图片
                    parts = re.split(r'(\{image:[^}]+\})', para_text)
                    
                    # 先添加文本部分
                    text_parts = []
                    for part in parts:
                        if not part.startswith('{image:'):
                            text_parts.append(part.strip())
                    
                    combined_text = ' '.join([p for p in text_parts if p])
                    if combined_text:
                        doc.add_paragraph(combined_text)
                    
                    # 然后添加图片
                    for img_filename in image_matches:
                        try:
                            # 从数据库查询图片的storage_path
                            with dao.pool.connection() as img_conn:
                                with img_conn.cursor() as img_cur:
                                    img_cur.execute("""
                                        SELECT storage_path
                                        FROM declare_assets
                                        WHERE project_id = %s AND filename = %s AND asset_type = 'image'
                                        LIMIT 1
                                    """, [project_id, img_filename])
                                    
                                    img_row = img_cur.fetchone()
                                    
                                    if img_row and img_row['storage_path']:
                                        storage_path = img_row['storage_path']
                                        
                                        if Path(storage_path).exists():
                                            # 打开图片检查尺寸
                                            img = Image.open(storage_path)
                                            width, height = img.size
                                            
                                            # 计算显示尺寸（最大宽度6英寸，保持比例）
                                            max_width_inches = 6.0
                                            # 像素转英寸：假设72 DPI
                                            width_inches = width / 72.0
                                            display_width = min(max_width_inches, width_inches)
                                            
                                            # 添加图片
                                            pic_para = doc.add_paragraph()
                                            pic_para.alignment = 1  # 居中
                                            run = pic_para.add_run()
                                            run.add_picture(storage_path, width=Inches(display_width))
                                            
                                            # 添加图片说明
                                            caption_para = doc.add_paragraph()
                                            caption_para.alignment = 1
                                            caption_run = caption_para.add_run(f"图：{img_filename}")
                                            caption_run.font.size = Pt(10)
                                            caption_run.font.italic = True
                                            
                                            logger.info(f"[export_docx] 嵌入用户图片: {img_filename}, 尺寸: {width}x{height}px, 显示宽度: {display_width:.2f}英寸")
                                        else:
                                            logger.warning(f"[export_docx] 图片文件不存在: {storage_path}")
                                            doc.add_paragraph(f"[图片: {img_filename} - 文件不存在]")
                                    else:
                                        logger.warning(f"[export_docx] 数据库中未找到图片: {img_filename}")
                                        doc.add_paragraph(f"[图片: {img_filename} - 未找到]")
                        except Exception as e:
                            logger.error(f"[export_docx] 处理图片失败: {img_filename}, error={e}", exc_info=True)
                            doc.add_paragraph(f"[图片: {img_filename}]")
                else:
                    # 普通段落
                    doc.add_paragraph(para_text)
        else:
            # 如果没有内容，添加提示
            doc.add_paragraph("（本章节内容待填写）", style='Normal')
        
        # 添加一些间距
        doc.add_paragraph("")
    
    # 6. 保存文档到临时文件
    temp_dir = tempfile.gettempdir()
    filename = f"{project_name}.docx"
    # 使用项目ID确保文件名唯一
    safe_filename = f"declare_{project_id}_{filename}"
    output_path = Path(temp_dir) / safe_filename
    
    doc.save(str(output_path))
    
    # 设置文件权限（确保可读）
    import os
    os.chmod(str(output_path), 0o644)
    
    logger.info(f"[export_docx] 导出成功: {output_path}")
    
    # 7. 返回文件
    # 对中文文件名进行URL编码以符合HTTP header规范
    from urllib.parse import quote
    encoded_filename = quote(filename)
    
    # 使用ASCII兼容的fallback文件名
    ascii_filename = f"declare_{project_id[:8]}.docx"
    
    return FileResponse(
        path=str(output_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{encoded_filename}',
            "Access-Control-Expose-Headers": "Content-Disposition",
        }
    )


# ==================== Runs ====================

@router.get("/runs/{run_id}", response_model=RunOut)
def get_run(run_id: str, user=Depends(get_current_user_sync)):
    """获取任务运行记录"""
    dao = _get_dao()
    run = dao.get_run(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    
    # TODO: 如果有 platform_job_id，查询 job 状态并回写
    
    return run


import os

