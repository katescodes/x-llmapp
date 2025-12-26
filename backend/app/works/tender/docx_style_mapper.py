"""
DOCX 导出增强 - 模板样式映射验证

确保导出的 DOCX 严格按照模板样式渲染
"""
import logging
from typing import Dict, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def validate_template_styles(template_path: str) -> Dict[int, str]:
    """
    验证模板中的样式并返回层级映射
    
    Args:
        template_path: 模板文件路径
    
    Returns:
        {level: style_name} 映射，例如：
        {
            1: "Heading 1",
            2: "Heading 2",
            3: "Heading 3",
            ...
        }
    """
    try:
        doc = Document(template_path)
        styles = doc.styles
        
        heading_map = {}
        
        # 检测可用的标题样式
        for style in styles:
            if style.name.startswith("Heading"):
                try:
                    level = int(style.name.split()[-1])
                    heading_map[level] = style.name
                except ValueError:
                    continue
        
        logger.info(f"Template styles detected: {heading_map}")
        
        # 如果没有找到标题样式，使用默认值
        if not heading_map:
            heading_map = {
                1: "Heading 1",
                2: "Heading 2",
                3: "Heading 3",
                4: "Heading 4",
                5: "Heading 5",
                6: "Heading 6"
            }
            logger.warning(f"No heading styles found in template, using defaults")
        
        return heading_map
    
    except Exception as e:
        logger.error(f"Failed to validate template styles: {e}")
        # 返回默认映射
        return {
            1: "Heading 1",
            2: "Heading 2",
            3: "Heading 3",
            4: "Heading 4",
            5: "Heading 5",
            6: "Heading 6"
        }


def insert_toc_field(doc: Document) -> None:
    """
    在文档开头插入可更新的目录（TOC）field
    
    用户在 Word 中按 F9 即可更新页码
    
    Args:
        doc: Document 对象
    """
    try:
        # 在第一个段落前插入目录
        if len(doc.paragraphs) == 0:
            # 如果文档为空，先添加一个段落
            doc.add_paragraph()
        
        # 创建目录标题段落
        toc_title = doc.paragraphs[0].insert_paragraph_before("目录")
        toc_title.style = doc.styles.get("Heading 1", "Heading 1")
        
        # 创建 TOC field 段落
        toc_para = doc.paragraphs[1].insert_paragraph_before()
        
        # 构建 TOC field XML
        # <w:fldSimple w:instr="TOC \o &quot;1-3&quot; \h \z \u">
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), r'TOC \o "1-6" \h \z \u')
        
        # 添加占位文本（在未更新前显示）
        run = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = "【右键点击此处，选择"更新域"以生成目录】"
        run.append(t)
        fld_simple.append(run)
        
        # 插入到段落
        toc_para._element.append(fld_simple)
        
        logger.info("TOC field inserted successfully")
    
    except Exception as e:
        logger.error(f"Failed to insert TOC field: {e}")


def apply_heading_style(paragraph, level: int, heading_map: Dict[int, str]) -> None:
    """
    应用标题样式
    
    Args:
        paragraph: Paragraph 对象
        level: 层级（1-6）
        heading_map: 样式映射
    """
    # 确保 level 在有效范围内
    level = max(1, min(6, level))
    
    # 获取对应的样式名称
    style_name = heading_map.get(level, f"Heading {level}")
    
    try:
        paragraph.style = style_name
    except KeyError:
        logger.warning(f"Style '{style_name}' not found, using default")
        paragraph.style = f"Heading {level}"


def ensure_template_compatibility(template_path: str) -> Dict[str, any]:
    """
    确保模板兼容性，返回配置信息
    
    Args:
        template_path: 模板文件路径
    
    Returns:
        {
            "heading_map": {1: "Heading 1", ...},
            "has_toc": True/False,
            "normal_style": "Normal" | "正文" | ...
        }
    """
    try:
        doc = Document(template_path)
        styles = doc.styles
        
        heading_map = validate_template_styles(template_path)
        
        # 检测正文样式
        normal_style = "Normal"
        for style in styles:
            if style.name in ["Normal", "正文", "Body Text"]:
                normal_style = style.name
                break
        
        # 检测是否已有 TOC
        has_toc = any("TOC" in p.text for p in doc.paragraphs if p.text)
        
        config = {
            "heading_map": heading_map,
            "has_toc": has_toc,
            "normal_style": normal_style
        }
        
        logger.info(f"Template compatibility config: {config}")
        
        return config
    
    except Exception as e:
        logger.error(f"Failed to ensure template compatibility: {e}")
        return {
            "heading_map": {i: f"Heading {i}" for i in range(1, 7)},
            "has_toc": False,
            "normal_style": "Normal"
        }

