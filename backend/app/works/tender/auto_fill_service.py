"""
AutoFillService - 范本自动填充服务
整合字段提取和占位符检测，实现自动填充功能
"""
import logging
from typing import Dict, Optional, List
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from io import BytesIO

from .field_extractor import TenderFieldExtractor, FIELD_DISPLAY_NAMES
from .placeholder_detector import PlaceholderDetector, get_placeholder_detector

logger = logging.getLogger(__name__)


class AutoFillService:
    """范本自动填充服务"""
    
    def __init__(self):
        self.field_extractor = TenderFieldExtractor()
        self.placeholder_detector = get_placeholder_detector()
    
    def auto_fill_snippet(
        self, 
        project_id: str,
        snippet_docx_bytes: bytes,
        extracted_fields: Optional[Dict[str, str]] = None
    ) -> bytes:
        """
        自动填充范本文档
        
        Args:
            project_id: 项目ID
            snippet_docx_bytes: 范本docx文件的字节内容
            extracted_fields: 已提取的字段（如果为None则从项目提取）
            
        Returns:
            填充后的docx文件字节内容
        """
        try:
            # 1. 如果没有提供字段，则从项目提取
            if extracted_fields is None:
                logger.info(f"从项目 {project_id} 提取字段...")
                extracted_fields = self.field_extractor.extract_from_project(project_id)
            
            if not extracted_fields:
                logger.warning("未提取到任何字段，跳过自动填充")
                return snippet_docx_bytes
            
            logger.info(f"提取到 {len(extracted_fields)} 个字段: {list(extracted_fields.keys())}")
            
            # 2. 加载范本文档
            doc = Document(BytesIO(snippet_docx_bytes))
            
            # 3. 遍历所有段落，检测并填充占位符
            filled_count = 0
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                if not original_text or len(original_text) < 3:
                    continue
                
                # 检测占位符
                placeholders = self.placeholder_detector.detect_placeholders(original_text)
                
                if placeholders:
                    # 填充占位符
                    filled_text = self.placeholder_detector.fill_placeholders(
                        original_text, 
                        extracted_fields, 
                        placeholders
                    )
                    
                    if filled_text != original_text:
                        # 保留原有格式，逐run替换
                        self._replace_paragraph_text(paragraph, original_text, filled_text)
                        filled_count += 1
            
            # 4. 遍历所有表格，检测并填充占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            
                            if not original_text or len(original_text) < 3:
                                continue
                            
                            # 检测占位符
                            placeholders = self.placeholder_detector.detect_placeholders(original_text)
                            
                            if placeholders:
                                # 填充占位符
                                filled_text = self.placeholder_detector.fill_placeholders(
                                    original_text, 
                                    extracted_fields, 
                                    placeholders
                                )
                                
                                if filled_text != original_text:
                                    self._replace_paragraph_text(paragraph, original_text, filled_text)
                                    filled_count += 1
            
            logger.info(f"✅ 自动填充完成，共处理 {filled_count} 个段落/单元格")
            
            # 5. 保存填充后的文档
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            return output_buffer.read()
        
        except Exception as e:
            logger.error(f"自动填充失败: {e}", exc_info=True)
            # 失败时返回原文档
            return snippet_docx_bytes
    
    def _replace_paragraph_text(
        self, 
        paragraph, 
        old_text: str, 
        new_text: str
    ):
        """
        替换段落文本，尽量保留原有格式
        
        Args:
            paragraph: docx段落对象
            old_text: 原文本
            new_text: 新文本
        """
        # 简单策略：清空所有run，创建一个新run
        # （更复杂的策略是保留每个run的格式，但实现复杂）
        
        # 保存第一个run的格式（如果有）
        first_run_format = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
            }
        
        # 清空所有run
        for run in paragraph.runs:
            run.text = ""
        
        # 创建新run
        if paragraph.runs:
            new_run = paragraph.runs[0]
        else:
            new_run = paragraph.add_run()
        
        new_run.text = new_text
        
        # 应用格式
        if first_run_format:
            if first_run_format['font_name']:
                new_run.font.name = first_run_format['font_name']
            if first_run_format['font_size']:
                new_run.font.size = first_run_format['font_size']
            if first_run_format['bold']:
                new_run.bold = True
            if first_run_format['italic']:
                new_run.italic = True
            if first_run_format['underline']:
                new_run.underline = True
    
    def extract_project_fields(self, project_id: str) -> Dict[str, str]:
        """
        从项目中提取所有字段（供前端预览或调试）
        
        Args:
            project_id: 项目ID
            
        Returns:
            提取的字段字典
        """
        return self.field_extractor.extract_from_project(project_id)
    
    def get_fill_summary(
        self, 
        project_id: str,
        snippet_text: str
    ) -> Dict[str, any]:
        """
        获取自动填充摘要（不实际填充，只返回统计信息）
        
        Args:
            project_id: 项目ID
            snippet_text: 范本文本
            
        Returns:
            {
                "total_placeholders": 5,
                "can_fill": 3,
                "cannot_fill": 2,
                "extracted_fields": {...},
                "placeholders": [...]
            }
        """
        try:
            # 提取字段
            extracted_fields = self.field_extractor.extract_from_project(project_id)
            
            # 检测占位符
            placeholders = self.placeholder_detector.detect_placeholders(snippet_text)
            
            # 统计
            can_fill = 0
            cannot_fill = 0
            
            for ph in placeholders:
                field_name = ph["field_name"]
                if extracted_fields.get(field_name):
                    can_fill += 1
                else:
                    cannot_fill += 1
            
            return {
                "total_placeholders": len(placeholders),
                "can_fill": can_fill,
                "cannot_fill": cannot_fill,
                "extracted_fields": extracted_fields,
                "placeholders": placeholders,
            }
        
        except Exception as e:
            logger.error(f"获取填充摘要失败: {e}", exc_info=True)
            return {
                "total_placeholders": 0,
                "can_fill": 0,
                "cannot_fill": 0,
                "extracted_fields": {},
                "placeholders": [],
            }


# 全局实例
_service_instance = None

def get_auto_fill_service() -> AutoFillService:
    """获取全局自动填充服务实例"""
    global _service_instance
    if _service_instance is None:
        _service_instance = AutoFillService()
    return _service_instance

