#!/usr/bin/env python3
"""
测试脚本：验证 fragment_extractor 逻辑
直接在容器内运行，测试核心逻辑
"""
import sys
import os

# 添加应用路径
sys.path.insert(0, '/app')

def test_extract_body_items():
    """测试 _extract_body_items 的索引生成"""
    print("=" * 60)
    print("测试1: _extract_body_items() 索引生成")
    print("=" * 60)
    
    from docx import Document
    from io import BytesIO
    from app.services.fragment.fragment_extractor import TenderSampleFragmentExtractor
    from app.services.dao.tender_dao import TenderDAO
    
    # 创建测试文档
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("第三段")
    doc.add_table(2, 2)
    doc.add_paragraph("第五段")
    
    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # 重新加载
    doc = Document(buffer)
    
    # 测试 extractor
    class MockDAO:
        pass
    
    extractor = TenderSampleFragmentExtractor(MockDAO())
    body_items = extractor._extract_body_items(doc)
    
    print(f"body_items 数量: {len(body_items)}")
    for i, item in enumerate(body_items):
        print(f"  [{i}] bodyIndex={item.get('bodyIndex')} type={item.get('type')} text={item.get('text', '')[:20]}")
    
    # 验证索引连续性
    indices = [item.get('bodyIndex') for item in body_items]
    expected = list(range(len(body_items)))
    
    if indices == expected:
        print("\n✓ 索引连续且正确")
        return True
    else:
        print(f"\n✗ 索引不连续！")
        print(f"  预期: {expected}")
        print(f"  实际: {indices}")
        return False

def test_identify_fragments_logic():
    """测试 _identify_fragments_by_rules 的基本逻辑"""
    print("\n" + "=" * 60)
    print("测试2: _identify_fragments_by_rules() 基本逻辑")
    print("=" * 60)
    
    from app.services.fragment.fragment_extractor import TenderSampleFragmentExtractor
    
    class MockDAO:
        pass
    
    extractor = TenderSampleFragmentExtractor(MockDAO())
    
    # 模拟 body_items
    body_items = [
        {"bodyIndex": 0, "type": "paragraph", "text": "第一章 项目概述"},
        {"bodyIndex": 1, "type": "paragraph", "text": "投标文件格式"},  # 候选起点
        {"bodyIndex": 2, "type": "paragraph", "text": "一、投标函"},
        {"bodyIndex": 3, "type": "table", "text": ""},
        {"bodyIndex": 4, "type": "paragraph", "text": "二、开标一览表"},
        {"bodyIndex": 5, "type": "table", "text": ""},
        {"bodyIndex": 6, "type": "paragraph", "text": "三、授权委托书"},
        {"bodyIndex": 7, "type": "paragraph", "text": "第二章 评标办法"},  # 结束点
    ]
    
    try:
        fragments = extractor._identify_fragments_by_rules(body_items)
        print(f"\n检测到 {len(fragments)} 个片段:")
        
        for frag in fragments:
            print(f"\n  标题: {frag.get('title')}")
            print(f"  norm_key: {frag.get('norm_key')}")
            print(f"  范围: {frag.get('start_body_index')} -> {frag.get('end_body_index')}")
            print(f"  策略: {frag.get('strategy')}")
        
        # 检查诊断信息
        diag = getattr(extractor, '_last_rules_diag', {})
        print(f"\n诊断信息:")
        print(f"  candidates: {diag.get('candidates')}")
        print(f"  chosen_start: {diag.get('chosen_start')}")
        print(f"  chosen_end: {diag.get('chosen_end')}")
        print(f"  best_score: {diag.get('best_score')}")
        print(f"  heads_found: {diag.get('heads_found')}")
        
        if len(fragments) > 0:
            print("\n✓ 规则逻辑正常工作")
            return True
        else:
            print("\n✗ 未检测到任何片段")
            return False
    except Exception as e:
        print(f"\n✗ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_sample_data():
    """检查数据库中是否有实际的样本数据"""
    print("\n" + "=" * 60)
    print("测试3: 检查数据库中的样本数据")
    print("=" * 60)
    
    try:
        from app.services.dao.tender_dao import TenderDAO
        from app.db import get_db_pool
        
        pool = get_db_pool()
        dao = TenderDAO(pool)
        
        # 查询最近的一个项目
        with pool.connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT id, title, created_at 
                    FROM tender_project 
                    ORDER BY created_at DESC 
                    LIMIT 1
                """)
                project = cur.fetchone()
                
                if not project:
                    print("✗ 数据库中没有项目")
                    return False
                
                project_id = project[0]
                print(f"\n最新项目: {project[1]}")
                print(f"项目ID: {project_id}")
                
                # 查询 fragments
                fragments = dao.list_fragments("PROJECT", project_id)
                print(f"\nfragments 数量: {len(fragments)}")
                
                if len(fragments) > 0:
                    print("\n前3个 fragments:")
                    for frag in fragments[:3]:
                        print(f"\n  ID: {frag.get('id')}")
                        print(f"  标题: {frag.get('title')}")
                        print(f"  类型: {frag.get('fragment_type')}")
                        print(f"  范围: {frag.get('start_body_index')} -> {frag.get('end_body_index')}")
                        print(f"  置信度: {frag.get('confidence')}")
                        
                        # 检查诊断信息
                        diag_json = frag.get('diagnostics_json')
                        if diag_json:
                            import json
                            diag = json.loads(diag_json)
                            print(f"  诊断: {diag.get('reason')}")
                    
                    print("\n✓ 数据库中有 fragments 数据")
                else:
                    print("\n⚠ 数据库中 fragments 为空（可能是真的提取了0个）")
                
                return True
                
    except Exception as e:
        print(f"\n✗ 数据库查询失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("开始测试 fragment_extractor 逻辑...\n")
    
    results = []
    
    # 测试1: 索引生成
    try:
        results.append(("索引生成", test_extract_body_items()))
    except Exception as e:
        print(f"测试1失败: {e}")
        results.append(("索引生成", False))
    
    # 测试2: 规则逻辑
    try:
        results.append(("规则逻辑", test_identify_fragments_logic()))
    except Exception as e:
        print(f"测试2失败: {e}")
        results.append(("规则逻辑", False))
    
    # 测试3: 数据库数据
    try:
        results.append(("数据库数据", check_sample_data()))
    except Exception as e:
        print(f"测试3失败: {e}")
        results.append(("数据库数据", False))
    
    # 汇总结果
    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    for name, passed in results:
        status = "✓" if passed else "✗"
        print(f"{status} {name}")
    
    all_passed = all(r[1] for r in results)
    if all_passed:
        print("\n✓ 所有测试通过")
    else:
        print("\n✗ 部分测试失败")
    
    return 0 if all_passed else 1

if __name__ == "__main__":
    sys.exit(main())

