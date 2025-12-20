from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.services.template.docx_extractor import DocxBlockExtractor
from app.services.template.style_hints_fallback import build_style_hints_fallback


def _build_docx_bytes() -> bytes:
    doc = Document()

    # 自定义样式：+标题1~+标题5（base_style 指向 Heading 1..5）
    for lvl in range(1, 6):
        s = doc.styles.add_style(f"+标题{lvl}", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = doc.styles[f"Heading {lvl}"]
        doc.add_paragraph(f"这是标题{lvl}", style=s)

    # 正文样式：++正文
    body = doc.styles.add_style("++正文", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = doc.styles["Normal"]
    for i in range(8):
        doc.add_paragraph(f"这是正文段落 {i}", style=body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_does_not_crash_without_outlineLvl_and_fallback_style_hints():
    docx_bytes = _build_docx_bytes()

    extractor = DocxBlockExtractor()
    extract_result = extractor.extract(docx_bytes, max_blocks=200, max_chars_per_block=200)

    # 断言：extract 不抛异常且有块
    assert extract_result.blocks

    # 断言：fallback style_hints 能识别 +标题1 和 ++正文
    hints = build_style_hints_fallback(extract_result)
    assert hints.heading1_style is not None and "标题1" in hints.heading1_style
    assert hints.body_style is not None and "正文" in hints.body_style


