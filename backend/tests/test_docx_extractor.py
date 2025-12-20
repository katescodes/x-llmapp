"""
DocxBlockExtractor 单元测试
"""
import pytest
from io import BytesIO
from docx import Document
from app.services.template.docx_extractor import (
    DocxBlockExtractor,
    BlockType
)


@pytest.fixture
def sample_docx_bytes():
    """创建测试用的 docx 字节内容"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是第一章的内容")
    
    doc.add_heading("1.1 背景", level=2)
    doc.add_paragraph("背景描述")
    
    # 添加表格
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "列2"
    table.cell(0, 2).text = "列3"
    table.cell(1, 0).text = "数据1"
    table.cell(1, 1).text = "数据2"
    table.cell(1, 2).text = "数据3"
    
    # 保存到字节流
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_extractor_basic(sample_docx_bytes):
    """测试基本提取功能"""
    extractor = DocxBlockExtractor()
    result = extractor.extract(sample_docx_bytes, max_blocks=100, max_chars_per_block=300)
    
    # 验证提取到了块
    assert len(result.blocks) > 0
    
    # 验证有段落块
    paragraphs = [b for b in result.blocks if b.type == BlockType.PARAGRAPH]
    assert len(paragraphs) > 0
    
    # 验证有表格块
    tables = [b for b in result.blocks if b.type == BlockType.TABLE]
    assert len(tables) > 0
    
    # 验证样式统计
    assert result.style_stats["total_blocks"] == len(result.blocks)
    assert len(result.style_stats["heading_styles"]) > 0


def test_docx_extractor_max_blocks():
    """测试 max_blocks 限制"""
    doc = Document()
    
    # 添加很多段落
    for i in range(100):
        doc.add_paragraph(f"段落 {i}")
    
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    
    extractor = DocxBlockExtractor()
    result = extractor.extract(docx_bytes, max_blocks=20, max_chars_per_block=300)
    
    # 验证块数量被限制
    assert len(result.blocks) <= 20


def test_docx_extractor_max_chars():
    """测试 max_chars_per_block 限制"""
    doc = Document()
    
    # 添加一个很长的段落
    long_text = "这是一个非常长的段落。" * 100
    doc.add_paragraph(long_text)
    
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    
    extractor = DocxBlockExtractor()
    result = extractor.extract(docx_bytes, max_blocks=100, max_chars_per_block=50)
    
    # 验证文本被截断
    if len(result.blocks) > 0:
        first_block = result.blocks[0]
        assert len(first_block.text) <= 50


def test_docx_extractor_style_stats():
    """测试样式统计"""
    doc = Document()
    
    doc.add_heading("标题1", level=1)
    doc.add_heading("标题2", level=2)
    doc.add_paragraph("正文内容")
    
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    
    extractor = DocxBlockExtractor()
    result = extractor.extract(docx_bytes)
    
    # 验证样式统计
    assert "style_count" in result.style_stats
    assert "heading_styles" in result.style_stats
    assert len(result.style_stats["heading_styles"]) > 0


def test_docx_extractor_table_meta():
    """测试表格元数据提取"""
    doc = Document()
    
    # 添加表格
    table = doc.add_table(rows=5, cols=3)
    table.cell(0, 0).text = "表头1"
    table.cell(0, 1).text = "表头2"
    table.cell(0, 2).text = "表头3"
    
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    
    extractor = DocxBlockExtractor()
    result = extractor.extract(docx_bytes)
    
    # 查找表格块
    tables = [b for b in result.blocks if b.type == BlockType.TABLE]
    assert len(tables) > 0
    
    table_block = tables[0]
    assert table_block.table_meta is not None
    assert table_block.table_meta["rows"] == 5
    assert table_block.table_meta["cols"] == 3


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
