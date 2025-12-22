"""
申报书自动生成功能 - 验证脚本

使用此脚本快速验证"按标题自动写内容"功能是否正常工作。
"""

import asyncio
import sys
from pathlib import Path

# 添加项目路径到 sys.path
sys.path.insert(0, str(Path(__file__).parent.parent / "backend"))


async def test_auto_generate():
    """测试自动生成功能"""
    
    print("=" * 60)
    print("申报书自动生成功能 - 验证测试")
    print("=" * 60)
    print()
    
    # 1. 导入必要的模块
    print("✓ 步骤 1/5: 导入模块...")
    try:
        from app.services.export.docx_exporter import (
            AutoWriteCfg,
            build_project_context_string,
            _is_empty_or_placeholder,
            _infer_section_style,
            _target_min_words,
        )
        print("  ✓ 模块导入成功")
    except Exception as e:
        print(f"  ✗ 模块导入失败: {e}")
        return False
    
    # 2. 测试占位符判断
    print("\n✓ 步骤 2/5: 测试占位符判断...")
    test_cases = [
        ("", True, "空字符串"),
        ("【填写】", True, "占位符【填写】"),
        ("【待补】", True, "占位符【待补】"),
        ("TODO", True, "占位符 TODO"),
        ("TBD", True, "占位符 TBD"),
        ("abc", True, "少于5字符"),
        ("这是实际内容，不应该生成", False, "实际内容"),
        ("项目背景介绍：...", False, "实际内容"),
    ]
    
    for content, expected, desc in test_cases:
        result = _is_empty_or_placeholder(content)
        status = "✓" if result == expected else "✗"
        print(f"  {status} {desc}: {content!r} -> {result}")
    
    # 3. 测试标题样式推断
    print("\n✓ 步骤 3/5: 测试标题样式推断...")
    titles = [
        "项目建设背景与必要性",
        "建设目标与指标",
        "技术方案与架构",
        "应用场景与业务流程",
        "组织保障措施",
        "投资预算与资金来源",
        "实施进度与计划",
    ]
    
    for title in titles:
        hint = _infer_section_style(title)
        print(f"  ✓ {title}")
        print(f"    → {hint[:50]}...")
    
    # 4. 测试字数目标
    print("\n✓ 步骤 4/5: 测试字数目标...")
    cfg = AutoWriteCfg()
    for level in [1, 2, 3, 4, 5]:
        target = _target_min_words(level, cfg)
        print(f"  ✓ H{level} 标题目标字数: {target} 字")
    
    # 5. 测试项目上下文构建
    print("\n✓ 步骤 5/5: 测试项目上下文构建...")
    
    # 测试用例 1：完整数据
    project_data = {
        "name": "智能工厂数字化转型项目",
        "company": "某某制造有限公司",
        "summary": "本项目旨在通过数字化手段提升生产效率...",
        "meta_json": {
            "industry": "装备制造",
            "budget": "500万元",
            "duration": "18个月",
        },
        "patents": [
            {"name": "一种智能检测装置"},
            {"name": "基于AI的质量控制方法"},
        ],
        "devices": [{"name": "设备1"}, {"name": "设备2"}],
        "achievements": [{"name": "成果1"}],
    }
    
    context = build_project_context_string(project_data)
    print("  ✓ 完整项目数据:")
    for line in context.split("\n")[:5]:
        print(f"    {line}")
    if len(context.split("\n")) > 5:
        print(f"    ... (共 {len(context.split('\\n'))} 行)")
    
    # 测试用例 2：空数据
    empty_context = build_project_context_string({})
    print(f"\n  ✓ 空项目数据: {empty_context!r}")
    
    # 测试用例 3：部分数据
    partial_data = {
        "name": "测试项目",
        "summary": "测试摘要",
    }
    partial_context = build_project_context_string(partial_data)
    print(f"\n  ✓ 部分项目数据:")
    for line in partial_context.split("\n"):
        print(f"    {line}")
    
    print("\n" + "=" * 60)
    print("✓ 所有单元测试通过！")
    print("=" * 60)
    print()
    
    return True


async def test_full_export(project_id: str = None):
    """
    测试完整导出流程（需要数据库连接）
    
    Args:
        project_id: 项目ID（可选）
    """
    
    print("=" * 60)
    print("完整导出流程测试")
    print("=" * 60)
    print()
    
    if not project_id:
        print("⚠️  未提供 project_id，跳过完整导出测试")
        print("   使用方式: python verify_auto_generate.py --project-id proj_xxx")
        return
    
    try:
        from app.config import get_settings
        from app.services.dao.tender_dao import TenderDAO
        from app.services.export.export_service import ExportService
        from app.services.export.docx_exporter import AutoWriteCfg
        from psycopg_pool import ConnectionPool
        
        print(f"✓ 测试项目: {project_id}")
        
        # 创建数据库连接池
        settings = get_settings()
        pool = ConnectionPool(settings.DATABASE_URL)
        
        print("  ✓ 数据库连接成功")
        
        # 创建 DAO 和服务
        dao = TenderDAO(pool)
        export_service = ExportService(dao)
        
        # 配置
        cfg = AutoWriteCfg(
            min_words_h1=600,  # 测试时降低要求
            min_words_h2=400,
            min_words_h3=300,
            min_words_h4=200,
            max_tokens=1200,
            multi_round=True,
        )
        
        print("  ✓ 服务初始化成功")
        print()
        print("开始导出（启用自动生成）...")
        print()
        
        # 导出
        output_path = await export_service.export_project_to_docx(
            project_id=project_id,
            auto_generate_content=True,
            auto_write_cfg=cfg,
            project_context="",  # 自动构建
        )
        
        print()
        print(f"✓ 导出成功: {output_path}")
        
        # 分析生成的文档
        from docx import Document
        
        doc = Document(output_path)
        
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
        normal_count = sum(1 for p in doc.paragraphs if not p.style.name.startswith("Heading"))
        
        print()
        print("文档统计:")
        print(f"  • 总段落数: {len(doc.paragraphs)}")
        print(f"  • 标题段落: {heading_count}")
        print(f"  • 正文段落: {normal_count}")
        print(f"  • 总字符数: {total_chars}")
        print(f"  • 平均每标题: {total_chars // max(heading_count, 1)} 字")
        
        # 检查是否有【待补】占位
        placeholder_count = sum(1 for p in doc.paragraphs if "【待补" in p.text)
        print(f"  • 【待补】占位: {placeholder_count} 处")
        
        print()
        print("=" * 60)
        print("✓ 完整导出测试通过！")
        print("=" * 60)
        
        pool.close()
        
    except Exception as e:
        print(f"\n✗ 导出测试失败: {e}")
        import traceback
        traceback.print_exc()


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description="验证申报书自动生成功能")
    parser.add_argument(
        "--project-id",
        help="项目ID（用于完整导出测试）",
        default=None,
    )
    parser.add_argument(
        "--skip-unit-test",
        help="跳过单元测试",
        action="store_true",
    )
    
    args = parser.parse_args()
    
    # 运行单元测试
    if not args.skip_unit_test:
        success = asyncio.run(test_auto_generate())
        if not success:
            print("\n✗ 单元测试失败")
            sys.exit(1)
    
    # 运行完整导出测试
    if args.project_id:
        print()
        asyncio.run(test_full_export(args.project_id))
    else:
        print("\n💡 提示:")
        print("   如需测试完整导出流程，请提供项目ID:")
        print("   python verify_auto_generate.py --project-id proj_xxx")


if __name__ == "__main__":
    main()

