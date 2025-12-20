#!/usr/bin/env python3
"""
E2E: 自动套用格式（真实项目 + 真实格式模板）

Steps
- 选择一个项目与一个已分析且 style_hints 合格的格式模板
- 调用 apply-format-template
- 再调用 directory/meta 确认持久化（applied_format_template_id + style_hints）
- 调用 export/docx 导出并保存 /tmp/out.docx
- 用 python-docx 校验：一级标题段落 style.name == spec.style_hints.heading1_style（至少对上一级）

Env
- BASE_URL: default http://localhost:8000
- PROJECT_ID: optional
- TEMPLATE_ID: optional
- AUTH_TOKEN: optional (Bearer)
"""

from __future__ import annotations

import os
import sys
from typing import Any, Dict, List, Optional, Tuple

import httpx
from docx import Document


def _env(name: str, default: Optional[str] = None) -> Optional[str]:
    v = os.getenv(name)
    if v is None:
        return default
    v = v.strip()
    return v or default


def _fail(msg: str) -> None:
    print("FAIL:", msg)
    raise SystemExit(1)


def _pass(msg: str) -> None:
    print("PASS:", msg)
    raise SystemExit(0)


def _headers(auth_token: Optional[str]) -> Dict[str, str]:
    h: Dict[str, str] = {"Accept": "application/json"}
    if auth_token:
        h["Authorization"] = f"Bearer {auth_token}"
    return h


def _json_get(client: httpx.Client, url: str, auth_token: Optional[str]) -> Any:
    r = client.get(url, headers=_headers(auth_token))
    if r.status_code in (401, 403):
        _fail(f"Auth failed for GET {url}: {r.status_code} {r.text} (set AUTH_TOKEN)")
    r.raise_for_status()
    return r.json()


def _json_post(client: httpx.Client, url: str, payload: Any, auth_token: Optional[str]) -> Any:
    r = client.post(url, headers=_headers(auth_token), json=payload)
    if r.status_code in (401, 403):
        _fail(f"Auth failed for POST {url}: {r.status_code} {r.text} (set AUTH_TOKEN)")
    r.raise_for_status()
    return r.json()


def _pick_project_id(client: httpx.Client, base_url: str, auth_token: Optional[str], env_project_id: Optional[str]) -> str:
    if env_project_id:
        return env_project_id
    rows = _json_get(client, f"{base_url}/api/apps/tender/projects", auth_token)
    if not isinstance(rows, list) or not rows:
        _fail("No projects found via GET /api/apps/tender/projects")
    pid = str(rows[0].get("id") or "")
    if not pid:
        _fail("First project has empty id")
    return pid


def _list_format_templates(client: httpx.Client, base_url: str, auth_token: Optional[str]) -> List[Dict[str, Any]]:
    rows = _json_get(client, f"{base_url}/api/apps/tender/format-templates", auth_token)
    if not isinstance(rows, list):
        return []
    out: List[Dict[str, Any]] = []
    for r in rows:
        if isinstance(r, dict) and r.get("id"):
            out.append(r)
    return out


def _get_template_spec(client: httpx.Client, base_url: str, auth_token: Optional[str], template_id: str) -> Dict[str, Any]:
    spec = _json_get(client, f"{base_url}/api/apps/tender/format-templates/{template_id}/spec", auth_token)
    if not isinstance(spec, dict):
        _fail(f"Template spec is not object: template_id={template_id}")
    return spec


def _get_template_analysis_summary(client: httpx.Client, base_url: str, auth_token: Optional[str], template_id: str) -> Dict[str, Any]:
    s = _json_get(client, f"{base_url}/api/apps/tender/format-templates/{template_id}/analysis-summary", auth_token)
    return s if isinstance(s, dict) else {}


def _style_hints_ok(spec: Dict[str, Any]) -> Tuple[bool, str]:
    hints = spec.get("style_hints") or {}
    if not isinstance(hints, dict):
        return False, "spec.style_hints not dict"
    # 最低要求：heading1_style 非空
    h1 = hints.get("heading1_style")
    if not isinstance(h1, str) or not h1.strip():
        return False, "spec.style_hints.heading1_style empty"
    return True, ""


def _pick_template_id(
    client: httpx.Client,
    base_url: str,
    auth_token: Optional[str],
    env_template_id: Optional[str],
) -> str:
    if env_template_id:
        return env_template_id

    templates = _list_format_templates(client, base_url, auth_token)
    if not templates:
        _fail("No format templates found via GET /api/apps/tender/format-templates")

    best: Optional[str] = None
    for t in templates:
        tid = str(t.get("id") or "")
        if not tid:
            continue
        summary = _get_template_analysis_summary(client, base_url, auth_token, tid)
        analyzed = bool(summary.get("analyzed", False))
        confidence = float(summary.get("confidence", 0) or 0)
        style_hints_count = int(summary.get("style_hints_count", 0) or 0)
        if analyzed and confidence > 0 and style_hints_count > 0:
            best = tid
            break

    if not best:
        _fail("No analyzed+confident format template found (need analyzed==true, confidence>0, style_hints_count>0)")
    return best


def _save_docx(client: httpx.Client, url: str, auth_token: Optional[str], out_path: str) -> None:
    r = client.get(url, headers=_headers(auth_token))
    if r.status_code in (401, 403):
        _fail(f"Auth failed for GET {url}: {r.status_code} {r.text} (set AUTH_TOKEN)")
    r.raise_for_status()
    ct = (r.headers.get("content-type") or "").lower()
    if "application/vnd.openxmlformats" not in ct and "application/octet-stream" not in ct:
        _fail(f"export/docx content-type unexpected: {ct}, body={r.text[:300]}")
    with open(out_path, "wb") as w:
        w.write(r.content)


def _doc_styles_top(doc: Document, n: int = 30) -> List[str]:
    names: List[str] = []
    for s in doc.styles:
        try:
            names.append(str(s.name))
        except Exception:
            continue
        if len(names) >= n:
            break
    return names


def _paragraph_style_dump(doc: Document, limit: int = 25) -> List[Tuple[str, str]]:
    out: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        try:
            sname = p.style.name if p.style else ""
        except Exception:
            sname = ""
        out.append((txt[:120], sname))
        if len(out) >= limit:
            break
    return out


def main() -> None:
    base_url = _env("BASE_URL", "http://localhost:8000")
    project_id = _env("PROJECT_ID")
    template_id = _env("TEMPLATE_ID")
    auth_token = _env("AUTH_TOKEN")
    username = _env("USERNAME", "admin")
    password = _env("PASSWORD", "admin123")
    auto_register = (_env("AUTO_REGISTER", "true") or "true").lower() == "true"

    out_path = "/tmp/out.docx"

    # trust_env=False: 避免宿主机环境变量里配置了 SOCKS 代理但缺少 socksio 依赖导致 httpx 初始化失败
    with httpx.Client(timeout=60.0, trust_env=False) as client:
        # 如果未显式提供 AUTH_TOKEN，尝试用默认管理员账号登录获取 token
        if not auth_token:
            try:
                r = client.post(f"{base_url}/api/auth/login", json={"username": username, "password": password})
                if r.status_code == 200:
                    data = r.json()
                    auth_token = str(data.get("access_token") or "").strip() or None
            except Exception:
                auth_token = None

        # 若默认账号不可用，尝试自动注册一个 customer 账号再登录
        if not auth_token and auto_register:
            try:
                reg_user = username
                reg_pass = password
                if reg_user == "admin":
                    reg_user = f"e2e_customer_{os.getpid()}"
                    reg_pass = "e2e_pass_123456"

                client.post(
                    f"{base_url}/api/auth/register",
                    json={"username": reg_user, "password": reg_pass, "role": "customer"},
                )
                r = client.post(f"{base_url}/api/auth/login", json={"username": reg_user, "password": reg_pass})
                if r.status_code == 200:
                    data = r.json()
                    auth_token = str(data.get("access_token") or "").strip() or None
            except Exception:
                auth_token = None
        if not auth_token:
            _fail("Not authenticated (set AUTH_TOKEN, or provide USERNAME/PASSWORD; default is admin/admin123)")

        pid = _pick_project_id(client, base_url, auth_token, project_id)
        tid = _pick_template_id(client, base_url, auth_token, template_id)

        spec = _get_template_spec(client, base_url, auth_token, tid)
        ok, reason = _style_hints_ok(spec)
        if not ok:
            _fail(f"模板分析不合格: template_id={tid}: {reason}")

        expected_h1 = str((spec.get("style_hints") or {}).get("heading1_style") or "").strip()

        # apply-format-template
        apply_resp = _json_post(
            client,
            f"{base_url}/api/apps/tender/projects/{pid}/directory/apply-format-template",
            {"format_template_id": tid},
            auth_token,
        )
        if not isinstance(apply_resp, dict):
            _fail("apply-format-template resp not object")
        if not apply_resp.get("nodes"):
            _fail("apply-format-template resp.nodes empty")
        if not apply_resp.get("style_hints"):
            _fail("apply-format-template resp.style_hints empty")
        if str(apply_resp.get("applied_format_template_id") or "") != tid:
            _fail(f"apply-format-template applied_format_template_id mismatch: got={apply_resp.get('applied_format_template_id')} expected={tid}")

        # verify persistence via directory/meta
        meta = _json_get(client, f"{base_url}/api/apps/tender/projects/{pid}/directory/meta", auth_token)
        if not isinstance(meta, dict):
            _fail("directory/meta resp not object")
        nodes = meta.get("nodes")
        if not isinstance(nodes, list) or not nodes:
            _fail("directory/meta resp.nodes empty")
        if not meta.get("style_hints"):
            _fail("directory/meta resp.style_hints empty")
        if str(meta.get("applied_format_template_id") or "") != tid:
            _fail(f"directory/meta applied_format_template_id mismatch: got={meta.get('applied_format_template_id')} expected={tid}")

        # export docx
        _save_docx(
            client,
            f"{base_url}/api/apps/tender/projects/{pid}/export/docx?format_template_id={tid}",
            auth_token,
            out_path,
        )

    # validate docx
    doc = Document(out_path)
    # pick first level-1 title from directory nodes
    first_h1_title = None
    for n in nodes:
        try:
            if int(n.get("level") or 0) == 1:
                first_h1_title = str(n.get("title") or "").strip()
                if first_h1_title:
                    break
        except Exception:
            continue
    if not first_h1_title:
        _fail("Cannot find level==1 node title from directory/meta")

    matched_styles: List[str] = []
    for p in doc.paragraphs:
        if (p.text or "").strip() != first_h1_title:
            continue
        try:
            matched_styles.append(p.style.name if p.style else "")
        except Exception:
            matched_styles.append("")

    if not matched_styles:
        dump = _paragraph_style_dump(doc, limit=25)
        _fail(f"Cannot find exported paragraph matching first H1 title '{first_h1_title}'. First paragraphs={dump}")

    if expected_h1 not in matched_styles:
        dump = _paragraph_style_dump(doc, limit=25)
        styles = _doc_styles_top(doc, n=30)
        _fail(
            "Heading1 style mismatch.\n"
            f"  template_id={tid}\n"
            f"  project_id={pid}\n"
            f"  expected spec.style_hints.heading1_style={expected_h1!r}\n"
            f"  matched paragraph styles for title={first_h1_title!r}: {matched_styles}\n"
            f"  first_paragraphs(text, style)={dump}\n"
            f"  out.docx styles (top30)={styles}"
        )

    _pass(f"project_id={pid} template_id={tid} heading1_style OK ({expected_h1}) out={out_path}")


if __name__ == "__main__":
    try:
        main()
    except httpx.HTTPError as e:
        _fail(f"HTTP error: {type(e).__name__}: {e}")

